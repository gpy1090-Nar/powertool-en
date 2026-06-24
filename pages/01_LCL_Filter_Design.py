import streamlit as st
import db_utils
from streamlit_option_menu import option_menu
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import numpy as np
from scipy import signal
import matplotlib.pyplot as plt
import textwrap
import pandas as pd

from utils.lcl_algo import calculate_lcl_parameters, get_simulation_waveform
from utils.inverter_algo import calculate_max_ripple, get_voltage_limits, calculate_modulation_info, generate_modulation_waveforms
from utils.vector_algo import calculate_vector_analysis

# ==========================================
# 0. Page config - must be first Streamlit command
# ==========================================
st.set_page_config(page_title="LCL Filter & Inverter Design", page_icon="⚡", layout="wide")

# ======== Authentication & Trial Lock ========
auth_status = st.session_state.get("authentication_status")

if not auth_status:
    st.warning("🔒 You are not logged in or your account is not activated.")
    st.error("👉 Please click 'Home' in the sidebar to log in or start a free trial.")
    st.stop()

is_trial = (auth_status == "trial")

if not is_trial:
    current_username = st.session_state.get("username")
    local_token = st.session_state.get("session_token")
    try:
        db_token = db_utils.get_session_token(current_username)
        if local_token and db_token and local_token != db_token:
            st.session_state["authentication_status"] = False
            st.session_state["session_token"] = None
            st.error("🚨 Warning: Your account has been logged in on another device!")
            st.stop()
    except Exception:
        st.error("System busy, please try again later.")
        st.stop()
# ======== Authentication & Trial Lock End ========


def trial_lock(label="This feature requires a paid account", compact=False):
    """Trial mode lock component - shows a locked card instead of actual results."""
    pad = "14px 16px" if compact else "24px 18px"
    icon_size = "20px" if compact else "26px"
    st.markdown(
        f"""
        <div style="
            background: repeating-linear-gradient(
                45deg, #f5f5f5, #f5f5f5 10px, #ececec 10px, #ececec 20px);
            border: 1.5px dashed #faad14;
            border-radius: 10px;
            padding: {pad};
            text-align: center;
            margin: 6px 0 14px 0;
        ">
            <div style="font-size:{icon_size}; margin-bottom:6px;">🔒</div>
            <div style="font-weight:bold; font-size:14px; color:#333; margin-bottom:4px;">{label}</div>
            <div style="color:#888; font-size:12px; margin-bottom:12px;">
                Purchase to unlock all features — one-time payment, lifetime access
            </div>
            <a href="https://www.creem.io/payment/prod_3BMbIvss0TmmfCVLwUX46Z"
               target="_blank"
               style="display:inline-block; background:#a855f7; color:white;
                      padding:8px 24px; border-radius:8px; text-decoration:none;
                      font-size:13px; font-weight:bold; margin-top:8px;">
                🛒 Unlock Now — $39
            </a>
        </div>
        """,
        unsafe_allow_html=True
    )


if is_trial:
    st.info("🎁 You are in Free Trial mode. Some features are locked. [Purchase to unlock all features →](mailto:gpy1090@gmail.com)")
else:
    st.success(f"⚡ Welcome back, {st.session_state.get('name')}!")

# ==========================================
# Global state initialization
# ==========================================
default_params = {
    'sys_ug': 1140.0, 'sys_fg': 50.0, 'sys_fsw': 2500.0,
    'lcl_l1': 0.500,  'lcl_c': 100.0, 'lcl_l2': 0.250,
    'vec_ug': 690.0,  'vec_udc': 1100.0, 'vec_i': 500.0, 'vec_ang': 0.0, 'vec_l': 0.5
}
for key, val in default_params.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ==========================================
# 1. Sidebar navigation
# ==========================================
with st.sidebar:
    st.markdown("## ⚡ LCL Design Platform")
    st.caption("Ver 1.0 | PowerTool")
    
    # --- 1. Define nav options list ---
    nav_options = [
        "1. Theory Fundamentals",
        "2. LCL Parameter Design", 
        "3. Resonance Frequency Analysis",
        "4. Four-Quadrant Operation",
        "5. Modulation Index & Voltage Utilization",
        "6. Ripple Current Analysis",
        "7. Harmonic Standard Verification",
        "8. Passive Harmonic Filter Design",
        "9. Magnetic Component Design",
        "10. Filter Capacitor Design"
    ]
    
    # --- 2. Create styled menu ---
    selection = option_menu(
        menu_title="Design Workflow",
        options=nav_options, 
        icons=[
            "book", "sliders", "activity", "compass", 
            "bar-chart-steps", "water", "list-check", 
            "pie-chart", "lightning", "bricks"
        ],
        menu_icon="cast",
        default_index=0,
        styles={
            "container": {"padding": "0!important", "background-color": "#fafafa"},
            "icon": {"color": "#333", "font-size": "16px"}, 
            "nav-link": {
                "font-size": "15px", 
                "text-align": "left", 
                "margin": "0px", 
                "--hover-color": "#eee",
                "padding-top": "10px",    # increase vertical spacing
                "padding-bottom": "10px", 
            },
            "nav-link-selected": {"background-color": "#007bff", "font-weight": "normal"},
        }
    )

# Dynamically display current title
st.markdown(f"### {selection}")
st.divider()






# ==========================================
# Section 1: Theory Fundamentals
# ==========================================

if selection == nav_options[0]:
    st.header("📖 LCL Filter Theory Fundamentals")
    
    # --- 0. Helper: L vs LCL Bode plot comparison ---
    def plot_bode_comparison():
        L_total = 2.0e-3
        L_only = L_total
        L1 = 1.0e-3
        L2 = 1.0e-3
        Cf = 20.0e-6
        
        w = np.logspace(1, 5, 1000) * 2 * np.pi
        
        mag_L = -20 * np.log10(w * L_only)
        
        sys_lcl = signal.TransferFunction([1], [L1*L2*Cf, 0, L1+L2, 0])
        w_scipy, mag_lcl, phase_lcl = signal.bode(sys_lcl, w)
        
        fig = go.Figure()
        
        fig.add_trace(go.Scatter(
            x=w/(2*np.pi), y=mag_L, 
            name='L Filter (-20dB/dec)',
            line=dict(color='gray', width=2, dash='dash')
        ))
        
        fig.add_trace(go.Scatter(
            x=w_scipy/(2*np.pi), y=mag_lcl, 
            name='LCL Filter (-60dB/dec)',
            line=dict(color='#0d6efd', width=3)
        ))
        
        f_res = 1/(2*np.pi) * np.sqrt((L1+L2)/(L1*L2*Cf))
        fig.add_vline(x=f_res, line_dash="dot", line_color="red", annotation_text="Resonance Point (Fres)")
        
        fig.update_layout(
            title="L vs LCL Filter Attenuation Comparison (Same Total Inductance)",
            xaxis_title="Frequency (Hz)",
            yaxis_title="Gain (dB)",
            xaxis_type="log",
            height=400,
            template="plotly_white",
            legend=dict(yanchor="top", y=0.99, xanchor="right", x=0.99)
        )
        return fig

    # --- 1. SVG topology diagram ---
    def render_lcl_schematic():
        svg_lines = [
            '<svg width="100%" height="240" xmlns="http://www.w3.org/2000/svg">',
            '  <rect width="100%" height="100%" fill="white"/>',
            '  <!-- 1. Inverter Source -->',
            '  <rect x="20" y="80" width="50" height="80" fill="#f8f9fa" stroke="#333" stroke-width="2"/>',
            '  <text x="25" y="115" font-family="sans-serif" font-size="12" fill="#333">Vinv</text>',
            '  <text x="25" y="135" font-family="sans-serif" font-size="10" fill="#666">(PWM)</text>',
            '  <line x1="70" y1="120" x2="130" y2="120" stroke="black" stroke-width="2"/>',
            '  <!-- 2. L1 Inductor -->',
            '  <path d="M130,120 Q140,100 150,120 Q160,100 170,120 Q180,100 190,120 Q200,100 210,120" fill="none" stroke="#d63384" stroke-width="2.5"/>',
            '  <text x="155" y="90" font-family="sans-serif" font-size="14" font-weight="bold" fill="#d63384">L1</text>',
            '  <text x="145" y="150" font-family="sans-serif" font-size="11" fill="#666">Inverter Side</text>',
            '  <line x1="210" y1="120" x2="310" y2="120" stroke="black" stroke-width="2"/>',
            '  <!-- 3. Capacitor Cf Branch -->',
            '  <line x1="260" y1="120" x2="260" y2="160" stroke="black" stroke-width="2"/>',
            '  <line x1="245" y1="160" x2="275" y2="160" stroke="black" stroke-width="2"/>',
            '  <line x1="245" y1="170" x2="275" y2="170" stroke="black" stroke-width="2"/>',
            '  <line x1="260" y1="170" x2="260" y2="190" stroke="black" stroke-width="2"/>',
            '  <line x1="250" y1="190" x2="270" y2="190" stroke="black" stroke-width="2"/>',
            '  <line x1="254" y1="194" x2="266" y2="194" stroke="black" stroke-width="2"/>',
            '  <text x="280" y="170" font-family="sans-serif" font-size="14" font-weight="bold" fill="#0d6efd">Cf</text>',
            '  <!-- 4. L2 Inductor -->',
            '  <path d="M310,120 Q320,100 330,120 Q340,100 350,120 Q360,100 370,120 Q380,100 390,120" fill="none" stroke="#d63384" stroke-width="2.5"/>',
            '  <text x="335" y="90" font-family="sans-serif" font-size="14" font-weight="bold" fill="#d63384">L2</text>',
            '  <text x="330" y="150" font-family="sans-serif" font-size="11" fill="#666">Grid Side</text>',
            '  <line x1="390" y1="120" x2="450" y2="120" stroke="black" stroke-width="2"/>',
            '  <!-- 5. Grid Source -->',
            '  <circle cx="480" cy="120" r="30" fill="none" stroke="#198754" stroke-width="2"/>',
            '  <path d="M465,120 Q472,105 480,120 T495,120" fill="none" stroke="#198754" stroke-width="1.5"/>',
            '  <text x="465" y="170" font-family="sans-serif" font-size="12" font-weight="bold" fill="#198754">Grid (Ug)</text>',
            '  <text x="90" y="110" font-family="sans-serif" font-size="12" fill="red">i1 →</text>',
            '  <text x="410" y="110" font-family="sans-serif" font-size="12" fill="green">i2 →</text>',
            '</svg>'
        ]
        st.markdown("".join(svg_lines), unsafe_allow_html=True)

    # --- Page layout ---
    with st.container():
        c1, c2 = st.columns([3, 2])
        with c1:
            st.markdown("##### 📐 LCL Filter Single-Phase Topology")
            render_lcl_schematic()
        with c2:
            st.info("📌 **Core Component Definitions**")
            st.markdown("""
            *   **$L_1$ (Inverter-Side Inductor)**: The primary buffer. Limits $di/dt$ during IGBT switching and suppresses high-frequency ripple.
            *   **$C_f$ (Filter Capacitor)**: The **bypass path**. Provides a low-impedance route for high-frequency harmonics to prevent them from reaching the grid.
            *   **$L_2$ (Grid-Side Inductor)**: Secondary filtering stage; also limits fault current during grid short-circuit events.
            """)

    st.divider()

    # --- Deep-dive teaching tabs ---
    tab_why, tab_vs, tab_math = st.tabs(["1. Physical Fundamentals: Why Inductors Are Essential", "2. Deep Comparison: L vs LCL", "3. Mathematical Model & Transfer Function"])
    
    # === Tab 1: Physical Fundamentals ===
    with tab_why:
        st.subheader("💡 The Core Problem: Two Voltage Sources in Direct Conflict")
        
        col_phy1, col_phy2 = st.columns([1, 1])
        with col_phy1:
            st.markdown(r"""
            **1. Voltage Source Properties**
            *   **Grid ($V_g$)**: An extremely stiff voltage source with near-zero impedance. Its voltage waveform is a pure sinusoid.
            *   **Converter ($V_{inv}$)**: The DC bus capacitor acts as a voltage source. Through IGBT switching, it outputs a **PWM square wave** with amplitude $\\pm V_{dc}$.

            **2. Why Can't They Be Connected Directly?**
            By Ohm's Law: $I = \Delta V / Z$. If two ideal voltage sources of different values are connected directly ($Z \approx 0$):
            *   **Voltage difference $\Delta V$**: Since $V_{inv}$ is a square wave, the instantaneous voltage difference is enormous.
            *   **Current $I$**: $I \to \infty$. This causes instantaneous current surge that **destroys the IGBTs**.
            
            **3. The Role of the Inductor: Elastic Buffer**
            The inductor obeys **$V_L = L \cdot \frac{di}{dt}$**.
            Like a spring or hydraulic damper, it absorbs the instantaneous voltage difference between $V_{inv}$ and $V_g$, limiting $di/dt$ and protecting the switching devices.
            """)
        
        with col_phy2:
            t = np.linspace(0, 0.02, 1000)
            vg = 311 * np.sin(2*np.pi*50*t)
            carrier = 350 * signal.sawtooth(2*np.pi*1000*t)
            vinv = np.where(vg > carrier, 350, -350)
            
            fig_conflict = go.Figure()
            fig_conflict.add_trace(go.Scatter(
                x=t*1000, y=vinv,
                name='Converter Output (PWM)',
                line=dict(color='#dc3545', width=1),
                fill='tozeroy',
                fillcolor='rgba(220, 53, 69, 0.1)'
            ))
            fig_conflict.add_trace(go.Scatter(
                x=t*1000, y=vg,
                name='Grid Voltage (Sin)',
                line=dict(color='#198754', width=3)
            ))
            
            fig_conflict.update_layout(
                title="Voltage Difference Between Square Wave and Sine Wave", 
                xaxis_title="Time (ms)",
                yaxis_title="Voltage (V)",
                height=300,
                margin=dict(l=20, r=20, t=40, b=20),
                showlegend=True
            )
            st.plotly_chart(fig_conflict, use_container_width=True)
            st.caption("The voltage jump in the red zone must be absorbed by inductor L — without it, current becomes uncontrollable.")

    # === Tab 2: L vs LCL Deep Comparison ===
    with tab_vs:
        st.subheader("⚔️ L vs LCL: The Definitive Comparison")
        
        c_vs1, c_vs2 = st.columns([1, 1])
        
        with c_vs1:
            st.markdown("#### 🔴 L Filter (First-Order System)")
            st.markdown("""
            *   **System Order**: 1st order
            *   **High-Frequency Attenuation**: **-20dB/dec**.
                *   *Meaning*: Each decade increase in frequency reduces ripple amplitude by 10×.
            *   **Engineering Problem**:
                To meet IEEE 519 harmonic limits in MW-scale, low switching frequency applications, the required inductance is **enormous**.
                > **Consequence**: Inductor the size of a refrigerator, very high copper loss, and extremely slow dynamic response.
            """)
            
            st.markdown("#### 🔵 LCL Filter (Third-Order Oscillatory System)")
            st.markdown("""
            *   **System Order**: 3rd order
            *   **High-Frequency Attenuation**: **-60dB/dec**.
                *   *Meaning*: Each frequency decade reduces ripple amplitude by an impressive factor of **1000×**.
            *   **Engineering Advantage**:
                Leverages the low-impedance bypass of $C_f$ for outsized filtering performance. Total inductance is only **1/3 to 1/4** of an equivalent L filter.
            *   **Critical Weakness**:
                **Resonance**. At $f_{res}$, system gain spikes dramatically (see the red peak in the right chart).
                > **Analogy**: LCL is like a precision scalpel — extraordinarily effective at eliminating harmonics, but without damping control it can destabilize the system (oscillation).
            """)

        with c_vs2:
            st.plotly_chart(plot_bode_comparison(), use_container_width=True)
            st.caption("Note: At the same total inductance, LCL provides far superior high-frequency harmonic attenuation than L filter (e.g., at switching frequency 2.5kHz).")

    # === Tab 3:  ===
    with tab_math:
        st.subheader("📐 Mathematical Model & Transfer Function Derivation")
        st.markdown("To design control parameters rigorously, we need to establish the transfer function from **Converter Voltage $V_{inv}$** to **grid-injected current $I_2$**.transfer function.")
        
        st.markdown("##### 1. L Filter Model")
        st.latex(r"G_L(s) = \frac{i(s)}{v_{inv}(s)} = \frac{1}{sL + R} \approx \frac{1}{sL}")
        st.caption("Characteristics: single-pole system, phase lagging 90°.")
        
        st.divider()
        
        st.markdown("##### 2. LCL Filter Model (ignoring parasitic resistance)")
        st.markdown("Derivation using Kirchhoff's Voltage and Current Laws (KVL/KCL):")
        st.latex(r"""
        \begin{cases}
        v_{inv} - v_{cf} = L_1 \frac{di_1}{dt} \\
        v_{cf} - v_g = L_2 \frac{di_2}{dt} \\
        i_1 - i_2 = C_f \frac{dv_{cf}}{dt}
        \end{cases}
        \quad \xrightarrow{\mathscr{L}\text{-transform}} \quad
        G_{LCL}(s) = \frac{i_2(s)}{v_{inv}(s)} = \frac{1}{s^3 L_1 L_2 C_f + s(L_1 + L_2)}
        """)
        
        col_math1, col_math2 = st.columns(2)
        with col_math1:
            st.markdown("**🔍 Low-Frequency Behavior ($s \to 0$)**")
            st.latex(r"G_{LCL}(s) \approx \frac{1}{s(L_1 + L_2)}")
            st.markdown("At low frequencies, the LCL filter behaves as a simple inductor with total inductance $L_1+L_2$.")
            
        with col_math2:
            st.markdown("**🔍 Resonance Frequency ($Denominator = 0$)**")
            st.latex(r"\omega_{res} = \sqrt{\frac{L_1 + L_2}{L_1 L_2 C_f}} \quad \Rightarrow \quad f_{res} = \frac{\omega_{res}}{2\pi}")
            st.markdown("This is the most unstable operating point. Ensure $f_{res}$ is kept away from multiples of $f_{sw}$.")




# ==========================================
#  2: LCL Parameter Design ( SOA )
# ==========================================
elif selection == nav_options[1]:
    st.header("🛠️ LCL Parameter Design & SOA Feasibility Optimization")
    st.caption("🔥 Core feature: Provides white-box mathematical derivations and maps 5 non-linear constraints onto a 2D plane — exposing the absolute safe parameter zone at a glance.")

    # =========================================================
    # ( + )
    # =========================================================
    import io
    import datetime
    import subprocess as _subprocess
    import importlib as _importlib
    import sys as _sys_top
    import os as _os_top


    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib import font_manager as _fm

    # ── matplotlib ()──
    def _find_cn_font():
        if _sys_top.platform.startswith('win'):
            fd = _os_top.path.join(_os_top.environ.get('WINDIR','C:\\Windows'), 'Fonts')
            for fname in ['msyh.ttc','msyhbd.ttc','simhei.ttf','simsun.ttc']:
                fp = _os_top.path.join(fd, fname)
                if _os_top.path.isfile(fp):
                    try: return _fm.FontProperties(fname=fp)
                    except: pass
        elif _sys_top.platform.startswith('darwin'):
            for fp in ['/System/Library/Fonts/PingFang.ttc']:
                if _os_top.path.isfile(fp):
                    try: return _fm.FontProperties(fname=fp)
                    except: pass
        else:
            for fp in ['/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
                       '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc']:
                if _os_top.path.isfile(fp):
                    try: return _fm.FontProperties(fname=fp)
                    except: pass
        return _fm.FontProperties()

    _CN_FONT = _find_cn_font()

    # ── (LaTeX → PNG bytes)──
    def _formula_png(latex_str, fontsize=13, color='black', dpi=150, indent=True):
        """Renders LaTeX string to a transparent-background PNG suitable for embedding in Word/PDF."""
        fig = plt.figure(figsize=(0.01, 0.01))
        fig.patch.set_alpha(0)
        renderer = fig.canvas.get_renderer()
        t = fig.text(0, 0, f'${latex_str}$', fontsize=fontsize, color=color)
        bbox = t.get_window_extent(renderer=renderer)
        w_inch = (bbox.width + 30) / dpi
        h_inch = (bbox.height + 14) / dpi
        plt.close(fig)
        fig = plt.figure(figsize=(max(w_inch, 1.5), max(h_inch, 0.38)))
        fig.patch.set_alpha(0)
        x0 = 0.06 if indent else 0.02
        fig.text(x0, 0.18, f'${latex_str}$', fontsize=fontsize, color=color, va='bottom')
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                    transparent=True, pad_inches=0.03)
        plt.close(fig)
        buf.seek(0)
        return buf.read()

    # ── Word  ──
    def _gen_word_report(params: dict, soa_png: bytes, bode_png: bytes, fig_soa=None) -> bytes:
        import sys as _sw, os as _ow
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # 
        def _cn():
            if _sw.platform.startswith('win'):
                fd = _ow.path.join(_ow.environ.get('WINDIR','C:\\Windows'), 'Fonts')
                if _ow.path.isfile(_ow.path.join(fd,'msyh.ttc')): return 'Microsoft YaHei'
                if _ow.path.isfile(_ow.path.join(fd,'simhei.ttf')): return 'SimHei'
                return 'SimSun'
            elif _sw.platform.startswith('darwin'): return 'PingFang SC'
            return 'Noto Sans CJK SC'

        CNF = _cn()

        def _inject(doc):
            """Injects font settings at docDefaults and Normal style levels to prevent encoding issues."""
            normal = doc.styles['Normal']
            rPr_n = normal.element.get_or_add_rPr()
            for x in rPr_n.findall(qn('w:rFonts')): rPr_n.remove(x)
            rf = OxmlElement('w:rFonts')
            for a,v in [('w:ascii','Arial'),('w:hAnsi','Arial'),('w:eastAsia',CNF),('w:cs',CNF)]:
                rf.set(qn(a), v)
            rPr_n.insert(0, rf); normal.font.size = Pt(11)
            dd = doc.styles.element.find(qn('w:docDefaults'))
            if dd is None: return
            rpd = dd.find('.//' + qn('w:rPrDefault'))
            if rpd is None: return
            rPr_d = rpd.find(qn('w:rPr'))
            if rPr_d is None: rPr_d = OxmlElement('w:rPr'); rpd.append(rPr_d)
            for x in rPr_d.findall(qn('w:rFonts')): rPr_d.remove(x)
            rf2 = OxmlElement('w:rFonts')
            for a,v in [('w:ascii','Arial'),('w:hAnsi','Arial'),('w:eastAsia',CNF),('w:cs',CNF)]:
                rf2.set(qn(a), v)
            rPr_d.insert(0, rf2)

        def _fr(run, size=11, bold=False, color=None):
            run.font.size = Pt(size); run.font.bold = bold
            if color: run.font.color.rgb = RGBColor(*color)
            rPr = run._r.get_or_add_rPr()
            for x in rPr.findall(qn('w:rFonts')): rPr.remove(x)
            rf = OxmlElement('w:rFonts')
            for a,v in [('w:ascii','Arial'),('w:hAnsi','Arial'),('w:eastAsia',CNF),('w:cs',CNF)]:
                rf.set(qn(a), v)
            rPr.insert(0, rf)

        def _pp(para):
            pPr = para._p.get_or_add_pPr()
            rPr = pPr.find(qn('w:rPr'))
            if rPr is None: rPr = OxmlElement('w:rPr'); pPr.append(rPr)
            for x in rPr.findall(qn('w:rFonts')): rPr.remove(x)
            rf = OxmlElement('w:rFonts')
            for a,v in [('w:ascii','Arial'),('w:hAnsi','Arial'),('w:eastAsia',CNF),('w:cs',CNF)]:
                rf.set(qn(a), v)
            rPr.insert(0, rf)

        def H(doc, text, level=1):
            para = doc.add_heading('', level=level); para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _pp(para)
            clr = (31,73,125) if level==1 else (46,116,181)
            sz  = 16 if level==1 else 13
            run = para.add_run(text); _fr(run, size=sz, bold=True, color=clr)
            return para

        def P(doc, text, size=11, bold=False, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, indent_cm=0):
            para = doc.add_paragraph(); para.alignment = align; _pp(para)
            if indent_cm: para.paragraph_format.left_indent = Cm(indent_cm)
            run = para.add_run(text); _fr(run, size=size, bold=bold, color=color)
            return para

        def F(doc, latex, color='black', fontsize=13, width_inches=5.6):
            """Embeds a single-line LaTeX formula image, left-aligned."""
            png = _formula_png(latex, fontsize=fontsize, color=color)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run()
            run.add_picture(io.BytesIO(png), width=Inches(min(width_inches, 5.6)))
            return para

        def shade(cell, fill):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            shd.set(qn('w:fill'), fill); tcPr.append(shd)

        def KV(doc, rows, cw=(7,9)):
            t = doc.add_table(rows=len(rows), cols=2)
            t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
            for i,(k,v) in enumerate(rows):
                for j,txt in enumerate([k,v]):
                    cell = t.rows[i].cells[j]; cell.width = Cm(cw[j])
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    para = cell.paragraphs[0]; _pp(para)
                    run = para.add_run(str(txt)); _fr(run, size=10, bold=(i==0))
                    if i==0: shade(cell, '1F4973')
                    elif i%2==0: shade(cell, 'EBF3FB')
            for j in range(2):
                _fr(t.rows[0].cells[j].paragraphs[0].runs[0],
                    size=10, bold=True, color=(255,255,255))
            doc.add_paragraph(); return t

        def T3(doc, rows, cw=(5,4.5,6.5)):
            t = doc.add_table(rows=len(rows), cols=3)
            t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
            for i,row_d in enumerate(rows):
                for j,txt in enumerate(row_d):
                    cell = t.rows[i].cells[j]; cell.width = Cm(cw[j])
                    para = cell.paragraphs[0]; _pp(para)
                    run = para.add_run(str(txt)); _fr(run, size=10, bold=(i==0))
                    if i==0: shade(cell, '1F4973')
                    elif i%2==0: shade(cell, 'EBF3FB')
            for j in range(3):
                _fr(t.rows[0].cells[j].paragraphs[0].runs[0],
                    size=10, bold=True, color=(255,255,255))
            doc.add_paragraph(); return t

        # ════════════════ SOA PNG: plotly ════════════════
        if fig_soa is not None:
            try:
                soa_png = fig_soa.to_image(format="png", width=950, height=580, scale=1.5)
            except Exception:
                pass  # fallback to passed matplotlib soa_png
        # ════════════════  ════════════════
        doc = Document()
        _inject(doc)
        sec = doc.sections[0]
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.left_margin=sec.right_margin=Cm(2.5)
        sec.top_margin=sec.bottom_margin=Cm(2.5)
        prm = params

        # ──  ──
        doc.add_paragraph()
        tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER; _pp(tp)
        _fr(tp.add_run('LCL Filter Design Report'), size=26, bold=True, color=(31,73,125))
        sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER; _pp(sp)
        _fr(sp.add_run('PowerTool  Power Electronics Design Tool'), size=13, color=(89,89,89))
        dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER; _pp(dp)
        _fr(dp.add_run(f'Generated: {datetime.datetime.now().strftime("%Y-%m-%d  %H:%M")}'),
            size=11, color=(89,89,89))
        doc.add_page_break()

        # ── 1 ──
        H(doc, '1. Operating Conditions & Boundary Constraints')
        P(doc, 'The following are the system operating conditions and design limit specifications used in this design.')
        doc.add_paragraph()
        H(doc, '1.1  System Operating Parameters', level=2)
        KV(doc, [
            ('Parameter',             'Value'),
            ('Grid Line Voltage Vg_line',  f'{prm["vg_line"]:.1f} Vrms'),
            ('Grid Phase Voltage Vg_phase', f'{prm["vg_phase"]:.2f} Vrms'),
            ('Grid Frequency fg',        f'{prm["fg"]:.1f} Hz'),
            ('Rated Phase Current Irated',  f'{prm["i_rated"]:.1f} Arms'),
            ('DC Bus Voltage Vdc',   f'{prm["vdc"]:.1f} V'),
            ('Switching Frequency fsw',       f'{prm["fsw"]:.0f} Hz'),
            ('Topology',           prm['topo']),
            ('Modulation Strategy',           prm['mod']),
            ('Minimum Power Factor',       f'{prm["pf"]:.2f}  (Lagging)'),
            ('Damping Method',           prm['damp_mode']),
        ])
        H(doc, '1.2  Design Limit Specifications', level=2)
        KV(doc, [
            ('Specification',            'Limit Value'),
 ('Maximum Allowable Ripple Current', f'{prm["ripple_pct"]:.1f}% -> ΔIpp = {prm["delta_i_pp"]:.2f} A (peak-peak)'),
            ('Max Capacitor Reactive Power Ratio', f'{prm["reactive_pct"]:.1f}%'),
            ('Resonance Safety Island',    f'{prm["f_safe_low"]:.0f} Hz  ~  {prm["f_safe_high"]:.0f} Hz'),
        ])
        doc.add_page_break()

        # ── 2 ──
        H(doc, '2. Rigorous Physical Limit Derivation (with numerical substitution)')

        H(doc, '2.1  Constraint 1: Total Inductance Upper Limit (Four-Quadrant Vector Limit)', level=2)
        P(doc, 'The converter output phase voltage vector must satisfy:')
        F(doc, r'\vec{V}_{inv} = \vec{V}_g + j\omega_g L_{total} \vec{I}')
        P(doc, 'At current Lagging Power Factor Angle phi, using the law of cosines:')
        F(doc, r'V_L^2 + 2\,V_g\sin\varphi\cdot V_L + (V_g^2 - V_{max}^2) \leq 0')
        P(doc, 'Substituting current operating values:')

        if prm['mod'] == 'SPWM':
            F(doc, rf'V_{{max\_out}} = \frac{{V_{{dc}}}}{{2\sqrt{{2}}}} = \frac{{{prm["vdc"]:.1f}}}{{2.828}} = {prm["v_max_out"]:.2f}\ \mathrm{{V}}', color='darkred')
        else:
            F(doc, rf'V_{{max\_out}} = \frac{{V_{{dc}}}}{{\sqrt{{6}}}} = \frac{{{prm["vdc"]:.1f}}}{{2.449}} = {prm["v_max_out"]:.2f}\ \mathrm{{V}}', color='darkred')

        F(doc, rf'\varphi = \arccos({prm["pf"]}) = {prm["phi_deg"]:.2f}^\circ \quad \sin\varphi = {prm["sin_phi"]:.4f}')
        F(doc, rf'V_{{g,ph}} = \frac{{{prm["vg_line"]:.1f}}}{{\sqrt{{3}}}} = {prm["vg_phase"]:.2f}\ \mathrm{{V}}')

        if prm['l_max_mh'] > 0:
            F(doc, rf'V_{{L\_max}} = {prm["v_l_max"]:.2f}\ \mathrm{{V}}', color='darkred')
            F(doc, rf'L_{{total\_max}} = \frac{{V_{{L\_max}}}}{{\omega_g \cdot I_{{rated}}}} = \frac{{{prm["v_l_max"]:.2f}}}{{{prm["omega_g"]:.2f} \times {prm["i_rated"]:.1f}}}',
              color='darkred')
            F(doc, r'\mathbf{\Rightarrow\ L_{total\_max} = ' + f'{prm["l_max_mh"]:.4f}' + r'\ mH}', color='darkred', fontsize=14)
        else:
            P(doc, '⚠️  Insufficient DC Bus Voltage — Grid phase voltage exceeds the converter\'s maximum output. Grid connection is not possible!',
              color=(200,0,0), bold=True)
        doc.add_paragraph()

        H(doc, '2.2  Constraint 2: Inverter-Side Inductance L1 Lower Limit (Ripple Limit)', level=2)
        F(doc, r'L_{1} \geq \frac{V_{dc}}{K \cdot f_{sw} \cdot \Delta I_{pp}}')
        P(doc, f'Topology Ripple Coefficient K = {prm["k_ripple"]}({prm["topo"]} + {prm["mod"]})')
        F(doc, rf'\Delta I_{{pp}} = {prm["ripple_pct"]}\%\times\sqrt{{2}}\times {prm["i_rated"]:.1f} = {prm["delta_i_pp"]:.2f}\ \mathrm{{A}}')
        F(doc, rf'L_{{1\_min}} = \frac{{{prm["vdc"]:.1f}}}{{{prm["k_ripple"]} \times {prm["fsw"]:.0f} \times {prm["delta_i_pp"]:.2f}}}',
          color='darkblue')
        _v = prm["l1_min_mh"]
        F(doc, r'\mathbf{\Rightarrow\ L_{1\_min} = ' + f'{_v:.4f}' + r'\ mH}', color='darkblue', fontsize=14)
        doc.add_paragraph()

        H(doc, '2.3  Constraint 3: Filter Capacitor Cf Upper Limit (Reactive Power Limit)', level=2)
        F(doc, r'C_f \leq \frac{\lambda\% \cdot I_{rated}}{\omega_g \cdot V_{g,ph}}')
        F(doc, rf'C_{{f\_max}} = \frac{{{prm["reactive_pct"]}\%\times {prm["i_rated"]:.1f}}}{{{prm["omega_g"]:.2f}\times {prm["vg_phase"]:.2f}}}',
          color='purple')
        _v = prm["c_max_uf"]
        F(doc, r'\mathbf{\Rightarrow\ C_{f\_max} = ' + f'{_v:.2f}' + r'\ \mu F}', color='purple', fontsize=14)
        doc.add_page_break()

        # ── Chapter 3 ──
        H(doc, '3. External Grid-Side Leakage Inductance Derivation')
        P(doc, f'Calculation Mode: {prm["grid_mode"]}')
        doc.add_paragraph()

        if 'transformer' in prm['grid_mode']:
            H(doc, '3.1  Based on Transformer Nameplate Parameters (independent of converter rating)', level=2)
            P(doc, f'Transformer rated capacity Stx = {prm["tx_kva"]:.0f} kVA, short-circuit impedance uk = {prm["tx_uk"]:.1f}%')
            F(doc, rf'Z_{{base,tx}} = \frac{{U_{{line}}^2}}{{S_{{tx}}}} = \frac{{{prm["vg_line"]:.1f}^2}}{{{prm["tx_kva"]*1000:.0f}}} = {prm["z_base_tx"]:.4f}\ \Omega')
            F(doc, rf'Z_T = u_k\%\times Z_{{base,tx}} = {prm["tx_uk"]/100:.3f}\times {prm["z_base_tx"]:.4f} = {prm["z_tx"]:.4f}\ \Omega')
            F(doc, rf'L_g = \frac{{Z_T}}{{\omega_g}} = \frac{{{prm["z_tx"]:.4f}}}{{{prm["omega_g"]:.2f}}}', color='darkgreen')
            F(doc, r'\mathbf{\Rightarrow\ L_g = ' + f'{prm["lg_mh"]:.4f}' + r'\ mH,\ SCR_{eq}\approx ' + f'{prm["eq_scr"]:.1f}' + r'}', color='darkgreen', fontsize=14)
        else:
            H(doc, '3.1  Based on Grid Short Circuit Ratio (SCR)', level=2)
            F(doc, rf'S_{{inv}} = \sqrt{{3}}\times {prm["vg_line"]:.1f}\times {prm["i_rated"]:.1f} = {prm["s_inv_kva"]:.1f}\ \mathrm{{kVA}}')
            F(doc, rf'Z_{{base}} = \frac{{U_{{line}}^2}}{{S_{{inv}}}} = \frac{{{prm["vg_line"]:.1f}^2}}{{{prm["s_inv_kva"]*1000:.0f}}} = {prm["z_base"]:.4f}\ \Omega')
            F(doc, rf'Z_g = \frac{{Z_{{base}}}}{{SCR}} = \frac{{{prm["z_base"]:.4f}}}{{{prm["scr"]:.1f}}} = {prm["z_g"]:.4f}\ \Omega')
            F(doc, rf'L_g = \frac{{Z_g}}{{2\pi f_g}} = \frac{{{prm["z_g"]:.4f}}}{{{prm["omega_g"]:.2f}}}', color='darkgreen')
            F(doc, r'\mathbf{\Rightarrow\ L_g = ' + f'{prm["lg_mh"]:.4f}' + r'\ mH}', color='darkgreen', fontsize=14)
        doc.add_page_break()

        # ── Chapter 4 ──
        H(doc, '4. Finalized Parameter Summary')
        T3(doc, [
            ('Parameter', 'Locked Value', 'Description'),
            ('L1  Inverter-Side Inductor',     f'{prm["final_l1"]:.3f} mH',                     'Locked from SOA map, satisfies ripple constraint'),
            ('L2  Internal Grid-Side Inductor',
             f'{prm["final_l2"]:.3f} mH' if prm['has_l2'] else '0  (CL Topology)',
             'L2 = L1 × k' if prm['has_l2'] else 'Uses external leakage inductance as L2'),
            ('Lg  External Leakage Inductance',       f'{prm["lg_mh"]:.4f} mH',                        'Transformer leakage or grid impedance equivalent'),
            ('L_total  Total System Inductance',f'{prm["l_total_mh"]:.3f} mH',                   'L1 + L2 + Lg'),
            ('Cf  Filter Capacitor (Y equivalent)',f'{prm["final_c"]:.1f} uF',                      'Delta config physical value = Y value / 3'),
            ('Rd  Damping Resistor',
             f'{prm["final_rd"]:.3f} Ohm' if prm['has_damping'] else '-- (no damping)',
             '1/(3 x omega_res x Cf)'),
            ('f_res  Resonance Frequency',    f'{prm["f_res"]:.1f} Hz',
             f'Safe Island {prm["f_safe_low"]:.0f}~{prm["f_safe_high"]:.0f} Hz'),
            ('Compliance Status',
             '[OK] All Passed' if prm['all_pass'] else '[!!] Constraint Violated',
             'See SOA map for details'),
        ])
        doc.add_page_break()

        # ── Chapter 5: SOA Map ──
        H(doc, '5. SOA Parameter Feasibility Map')
        P(doc, f'Green region represents the safe parameter space satisfying all five constraints. The star marks the current design point.'
               f'X-axis = L_total (Lg = {prm["lg_mh"]:.4f} mH included), Y-axis = Cf.')
        doc.add_paragraph()
        doc.add_picture(io.BytesIO(soa_png), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        P(doc, 'Figure 5-1  SOA Five-Constraint Feasibility Map', size=9, color=(100,100,100),
          align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_page_break()

        # ── Chapter 6: Bode Plot ──
        H(doc, '6. Resonance Frequency Bode Plot')
        P(doc, 'Gray dashed: ideal grid (Lg=0); Red solid: actual response with leakage inductance; Green shading: safe island frequency band.')
        doc.add_paragraph()
        doc.add_picture(io.BytesIO(bode_png), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        P(doc, 'Figure 6-1  LCL Filter Bode Plot (with external leakage inductance)', size=9, color=(100,100,100),
          align=WD_ALIGN_PARAGRAPH.CENTER)

        # Footer
        footer = doc.sections[0].footer
        fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; _pp(fp)
        _fr(fp.add_run('This report was automatically generated by PowerTool — For engineering reference only'),
            size=9, color=(128,128,128))

        out = io.BytesIO(); doc.save(out); out.seek(0)
        return out.read()

    # ── PDF  ──
    def _gen_pdf_report(params: dict, soa_png: bytes, bode_png: bytes, fig_soa=None) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm as rl_cm
        from reportlab.lib import colors
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Table, TableStyle, Image as RLImage,
                                        PageBreak, HRFlowable)
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
        CN = 'STSong-Light'
        MG = 2.2 * rl_cm
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=MG, rightMargin=MG,
                                topMargin=MG, bottomMargin=MG)
        S = lambda n, **kw: ParagraphStyle(n, fontName=CN, **kw)
        sty = {
            'title':   S('t',  fontSize=20, leading=28, alignment=TA_CENTER,
                          textColor=colors.HexColor('#1F4973'), spaceAfter=6),
            'sub':     S('s',  fontSize=12, leading=18, alignment=TA_CENTER,
                          textColor=colors.HexColor('#595959'), spaceAfter=4),
            'date':    S('d',  fontSize=10, leading=16, alignment=TA_CENTER,
                          textColor=colors.HexColor('#888888')),
            'h1':      S('h1', fontSize=14, leading=22, spaceBefore=14, spaceAfter=4,
                          textColor=colors.HexColor('#1F4973')),
            'h2':      S('h2', fontSize=12, leading=18, spaceBefore=10, spaceAfter=3,
                          textColor=colors.HexColor('#2E74B5')),
            'body':    S('b',  fontSize=10, leading=16, spaceAfter=3),
            'caption': S('c',  fontSize=9,  leading=14, alignment=TA_CENTER,
                          textColor=colors.HexColor('#666666')),
        }

        def HR(): return HRFlowable(width='100%', thickness=0.6,
                                    color=colors.HexColor('#2E74B5'),
                                    spaceAfter=6, spaceBefore=2)

        def FI(latex, color='black', fontsize=13, left_cm=1.0):
            """PDF formula image flowable."""
            png = _formula_png(latex, fontsize=fontsize, color=color)
            img = RLImage(io.BytesIO(png))
            # Maintain aspect ratio, limit max width
            max_w = (A4[0] - 2*MG - left_cm*rl_cm)
            img.drawWidth  = min(img.imageWidth / 150 * rl_cm, max_w)
            img.drawHeight = img.imageHeight / 150 * rl_cm
            return img

        def KVT(rows, cw=None):
            if cw is None: cw = [7*rl_cm, 9*rl_cm]
            ts = TableStyle([
                ('BACKGROUND',(0,0),(-1,0),  colors.HexColor('#1F4973')),
                ('TEXTCOLOR', (0,0),(-1,0),  colors.white),
                ('FONTNAME',  (0,0),(-1,-1), CN), ('FONTSIZE',(0,0),(-1,-1),9),
                ('GRID',      (0,0),(-1,-1), 0.4, colors.HexColor('#AAAAAA')),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white, colors.HexColor('#EBF3FB')]),
                ('TOPPADDING',(0,0),(-1,-1),4), ('BOTTOMPADDING',(0,0),(-1,-1),4),
                ('LEFTPADDING',(0,0),(-1,-1),6),
            ])
            td = [[Paragraph(str(c), sty['body']) for c in r] for r in rows]
            t = Table(td, colWidths=cw); t.setStyle(ts); return t

        def T3T(rows):
            cw = [5*rl_cm, 4.5*rl_cm, 6.5*rl_cm]
            ts = TableStyle([
                ('BACKGROUND',(0,0),(-1,0),  colors.HexColor('#1F4973')),
                ('TEXTCOLOR', (0,0),(-1,0),  colors.white),
                ('FONTNAME',  (0,0),(-1,-1), CN), ('FONTSIZE',(0,0),(-1,-1),9),
                ('GRID',      (0,0),(-1,-1), 0.4, colors.HexColor('#AAAAAA')),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white, colors.HexColor('#EBF3FB')]),
                ('TOPPADDING',(0,0),(-1,-1),4), ('BOTTOMPADDING',(0,0),(-1,-1),4),
                ('LEFTPADDING',(0,0),(-1,-1),6),
            ])
            td = [[Paragraph(str(c), sty['body']) for c in r] for r in rows]
            t = Table(td, colWidths=cw); t.setStyle(ts); return t

        prm = params
        # SOA PNG: prefer plotly original (matches web view exactly)
        if fig_soa is not None:
            try:
                soa_png = fig_soa.to_image(format="png", width=950, height=580, scale=1.5)
            except Exception:
                pass  # fallback to passed matplotlib soa_png
        story = []

        story += [Spacer(1,1.5*rl_cm),
                  Paragraph('LCL Filter Design Report', sty['title']),
                  Spacer(1,0.3*rl_cm),
                  Paragraph('PowerTool  Power Electronics Design Tool', sty['sub']),
                  Paragraph(f'Generated: {datetime.datetime.now().strftime("%Y-%m-%d  %H:%M")}', sty['date']),
                  PageBreak()]

        story += [Paragraph('1. Operating Conditions & Boundary Constraints', sty['h1']), HR(),
                  Paragraph('1.1  System Operating Parameters', sty['h2'])]
        story.append(KVT([
            ('Parameter','Value'),
            ('Grid Line Voltage Vg_line',  f'{prm["vg_line"]:.1f} Vrms'),
            ('Grid Phase Voltage Vg_phase', f'{prm["vg_phase"]:.2f} Vrms'),
            ('Grid Frequency fg',         f'{prm["fg"]:.1f} Hz'),
            ('Rated Phase Current Irated',   f'{prm["i_rated"]:.1f} Arms'),
            ('DC Bus Voltage Vdc',    f'{prm["vdc"]:.1f} V'),
            ('Switching Frequency fsw',        f'{prm["fsw"]:.0f} Hz'),
            ('Topology', prm['topo']), ('Modulation Strategy', prm['mod']),
            ('Minimum Power Factor', f'{prm["pf"]:.2f}  (Lagging)'),
            ('Damping Method', prm['damp_mode']),
        ]))
        story += [Spacer(1,0.4*rl_cm), Paragraph('1.2  Design Limit Specifications', sty['h2'])]
        story.append(KVT([
            ('Specification','Limit Value'),
            ('Maximum Allowable Ripple Current', f'{prm["ripple_pct"]:.1f}% -> ΔIpp = {prm["delta_i_pp"]:.2f} A'),
            ('Max Capacitor Reactive Power Ratio', f'{prm["reactive_pct"]:.1f}%'),
            ('Resonance Safety Island',    f'{prm["f_safe_low"]:.0f} ~ {prm["f_safe_high"]:.0f} Hz'),
        ]))
        story.append(PageBreak())

        # Chapter 2
        story += [Paragraph('2. Rigorous Physical Limit Derivation (with numerical substitution)', sty['h1']), HR(),
                  Paragraph('2.1  Constraint 1: Total Inductance Upper Limit (Four-Quadrant Vector Limit)', sty['h2']),
                  Paragraph('Vector relationship constraint:', sty['body']),
                  FI(r'\vec{V}_{inv} = \vec{V}_g + j\omega_g L_{total}\,\vec{I}'),
                  FI(r'V_L^2 + 2V_g\sin\varphi\cdot V_L + (V_g^2 - V_{max}^2)\leq 0'),
                  Paragraph('Substituting values:', sty['body'])]

        if prm['mod'] == 'SPWM':
            story.append(FI(rf'V_{{max\_out}} = V_{{dc}}/(2\sqrt{{2}}) = {prm["vdc"]:.1f}/2.828 = {prm["v_max_out"]:.2f}\ V', color='darkred'))
        else:
            story.append(FI(rf'V_{{max\_out}} = V_{{dc}}/\sqrt{{6}} = {prm["vdc"]:.1f}/2.449 = {prm["v_max_out"]:.2f}\ V', color='darkred'))

        story += [
            FI(rf'\varphi=\arccos({prm["pf"]})={prm["phi_deg"]:.2f}^\circ,\quad\sin\varphi={prm["sin_phi"]:.4f}'),
            FI(rf'V_{{g,ph}}={prm["vg_phase"]:.2f}\ V'),
        ]
        if prm['l_max_mh'] > 0:
            story += [
                FI(rf'V_{{L\_max}}={prm["v_l_max"]:.2f}\ V', color='darkred'),
                FI(rf'L_{{total\_max}} = \frac{{{prm["v_l_max"]:.2f}}}{{{prm["omega_g"]:.2f}\times{prm["i_rated"]:.1f}}}', color='darkred'),
                FI(r'\mathbf{\Rightarrow\ L_{total\_max} = ' + f'{prm["l_max_mh"]:.4f}' + r'\ mH}', color='darkred', fontsize=15),
            ]
        else:
            story.append(Paragraph('⚠️ Insufficient DC Bus Voltage — Grid phase voltage exceeds converter maximum output. Grid connection impossible!', sty['body']))

        story += [
            Spacer(1,0.3*rl_cm),
            Paragraph('2.2  Constraint 2: Inverter-Side Inductance L1 Lower Limit (Ripple Limit)', sty['h2']),
            FI(r'L_{1} \geq \frac{V_{dc}}{K\cdot f_{sw}\cdot\Delta I_{pp}}'),
            FI(rf'\Delta I_{{pp}}={prm["ripple_pct"]}\%\times\sqrt{{2}}\times{prm["i_rated"]:.1f}={prm["delta_i_pp"]:.2f}\ A'),
            FI(rf'L_{{1\_min}}=\frac{{{prm["vdc"]:.1f}}}{{{prm["k_ripple"]}\times{prm["fsw"]:.0f}\times{prm["delta_i_pp"]:.2f}}}', color='darkblue'),
            FI(r'\mathbf{\Rightarrow\ L_{1\_min} = ' + f'{prm["l1_min_mh"]:.4f}' + r'\ mH}', color='darkblue', fontsize=14),
            Spacer(1,0.3*rl_cm),
            Paragraph('2.3  Constraint 3: Filter Capacitor Cf Upper Limit (Reactive Power Limit)', sty['h2']),
            FI(r'C_f\leq\frac{\lambda\%\cdot I_{rated}}{\omega_g\cdot V_{g,ph}}'),
            FI(rf'C_{{f\_max}}=\frac{{{prm["reactive_pct"]}\%\times{prm["i_rated"]:.1f}}}{{{prm["omega_g"]:.2f}\times{prm["vg_phase"]:.2f}}}', color='purple'),
            FI(r'\mathbf{\Rightarrow\ C_{f\_max} = ' + f'{prm["c_max_uf"]:.2f}' + r'\ \mu F}', color='purple', fontsize=14),
            PageBreak(),
        ]

        # Chapter 3
        story += [Paragraph('3. External Grid-Side Leakage Inductance Derivation', sty['h1']), HR(),
                  Paragraph(f'Calculation Mode: {prm["grid_mode"]}', sty['body']),
                  Spacer(1,0.3*rl_cm)]
        if 'transformer' in prm['grid_mode']:
            story += [
                Paragraph('3.1  Based on Transformer Nameplate Parameters (independent of converter rating)', sty['h2']),
                FI(rf'Z_{{base,tx}}=U_{{line}}^2/S_{{tx}}={prm["vg_line"]:.1f}^2/{prm["tx_kva"]*1000:.0f}={prm["z_base_tx"]:.4f}\ \Omega'),
                FI(rf'Z_T=u_k\%\times Z_{{base,tx}}={prm["tx_uk"]/100:.3f}\times{prm["z_base_tx"]:.4f}={prm["z_tx"]:.4f}\ \Omega'),
                FI(rf'L_g=Z_T/\omega_g={prm["z_tx"]:.4f}/{prm["omega_g"]:.2f}', color='darkgreen'),
                FI(r'\mathbf{\Rightarrow\ L_g = ' + f'{prm["lg_mh"]:.4f}' + r'\ mH,\ SCR_{eq}\approx ' + f'{prm["eq_scr"]:.1f}' + r'}', color='darkgreen', fontsize=14),
            ]
        else:
            story += [
                Paragraph('3.1  Based on Grid Short Circuit Ratio (SCR)', sty['h2']),
                FI(rf'S_{{inv}}=\sqrt{{3}}\times{prm["vg_line"]:.1f}\times{prm["i_rated"]:.1f}={prm["s_inv_kva"]:.1f}\ \mathrm{{kVA}}'),
                FI(rf'Z_{{base}}=U_{{line}}^2/S_{{inv}}={prm["vg_line"]:.1f}^2/{prm["s_inv_kva"]*1000:.0f}={prm["z_base"]:.4f}\ \Omega'),
                FI(rf'Z_g=Z_{{base}}/SCR={prm["z_base"]:.4f}/{prm["scr"]:.1f}={prm["z_g"]:.4f}\ \Omega'),
                FI(rf'L_g=Z_g/(2\pi f_g)={prm["z_g"]:.4f}/{prm["omega_g"]:.2f}', color='darkgreen'),
                FI(r'\mathbf{\Rightarrow\ L_g = ' + f'{prm["lg_mh"]:.4f}' + r'\ mH}', color='darkgreen', fontsize=14),
            ]
        story.append(PageBreak())

        # 4
        story += [Paragraph('4. Finalized Parameter Summary', sty['h1']), HR()]
        story.append(T3T([
            ('Parameter','Locked Value','Description'),
            ('L1  Inverter-Side Inductor',     f'{prm["final_l1"]:.3f} mH', 'Locked from SOA map'),
            ('L2  Internal Grid-Side Inductor',
             f'{prm["final_l2"]:.3f} mH' if prm['has_l2'] else '0  (CL Topology)',
             'L2 = L1 × k' if prm['has_l2'] else 'Uses external leakage as L2'),
            ('Lg  External Leakage Inductance',       f'{prm["lg_mh"]:.4f} mH', 'Transformer / grid impedance equivalent'),
            ('L_total  Total System Inductance',f'{prm["l_total_mh"]:.3f} mH', 'L1+L2+Lg'),
            ('Cf  Filter Capacitor (Y equivalent)',f'{prm["final_c"]:.1f} uF', 'Delta config physical value = Y value / 3'),
            ('Rd  Damping Resistor',
             f'{prm["final_rd"]:.3f} Ohm' if prm['has_damping'] else '-- (no damping)',
             '1/(3 x omega_res x Cf)'),
            ('f_res  Resonance Frequency',    f'{prm["f_res"]:.1f} Hz',
             f'Island {prm["f_safe_low"]:.0f}~{prm["f_safe_high"]:.0f} Hz'),
            ('Compliance Status',
             'All Passed OK' if prm['all_pass'] else 'Constraint Violated !!', 'See SOA map'),
        ]))
        story.append(PageBreak())

        story += [
            Paragraph('5. SOA Parameter Feasibility Map', sty['h1']), HR(),
            Paragraph(f'Green region = safe parameter space. Star = current design point. X-axis = L_total (Lg={prm["lg_mh"]:.4f} mH included).', sty['body']),
            Spacer(1,0.3*rl_cm),
            RLImage(io.BytesIO(soa_png), width=15*rl_cm, height=9.5*rl_cm),
            Paragraph('Figure 5-1  SOA Five-Constraint Feasibility Map', sty['caption']),
            PageBreak(),
            Paragraph('6. Resonance Frequency Bode Plot', sty['h1']), HR(),
            Paragraph('Gray dashed: ideal grid (Lg=0); Red solid: actual response with leakage; Green shading: safety island.', sty['body']),
            Spacer(1,0.3*rl_cm),
            RLImage(io.BytesIO(bode_png), width=15*rl_cm, height=11*rl_cm),
            Paragraph('Figure 6-1  LCL Filter Bode Plot (with external leakage inductance)', sty['caption']),
        ]
        doc.build(story)
        buf.seek(0)
        return buf.read()

    # ── SOA matplotlib  ──
    def _make_soa_fig(l1_min, l_max, c_max, p_fsw, p_fg,
                      ltot_ripple_limit, lg, has_l2, k_int,
                      l_tot_design_mh, final_c, ltot_axis_max, cf_axis_max):
        fig, ax = plt.subplots(figsize=(8, 5))
        Ltot = np.linspace(0.001, ltot_axis_max, 300)
        Cf   = np.linspace(0.1,   cf_axis_max,   300)
        Lg2, Cg2 = np.meshgrid(Ltot, Cf)
        Ls = Lg2*1e-3; Cs = Cg2*1e-6
        if has_l2:
            L1g = np.maximum((Ls - lg)/(1 + k_int), 1e-9)
            L2g = L1g * k_int + lg
        else:
            L1g = np.maximum(Ls - lg, 1e-9)
            L2g = np.full_like(L1g, lg)
        Lp  = (L1g*L2g)/(L1g+L2g+1e-18)
        fr  = 1/(2*np.pi*np.sqrt(Lp*Cs+1e-30))
        safe = (L1g>=l1_min)&(Ls<=l_max)&(Cs<=c_max)&(fr<=0.5*p_fsw)&(fr>=10*p_fg)
        ax.contourf(Ltot, Cf, safe.astype(float),
                    levels=[0.5,1.5], colors=['#28a74555'], alpha=0.6)
        ax.axvline(ltot_ripple_limit, color='blue', ls='--', lw=1.8, label='① Ripple Limit L1_min')
        if l_max*1000 > 0:
            ax.axvline(l_max*1000, color='red', ls='--', lw=1.8, label='② Voltage Drop Limit L_max')
        ax.axhline(c_max*1e6, color='purple', ls='--', lw=1.8, label='③ Reactive Power Limit Cf_max')
        L1l = np.linspace(max(lg+1e-6, 0.001e-3), ltot_axis_max*1e-3, 200)
        if has_l2:
            L2l = np.maximum((L1l-lg)/(1+k_int), 1e-9)*k_int + lg
        else:
            L2l = np.full_like(L1l, lg)
        Lpl = (L1l*L2l)/(L1l+L2l+1e-18)
        cfmx = (1/(Lpl*(2*np.pi*0.5*p_fsw)**2))*1e6
        cfmn = (1/(Lpl*(2*np.pi*10*p_fg)**2))*1e6
        cfmx[cfmx>cf_axis_max*2]=np.nan; cfmn[cfmn>cf_axis_max*2]=np.nan
        ax.plot(L1l*1000, cfmx, color='orange', lw=1.8, label='④ fres=0.5fsw upper limit')
        ax.plot(L1l*1000, cfmn, color='cyan',   lw=1.8, label='⑤ fres=10fg lower limit')
        ax.plot(l_tot_design_mh, final_c, marker='*', ms=16,
                color='gold', mec='black', zorder=10, label='Current Design Point')
        ax.set_xlabel('Total System Inductance L_total = L1+L2+Lg (mH)', fontsize=10, fontproperties=_CN_FONT)
        ax.set_ylabel('Filter Capacitor Cf (uF)', fontsize=10, fontproperties=_CN_FONT)
        ax.set_title('SOA Five-Constraint Feasibility Map', fontsize=11, fontproperties=_CN_FONT)
        ax.legend(fontsize=8, prop=_CN_FONT)
        ax.set_xlim(0, ltot_axis_max); ax.set_ylim(0, cf_axis_max)
        ax.grid(alpha=0.3)
        buf = io.BytesIO(); fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig); buf.seek(0); return buf.read()

    # ──  ──
    def _make_bode_fig(l1_si, l2_si, lg, cf_si, rd_si, p_fg, p_fsw, has_l2):
        w = np.logspace(1, 5, 800)*2*np.pi; freqs = w/(2*np.pi); s = 1j*w
        l2t = l2_si + lg; cf = cf_si; rd = rd_si
        if has_l2 and l2_si > 0:
            num_i = [rd*cf, 1]; den_i = [l1_si*l2_si*cf, rd*cf*(l1_si+l2_si), l1_si+l2_si, 0]
            li = 'LCL (Lg=0)'
        else:
            num_i = [1]; den_i = [l1_si, 0]; li = 'L (Lg=0)'
        Hi  = np.polyval(num_i, s)/np.polyval(den_i, s)
        num_a = [rd*cf,1]; den_a = [l1_si*l2t*cf, rd*cf*(l1_si+l2t), l1_si+l2t, 0]
        Ha  = np.polyval(num_a, s)/np.polyval(den_a, s)
        mag_i = 20*np.log10(np.abs(Hi)+1e-30)
        mag_a = 20*np.log10(np.abs(Ha)+1e-30)
        ph_a  = np.angle(Ha, deg=True)
        lp = (l1_si*l2t)/(l1_si+l2t+1e-18)
        fr = 1/(2*np.pi*np.sqrt(lp*cf)) if cf > 0 else 0
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(8,6), sharex=True)
        ax1.semilogx(freqs, mag_i, color='gray', ls='--', lw=1.5, alpha=0.7, label=f'Ideal Grid Lg=0 ({li})')
        ax1.semilogx(freqs, mag_a, color='#dc3545', lw=2, label=f'Actual with leakage Lg={lg*1000:.3f}mH')
        if fr > 0:
            ax1.axvline(fr, color='#dc3545', ls=':', lw=1.2)
            ax1.text(fr*1.05, np.nanmax(mag_a)-5, f'fres={fr:.0f}Hz',
                     color='#dc3545', fontsize=8, fontproperties=_CN_FONT)
        ax1.axvspan(10*p_fg, 0.5*p_fsw, alpha=0.07, color='green', label='Safety Island')
        ax1.set_ylabel('Gain (dB)', fontsize=9, fontproperties=_CN_FONT)
        ax1.set_title('LCL Filter Bode Plot (with external leakage inductance)', fontsize=11, fontproperties=_CN_FONT)
        ax1.legend(fontsize=8, prop=_CN_FONT); ax1.grid(True, which='both', alpha=0.3)
        ax2.semilogx(freqs, ph_a, color='#0d6efd', lw=2, label='Phase (with leakage)')
        ax2.set_ylabel('Phase (deg)', fontsize=9, fontproperties=_CN_FONT)
        ax2.set_xlabel('Frequency (Hz)', fontsize=9, fontproperties=_CN_FONT)
        ax2.legend(fontsize=8, prop=_CN_FONT); ax2.grid(True, which='both', alpha=0.3)
        plt.tight_layout()
        buf = io.BytesIO(); fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig); buf.seek(0); return buf.read()

    # =========================================================
    # 0. Parameter
    # =========================================================
    default_values = {
        'sys_ug': 690.0, 'sys_fg': 50.0, 'vec_i': 500.0, 'vec_udc': 1100.0, 'sys_fsw': 2500.0,
        'sys_topo': "2-Level", 'sys_mod': "SVPWM", 'design_ripple': 20.0, 'design_q': 5.0,
        'target_pf': 0.90, 'lcl_l1': 0.50, 'lcl_l2': 0.25, 'lcl_c': 100.0, 'lcl_rd': 0.05,
        'damp_mode': "Series Resistor Damping"
    }
    for key, val in default_values.items():
        if key not in st.session_state:
            st.session_state[key] = val

    # =========================================================
    # 1. :  + 
    # =========================================================
    damping_option = st.radio("Damping Configuration", ["No Damping", "Series Resistor Damping"], horizontal=True, key='damp_mode')
    has_damping = (damping_option == "Series Resistor Damping")

    c_L, c_C, c_R = "#d63384", "#0d6efd", "#fd7e14"
    if has_damping:
        cap_branch = [
            f'  <line x1="220" y1="90" x2="220" y2="115" stroke="black" stroke-width="2"/>',
            f'  <line x1="205" y1="115" x2="235" y2="115" stroke="black" stroke-width="2"/>',
            f'  <line x1="205" y1="125" x2="235" y2="125" stroke="black" stroke-width="2"/>',
            f'  <text x="240" y="125" font-family="sans-serif" font-size="14" fill="{c_C}" font-weight="bold">Cf</text>',
            f'  <line x1="220" y1="125" x2="220" y2="140" stroke="black" stroke-width="2"/>',
            f'  <polyline points="220,140 215,145 225,150 215,155 225,160 220,165" fill="none" stroke="{c_R}" stroke-width="2"/>',
            f'  <text x="240" y="160" font-family="sans-serif" font-size="14" fill="{c_R}" font-weight="bold">Rd</text>',
            f'  <line x1="220" y1="165" x2="220" y2="180" stroke="black" stroke-width="2"/>',
        ]
    else:
        cap_branch = [
            f'  <line x1="220" y1="90" x2="220" y2="130" stroke="black" stroke-width="2"/>',
            f'  <line x1="205" y1="130" x2="235" y2="130" stroke="black" stroke-width="2"/>',
            f'  <line x1="205" y1="140" x2="235" y2="140" stroke="black" stroke-width="2"/>',
            f'  <text x="240" y="140" font-family="sans-serif" font-size="14" fill="{c_C}" font-weight="bold">Cf</text>',
            f'  <line x1="220" y1="140" x2="220" y2="180" stroke="black" stroke-width="2"/>',
        ]
    svg_topo = [
        '<svg width="100%" height="180" xmlns="http://www.w3.org/2000/svg">',
        '  <rect x="10" y="50" width="50" height="80" fill="#f8f9fa" stroke="#333" stroke-width="2"/>',
        '  <text x="20" y="95" font-family="sans-serif" font-size="12" fill="#333">Vinv</text>',
        '  <line x1="60" y1="90" x2="100" y2="90" stroke="black" stroke-width="2"/>',
        f'  <path d="M100,90 Q110,70 120,90 Q130,70 140,90 Q150,70 160,90" fill="none" stroke="{c_L}" stroke-width="2.5"/>',
        f'  <text x="115" y="60" font-family="sans-serif" font-size="14" fill="{c_L}" font-weight="bold">L1</text>',
        '  <line x1="160" y1="90" x2="220" y2="90" stroke="black" stroke-width="2"/>',
        *cap_branch,
        '  <line x1="210" y1="180" x2="230" y2="180" stroke="black" stroke-width="2"/>',
        f'  <path d="M220,90 Q230,70 240,90 Q250,70 260,90 Q270,70 280,90" fill="none" stroke="{c_L}" stroke-width="2.5"/>',
        f'  <text x="240" y="60" font-family="sans-serif" font-size="14" fill="{c_L}" font-weight="bold">L2</text>',
        '  <line x1="280" y1="90" x2="330" y2="90" stroke="black" stroke-width="2"/>',
        '  <circle cx="360" cy="90" r="30" fill="none" stroke="#198754" stroke-width="2"/>',
        '  <text x="348" y="95" font-family="sans-serif" font-size="12" font-weight="bold" fill="#198754">Grid</text>',
        '</svg>',
    ]
    st.markdown("".join(svg_topo), unsafe_allow_html=True)
    st.divider()

    # =========================================================
    # 1: 
    # =========================================================
    st.markdown("#### 1. Operating Conditions & Boundary Constraints")
    c1, c2, c3 = st.columns(3)
    with c1:
        p_vg_line = st.number_input("Grid Line Voltage (Vrms)", min_value=10.0,  step=10.0,  key='sys_ug')
        p_in      = st.number_input("Rated Phase Current (Arms)", min_value=1.0,   step=10.0,  key='vec_i')
        p_vdc     = st.number_input("DC Bus Voltage (V)",  min_value=100.0, step=10.0,  key='vec_udc')
    with c2:
        p_fg      = st.number_input("Grid Frequency (Hz)",     min_value=1.0,   step=1.0,   key='sys_fg')
        p_fsw     = st.number_input("Switching Frequency (Hz)",     min_value=100.0, step=100.0, key='sys_fsw')
        topo_type = st.selectbox("Topology", ["2-Level", "3-Level (NPC/T-Type)"], key='sys_topo')
        mod_type  = st.selectbox("Modulation Strategy", ["SVPWM", "SPWM"], key='sys_mod')
    with c3:
        target_pf         = st.slider("Minimum Power Factor (Lagging)", 0.0, 1.0, step=0.05, key='target_pf')
        target_ripple_pct = st.number_input("Maximum Allowable Ripple (%)", min_value=1.0, step=1.0, key='design_ripple', help="Typical range: 15%~25%")
        max_reactive_pct = st.number_input("Capacitor Reactive Power Limit (%)", min_value=0.5, step=0.5, key='design_q', help="Typically 5% of rated power")
    st.divider()

    # =========================================================
    # 
    # =========================================================
    vg_phase = p_vg_line / np.sqrt(3)
    omega_g  = 2 * np.pi * p_fg
    v_max_factor = 1/(2*np.sqrt(2)) if mod_type == "SPWM" else 1/np.sqrt(6)
    v_max_out    = p_vdc * v_max_factor
    phi_rad      = np.arccos(target_pf)
    sin_phi      = np.sin(phi_rad)
    a_quad = 1.0; b_quad = 2*vg_phase*sin_phi; c_quad = vg_phase**2 - v_max_out**2
    delta_quad = b_quad**2 - 4*a_quad*c_quad
    l_max_val = 0.0; v_l_max = 0.0
    if delta_quad >= 0 and c_quad <= 0:
        v_l_max   = (-b_quad + np.sqrt(delta_quad)) / (2*a_quad)
        l_max_val = v_l_max / (omega_g * p_in)
    k_ripple   = 4 if topo_type == "2-Level" else (8 if mod_type == "SPWM" else 12)
    delta_i_pp = p_in * np.sqrt(2) * (target_ripple_pct / 100.0)
    l1_min_val = p_vdc / (k_ripple * p_fsw * delta_i_pp) if delta_i_pp > 0 else 0.0
    c_max_reactive = (max_reactive_pct/100.0) * p_in / (omega_g * vg_phase)

    # =========================================================
    # 2: 
    # =========================================================
    st.markdown("#### 2. Physical Limit Derivation & Calculation Results")
    rd1, rd2 = st.columns(2)
    with rd1:
        st.markdown("##### ① Total Inductance Upper Limit (Four-Quadrant Vector Limit) 👉 **[See Navigation: 4. Four-Quadrant Operation]**")
        st.markdown("$\\vec{V}_{inv} = \\vec{V}_g + j\\omega L \\vec{I}$, at power factor Lagging $\\phi$:")
        st.latex(r"|\vec{V}_{inv}|^2 = V_g^2 + V_L^2 + 2 V_g V_L \sin(\phi) \le V_{max\_out}^2")
        if is_trial:
            trial_lock("Unlock with paid access: total inductance limit result", compact=True)
        elif l_max_val > 0:
            st.markdown(f"Substituting $V_{{max\\_out}} = {v_max_out:.1f}$ V, $V_g = {vg_phase:.1f}$ V, we get **$L_{{total}} \\le {l_max_val*1000:.3f}$ mH**.")
        else:
            st.error(f"🚨 **Grid voltage exceeds converter maximum output** (Vg_phase={vg_phase:.1f}V > V_max_out={v_max_out:.1f}V). These parameters cannot achieve grid connection. Increase Vdc or lower grid voltage.")
        st.divider()
        st.markdown("##### ③ Cf Upper Limit (Reactive Power Limit)")
        st.latex(r"C_f \le \frac{\lambda \% \cdot I_{rated}}{\omega_g \cdot V_{g,ph}}")
        if is_trial:
            trial_lock("Unlock with paid access: Cf upper limit result", compact=True)
        else:
            st.markdown(f"Current Result: **$C_f \\le {c_max_reactive*1e6:.1f}$ μF**")
    with rd2:
        st.markdown("##### ② L1 Lower Limit (Ripple Limit) 👉 **[See Navigation: 6. Ripple Current Analysis]**")
        st.latex(r"L_1 \ge \frac{V_{dc}}{K \cdot f_{sw} \cdot \Delta I_{pp}}")
        if is_trial:
            trial_lock("Unlock with paid access: L1 lower limit result", compact=True)
        else:
            st.markdown(f"$K={k_ripple}$, $\\Delta I_{{pp}}={delta_i_pp:.1f}$ A, giving **$L_1 \\ge {l1_min_val*1000:.3f}$ mH**.")
        if has_damping:
            st.divider()
            st.markdown("##### ④ Passive Damping Resistor Rd Design Basis")
            st.latex(r"R_d \approx \frac{1}{3 \cdot 2\pi f_{res} \cdot C_f}")
            st.caption("📚 Liserre et al., IEEE Trans. Ind. Electron., 2005, 52(5):1199-1206")
        st.divider()
        st.markdown("##### 🏝️ Safety Island Frequency Band (Constraints ④⑤)")
        st.latex(r"10 f_g < f_{res} < 0.5 f_{sw}")
        if is_trial:
            trial_lock("Unlock with paid access: safety island frequency band result", compact=True)
        else:
            st.markdown(f"Current safety frequency band: **{10*p_fg:.0f} Hz ~ {0.5*p_fsw:.0f} Hz**")
    st.divider()

    # =========================================================
    # 3. Parameter SOA 
    # =========================================================
    st.markdown("### 🎯 3. Parameter Lock & SOA Feasibility Analysis")
    col_lock, col_soa = st.columns([1, 2.5])

    with col_lock:
        st.markdown("#### 🌐 External Grid-Side Leakage Definition")
        grid_mode_b2 = st.radio("Grid Impedance Calculation Method:",
                                ["Specify Grid SCR directly", "Estimate via Transformer Parameters"],
                                key="b2_grid_mode")
        safe_ug_b2 = float(st.session_state.get('sys_ug', 690.0))
        safe_fg_b2 = float(st.session_state.get('sys_fg', 50.0))
        safe_i_b2  = float(st.session_state.get('vec_i', 500.0))
        s_inv_b2   = np.sqrt(3) * safe_ug_b2 * safe_i_b2

        tx_s_kva_b2 = 1000.0; tx_uk_b2 = 6.0
        z_base_tx_b2 = 1.0;   z_tx_b2  = 0.0
        z_base_b2    = 1.0;   z_g_b2   = 0.0
        scr_val_b2   = 15.0;  eq_scr_b2 = 15.0

        if grid_mode_b2 == "Specify Grid SCR directly":
            scr_val_b2 = st.slider("Grid Point SCR", min_value=1.5, max_value=50.0,
                                   value=15.0, step=0.5, key="b2_scr")
            z_base_b2  = (safe_ug_b2**2) / s_inv_b2 if s_inv_b2 > 0 else 1.0
            z_g_b2     = z_base_b2 / scr_val_b2
            lg_b2      = z_g_b2 / (2*np.pi*safe_fg_b2)
            eq_scr_b2  = scr_val_b2
        else:
            cb2_1, cb2_2 = st.columns(2)
            with cb2_1:
                tx_s_kva_b2 = st.number_input("Transformer Rating (kVA)", min_value=10.0, value=1000.0,
                                              step=100.0, key="b2_tx_kva")
            with cb2_2:
                tx_uk_b2 = st.number_input("Short-Circuit Impedance uk (%)", min_value=1.0, max_value=20.0,
                                           value=6.0, step=0.5, key="b2_tx_uk")
            tx_s_b2      = tx_s_kva_b2 * 1000.0
            z_base_tx_b2 = (safe_ug_b2**2) / tx_s_b2 if tx_s_b2 > 0 else 1.0
            z_tx_b2      = (tx_uk_b2/100.0) * z_base_tx_b2
            lg_b2        = z_tx_b2 / (2*np.pi*safe_fg_b2)
            eq_scr_b2    = tx_s_b2 / ((tx_uk_b2/100.0)*s_inv_b2) if s_inv_b2 > 0 else 999.0

        if is_trial:
            trial_lock("Unlock with paid access: external leakage Lg and equivalent SCR result", compact=True)
        else:
            col_m1, col_m2 = st.columns(2)
            with col_m1: st.metric("📐 External Leakage $L_g$", f"{lg_b2*1000:.2f} mH")
            with col_m2: st.metric("📊 Equivalent SCR", f"{eq_scr_b2:.1f}")
        st.divider()

        st.markdown("#### 🔩 Converter Internal Structure")
        has_l2_b2 = st.checkbox("Converter includes internal grid-side inductor L2", value=True, key="b2_has_l2")
        l_inner_max = max(l_max_val - lg_b2, 0.0)
        l_inner_min = l1_min_val
        st.divider()

        if has_l2_b2:
            st.markdown("#### 🎛️ L1 / L2 Parameter Lock")
            if is_trial:
                trial_lock("Unlock with paid access: internal inductor recommended range", compact=True)
            else:
                st.info(f"Internal inductor range (total limit {l_max_val*1000:.3f} mH minus leakage {lg_b2*1000:.3f} mH):\n\n"
                        f"Recommended L1+L2: **{l_inner_min*1000:.3f} ~ {l_inner_max*1000:.3f}** mH\n\nSuggested L2/L1 ratio: 0.3 ~ 0.6")
            l2_ratio_b2 = st.slider("L2/L1 Ratio k", 0.1, 1.0, 0.5, 0.05, key="b2_l2_ratio")
            l1_suggest_max = l_inner_max/(1+l2_ratio_b2) if (1+l2_ratio_b2) > 0 else l_inner_max
            if not is_trial:
                st.caption(f"With k={l2_ratio_b2}, recommended L1: {l_inner_min*1000:.3f} ~ {l1_suggest_max*1000:.3f} mH")
            user_l1_b2 = st.number_input("Lock L1 (mH)", min_value=0.001,
                                         value=max(float(st.session_state.get('lcl_l1_with_l2', 0.5)), 0.001),
                                         step=0.01, format="%.3f", key='lcl_l1_with_l2')
            l2_locked = user_l1_b2 * l2_ratio_b2
            st.session_state['lcl_l2'] = l2_locked
            if is_trial:
                trial_lock("Unlock with paid access: L2 linked value and total system inductance", compact=True)
            else:
                st.markdown(f"Linked setting **L2 = {l2_locked:.3f} mH**\n\n"
                            f"Total system inductance L1+L2+Lg = **{(user_l1_b2+l2_locked+lg_b2*1000):.3f} mH**")
        else:
            st.markdown("#### 🎛️ L1 Parameter Lock (CL Topology)")
            if is_trial:
                trial_lock("Unlock with paid access: CL topology recommended range", compact=True)
            else:
                st.info(f"CL topology: external leakage {lg_b2*1000:.3f} mH serves as L2.\n\nRecommended L1: **{l_inner_min*1000:.3f} ~ {l_inner_max*1000:.3f}** mH")
            l2_ratio_b2 = 0.0
            user_l1_b2 = st.number_input("Lock L1 (mH)", min_value=0.001,
                                         value=max(float(st.session_state.get('lcl_l1_no_l2', 0.5)), 0.001),
                                         step=0.01, format="%.3f", key='lcl_l1_no_l2')
            st.session_state['lcl_l2'] = 0.0
            l2_locked = 0.0
            if is_trial:
                trial_lock("Unlock with paid access: total system inductance summary", compact=True)
            else:
                st.markdown(f"Total system inductance L1+Lg = **{(user_l1_b2+lg_b2*1000):.3f} mH**")

        st.divider()
        cf_col1, cf_col2 = st.columns([2,1])
        with cf_col1:
            final_c = st.number_input("Lock Cf (μF)  ⚠️ Y-equivalent value", min_value=0.1,
                                      value=float(st.session_state.get('lcl_c', 100.0)),
                                      step=1.0, format="%.1f", key='lcl_c',
                                      help="For delta (Δ) configuration, physical capacitor value = this value ÷ 3.")
        with cf_col2:
            st.markdown("<br>", unsafe_allow_html=True)
            st.caption("Delta config:\nPhysical = Y÷3")

        l1_si    = user_l1_b2 * 1e-3
        l2_si    = l2_locked  * 1e-3 + lg_b2
        cf_si    = final_c * 1e-6
        l_tot_si = l1_si + l2_si
        l_par_b2 = (l1_si*l2_si)/(l1_si+l2_si+1e-18)
        cur_f_res_b2 = 1/(2*np.pi*np.sqrt(l_par_b2*cf_si)) if (l_par_b2*cf_si) > 0 else 0

        if has_damping:
            recommended_rd = 1/(3*2*np.pi*cur_f_res_b2*cf_si) if cur_f_res_b2 > 0 else 0
            if is_trial:
                pass  # hide Rd recommendation in trial mode
            else:
                st.markdown(f"Theoretical optimal Rd ≈ **{recommended_rd:.3f}** Ω")
            final_rd = st.number_input("Lock Rd (Ω)", min_value=0.0,
                                       value=float(st.session_state.get('lcl_rd', 0.05)),
                                       step=0.01, format="%.3f")
        else:
            final_rd = 0.0

    with col_soa:
        st.markdown("#### 📊 SOA Five-Constraint Feasibility Map")
        k_int = l2_ratio_b2
        ltot_axis_max  = max(l_max_val, (l1_min_val+lg_b2)*2) * 1.3 * 1000
        cf_axis_max_b2 = c_max_reactive * 1.5 * 1e6
        if ltot_axis_max < 0.5:  ltot_axis_max  = 5.0
        if cf_axis_max_b2 < 1.0: cf_axis_max_b2 = 200.0

        Ltot_arr = np.linspace(0.001, ltot_axis_max, 300)
        Cf_arr   = np.linspace(0.1,   cf_axis_max_b2, 300)
        Ltot_grid, Cf_grid = np.meshgrid(Ltot_arr, Cf_arr)
        Ltot_si = Ltot_grid*1e-3; Cf_si_g = Cf_grid*1e-6
        if has_l2_b2:
            L1_si_g    = np.maximum((Ltot_si-lg_b2)/(1+k_int), 1e-9)
            L2eff_si_g = L1_si_g*k_int + lg_b2
        else:
            L1_si_g    = np.maximum(Ltot_si-lg_b2, 1e-9)
            L2eff_si_g = np.full_like(L1_si_g, lg_b2)
        L_par_si_g = (L1_si_g*L2eff_si_g)/(L1_si_g+L2eff_si_g+1e-18)
        f_res_g    = 1/(2*np.pi*np.sqrt(L_par_si_g*Cf_si_g+1e-30))
        if has_l2_b2:
            ltot_ripple_limit = (l1_min_val*(1+k_int)+lg_b2)*1000
        else:
            ltot_ripple_limit = (l1_min_val+lg_b2)*1000
        cond_r = L1_si_g>=l1_min_val; cond_v = Ltot_si<=l_max_val
        cond_q = Cf_si_g<=c_max_reactive
        safe_zone_g = cond_r & cond_v & cond_q & (f_res_g<=0.5*p_fsw) & (f_res_g>=10*p_fg)

        fig_soa = go.Figure()
        fig_soa.add_trace(go.Contour(x=Ltot_arr, y=Cf_arr, z=safe_zone_g.astype(int),
                                     colorscale=[[0,'rgba(0,0,0,0)'],[1,'rgba(40,167,69,0.35)']],
                                     showscale=False, hoverinfo='skip'))
        fig_soa.add_vline(x=ltot_ripple_limit, line_width=2, line_dash="dash", line_color="blue",
                          annotation_text="① Ripple Limit", annotation_position="bottom right")
        if l_max_val > 0:
            fig_soa.add_vline(x=l_max_val*1000, line_width=2, line_dash="dash", line_color="red",
                              annotation_text="② Voltage Drop Limit", annotation_position="top left")
        fig_soa.add_hline(y=c_max_reactive*1e6, line_width=2, line_dash="dash", line_color="purple",
                          annotation_text="③ Reactive Power Limit", annotation_position="bottom left")
        Ltot_line = np.linspace(max(lg_b2+1e-6, 0.001e-3), ltot_axis_max*1e-3, 300)
        if has_l2_b2:
            L1_line    = np.maximum((Ltot_line-lg_b2)/(1+k_int), 1e-9)
            L2eff_line = L1_line*k_int + lg_b2
        else:
            L1_line    = np.maximum(Ltot_line-lg_b2, 1e-9)
            L2eff_line = np.full_like(L1_line, lg_b2)
        L_par_line = (L1_line*L2eff_line)/(L1_line+L2eff_line+1e-18)
        Cf_fmax_line = (1/(L_par_line*(2*np.pi*0.5*p_fsw)**2))*1e6
        Cf_fmin_line = (1/(L_par_line*(2*np.pi*10*p_fg)**2))*1e6
        Cf_fmax_line[Cf_fmax_line>cf_axis_max_b2*2]=np.nan
        Cf_fmin_line[Cf_fmin_line>cf_axis_max_b2*2]=np.nan
        fig_soa.add_trace(go.Scatter(x=Ltot_line*1000, y=Cf_fmax_line, mode='lines',
                                     name='④ Resonance upper limit (0.5fsw)', line=dict(color='orange',width=2)))
        fig_soa.add_trace(go.Scatter(x=Ltot_line*1000, y=Cf_fmin_line, mode='lines',
                                     name='⑤ Resonance lower limit (10fg)',   line=dict(color='cyan',  width=2)))
        l_tot_design_mh = (l1_si+l2_si)*1000
        fig_soa.add_trace(go.Scatter(x=[l_tot_design_mh], y=[final_c], mode='markers',
                                     name='Current Design Point ⭐',
                                     marker=dict(size=18, symbol='star', color='gold',
                                                 line=dict(width=2, color='black'))))
        fig_soa.update_layout(
            title=f"Five-Constraint Feasibility Map (X-axis=L_total, Lg={lg_b2*1000:.3f} mH)",
            xaxis_title="Total System Inductance L_total = L1+L2+Lg (mH)",
            yaxis_title="Filter Capacitor Cf (μF)",
            xaxis=dict(range=[0,ltot_axis_max]), yaxis=dict(range=[0,cf_axis_max_b2]),
            height=520, margin=dict(l=40,r=40,t=50,b=40),
            legend=dict(x=0.6, y=0.98, bgcolor="rgba(255,255,255,0.8)"),
            plot_bgcolor="#fafafa"
        )
        st.plotly_chart(fig_soa, use_container_width=True)
        st.divider()

        # — Compute compliance flags (always calculated, used in report even for trial) —
        l1_pass    = l1_si >= l1_min_val
        ltot_pass  = l_tot_si <= l_max_val
        c_pass     = cf_si <= c_max_reactive
        f_pass     = (10*p_fg <= cur_f_res_b2 <= 0.5*p_fsw)
        all_pass   = l1_pass and ltot_pass and c_pass and f_pass

        if is_trial:
            # Trial mode: show constraint names only, hide exact values
            st.markdown(
                "**📌 Five-Constraint Boundary Quick Reference (Trial Mode):**\n\n"
                "* Constraint ① Blue line: L1 lower limit\n"
                "* Constraint ② Red line: L_total upper limit\n"
                "* Constraint ③ Purple line: Cf upper limit\n"
                "* Constraints ④⑤ Safety island frequency band"
            )
            trial_lock("Unlock with paid access: exact constraint boundary values and compliance check results")
        else:
            st.markdown("#### 🧠 Operating Point Constraint Summary & Safety Island Verification")
            st.info(
                f"**📌 Five-Constraint Boundary Summary:**\n\n"
                f"*   **Constraint ① Blue line**: L1 lower limit ≥ **{l1_min_val*1000:.3f} mH**\n"
                f"*   **Constraint ② Red line**: L_total upper limit ≤ **{l_max_val*1000:.3f} mH**\n"
                f"*   **Constraint ③ Purple line**: Cf upper limit ≤ **{c_max_reactive*1e6:.1f} μF**\n"
                f"*   **Constraints ④⑤ Safety Island**: **{10*p_fg:.0f} ~ {0.5*p_fsw:.0f} Hz**"
            )

        if l_max_val == 0:
            st.error("🚨 DC Bus Voltage too low (Vg_phase > V_max_out) — Grid connection impossible. Please increase Vdc.")
        elif is_trial:
            pass  # trial mode already handled above
        elif all_pass:
            st.success(
                f"✅ **Current design point is within the safety island.**\n\n"
                f"*   Constraint ① L1={user_l1_b2:.3f} mH ≥ {l1_min_val*1000:.3f} mH ✅\n"
                f"*   Constraint ② L_total={l_tot_si*1000:.3f} mH ≤ {l_max_val*1000:.3f} mH ✅\n"
                f"*   Constraint ③ Cf={final_c:.1f} μF ≤ {c_max_reactive*1e6:.1f} μF ✅\n"
                f"*   Constraints ④⑤ {cur_f_res_b2:.0f} Hz ∈ [{10*p_fg:.0f}, {0.5*p_fsw:.0f}] Hz ✅"
            )
            st.metric("Actual Resonance Frequency", f"{cur_f_res_b2:.1f} Hz")
        else:
            st.error("🚨 **Parameter Out of Bounds** — Design point is outside the safe zone!")
            cw1, cw2 = st.columns(2)
            with cw1:
                if not l1_pass:   st.markdown(f"❌ Ripple violated: L1={user_l1_b2:.3f} < {l1_min_val*1000:.3f} mH")
                if not ltot_pass: st.markdown(f"❌ Over-modulation: L_total={l_tot_si*1000:.3f} mH exceeds limit {l_max_val*1000:.3f} mH")
            with cw2:
                if not c_pass: st.markdown(f"❌ Reactive power exceeded: Cf={final_c:.1f} > {c_max_reactive*1e6:.1f} μF")
                if not f_pass: st.markdown(f"❌ Safety island violated: {cur_f_res_b2:.0f} Hz not in [{10*p_fg:.0f},{0.5*p_fsw:.0f}] Hz")
            st.metric("Current Resonance Frequency", f"{cur_f_res_b2:.1f} Hz")

    st.divider()

    # =========================================================
    # 4. One-click report generation
    # =========================================================
    st.markdown("### 📄 4. Generate LCL Design Report")
    st.caption("Packages all current design parameters, LaTeX formula derivations, SOA map, and Bode plot into a professional report for engineering documentation or design review.")

    rep_col1, rep_col2 = st.columns([1, 2])
    with rep_col1:
        report_fmt = st.radio("Report Format", ["Word (.docx)", "PDF (.pdf)"],
                              horizontal=False, key="report_fmt")
        if is_trial:
            st.button("🔒 Generate Report (Paid Access Required)", type="primary",
                      use_container_width=True, disabled=True)
            trial_lock("Unlock with paid access: generate full LCL design report", compact=True)
            gen_btn = False
        else:
            gen_btn = st.button("🚀 Generate Report", type="primary", use_container_width=True)
    if not is_trial:
        with rep_col2:
            st.info(
                "**Report sections include:**\n\n"
                "1. Operating conditions & boundary constraint parameter summary\n"
                "2. LaTeX formula derivation for the three physical limits (rendered as math images)\n"
                "3. External grid-side leakage inductance derivation (SCR or transformer mode)\n"
                "4. Finalized parameter summary table (including resonance frequency and compliance status)\n"
                "5. SOA five-constraint feasibility map\n"
                "6. Resonance frequency Bode plot (including external leakage inductance response)"
            )

        if gen_btn:
            with st.spinner("⏳ Rendering formulas and generating report, please wait..."):
                try:
                    rp = {
                        'vg_line': p_vg_line, 'vg_phase': vg_phase,
                        'fg': p_fg, 'omega_g': omega_g,
                        'i_rated': p_in, 'vdc': p_vdc,
                        'fsw': p_fsw, 'topo': topo_type,
                        'mod': mod_type, 'pf': target_pf,
                        'damp_mode': damping_option,
                        'ripple_pct': target_ripple_pct,
                        'reactive_pct': max_reactive_pct,
                        'delta_i_pp': delta_i_pp,
                        'k_ripple': k_ripple,
                        'f_safe_low': 10*p_fg, 'f_safe_high': 0.5*p_fsw,
                        'v_max_out': v_max_out,
                        'phi_deg': float(np.degrees(phi_rad)),
                        'sin_phi': sin_phi,
                        'v_l_max': v_l_max,
                        'l_max_mh': l_max_val*1000,
                        'l1_min_mh': l1_min_val*1000,
                        'c_max_uf': c_max_reactive*1e6,
                        'grid_mode': grid_mode_b2,
                        'lg_mh': lg_b2*1000,
                        'eq_scr': eq_scr_b2,
                        's_inv_kva': s_inv_b2/1000,
                        'z_base': z_base_b2, 'z_g': z_g_b2, 'scr': scr_val_b2,
                        'tx_kva': tx_s_kva_b2, 'tx_uk': tx_uk_b2,
                        'z_base_tx': z_base_tx_b2, 'z_tx': z_tx_b2,
                        'has_l2': has_l2_b2, 'has_damping': has_damping,
                        'final_l1': user_l1_b2, 'final_l2': l2_locked,
                        'final_c': final_c, 'final_rd': final_rd,
                        'l_total_mh': l_tot_si*1000,
                        'f_res': cur_f_res_b2, 'all_pass': all_pass,
                    }

                    soa_png  = _make_soa_fig(
                        l1_min=l1_min_val, l_max=l_max_val,
                        c_max=c_max_reactive, p_fsw=p_fsw, p_fg=p_fg,
                        ltot_ripple_limit=ltot_ripple_limit, lg=lg_b2,
                        has_l2=has_l2_b2, k_int=k_int,
                        l_tot_design_mh=l_tot_design_mh, final_c=final_c,
                        ltot_axis_max=ltot_axis_max, cf_axis_max=cf_axis_max_b2
                    )
                    bode_png = _make_bode_fig(
                        l1_si=l1_si, l2_si=l2_locked*1e-3,
                        lg=lg_b2, cf_si=cf_si,
                        rd_si=final_rd if has_damping else 0.0,
                        p_fg=p_fg, p_fsw=p_fsw, has_l2=has_l2_b2
                    )

                    if report_fmt == "Word (.docx)":
                        report_bytes = _gen_word_report(rp, soa_png, bode_png, fig_soa=fig_soa)
                        fname = f"LCL_Design_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                        mime  = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    else:
                        report_bytes = _gen_pdf_report(rp, soa_png, bode_png, fig_soa=fig_soa)
                        fname = f"LCL_Design_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
                        mime  = "application/pdf"

                    st.success("✅ Report generated successfully!")
                    st.download_button(
                        label=f"⬇️ Download {report_fmt.split()[0]} Report",
                        data=report_bytes, file_name=fname, mime=mime,
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"❌ Report generation failed: {str(e)}")
                    st.exception(e)







# ==========================================
# Section 3: Resonance Frequency Analysis (Weak Grid / Leakage Inductance & Robustness)
# ==========================================
elif selection == nav_options[2]:
    st.header("📈 Weak Grid Adaptability & Resonance Drift Analysis")
    st.caption("🔥 Core feature: Reveals the actual frequency-domain response of CL/LCL topologies under real transformer leakage inductance and weak grid conditions.")

    # --- 1. Parameter input and control panel ---
    c_param, c_plot = st.columns([1, 2.5])

    with c_param:
        st.markdown("#### 🎛️ 1. Converter Hardware Parameters")

        safe_ug = max(10.0, float(st.session_state.get('sys_ug', 690.0)))
        safe_fg = max(1.0, float(st.session_state.get('sys_fg', 50.0)))
        safe_i = max(1.0, float(st.session_state.get('vec_i', 500.0)))
        safe_fsw = max(100.0, float(st.session_state.get('sys_fsw', 2500.0)))

        with st.expander("Base Operating Parameters", expanded=False):
            sys_ug = st.number_input("Grid Line Voltage (V)", min_value=10.0, value=safe_ug, step=10.0, key="m3_ug")
            sys_fg = st.number_input("Grid Frequency (Hz)", min_value=1.0, value=safe_fg, step=1.0, key="m3_fg")
            vec_i = st.number_input("Converter Rated Phase Current (A)", min_value=1.0, value=safe_i, step=10.0, key="m3_i")
            fsw = st.number_input("Switching Frequency (Hz)", min_value=100.0, value=safe_fsw, step=100.0, key="m3_fsw")

        s_inv = np.sqrt(3) * sys_ug * vec_i

        safe_l1 = max(0.001, float(st.session_state.get('lcl_l1', 0.5)))
        safe_c = max(0.1, float(st.session_state.get('lcl_c', 100.0)))
        safe_rd = max(0.0, float(st.session_state.get('lcl_rd', 0.05)))

        l1_mh = st.number_input("L1 Inverter-Side Inductance (mH)", min_value=0.001, value=safe_l1, step=0.01, format="%.3f", key="m3_l1")
        cf_uf = st.number_input("Cf Filter Capacitor (μF) [Y-equivalent]", min_value=0.1, value=safe_c, step=1.0, format="%.1f", key="m3_c",
                                help="This is the star (Y) equivalent capacitance. For delta (Δ) configuration, the physical capacitor value is 1/3 of this value.")
        rd_ohm = st.number_input("Rd Damping Resistor (Ω)", min_value=0.0, value=safe_rd, step=0.01, format="%.3f", key="m3_rd")

        has_l2 = st.checkbox("Converter includes internal grid-side inductor L2 (full LCL topology)", value=True,
                             help="High-power converters often omit internal L2, using external transformer leakage inductance instead (CL topology).")
        if has_l2:
            safe_l2 = max(0.001, float(st.session_state.get('lcl_l2', 0.25)))
            l2_mh = st.number_input("L2 Internal Grid-Side Inductance (mH)", min_value=0.001, value=safe_l2, step=0.01, format="%.3f", key="m3_l2")
        else:
            l2_mh = 0.0

        st.divider()
        st.markdown("#### 🌪️ 2. External Grid Environment Definition")
        grid_mode = st.radio("Grid Impedance Calculation Method:", ["Specify Grid SCR directly", "Estimate via Transformer Parameters"])

        if grid_mode == "Specify Grid SCR directly":
            scr_val = st.slider("Grid Point Short Circuit Ratio (SCR)", min_value=1.5, max_value=50.0, value=15.0, step=0.5)
            z_base_inv = (sys_ug ** 2) / s_inv if s_inv > 0 else 1.0
            l_base_inv = z_base_inv / (2 * np.pi * sys_fg)
            lg = l_base_inv / scr_val
            eq_scr = scr_val

            # ✅ Optimization 2: show results directly in SCR mode
            st.divider()
            if is_trial:
                trial_lock("Unlock with paid access: equivalent external leakage Lg and grid point SCR", compact=True)
            else:
                col_lg1, col_lg2 = st.columns([3, 2])
                with col_lg1:
                    st.metric(label="📐 Equivalent External Leakage $L_g$", value=f"{lg * 1000:.2f} mH")
                with col_lg2:
                    st.metric(label="📊 Grid Point SCR", value=f"{eq_scr:.1f}")

        else:
            c_tx1, c_tx2 = st.columns(2)
            with c_tx1:
                tx_s_kva = st.number_input("Transformer Rating (kVA)", min_value=10.0, value=1000.0, step=100.0,
                                           help="Enter the rated capacity from the step-up transformer nameplate. Does not need to match the converter rating.")
            with c_tx2:
                tx_uk = st.number_input("Short-Circuit Impedance uk (%)", min_value=1.0, max_value=20.0, value=6.0, step=0.5,
                                        help="Transformer nameplate short-circuit impedance percentage, typically 4%–8%.")

            # ✅ Optimization 1: leakage inductance based purely on transformer parameters
            tx_s = tx_s_kva * 1000.0
            z_base_tx = (sys_ug ** 2) / tx_s if tx_s > 0 else 1.0
            z_tx_actual = (tx_uk / 100.0) * z_base_tx
            lg = z_tx_actual / (2 * np.pi * sys_fg)

            # Equivalent SCR calculated separately as a system matching metric
            eq_scr = tx_s / ((tx_uk / 100.0) * s_inv) if s_inv > 0 else 999.0

            # ✅ Optimization 2: show results directly in transformer mode
            st.divider()
            if is_trial:
                trial_lock("Unlock with paid access: transformer leakage inductance Lg and equivalent SCR", compact=True)
            else:
                col_lg1, col_lg2 = st.columns(2)
                with col_lg1:
                    st.metric(label="📐 Transformer Leakage $L_g$", value=f"{lg * 1000:.3f} mH")
                with col_lg2:
                    st.metric(label="📊 Equivalent SCR", value=f"{eq_scr:.1f}",
                              help="Equivalent SCR = Transformer short-circuit capacity / (uk% × converter rating). This is a system matching metric, not an intrinsic transformer property.")

    # --- 2. Core mathematical calculations ---
    l1 = l1_mh * 1e-3
    l2 = l2_mh * 1e-3
    cf = cf_uf * 1e-6
    rd = rd_ohm

    l2_total = l2 + lg

    if has_l2:
        f_res_ideal = 1 / (2 * np.pi) * np.sqrt((l1 + l2) / (l1 * l2 * cf))
        ideal_type = "LCL Filter"
    else:
        f_res_ideal = 1 / (2 * np.pi) * np.sqrt(1 / (l1 * cf))
        ideal_type = "LC/L Filter"

    f_res_actual = 1 / (2 * np.pi) * np.sqrt((l1 + l2_total) / (l1 * l2_total * cf))

    with c_plot:
        # --- 3.  ---
        st.markdown("#### 🧠 Expert System Diagnostic")
        f_ctrl_bw_high = fsw / 6.0
        f_res_min_limit = 10 * sys_fg
        f_res_max_limit = 0.5 * fsw

        alert_container = st.container()

        if f_res_actual < f_res_min_limit:
            if is_trial:
                alert_container.error("🚨 **High Risk (Low-Frequency Oscillation)**: Excessive external leakage — actual resonance falls below the safe lower limit, risking low-order harmonic amplification and instability. (🔒 Exact frequency values require paid access)")
            else:
                alert_container.error(f"🚨 **High Risk (Low-Frequency Oscillation)**: Equivalent SCR={eq_scr:.1f}, excessive external leakage! Actual frequency ({f_res_actual:.0f} Hz) falls below lower limit ({f_res_min_limit:.0f} Hz) — risk of harmonic amplification.")
        elif f_res_actual < f_ctrl_bw_high:
            if is_trial:
                alert_container.warning("⚠️ **Moderate Risk (Control Coupling)**: Actual resonance frequency is approaching the control bandwidth — phase margin attenuation; active damping algorithm required. (🔒 Exact frequency values require paid access)")
            else:
                alert_container.warning(f"⚠️ **Moderate Risk (Control Coupling)**: Actual frequency ({f_res_actual:.0f} Hz) approaching control bandwidth ({f_ctrl_bw_high:.0f} Hz) — may cause phase attenuation; add active damping algorithm.")
        elif f_res_actual > f_res_max_limit:
            if is_trial:
                alert_container.error("🚨 **High Risk (High-Frequency Excitation)**: Actual resonance frequency exceeds the safe upper limit — will be excited by switching ripple, risking component destruction. (🔒 Exact frequency values require paid access)")
            else:
                alert_container.error(f"🚨 **High Risk (High-Frequency Excitation)**: Actual resonance frequency ({f_res_actual:.0f} Hz) exceeds upper limit ({f_res_max_limit:.0f} Hz) — risk of switching component destruction.")
        else:
            if not has_l2:
                if is_trial:
                    alert_container.success("✅ **CL system safely leveraging leakage**: No internal L2; transformer/grid leakage successfully forms LCL — actual resonance is within the Safe Zone. (🔒 Exact frequency and leakage values require paid access)")
                else:
                    alert_container.success(f"✅ **CL system safely leveraging leakage**: No internal L2; transformer/grid leakage ({lg*1000:.3f} mH) successfully forms LCL. Resonance frequency **{f_res_actual:.0f} Hz** is within the Safe Zone.")
            else:
                if is_trial:
                    alert_container.success('✅ **Firmly within Safety Island**: After frequency drift, resonance is safely bounded within $[10 f_g, 0.5 f_{sw}]$. (🔒 Exact frequency and leakage values require paid access)')
                else:
                    alert_container.success(f'✅ **Firmly within Safety Island**: External leakage = {lg*1000:.3f} mH. Frequency drifts to **{f_res_actual:.0f} Hz**, within $[10 f_g, 0.5 f_{{sw}}]$ Safe Zone.')

        # --- 4.  ---
        with st.expander("📝 Deep Dive: How is external grid impedance / transformer leakage derived?", expanded=False):
            if is_trial:
                trial_lock("Unlock with paid access: complete derivation of equivalent leakage inductance and SCR", compact=True)
            else:
                st.markdown(f"""
                **1. Converter Rated Apparent Power ($S_{{inv}}$)**:
                $$ S_{{inv}} = \\sqrt{{3}} \\times U_{{line}} \\times I_{{line}} = \\sqrt{{3}} \\times {sys_ug} \\times {vec_i} \\approx {s_inv/1000:.1f} \\text{{ kVA}} $$
                """)

                if grid_mode == "Specify Grid SCR directly":
                    st.markdown(f"""
                    **2. Converter Base Impedance ($Z_{{base}}$) and Grid Equivalent Inductance ($L_g$)**:
                    $$ Z_{{base}} = \\frac{{U_{{line}}^2}}{{S_{{inv}}}} = \\frac{{{sys_ug}^2}}{{{s_inv:.1f}}} \\approx {z_base_inv:.4f} \\text{{ }}\\Omega $$
                    $$ Z_g = \\frac{{Z_{{base}}}}{{SCR}} = \\frac{{{z_base_inv:.4f}}}{{{scr_val}}} \\approx {z_base_inv/scr_val:.4f} \\text{{ }}\\Omega $$
                    $$ L_g = \\frac{{Z_g}}{{2\\pi f_g}} = \\frac{{{z_base_inv/scr_val:.4f}}}{{2\\pi \\times {sys_fg}}} \\approx {lg*1000:.3f} \\text{{ mH}} $$
                    """)

                else:
                    # ✅ 1: Parameter,
                    st.markdown(f"""
                    **2. Deriving Leakage Inductance ($L_g$) from Transformer Nameplate Parameters**

                    Transformer rating $S_{{tx}} = {tx_s_kva:.0f}\\text{{ kVA}}$, line voltage $U_{{line}} = {sys_ug}\\text{{ V}}$, short-circuit impedance $u_k = {tx_uk}\\%$.

                    **① Transformer Base Impedance** (defined by transformer nameplate, independent of converter rating):
                    $$ Z_{{base,tx}} = \\frac{{U_{{line}}^2}}{{S_{{tx}}}} = \\frac{{{sys_ug}^2}}{{{tx_s:.0f}}} \\approx {z_base_tx:.4f} \\;\\Omega $$

                    **② Transformer Short-Circuit Impedance** (actual leakage reactance from nameplate $u_k\\%$, resistance neglected):
                    $$ Z_T = u_k\\% \\times Z_{{base,tx}} = {tx_uk/100.0:.3f} \\times {z_base_tx:.4f} \\approx {z_tx_actual:.4f} \\;\\Omega $$

                    **③ Convert to Leakage Inductance** ($\\omega_g = 2\\pi f_g = 2\\pi \\times {sys_fg}$ rad/s):
                    $$ L_g = \\frac{{Z_T}}{{\\omega_g}} = \\frac{{{z_tx_actual:.4f}}}{{2\\pi \\times {sys_fg}}} \\approx \\mathbf{{{lg*1000:.3f}\\text{{ mH}}}} $$

                    > 📌 The above derivation uses only transformer nameplate parameters ($S_{{tx}}$, $U_{{line}}$, $u_k\\%$). $L_g$ is an intrinsic transformer property, independent of the downstream converter rating.

                    **④ Equivalent SCR** (system matching metric, requires converter rating $S_{{inv}}$):
                    $$ SCR_{{eq}} = \\frac{{S_{{tx,sc}}}}{{S_{{inv}}}} = \\frac{{S_{{tx}} / u_k\\%}}{{S_{{inv}}}} = \\frac{{{tx_s:.0f} / {tx_uk/100.0:.3f}}}{{{s_inv:.1f}}} \\approx {eq_scr:.1f} $$

                    > ⚠️ Equivalent SCR reflects the matching between transformer short-circuit capacity and converter rating. It is a system-level metric and does not participate in the leakage inductance derivation.
                    """)

        # --- 5.  Bode  ---
        try:
            w_range = np.logspace(1, 5, 1000) * 2 * np.pi
            freqs_hz = w_range / (2 * np.pi)
            s = 1j * w_range

            if has_l2:
                num_ideal = [rd * cf, 1]
                den_ideal = [l1 * l2 * cf, rd * cf * (l1 + l2), l1 + l2, 0]
            else:
                num_ideal = [1]
                den_ideal = [l1, 0]

            H_ideal = np.polyval(num_ideal, s) / np.polyval(den_ideal, s)
            mag_ideal = 20 * np.log10(np.abs(H_ideal))

            num_actual = [rd * cf, 1]
            den_actual = [l1 * l2_total * cf, rd * cf * (l1 + l2_total), l1 + l2_total, 0]
            H_actual = np.polyval(num_actual, s) / np.polyval(den_actual, s)
            mag_actual = 20 * np.log10(np.abs(H_actual))
            phase_actual = np.angle(H_actual, deg=True)

            fig_bode = make_subplots(rows=2, cols=1, shared_xaxes=True, vertical_spacing=0.08, row_heights=[0.6, 0.4])

            fig_bode.add_trace(go.Scatter(x=freqs_hz, y=mag_ideal, name=f'Ideal Lab Grid ({ideal_type}, Lg=0)',
                                          line=dict(color='gray', dash='dash', width=2), opacity=0.6,
                                          hoverinfo='none' if is_trial else 'all'), row=1, col=1)
            _actual_name = 'Real Field Environment (with leakage)' if is_trial else f'Real Field Environment (with leakage, Lg={lg*1000:.3f}mH)'
            fig_bode.add_trace(go.Scatter(x=freqs_hz, y=mag_actual, name=_actual_name,
                                          line=dict(color='#dc3545', width=3),
                                          hoverinfo='none' if is_trial else 'all'), row=1, col=1)
            fig_bode.add_trace(go.Scatter(x=freqs_hz, y=phase_actual, name='Actual Field Phase',
                                          line=dict(color='#0d6efd', width=2),
                                          hoverinfo='none' if is_trial else 'all'), row=2, col=1)

            y_max_actual = np.nanmax(mag_actual)

            if has_l2:
                y_max_ideal = np.nanmax(mag_ideal)
                fig_bode.add_vline(x=f_res_ideal, line_dash="dot", line_color="gray", row=1, col=1)
                _ideal_txt = "Ideal" if is_trial else f"Ideal<br>{f_res_ideal:.0f}Hz"
                fig_bode.add_annotation(x=np.log10(f_res_ideal), y=y_max_ideal - 10, text=_ideal_txt,
                                        showarrow=True, arrowhead=2, ax=40, ay=30, font=dict(color="gray"), row=1, col=1)

            fig_bode.add_vline(x=f_res_actual, line_dash="dash", line_color="#dc3545", row=1, col=1)
            _actual_txt = "Actual Resonance" if is_trial else f"Actual Resonance<br>{f_res_actual:.0f}Hz"
            fig_bode.add_annotation(x=np.log10(f_res_actual), y=y_max_actual + 5, text=_actual_txt,
                                    showarrow=True, arrowhead=2, ax=-60, ay=-30,
                                    font=dict(color="#dc3545", size=13, weight="bold"), row=1, col=1)

            fig_bode.update_layout(
                height=480, margin=dict(t=50, b=20, l=20, r=20),
                hovermode=False if is_trial else 'x unified',
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="center", x=0.5)
            )
            fig_bode.update_xaxes(type="log", row=1, col=1)
            fig_bode.update_xaxes(type="log", title_text="Frequency (Hz)", row=2, col=1)
            fig_bode.update_yaxes(title_text="Gain Magnitude (dB)", row=1, col=1)
            fig_bode.update_yaxes(title_text="Phase (°)", row=2, col=1)

            if is_trial:
                st.markdown(
                    """
                    <style>
                    .bode-mask-wrap { position: relative; }
                    .bode-mask-overlay {
                        position: absolute; top: 0; left: 0; width: 100%; height: 100%;
                        z-index: 10; border-radius: 8px; pointer-events: none;
                        background: repeating-linear-gradient(
                            45deg, rgba(245,245,245,0.45), rgba(245,245,245,0.45) 12px,
                            rgba(200,200,200,0.45) 12px, rgba(200,200,200,0.45) 24px);
                        backdrop-filter: blur(2px);
                        display: flex; align-items: center; justify-content: center;
                    }
                    .bode-mask-badge {
                        background: rgba(255,255,255,0.92); border:1.5px dashed #faad14;
                        border-radius: 10px; padding: 14px 22px; text-align:center;
                    }
                    </style>
                    <div class="bode-mask-wrap"><div class="bode-mask-overlay">
                        <div class="bode-mask-badge">
                            <div style="font-size:22px;">🔒</div>
                            <div style="font-weight:bold; font-size:13px; color:#333;">Unlock full frequency response data with paid access</div>
                            <div style="color:#888; font-size:11px; margin-top:2px;">Preview only — one-time payment, lifetime access</div>
                        </div>
                    </div></div>
                    """,
                    unsafe_allow_html=True
                )
                st.plotly_chart(fig_bode, use_container_width=True)
                trial_lock("Unlock with paid access: complete resonance frequency response (including precise resonance point and leakage inductance values)", compact=True)
            else:
                st.plotly_chart(fig_bode, use_container_width=True)

        except Exception as e:
            st.error(f"❌ Chart rendering failed: {str(e)}")

    st.divider()

    # --- 6.  ---
    st.markdown("### 📖 Key Rules: Why Do High-Power Converters Prefer CL Topology?")
    c_edu1, c_edu2 = st.columns(2)
    with c_edu1:
        st.info('**1. The Art of "Free Lunch" with CL Topology**')
        st.markdown("""
        *   **Engineering Challenge**: For MW-scale wind/PV converters, even a fraction of a millihenry of inductance represents enormous volume and cost.
 * **Clever Design**: Since the converter will always connect to a step-up transformer (e.g. 690V to 35kV) which inherently has leakage reactance ($u_k\\%$), engineers simply **omit the internal $L_2$** and keep only $L_1$ $C_f$( CL topology).
        *   **Result**: When the converter pairs with the transformer, its leakage inductance perfectly substitutes for $L_2$, and the system naturally forms a complete LCL structure! The waveform in this tool beautifully demonstrates this physical transition from ideal monotonic attenuation to real-world resonant attenuation.
        """)
    with c_edu2:
        st.warning('**2. The LCL "Safety Island" Rule**')
        st.markdown(r"""
        Whether your LCL is formed via SCR or transformer leakage inductance, the actual resonance frequency $f_{res}'$ must be firmly bounded within:
        $$ 10 f_g < f_{res}' < 0.5 f_{sw} $$
        *   **Lower bound ($10 f_g$)**: Avoids superposition resonance with low-order background harmonics (5th, 7th, 11th, etc.) and ensures the grid current control loop has sufficient dynamic response bandwidth.
        *   **Upper bound ($0.5 f_{sw}$)**: Satisfies Nyquist sampling theorem. If too high, it is easily excited by high-frequency sideband harmonics from switching, causing uncontrollable high-frequency noise or even component destruction.
        """)










# ==========================================
#  4: Four-Quadrant Operation(v4 )
#
# ★ (2/4): 
#   : VSC — jX_L — Grid
#   KVL(I_gen  = VSC→Grid): 
#     V_s = V_on + V_L, V_L = jX_L · I_gen   ← , 
#
# ★ (2, ): 
#   O →  = V_on(Grid Voltage, )
#    → P  = V_L = jX_L·I_gen(Leading I_gen 90°)
#   O → P     = V_s(Converter Voltage)
#
#   ABCD(2): 
#     A = , P → V_L → I_gen(LeadingV_on 90°)
#     B = , P → V_L → I_gen(I_genV_on=0°)
#         2(b)=I_load=I_gen(I_genV_on)
#         → B""(I_load, I_gen)
#     C = , P → V_L → I_gen(LaggingV_on 90°)
#     D = , P → V_L → I_gen(I_genV_on→)
#
#   (2(b)): 
#     2(b): PB(), V_L = CTR→P = ()
#     I_gen = V_L / (jX_L) = (-j) / (j) = -1 → ()= V_on
#     I_load = -I_gen = ()= V_on →  ✓
#     2(a): PA(), V_L = CTR→P = ()
#     I_gen = (-1)/(j) = j → ()= LeadingV_on 90°
#     I_load =  ✓(i_s)
#
# ★ ABCD (: 4): 
#   C():  → VSC → Grid Voltage
#     I_gen LaggingV_on 90°(); Q_gen>0()
#     I_load LeadingV_on 90°(); Q_load<0
#   A():  → VSC → Grid Voltage
#     I_gen LeadingV_on 90°(); Q_gen<0()
#     I_load LaggingV_on 90°(); Q_load>0()
#   D():  → VSC
#     I_gen V_on (); P_gen>0()
#     I_load V_on (); P_load<0
#   Point B (bottom of circle): Motoring → VSC absorbs active power from grid
#     I_gen anti-phase with V_on (leftward); P_gen<0
#     I_load in-phase with V_on (rightward); P_load>0 (absorbs active power)
#
# ★ Power calculation (from S = V_on · I_gen*):
#   Let I_gen = |I|·e^(j·θ), θ = I_gen phase angle (positive = Leading V_on = CCW)
#   S_gen = V_on · I_gen* = |V_on|·|I|·e^(-jθ)
#   P_gen = |V_on|·|I|·cos(θ)
#   Q_gen = -|V_on|·|I|·sin(θ)    ← note the negative sign!
#   Verify point C: θ = -90° (I_gen Lagging V_on 90°)
#     Q_gen = -|V_on||I|·sin(-90°) = +|V_on||I| > 0 ✓ (generates reactive power)
#   Verify point A: θ = +90° (I_gen Leading V_on 90°)
#     Q_gen = -|V_on||I|·sin(+90°) = -|V_on||I| < 0 ✓ (absorbs reactive power)
#   Load convention: P_load = -P_gen, Q_load = -Q_gen
#
# ★ Physical meaning of phase angle θ and input convention:
#   θ > 0: I_gen Leading V_on → Q_gen < 0 (absorbs reactive, under-excited, lowers grid voltage)
#   θ < 0: I_gen Lagging V_on → Q_gen > 0 (generates reactive, over-excited, raises grid voltage)
#   Consistent with synchronous generator convention: over-excited (gen reactive) = current Lagging voltage
#
# ★ Load convention phasor diagram (key correction):
#   I_load = -I_gen, I_load is opposite to I_gen in the complex plane
#   In the phasor diagram:
#     Load convention: i_s arrow points from outside to point P (Grid→VSC, i.e. I_load direction)
#     I_load angle = θ + π (opposite to I_gen)
#     Arrow start = P - i_len·(I_load direction) = P + i_len·(I_gen direction)
#     Arrow end = P
#   Verify point C (Appendix Fig.1): P at right, I_gen downward (θ=-90°),
#     I_load upward, arrow points from above P to P ✓ (matches orange arrow in Appendix Fig.1)
#
# ★ Reactive power vs. Grid Voltage relationship (authoritative source: Appendix Fig.4;
#   Kundur, Power System Stability and Control, McGraw-Hill 1994, §10.2): 
#   Generate reactive (over-excited, I_gen Lagging, Q_gen>0) → raises grid voltage
#   Absorb reactive (under-excited, I_gen Leading, Q_gen<0) → lowers grid voltage
# ==========================================
elif selection == nav_options[3]:
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches
    import pandas as pd

    plt.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial', 'SimHei']
    plt.rcParams['axes.unicode_minus'] = False

    st.header("🔄 Four-Quadrant Operation & Reactive Power Limit Analysis")

    # =================================================================
    # 1. Operating parameter input
    # =================================================================
    with st.container():
        st.subheader("🎛️ Operating Parameter Input")

        def_ug  = st.session_state.get('sys_ug',   690.0)
        def_i   = st.session_state.get('vec_i',    500.0)
        def_l   = st.session_state.get('lcl_l1',   0.5)
        def_udc = st.session_state.get('vec_udc', 1100.0)

        c_in1, c_in2, c_in3 = st.columns(3)

        with c_in1:
            st.markdown("**1. Grid Settings**")
            u_line_in  = st.number_input("Grid Line Voltage (Vrms)", value=def_ug, step=10.0, key="m4_ug_line")
            sys_freq   = st.number_input("System Frequency (Hz)", value=50.0, step=1.0, key="m4_freq")
            u_grid_rms = u_line_in / np.sqrt(3)

        with c_in2:
            st.markdown("**2. Current Parameters**")
            i_rated    = st.number_input("Output Phase Current RMS (Arms)  [Y: phase = line current]", value=def_i, step=10.0, key="m4_i",
                                          help="Enter phase current RMS value (Arms). Phasor diagram uses RMS for both voltage and current. P = V_rms × I_rms × cosθ.")
            input_mode = st.radio("Phase Input Method", ["Current phase φ (°)", "Power factor PF"],
                                  horizontal=True, key="m4_input_mode")
            if input_mode == "Current phase φ (°)":
                pf_angle_deg = st.slider(
                    "I_gen Phase θ (°)", -180, 180, -30, 5, key="m4_phi",
                    help="θ = I_gen phase angle relative to V_on (CCW positive).\n"
                         "θ < 0: I_gen Lagging V_on → Q_gen > 0 (generates reactive, over-excited, raises grid voltage)\n"
                         "θ > 0: I_gen Leading V_on → Q_gen < 0 (absorbs reactive, under-excited, lowers grid voltage)\n"
                         "θ = 0: Point D, pure active power generation (P_gen > 0)\n"
                         "θ = ±180°: Point B, pure active power absorption (P_gen < 0)")
                pf_display = abs(np.cos(np.radians(pf_angle_deg)))
                if pf_angle_deg < -0.5:
                    ll_str = f"Lagging {abs(pf_angle_deg):.0f}° (over-excited, generates reactive, raises grid voltage)"
                elif pf_angle_deg > 0.5:
                    ll_str = f"Leading {pf_angle_deg:.0f}° (under-excited, absorbs reactive, lowers grid voltage)"
                else:
                    ll_str = "In-phase (Point D, pure active generation)"
                st.caption(f"PF = **{pf_display:.3f}**, I_gen {ll_str}")
            else:
                pf_input    = st.number_input("Power Factor PF (0~1)", value=0.866,
                    min_value=0.0, max_value=1.0, step=0.01, format="%.3f", key="m4_pf_val")
                pf_lead_lag = st.radio(
                    "I_gen relative to V_on (generator convention)",
                    ["Lagging", "Leading"],
                    horizontal=True, key="m4_leadlag")
                pf_angle_deg = np.degrees(np.arccos(np.clip(pf_input, 0, 1)))
                if "Leading" in pf_lead_lag:
                    pass           # θ > 0, Leading
                else:
                    pf_angle_deg = -pf_angle_deg   # θ < 0, Lagging
                st.caption(f"θ = **{pf_angle_deg:+.1f}°** (I_gen relative to V_on; negative = Lagging = over-excited)")

        with c_in3:
            st.markdown("**3. Hardware Constraints**")
            l_filter_mh = st.number_input("L1 Inductance (mH)", value=def_l, step=0.01,
                format="%.3f", key="m4_l", help="Inverter-side inductor")
            u_dc = st.number_input("DC Bus Voltage (V)", value=def_udc, step=50.0, key="m4_vdc")
            modulation = st.selectbox("Modulation Strategy",
                options=["SPWM (Sinusoidal PWM)", "SVPWM (Space Vector)", "Over-modulation (estimate only)"],
                index=1,
                help="• SPWM: Phase voltage limit = Vdc/(2√2) Vrms\n"
                     "• SVPWM: Limit = Vdc/√6 Vrms, 15.47% higher than SPWM\n"
                     "• Over-modulation: beyond linear region, estimate only")

        # ── L2 / Transformer Leakage ────────────────────────────────────
        st.markdown("**4. L2 / Transformer Leakage Settings**")
        l2_cols = st.columns([1, 1, 1])
        with l2_cols[0]:
            l2_method = st.radio("L2 Input Method",
                ["Enter L2 directly (mH)", "Calculate from transformer parameters"], key="m4_l2_method")
        with l2_cols[1]:
            if l2_method == "Enter L2 directly (mH)":
                l2_mh    = st.number_input("L2 Inductance (mH)", value=0.1, step=0.01,
                    format="%.3f", key="m4_l2_direct", help="Grid-side inductor or transformer leakage equivalent")
                l2_source = f"Direct input: {l2_mh:.3f} mH"
            else:
                tx_kva = st.number_input("Transformer Rating (kVA)", value=2000.0, step=100.0, key="m4_tx_kva")
                uk_pct = st.number_input("Short-Circuit Impedance uk (%)", value=6.0, step=0.5, key="m4_uk_pct",
                    help="Transformer nameplate short-circuit impedance, typically 4%–8%")
                xk_accurate = (uk_pct/100.0) * (u_line_in**2) / (tx_kva*1e3) if tx_kva > 0 else 1e-6
                omega_tmp   = 2 * np.pi * sys_freq
                l2_mh       = (xk_accurate / omega_tmp) * 1000
                l2_source   = f"From transformer: Xk={xk_accurate:.4f} Ω → L2={l2_mh:.3f} mH"
        with l2_cols[2]:
            if is_trial:
                trial_lock("Unlock with paid access: L2 calculation result and total inductance", compact=True)
            else:
                st.info(f"**L2 Calculation Result**\n\n{l2_source}\n\n> L_total = L1 + L2")

        # ── Core parameter calculations ─────────────────────────────────────
        omega    = 2 * np.pi * sys_freq
        l1_henry = l_filter_mh * 1e-3
        l2_henry = l2_mh * 1e-3
        l_total  = l1_henry + l2_henry
        xl       = omega * l_total if l_total > 0 else 1e-6
        xl1      = omega * l1_henry
        xl2      = omega * l2_henry

        # θ = pf_angle_deg (I_gen phase angle, θ<0=Lagging=over-excited=generates reactive=Q_gen>0)
        theta    = np.radians(pf_angle_deg)   # I_gen phase angle (complex plane, CCW positive)
        I_mag    = abs(i_rated)               # |I_gen| RMS value (same unit as V_on, P=V_rms×I_rms×cosθ)

        # Grid voltage V_on set as positive real axis reference
        vec_vg    = u_grid_rms + 0j
        # I_gen phasor = I_mag · e^(j·θ)
        vec_i_gen = I_mag * np.exp(1j * theta)
        # V_L = jX_L · I_gen (leads I_gen by 90°, i.e. V_L angle = θ + π/2)
        vec_vl    = 1j * xl * vec_i_gen
        # V_s = V_on + V_L(KVL)
        vec_vinv  = vec_vg + vec_vl

        req_v_phase = np.abs(vec_vinv)

        if "SPWM" in modulation and "SVPWM" not in modulation:
            limit_v, mod_label = u_dc / (2*np.sqrt(2)), "SPWM"
        elif "Over-modulation" in modulation:
            limit_v, mod_label = (4/np.pi)*u_dc/(2*np.sqrt(2)), "Over-modulation"
        else:
            limit_v, mod_label = u_dc / np.sqrt(6), "SVPWM"

        # ── Results preview ─────────────────────────────────────────────────
        delta_v = limit_v - req_v_phase
        st.divider()
        st.markdown("**5. Results Preview**")
        if is_trial:
            trial_lock("Calculation results require paid access")
        else:
            prev_cols = st.columns(3)
            with prev_cols[0]:
                st.metric("Required Converter Phase Voltage", f"{req_v_phase:.1f} Vrms",
                          help="Phase voltage Line-to-Neutral RMS")
            with prev_cols[1]:
                if delta_v >= 0:
                    st.metric(f"{mod_label} Voltage Margin", f"+{delta_v:.1f} V",
                              delta=f"Limit {limit_v:.1f} V", delta_color="normal")
                else:
                    st.metric(f"{mod_label} Voltage Margin", f"{delta_v:.1f} V",
                              delta=f"Limit {limit_v:.1f} V", delta_color="inverse")
            with prev_cols[2]:
                st.metric("Equivalent Total Reactance XL", f"{xl:.4f} Ω",
                          help=f"XL1={xl1:.4f} Ω + XL2={xl2:.4f} Ω")

        with st.expander("📐 View Calculation Details"):
            if is_trial:
                trial_lock("Unlock with paid access: complete four-quadrant calculation details and parameter table", compact=True)
            else:
              st.markdown(
                "**Physical Equation (KVL, I_gen = VSC→Grid):**\n\n"
                "$$V_s = V_{on} + jX_L \\cdot I_{gen}$$\n\n"
                "**Power (S = V_on · I_gen*):**\n\n"
                "$$P_{gen} = |V_{on}||I_{gen}|\\cos\\theta, "
                "\\quad Q_{gen} = -|V_{on}||I_{gen}|\\sin\\theta$$\n\n"
                "Note: θ<0 (I_gen Lagging) → Q_gen>0 (generates reactive);"
                "θ>0 (I_gen Leading) → Q_gen<0 (absorbs reactive)"
              )
              calc_data = {
                "Parameter": ["Grid Line Voltage V_line", "Grid Phase Voltage V_on", "L1 Filter Inductor", "L2/Transformer Leakage",
                         "L_total", "Total Reactance X_L",
                         "|I_gen| Phase Current RMS (=line current, Y)", "θ (I_gen phase, negative=Lagging)",
                         "V_L real part", "V_L imag part", "Required Converter Phase Voltage |V_s|",
                         f"{mod_label} Phase Voltage Limit", "Voltage Margin",
                         "3-phase Total Active Power P (=3×V_on×I×cosθ)",
                         "3-phase Total Reactive Power Q (=3×V_on×I×|sinθ|)",
                         "Power factor PF"],
                "Value": [f"{u_line_in:.1f} Vrms", f"{u_grid_rms:.2f} Vrms",
                         f"{l_filter_mh:.3f} mH", f"{l2_mh:.3f} mH",
                         f"{l_total*1000:.3f} mH", f"{xl:.4f} Ω",
                         f"{I_mag:.1f} Arms", f"{pf_angle_deg:+.1f}°",
                         f"{vec_vl.real:.3f} V", f"{vec_vl.imag:.3f} V",
                         f"{req_v_phase:.2f} Vrms", f"{limit_v:.2f} Vrms",
                         f"{delta_v:+.2f} V {chr(9989) if delta_v >= 0 else chr(128680)}",
                         f"{3*u_grid_rms*I_mag*np.cos(theta)/1000:+.3f} kW",
                         f"{-3*u_grid_rms*I_mag*np.sin(theta)/1000:+.3f} kvar",
                         f"{abs(np.cos(theta)):.4f}"]
              }
              st.table(pd.DataFrame(calc_data))
              st.markdown(
                f"> SPWM limit: Vdc/(2√2) = **{u_dc/(2*np.sqrt(2)):.2f} Vrms**\n\n"
                f"> SVPWM limit: Vdc/√6 = **{u_dc/np.sqrt(6):.2f} Vrms**\n\n"
                f"> Reference: Holmes & Lipo, *PWM for Power Converters*, IEEE Press 2003, Ch.3"
              )

    st.divider()

    # =================================================================
    # 2. 
    #
    # S_gen = V_on · I_gen* = |V_on|·|I_gen|·e^(-jθ)
    # P_gen =  |V_on|·|I_gen|·cos(θ)
    # Q_gen = -|V_on|·|I_gen|·sin(θ)
    #
    # ★ (4): 
    #   C θ=-90°: P_gen=0, Q_gen=-Vg·I·sin(-90°)=+Vg·I>0 ✓(, )
    #   A θ=+90°: P_gen=0, Q_gen=-Vg·I·sin(+90°)=-Vg·I<0 ✓(, )
    #   D θ=  0°: P_gen=+Vg·I>0, Q_gen=0 ✓()
    #   B θ=±180°: P_gen=-Vg·I<0, Q_gen=0 ✓()
    # =================================================================
    P_gen  =  u_grid_rms * I_mag * np.cos(theta)   # single-phaseactive(W)
    Q_gen  = -u_grid_rms * I_mag * np.sin(theta)   # single-phasereactive(var)
    P_load = -P_gen
    Q_load = -Q_gen
    pf_val = abs(np.cos(theta))
    # ── ()──────────────────────────────
    P_gen_3ph  = 3 * P_gen    # 3-phase total active power (W) = √3×V_line×I×cosθ
    Q_gen_3ph  = 3 * Q_gen    # 3-phase total reactive power (var)
    P_load_3ph = 3 * P_load
    Q_load_3ph = 3 * Q_load

    # (, )
    theta_deg = np.degrees(theta)
    eps = 2.0  # threshold (degrees)
    # _phys_state: 
    if I_mag < 1e-3:
        _phys_state = "zero"
    elif abs(theta_deg) < eps:
        _phys_state = "pure_p_gen"    # Point D: pure active generation
    elif abs(abs(theta_deg) - 180) < eps:
        _phys_state = "pure_p_load"   # Point B: pure active absorption
    elif theta_deg < 0:
        _phys_state = "over_exc"      # Over-excited: I_gen Lagging, generates reactive, raises voltage
    else:
        _phys_state = "under_exc"     # Under-excited: I_gen Leading, absorbs reactive, lowers voltage

    # =================================================================
    # 3.  Tabs
    # =================================================================
    tab1, tab2, tab3 = st.tabs(["📈 1. Waveform View", "📐 2. Voltage Triangle", "🕸️ 3. PQ Operating Circles"])

    # ─────────────────────────────────────────────────────────────────
    # Tab 1: 
    # ─────────────────────────────────────────────────────────────────
    with tab1:

        # ──  ──────────────────────────────────────────────────
        convention = st.radio(
            "📐 Display Convention (affects value signs, current direction, power flow diagnosis)",
            ["⚡ Generator Convention (IEEE 1547, I_gen: VSC→Grid)",
             "🔌 Load Convention (IEC 61400-21, I_load: Grid→VSC)"],
            horizontal=True, key="m4_convention"
        )
        use_gen = "Generator" in convention

        # 
        # 
        P_disp = P_gen_3ph  if use_gen else P_load_3ph
        Q_disp = Q_gen_3ph  if use_gen else Q_load_3ph
        p_disp_kw   = P_disp / 1000
        q_disp_kvar = Q_disp / 1000

        # ── ( i_s, )──────
        # : i_s = I_gen, θ<0=i_sLaggingV_on
        # :   i_s = I_load = -I_gen, i_s = θ+180°
        #   θ=-90° → i_s(load)LeadingV_on 90°(C, , )
        #   θ=+90° → i_s(load)LaggingV_on 90°(A, , )
        if _phys_state == "zero":
            phase_desc_phys = "i_s = 0 (no current)"
        elif _phys_state == "pure_p_gen":
            if use_gen:
                phase_desc_phys = "i_s in-phase with V_on (Point D, pure active generation, P_gen>0)"
            else:
                phase_desc_phys = "i_s anti-phase with V_on (Point D, load convention, P_load<0, VSC sends active to grid)"
        elif _phys_state == "pure_p_load":
            if use_gen:
                phase_desc_phys = "i_s anti-phase with V_on (Point B, pure active absorption, P_gen<0)"
            else:
                phase_desc_phys = "i_s in-phase with V_on (Point B, P_load>0, VSC absorbs active from grid)"
        elif _phys_state == "over_exc":
            # : I_genLaggingθ°, I_loadLeadingθ°()
            if use_gen:
                phase_desc_phys = (f"i_s Lagging V_on {abs(theta_deg):.1f}°"
                                   f"(over-excited, VSC generates reactive, **raises** inductive grid voltage)")
            else:
                phase_desc_phys = (f"i_s Leading V_on {abs(theta_deg):.1f}°"
                                   f"(load view: i_load Leading, over-excited, VSC generates reactive, **raises** inductive grid voltage)")
        else:  # under_exc
            # : I_genLeadingθ°, I_loadLaggingθ°
            if use_gen:
                phase_desc_phys = (f"i_s Leading V_on {abs(theta_deg):.1f}°"
                                   f"(under-excited, VSC absorbs reactive, **lowers** inductive grid voltage)")
            else:
                phase_desc_phys = (f"i_s Lagging V_on {abs(theta_deg):.1f}°"
                                   f"(load view: i_load Lagging, under-excited, VSC absorbs reactive, **lowers** inductive grid voltage)")

        if use_gen:
            st.info(
                "**Generator Convention (IEEE 1547-2018 §4.1)**\n\n"
                "| Sign | Meaning | Reactive Effect |\n"
                "|:---:|:---|:---:|\n"
                "| P > 0 | VSC **sends** active to grid | — |\n"
                "| Q > 0 | VSC **generates** reactive (I_gen **Lagging** V_on, over-excited) | **Raises** Grid Voltage |\n"
                "| Q < 0 | VSC **absorbs** reactive (I_gen **Leading** V_on, under-excited) | **Lowers** Grid Voltage |",
                icon="⚡"
            )
            st.caption("⚡ All power values are **3-phase total** (= 3 × V_phase × I_phase × cosθ = √3 × V_line × I_line × cosθ)")
        else:
            st.info(
                "**Load Convention (IEC 61400-21-1:2021 §4.2; GB/T 19963-2021 §6.3)**\n\n"
                "| Sign | Meaning | Reactive Effect |\n"
                "|:---:|:---|:---:|\n"
                "| P > 0 | VSC **absorbs active from grid** | — |\n"
                "| Q > 0 | VSC **absorbs** reactive from grid (I_load **Lagging** V_on, under-excited) | **Lowers** grid voltage |\n"
                "| Q < 0 | VSC **sends** reactive to grid (I_load **Leading** V_on, over-excited) | **Raises** grid voltage |",
                icon="🔌"
            )

        # ══════════════════════════════════════════════════════════════
        # ① 
        #
        # Core geometry (calibrated against Appendix Fig.2):
        # V_on = positive real axis (O→center)
        # V_L = jX_L·I_gen, angle = θ + π/2
        # Point P = CTR + R·[cos(θ+π/2), sin(θ+π/2)]
        #
        # ★ Load convention geometry correction (v4 core fix):
        # I_load = -I_gen, angle = θ + π
        # Arrow points toward P, start = P + i_len·[cos(θ), sin(θ)]
        #   = P + i_len·I_gen direction (from I_gen endpoint toward P)
        # i_s(load) points from I_gen side of P to P, direction = I_load ✓
        #
        # Verify point C (θ=-90°, Appendix Fig.1):
        #   P at right of circle, I_gen downward (-j direction)
        #   I_load upward (+j), arrow from above P to P ✓ (matches orange arrow Appendix Fig.1)
        # ══════════════════════════════════════════════════════════════
        st.markdown("#### 📡 Live Phasor Diagram (Dual Convention View)")
        st.caption(
            "**Key difference between conventions**: The underlying physics is identical — only the reference positive direction of current i_s differs."
            "Generator convention (left): i_s positive from converter to grid."
            "Load convention (right): i_s positive from grid to converter — so i_s arrows are reversed and P/Q signs are opposite."
            "Voltage vectors V_on and V_s are geometrically identical in both views; V_L arrow reverses due to KVL sign convention."
        )

        # Normalized geometry
        R   = 1.0
        CTR = np.array([0.0, 0.0])       # Circle center = V_on endpoint
        O   = np.array([-1.6*R, 0.0])    # Origin = V_on start point

        # Point P: V_L = jX_L·I_gen, angle(V_L) = θ + π/2
        theta_P = theta + np.pi / 2
        P = CTR + R * np.array([np.cos(theta_P), np.sin(theta_P)])

        i_len = 0.70 * R

        # I_gen unit direction vector (angle = θ)
        I_gen_dir = np.array([np.cos(theta), np.sin(theta)])

        # Generator convention: i_s from P, endpoint = P + i_len·I_gen_dir
        I_gen_end   = P + i_len * I_gen_dir

        # Load convention: I_load = -I_gen, i_s points toward P from outside
        # Start = P + i_len·I_gen_dir (coincides with generator arrow end, reversed direction to P)
        I_load_start = P + i_len * I_gen_dir   # Arrow start (shoots toward P)
        # Equivalent: I_load_start = I_gen_end

        # Unified viewport
        all_pts = np.vstack([O, CTR, P, I_gen_end,
                             CTR+[R,0], CTR+[-R,0], CTR+[0,R], CTR+[0,-R]])
        xs, ys   = all_pts[:,0], all_pts[:,1]
        cx = (xs.min()+xs.max())/2
        cy = (ys.min()+ys.max())/2
        half = max(xs.max()-xs.min(), ys.max()-ys.min())/2 * 1.38
        xlo, xhi = cx-half, cx+half
        ylo, yhi = cy-half, cy+half
        y_leg = ylo - 0.40*(yhi-ylo)
        span  = yhi - y_leg

        # Helper: draw arrow
        def _arr(fig, x0, y0, x1, y1, color, width=2.5):
            fig.add_annotation(
                x=x1, y=y1, ax=x0, ay=y0,
                xref='x', yref='y', axref='x', ayref='y',
                showarrow=True, arrowhead=3, arrowsize=1.5,
                arrowwidth=width, arrowcolor=color, text=""
            )
            fig.add_trace(go.Scatter(
                x=[x0, x1], y=[y0, y1], mode='lines',
                line=dict(color=color, width=width),
                hoverinfo='skip', showlegend=False
            ))

        # Helper: operating state label (physics-based, convention-independent)
        def _state_label():
            if I_mag < 1e-3:
                return "Zero current", "#888"
            S_base = u_grid_rms * I_mag
            thr = 0.05 * S_base
            if abs(Q_gen) < thr and P_gen > thr:
                return "Near Point D: Pure active generation (P>0, Q≈0)", "#198754"
            elif abs(Q_gen) < thr and P_gen < -thr:
                return "Near Point B: Pure active absorption (P<0, Q≈0)", "#dc3545"
            elif abs(P_gen) < thr and Q_gen > thr:
                return "Near Point C: Over-excited, generates reactive, raises grid voltage", "#0d6efd"
            elif abs(P_gen) < thr and Q_gen < -thr:
                return "Near Point A: Under-excited, absorbs reactive, lowers grid voltage", "#fd7e14"
            elif P_gen > 0 and Q_gen > 0:
                return "Gen active + Gen reactive (over-excited, raises voltage)", "#0d6efd"
            elif P_gen > 0 and Q_gen < 0:
                return "Gen active + Absorb reactive (under-excited, lowers voltage)", "#fd7e14"
            elif P_gen < 0 and Q_gen > 0:
                return "Absorb active + Gen reactive (over-excited, raises voltage)", "#6f42c1"
            else:
                return "Absorb active + Absorb reactive (under-excited, lowers voltage)", "#dc3545"

        state_txt, state_clr = _state_label()

        # Helper: draw shared elements (circle+ABCD+V_on+V_s+P point+state label)
        # V_L direction differs per convention, drawn separately
        def _draw_common(fig):
            # PWM limit circle
            ct = np.linspace(0, 2*np.pi, 300)
            fig.add_trace(go.Scatter(
                x=CTR[0]+R*np.cos(ct), y=CTR[1]+R*np.sin(ct),
                mode='lines', line=dict(color='#333', width=1.8),
                hoverinfo='skip', showlegend=False
            ))
            # ABCD labels (A=left, B=bottom, C=right, D=top)
            tk = 0.045 * R
            abcd = [
                (CTR[0]-R, CTR[1],    CTR[0]-R-0.13, CTR[1],         'A', 'right',  tk, 0),
                (CTR[0],   CTR[1]-R,  CTR[0],        CTR[1]-R-0.15,  'B', 'center', 0,  tk),
                (CTR[0]+R, CTR[1],    CTR[0]+R+0.13, CTR[1],         'C', 'left',   tk, 0),
                (CTR[0],   CTR[1]+R,  CTR[0],        CTR[1]+R+0.15,  'D', 'center', 0,  tk),
            ]
            for (px, py, tx, ty, lbl, anc, dx, dy) in abcd:
                fig.add_shape(type='line',
                    x0=px-dx, y0=py-dy, x1=px+dx, y1=py+dy,
                    line=dict(color='#555', width=1.5))
                fig.add_annotation(x=tx, y=ty, text=f"<b>{lbl}</b>",
                    showarrow=False, font=dict(size=13, color='#333'), xanchor=anc)
            # V_on: O → center (red, same in both views)
            _arr(fig, O[0], O[1], CTR[0], CTR[1], '#1a7a1a', width=2.8)
            fig.add_annotation(
                x=(O[0]+CTR[0])/2, y=(O[1]+CTR[1])/2 - 0.14,
                text="<b><i>V</i><sub>on</sub></b>",
                showarrow=False, font=dict(size=14, color='#1a7a1a'))
            # V_s: O → P (dark gray, same in both views)
            _arr(fig, O[0], O[1], P[0], P[1], '#555555', width=1.8)
            vs_mid = (np.array([O[0], O[1]]) + P) / 2
            vs_d   = P - np.array([O[0], O[1]])
            vs_n   = np.array([-vs_d[1], vs_d[0]]) / (np.linalg.norm(vs_d)+1e-9) * 0.14
            fig.add_annotation(
                x=vs_mid[0]+vs_n[0], y=vs_mid[1]+vs_n[1],
                text="<b><i>V</i><sub>s</sub></b>",
                showarrow=False, font=dict(size=14, color='#555555'))
            # Point P marker (same in both views)
            fig.add_trace(go.Scatter(
                x=[P[0]], y=[P[1]], mode='markers',
                marker=dict(size=9, color='#222'), hoverinfo='skip', showlegend=False
            ))
            # Operating state label (top-center, avoids overlap with P and info box)
            fig.add_annotation(
                x=0.5, y=1.0,
                xref='paper', yref='paper',
                text=f"<b>{state_txt}</b>",
                showarrow=False, xanchor='center', yanchor='top',
                font=dict(size=10, color=state_clr),
                bgcolor="rgba(255,255,255,0.93)", bordercolor=state_clr, borderwidth=1)
            return fig

        # V_L label offset (perpendicular to V_L direction)
        vl_d_gen = P - CTR                        # Generator: center→P direction
        vl_n_gen = np.array([-vl_d_gen[1], vl_d_gen[0]]) / (np.linalg.norm(vl_d_gen)+1e-9) * 0.14

        # ── Left: Generator Convention ──────────────────────────────────
        # V_L generator convention: center → P (= jX_L·I_gen, leading I_gen by 90°)
        # i_s: from P toward I_gen_end (VSC→Grid, angle=θ)
        fig_gen = go.Figure()
        _draw_common(fig_gen)
        # V_L (generator convention): center → P
        _arr(fig_gen, CTR[0], CTR[1], P[0], P[1], '#111111', width=2.2)
        vl_mid_gen = (CTR + P) / 2
        fig_gen.add_annotation(
            x=vl_mid_gen[0]+vl_n_gen[0], y=vl_mid_gen[1]+vl_n_gen[1],
            text="<b><i>V</i><sub>L</sub></b>",
            showarrow=False, font=dict(size=14, color='#111111'))
        # i_s (generator convention): P → I_gen_end (VSC→Grid, angle=θ)
        _arr(fig_gen, P[0], P[1], I_gen_end[0], I_gen_end[1], '#1155cc', width=2.5)
        off_gen = np.array([np.cos(theta), np.sin(theta)]) * 0.11
        fig_gen.add_annotation(
            x=I_gen_end[0]+off_gen[0], y=I_gen_end[1]+off_gen[1],
            text="<b><i>i</i><sub>s</sub></b>",
            showarrow=False, font=dict(size=14, color='#1155cc'))

        # Legend
        for idx, (c, lbl) in enumerate([
            ('#1155cc', 'i_s (P→out, VSC→Grid, angle=θ)'),
            ('#111111', 'V_L = jX_L·I_gen (center→P)'),
            ('#555555', 'V_s (O→P, Converter Output Voltage)'),
            ('#1a7a1a', 'V_on (O→center, Grid Voltage, reference)'),
        ]):
            fig_gen.add_annotation(
                x=xlo+0.03*(xhi-xlo),
                y=y_leg+0.03*span+idx*0.055*span,
                text=(f'<span style="color:{c};font-size:14px;font-weight:bold;">━</span>'
                      f'<span style="font-size:11px;color:#222;"> {lbl}</span>'),
                showarrow=False, xanchor='left', yanchor='bottom',
                font=dict(size=11), bgcolor='rgba(255,255,255,0)', borderwidth=0)

        fig_gen.add_annotation(
            x=0.01, y=0.32,
            xref='paper', yref='paper',
            text=("<b>Generator Convention</b><br>"
                  "P = 🔒 kW (3-phase)<br>"
                  "Q = 🔒 kvar (3-phase)<br>"
                  "KVL: V_s = V_on + V_L<br>"
                  "i_s: VSC→Grid") if is_trial else
                 (f"<b>Generator Convention</b><br>"
                  f"P = {P_gen_3ph/1000:+.3f} kW (3-phase)<br>"
                  f"Q = {Q_gen_3ph/1000:+.3f} kvar (3-phase)<br>"
                  f"KVL: V_s = V_on + V_L<br>"
                  f"i_s: VSC→Grid"),
            showarrow=False, xanchor='left', yanchor='bottom',
            font=dict(size=10, color='#0d4ea0'),
            bgcolor="rgba(230,241,251,0.95)", bordercolor='#1155cc', borderwidth=1,
            align='left')
        fig_gen.update_layout(
            xaxis=dict(range=[xlo, xhi], visible=False, showgrid=False, zeroline=False),
            yaxis=dict(range=[y_leg, yhi], visible=False, showgrid=False, zeroline=False,
                       scaleanchor='x', scaleratio=1),
            height=510, margin=dict(l=5, r=5, t=42, b=5),
            plot_bgcolor='rgba(245,249,255,1)', paper_bgcolor='rgba(255,255,255,0)',
            showlegend=False,
            title=dict(
                text=("<b>⚡ Generator Convention</b>"
                      "<span style='font-size:11px;color:#555;'>"
                      "  IEEE 1547-2018 | i_s positive: converter→grid</span>"),
                font=dict(size=13), x=0.02, xanchor='left'))

        # ── Right: Load Convention ────────────────────────────────────────
        # Load convention KVL: V_on = V_s + V_L_load, V_L_load = jX_L·I_load = -jX_L·I_gen = -V_L_gen
        # Load convention V_L direction: P → center (opposite to generator convention)
        # i_s (I_load = -I_gen): from P, angle = θ + π (opposite to I_gen)
        #   I_load endpoint = P + i_len·(-I_gen_dir) = P - i_len·I_gen_dir
        I_load_end = P - i_len * I_gen_dir   # I_load direction endpoint (shot from P)

        fig_load = go.Figure()
        _draw_common(fig_load)
        # V_L (load convention): P → center (opposite to generator convention)
        _arr(fig_load, P[0], P[1], CTR[0], CTR[1], '#111111', width=2.2)
        vl_mid_load = (P + CTR) / 2
        # Label offset same side as generator convention (perpendicular to V_L, negated to avoid overlap)
        fig_load.add_annotation(
            x=vl_mid_load[0]-vl_n_gen[0], y=vl_mid_load[1]-vl_n_gen[1],
            text="<b><i>V</i><sub>L</sub></b>",
            showarrow=False, font=dict(size=14, color='#111111'))
        # i_s (load convention): from P, direction = I_load = -I_gen (angle = θ+π)
        _arr(fig_load, P[0], P[1], I_load_end[0], I_load_end[1], '#cc6600', width=2.5)
        off_load = -np.array([np.cos(theta), np.sin(theta)]) * 0.11   # I_load direction
        fig_load.add_annotation(
            x=I_load_end[0]+off_load[0], y=I_load_end[1]+off_load[1],
            text="<b><i>i</i><sub>s</sub></b>",
            showarrow=False, font=dict(size=14, color='#cc6600'))

        # Legend()
        for idx, (c, lbl) in enumerate([
            ('#cc6600', 'i_s (P→out, Grid→VSC, angle=θ+180°)'),
            ('#111111', 'V_L_load (P→center, = -jX_L·I_gen, load convention)'),
            ('#555555', 'V_s (O→P, Converter Output Voltage, unchanged)'),
            ('#1a7a1a', 'V_on (O→center, Grid Voltage, unchanged)'),
        ]):
            fig_load.add_annotation(
                x=xlo+0.03*(xhi-xlo),
                y=y_leg+0.03*span+idx*0.055*span,
                text=(f'<span style="color:{c};font-size:14px;font-weight:bold;">━</span>'
                      f'<span style="font-size:11px;color:#222;"> {lbl}</span>'),
                showarrow=False, xanchor='left', yanchor='bottom',
                font=dict(size=11), bgcolor='rgba(255,255,255,0)', borderwidth=0)

        fig_load.add_annotation(
            x=0.01, y=0.32,
            xref='paper', yref='paper',
            text=("<b>Load Convention</b><br>"
                  "P = 🔒 kW (3-phase)<br>"
                  "Q = 🔒 kvar (3-phase)<br>"
                  "KVL: V_on = V_s + V_L_load<br>"
                  "i_s: Grid→VSC") if is_trial else
                 (f"<b>Load Convention</b><br>"
                  f"P = {P_load_3ph/1000:+.3f} kW (3-phase)<br>"
                  f"Q = {Q_load_3ph/1000:+.3f} kvar (3-phase)<br>"
                  f"KVL: V_on = V_s + V_L_load<br>"
                  f"i_s: Grid→VSC"),
            showarrow=False, xanchor='left', yanchor='bottom',
            font=dict(size=10, color='#0a5c3a'),
            bgcolor="rgba(225,245,238,0.95)", bordercolor='#198754', borderwidth=1,
            align='left')
        fig_load.update_layout(
            xaxis=dict(range=[xlo, xhi], visible=False, showgrid=False, zeroline=False),
            yaxis=dict(range=[y_leg, yhi], visible=False, showgrid=False, zeroline=False,
                       scaleanchor='x', scaleratio=1),
            height=510, margin=dict(l=5, r=5, t=42, b=5),
            plot_bgcolor='rgba(245,255,250,1)', paper_bgcolor='rgba(255,255,255,0)',
            showlegend=False,
            title=dict(
                text=("<b>🔌 Load Convention</b>"
                      "<span style='font-size:11px;color:#555;'>"
                      "  IEC 61400-21-1:2021 | i_s positive: grid→converter</span>"),
                font=dict(size=13), x=0.02, xanchor='left'))

        # Side-by-side display
        col_gen, col_load = st.columns(2)
        with col_gen:
            st.plotly_chart(fig_gen, use_container_width=True)
        with col_load:
            st.plotly_chart(fig_load, use_container_width=True)

        # Current operating point comparison table
        # Reactive power effect (based on Q_gen sign, convention-independent)
        if I_mag < 1e-3:
            q_eff = "No reactive power exchange"
        elif Q_gen > 1e-3 * u_grid_rms * I_mag:
            q_eff = "VSC **generates reactive power** (over-excited, I_gen Lagging, raises grid voltage)"
        elif Q_gen < -1e-3 * u_grid_rms * I_mag:
            q_eff = "VSC **absorbs reactive power** (under-excited, I_gen Leading, lowers grid voltage)"
        else:
            q_eff = "Reactive power negligible"

        if is_trial:
            st.info(
                "**📌 Current Operating Point**\n\n"
                "| | Generator Convention | Load Convention |\n"
                "|:---|:---:|:---:|\n"
                "| **i_s Direction** | P→out (VSC→Grid) | P→out (Grid→VSC, reversed vs left) |\n"
                "| **P (3-phase)** | 🔒 | 🔒 |\n"
                "| **Q (3-phase)** | 🔒 | 🔒 |\n"
                "| **Voltage triangle** | V_on+V_L=V_s | **Identical** |\n\n"
                "**Physical effect**: 🔒 (unlock with paid access for full operating diagnosis)"
            )
        else:
            st.info(
                f"**📌 Current Operating Point: {phase_desc_phys}**\n\n"
                f"| | Generator Convention | Load Convention |\n"
                f"|:---|:---:|:---:|\n"
                f"| **i_s Direction** | P→out (VSC→Grid) | P→out (Grid→VSC, reversed vs left) |\n"
                f"| **P (3-phase)** | {P_gen_3ph/1000:+.3f} kW | {P_load_3ph/1000:+.3f} kW |\n"
                f"| **Q (3-phase)** | {Q_gen_3ph/1000:+.3f} kvar | {Q_load_3ph/1000:+.3f} kvar |\n"
                f"| **Voltage triangle** | V_on+V_L=V_s | **Identical** |\n\n"
                f"**Physical effect**: {q_eff}\n\n"
                f"> Reference: IEEE Std 1547-2018 §4.1; IEC 61400-21-1:2021 §4.2;"
                f"Kundur, *Power System Stability and Control*, McGraw-Hill 1994, §10.2"
            )

        # ══════════════════════════════════════════════════════════════
        # ② ABCD four-point reference table (per Appendix Fig.4)
        # ══════════════════════════════════════════════════════════════
        with st.expander("📖 ABCD Four-Point Reference Table (per Standard Appendix)"):
            st.markdown(f"""
**Current Operating Point**: {phase_desc_phys}

---

#### Dual Convention Definitions

| | Generator Convention (IEEE 1547-2018 §4.1) | Load Convention (IEC 61400-21-1:2021 §4.2) |
|:---|:---|:---|
| **Current Reference** | I_gen: VSC → Grid (out = positive) | I_load = −I_gen: Grid → VSC (in = positive) |
| **P > 0** | VSC **sends** active power to grid | VSC **absorbs** active power from grid |
| **Q > 0** | VSC **generates** reactive; I_gen **Lagging** V_on (over-excited); **raises** grid voltage | VSC **absorbs** reactive (under-excited); I_load **Lagging** V_on; **lowers** inductive grid voltage |
| **Q < 0** | VSC **absorbs** reactive; I_gen **Leading** V_on (under-excited); **lowers** grid voltage | VSC **generates** reactive (over-excited); I_load **Leading** V_on; **raises** inductive grid voltage |
""")

            # Deep ABCD table locked for trial users
            if is_trial:
                trial_lock("Unlock with paid access: ABCD four-point operating condition table and reactive-voltage relationship analysis")
            else:
                st.markdown(f"""
---

#### ABCD Four-Point Operating Conditions (per Appendix Fig.4)

| Point | Physical Fact (invariant) | Generator Convention (IEEE 1547) | Load Convention (IEC 61400) |
|:---:|:---|:---|:---|
| **C** (right) | **Over-excited**: VSC generates reactive, **raises** grid voltage | I_gen **Lagging** 90°, **Q_gen > 0** (generates reactive) | I_load **Leading** 90°, **Q_load < 0** |
| **D** (top) | **Generation**: VSC sends active power to grid | I_gen **in-phase** 0°, **P_gen > 0** (active gen) | I_load **anti-phase** 180°, **P_load < 0** |
| **A** (left) | **Under-excited**: VSC absorbs reactive, **lowers** grid voltage | I_gen **Leading** 90°, **Q_gen < 0** (absorbs reactive) | I_load **Lagging** 90°, **Q_load > 0** (absorbs) |
| **B** (bottom) | **Motoring**: VSC absorbs active power from grid | I_gen **anti-phase** 180°, **P_gen < 0** | I_load **in-phase** 0°, **P_load > 0** (absorbs active) |

#### Grid Voltage Effect (Physical Mechanism)

| Physical State | Effect on Inductive Grid | Generator Convention Q | Load Convention Q |
|:---|:---:|:---:|:---:|
| Over-excited: VSC generates reactive (I_gen Lagging V_on) | **Raises** grid voltage | Q_gen **> 0** | Q_load **< 0** |
| Under-excited: VSC absorbs reactive (I_gen Leading V_on) | **Lowers** grid voltage | Q_gen **< 0** | Q_load **> 0** |

> References: Kundur, §10.2; IEEE 1547-2018 Annex B; Appendix Fig.4
""")

        st.divider()

        # ══════════════════════════════════════════════════════════════
        # ③ ()
        # ══════════════════════════════════════════════════════════════
        st.markdown("#### 📈 Instantaneous Waveform View")
        st.caption(
            "Grid Voltage V_on and instantaneous power are convention-independent;"
            "**Current i_s waveform reverses with convention** (generator and load conventions define opposite positive directions)"
        )

        T_period  = 1.0 / sys_freq if sys_freq > 0 else 0.02
        t_arr     = np.linspace(0, 2*T_period, 500)
        wt_arr    = omega * t_arr
        # V_on(t) = |V_on|·√2·sin(ωt)(V_ont=0)
        v_on_wave = u_grid_rms * np.sqrt(2) * np.sin(wt_arr)
        # I_gen(t) = |I|·√2·sin(ωt + θ)(θ < 0 = Lagging = )
        i_gen_wave  =  I_mag * np.sqrt(2) * np.sin(wt_arr + theta)
        # I_load(t) = -I_gen(t)
        i_load_wave = -i_gen_wave
        #  p(t) = V_on(t) · I_gen(t)(, =VSC)
        p_wave = v_on_wave * i_gen_wave
        # p(t) : V_on, 
        v_peak = u_grid_rms * np.sqrt(2)
        p_max  = np.max(np.abs(p_wave)) if np.max(np.abs(p_wave)) > 1 else 1.0
        p_wave_norm = p_wave / p_max * v_peak   # normalize to [-v_peak, +v_peak] range

        # Y: =V_on+p(t), =i_s(A)
        fig_wave = make_subplots(specs=[[{"secondary_y": True}]])

        # V_on(, )
        fig_wave.add_trace(go.Scatter(
            x=t_arr*1000, y=v_on_wave,
            name='V_on (Grid Voltage)' if is_trial else f'V_on (Grid Voltage, peak {v_peak:.0f}V)',
            line=dict(color='#1a7a1a', width=2.5),
            hovertemplate=None if is_trial else '%{y:.1f} V<extra>V_on</extra>',
            hoverinfo='none' if is_trial else 'all'
        ), secondary_y=False)

        # (, )
        if use_gen:
            i_wave_show = i_gen_wave
            i_name  = ('i_s Generator Convention (VSC→Grid)' if is_trial
                       else f'i_s Generator Convention (VSC→Grid, peak {I_mag*np.sqrt(2):.0f}A)')
            i_color = '#1155cc'
        else:
            i_wave_show = i_load_wave
            i_name  = ('i_s Load Convention (Grid→VSC, reversed vs left)' if is_trial
                       else f'i_s Load Convention (Grid→VSC, reversed vs left, peak {I_mag*np.sqrt(2):.0f}A)')
            i_color = '#cc6600'
        fig_wave.add_trace(go.Scatter(
            x=t_arr*1000, y=i_wave_show, name=i_name,
            line=dict(color=i_color, width=2.5),
            hovertemplate=None if is_trial else '%{y:.1f} A<extra>i_s</extra>',
            hoverinfo='none' if is_trial else 'all'
        ), secondary_y=True)

        # p(t) (, , V_on)
        p_avg_kw = np.mean(p_wave) / 1000
        fig_wave.add_trace(go.Scatter(
            x=t_arr*1000, y=p_wave_norm,
            name=f'p(t) normalized (positive=VSC generates active)' if is_trial else f'p(t) normalized (actual mean={p_avg_kw:.1f}kW, positive=VSC generates active)',
            line=dict(color='#fd7e14', width=1.5, dash='dot'), opacity=0.80,
            hoverinfo='none' if is_trial else 'all'
        ), secondary_y=False)

        # Leading/Lagging
        if abs(theta) > 0.05 and abs(abs(theta) - np.pi) > 0.05:
            t_i_zero_ms = (-theta / omega) * 1000 % (T_period * 1000)
            fig_wave.add_vline(x=0,            line_dash='dot', line_color='#1a7a1a', line_width=1.5)
            fig_wave.add_vline(x=t_i_zero_ms, line_dash='dot', line_color=i_color,   line_width=1.5)
            ll = "Lagging" if theta < 0 else "Leading"
            ann_x = t_i_zero_ms / 2 if theta < 0 else (t_i_zero_ms + T_period*1000) / 2
            _wave_ann_txt = f"I_gen {ll} V_on" if is_trial else f"I_gen {ll} V_on  |θ|={abs(theta_deg):.1f}°"
            fig_wave.add_annotation(
                x=ann_x % (T_period * 2000),
                y=v_peak * 1.15,
                text=_wave_ann_txt,
                showarrow=False, font=dict(size=11, color='#444'),
                bgcolor='rgba(255,255,255,0.90)')

        fig_wave.update_layout(
            title=dict(
                text=(f"<b>Instantaneous Waveform</b>  Convention: {'Generator (i_s=I_gen)' if use_gen else 'Load (i_s=I_load=-I_gen)'}"),
                font=dict(size=12)),
            xaxis_title="Time (ms)",
            legend=dict(orientation="h", y=1.18, x=0, font=dict(size=10)),
            height=320, hovermode=False if is_trial else "x unified",
            margin=dict(l=10, r=10, t=75, b=30),
            plot_bgcolor='rgba(250,250,250,1)'
        )
        # : (V_onp(t), )
        fig_wave.update_yaxes(
            title_text="Voltage (V) / p(t) normalized",
            range=[-v_peak*1.35, v_peak*1.35],
            showgrid=True, gridcolor='#eee', secondary_y=False)
        # : 
        i_peak = I_mag * np.sqrt(2)
        fig_wave.update_yaxes(
            title_text="Current (A)",
            range=[-i_peak*1.35, i_peak*1.35],
            showgrid=False, secondary_y=True)
        st.plotly_chart(fig_wave, use_container_width=True)

        st.divider()

        # ══════════════════════════════════════════════════════════════
        # ④ 
        # ══════════════════════════════════════════════════════════════
        st.markdown("#### 🧭 Power Flow Diagnosis")

        diag_c1, diag_c2 = st.columns([1.3, 1])

        with diag_c1:
            if is_trial:
                st.markdown(
                    """
                    <div style="display:flex;gap:24px;margin-bottom:8px;">
                      <div style="flex:1;background:#f0f4ff;border-radius:8px;padding:12px 16px;">
                        <div style="font-size:12px;color:#555;margin-bottom:4px;">3-Phase Total Active Power P</div>
                        <div style="font-size:22px;font-weight:700;color:#ccc;">🔒</div>
                      </div>
                      <div style="flex:1;background:#fff4e6;border-radius:8px;padding:12px 16px;">
                        <div style="font-size:12px;color:#555;margin-bottom:4px;">3-Phase Total Reactive Power Q</div>
                        <div style="font-size:22px;font-weight:700;color:#ccc;">🔒</div>
                      </div>
                      <div style="flex:1;background:#f0fff4;border-radius:8px;padding:12px 16px;">
                        <div style="font-size:12px;color:#555;margin-bottom:4px;">Power Factor PF</div>
                        <div style="font-size:22px;font-weight:700;color:#ccc;">🔒</div>
                      </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            else:
                st.markdown(
                    f"""
                    <div style="display:flex;gap:24px;margin-bottom:8px;">
                      <div style="flex:1;background:#f0f4ff;border-radius:8px;padding:12px 16px;">
                        <div style="font-size:12px;color:#555;margin-bottom:4px;">3-Phase Total Active Power P (generator convention: Q&gt;0=generates reactive)</div>
                        <div style="font-size:22px;font-weight:700;color:#0d47a1;">{p_disp_kw:+.3f} kW</div>
                      </div>
                      <div style="flex:1;background:#fff4e6;border-radius:8px;padding:12px 16px;">
                        <div style="font-size:12px;color:#555;margin-bottom:4px;">3-Phase Total Reactive Power Q (positive=generates reactive/over-excited)</div>
                        <div style="font-size:22px;font-weight:700;color:#e65100;">{q_disp_kvar:+.3f} kvar</div>
                      </div>
                      <div style="flex:1;background:#f0fff4;border-radius:8px;padding:12px 16px;">
                        <div style="font-size:12px;color:#555;margin-bottom:4px;">Power Factor PF</div>
                        <div style="font-size:22px;font-weight:700;color:#1b5e20;">{pf_val:.4f}</div>
                      </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

            # SVG power flow diagram
            p_to_grid = P_gen >= 0
            q_to_grid = Q_gen >= 0
            p_abs = abs(P_gen_3ph/1000)  # 3-phase
            q_abs = abs(Q_gen_3ph/1000)  # 3-phase

            if is_trial:
                svg_trial = """<svg width="100%" height="110" viewBox="0 0 330 110"
     xmlns="http://www.w3.org/2000/svg" style="font-family:sans-serif">
  <rect x="4" y="22" width="84" height="60" rx="8" fill="#f8f9fa" stroke="#adb5bd" stroke-width="1"/>
  <text x="46" y="47" font-size="13" font-weight="bold" text-anchor="middle" fill="#333">VSC</text>
  <text x="46" y="63" font-size="10" text-anchor="middle" fill="#666">Converter</text>
  <circle cx="284" cy="52" r="36" fill="none" stroke="#198754" stroke-width="1.5"/>
  <text x="284" y="48" font-size="11" font-weight="bold" text-anchor="middle" fill="#198754">GRID</text>
  <text x="284" y="63" font-size="10" text-anchor="middle" fill="#198754">Grid</text>
  <line x1="90" y1="36" x2="246" y2="36" stroke="#ccc" stroke-width="3" stroke-dasharray="6,4"/>
  <text x="168" y="28" font-size="10" font-weight="bold" text-anchor="middle" fill="#ccc">P = 🔒</text>
  <line x1="90" y1="68" x2="246" y2="68" stroke="#ccc" stroke-width="3" stroke-dasharray="6,4"/>
  <text x="168" y="95" font-size="10" font-weight="bold" text-anchor="middle" fill="#ccc">Q = 🔒</text>
</svg>"""
                st.markdown(svg_trial, unsafe_allow_html=True)
            else:
                if p_to_grid:
                    p_txt = f"P={p_abs:.2f}kW → Grid (3-phase, VSC generates active)"
                    px1, px2 = 90, 246
                else:
                    p_txt = f"P={p_abs:.2f}kW ← Grid (3-phase, VSC absorbs active)"
                    px1, px2 = 246, 90

                if q_to_grid:
                    q_txt = f"Q={q_abs:.2f}kvar → Grid (3-phase, over-excited, generates reactive, raises voltage)"
                    qx1, qx2 = 90, 246
                else:
                    q_txt = f"Q={q_abs:.2f}kvar ← Grid (3-phase, under-excited, absorbs reactive, lowers voltage)"
                    qx1, qx2 = 246, 90

                svg = f"""<svg width="100%" height="110" viewBox="0 0 330 110"
     xmlns="http://www.w3.org/2000/svg" style="font-family:sans-serif">
  <defs><marker id="fa2" markerWidth="8" markerHeight="6" refX="7" refY="3" orient="auto">
    <polygon points="0 0.5,7 3,0 5.5" fill="#444"/></marker></defs>
  <rect x="4" y="22" width="84" height="60" rx="8" fill="#f8f9fa" stroke="#adb5bd" stroke-width="1"/>
  <text x="46" y="47" font-size="13" font-weight="bold" text-anchor="middle" fill="#333">VSC</text>
  <text x="46" y="63" font-size="10" text-anchor="middle" fill="#666">Converter</text>
  <circle cx="284" cy="52" r="36" fill="none" stroke="#198754" stroke-width="1.5"/>
  <text x="284" y="48" font-size="11" font-weight="bold" text-anchor="middle" fill="#198754">GRID</text>
  <text x="284" y="63" font-size="10" text-anchor="middle" fill="#198754">Grid</text>
  <line x1="{px1}" y1="36" x2="{px2}" y2="36"
        stroke="#0d6efd" stroke-width="3" marker-end="url(#fa2)"/>
  <text x="168" y="28" font-size="10" font-weight="bold"
        text-anchor="middle" fill="#0d6efd">{p_txt}</text>
  <line x1="{qx1}" y1="68" x2="{qx2}" y2="68"
        stroke="#fd7e14" stroke-width="3" marker-end="url(#fa2)"/>
  <text x="168" y="95" font-size="10" font-weight="bold"
        text-anchor="middle" fill="#fd7e14">{q_txt}</text>
</svg>"""
                st.markdown(svg, unsafe_allow_html=True)

        with diag_c2:
            S_base = u_grid_rms * I_mag if I_mag > 0 else 1.0
            thr    = 0.05 * S_base
            if is_trial:
                trial_lock("Unlock with paid access: power flow diagnosis and operating quadrant analysis")
            else:
                # Operating state (per Appendix Fig.4, based on P_gen/Q_gen)
                if I_mag < 1e-3:
                    st.info("**🔵 Zero Current**\n\nNo power exchange.")
                elif abs(Q_gen) < thr and P_gen > thr:
                    st.success(
                        f"**✅ Near Point D: Pure active generation (PF≈1)**\n\n"
                        f"{phase_desc_phys}\n\n"
                        f"• Active: VSC delivers {P_gen_3ph/1000:.2f} kW to grid\n\n"
                        f"• Reactive: Q≈0, no effect on grid voltage"
                    )
                elif abs(Q_gen) < thr and P_gen < -thr:
                    st.warning(
                        f"**⚡ Near Point B: Pure active absorption (PF≈1)**\n\n"
                        f"{phase_desc_phys}\n\n"
                        f"• Active: VSC absorbs {abs(P_gen_3ph/1000):.2f} kW from grid\n\n"
                        f"• Reactive: Q≈0, no effect on grid voltage"
                    )
                elif abs(P_gen) < thr and Q_gen > thr:
                    st.info(
                        f"**🔵 Near Point C: Pure reactive generation (over-excited)**\n\n"
                        f"{phase_desc_phys}\n\n"
                        f"• Active: P≈0\n\n"
                        f"• Reactive: VSC generates {Q_gen_3ph/1000:.2f} kvar (over-excited, I_gen Lagging)\n\n"
                        f"• Effect: **raises** inductive grid voltage\n\n"
                        f"• {'Generator convention Q>0' if use_gen else 'Load convention Q<0'}"
                    )
                elif abs(P_gen) < thr and Q_gen < -thr:
                    st.warning(
                        f"**🟠 Near Point A: Pure reactive absorption (under-excited)**\n\n"
                        f"{phase_desc_phys}\n\n"
                        f"• Active: P≈0\n\n"
                        f"• Reactive: VSC absorbs {abs(Q_gen_3ph/1000):.2f} kvar (under-excited, I_gen Leading)\n\n"
                        f"• Effect: **lowers** inductive grid voltage\n\n"
                        f"• {'Generator convention Q<0' if use_gen else 'Load convention Q>0'}"
                    )
                elif P_gen > thr and Q_gen > thr:
                    st.success(
                        f"**✅ Quadrant I: Active generation + Reactive generation (over-excited)**\n\n"
                        f"{phase_desc_phys}\n\n"
                        f"• Active: VSC delivers {P_gen_3ph/1000:.2f} kW\n\n"
                        f"• Reactive: over-excited, generates {Q_gen_3ph/1000:.2f} kvar, **raises** grid voltage\n\n"
                        f"• {'Generator: P>0,Q>0' if use_gen else 'Load: P<0,Q<0'}"
                    )
                elif P_gen > thr and Q_gen < -thr:
                    st.warning(
                        f"**🟠 Quadrant IV: Active generation + Reactive absorption (under-excited)**\n\n"
                        f"{phase_desc_phys}\n\n"
                        f"• Active: VSC delivers {P_gen_3ph/1000:.2f} kW\n\n"
                        f"• Reactive: under-excited, absorbs {abs(Q_gen_3ph/1000):.2f} kvar, **lowers** grid voltage\n\n"
                        f"• {'Generator: P>0,Q<0' if use_gen else 'Load: P<0,Q>0'}"
                    )
                elif P_gen < -thr and Q_gen > thr:
                    st.info(
                        f"**🔵 Rectifier Mode: absorbs active power + generates reactive power (over-excited)**\n\n"
                        f"{phase_desc_phys}\n\n"
                        f"• Active: VSC absorbs {abs(P_gen_3ph/1000):.2f} kW from grid\n\n"
                        f"• Reactive: over-excited, generates {Q_gen_3ph/1000:.2f} kvar, **raises** grid voltage\n\n"
                        f"• {'Generator convention P<0,Q>0' if use_gen else 'Load convention P>0,Q<0'}"
                    )
                else:
                    st.warning(
                        f"**⚡ Rectifier Mode: absorbs active power + absorbs reactive power (under-excited)**\n\n"
                        f"{phase_desc_phys}\n\n"
                        f"• Active: VSC absorbs {abs(P_gen_3ph/1000):.2f} kW from grid\n\n"
                        f"• Reactive: under-excited, absorbs {abs(Q_gen_3ph/1000):.2f} kvar, **lowers** grid voltage\n\n"
                        f"• {'Generator convention P<0,Q<0' if use_gen else 'Load convention P>0,Q>0'}"
                    )

                st.caption(
                    "**Physical interpretation (Generator convention, see Appendix 4)**:\n\n"
                    "• **Over-excited** (I_gen Lagging V_on, Quadrant C, Q_gen > 0) → **Raises** inductive grid voltage\n\n"
                    "• **Under-excited** (I_gen Leading V_on, Quadrant A, Q_gen < 0) → **Lowers** inductive grid voltage\n\n"
                    "> Reference: Kundur, §10.2; Appendix 4; IEEE 1547-2018 Annex B"
                )



        with st.expander("📚 For Engineers: Understanding reactive/active power — amplitude and phase"):
            # ===== Part 1: Teaser (always visible) =====
            st.markdown(f"""
**Current Operating Point**: {phase_desc_phys}

---

#### 1. Reactive Power Q is Driven Mainly by Voltage Magnitude Difference

The reactive power of a grid-connected converter is fundamentally driven by the amplitude difference between V_s and V_on.
For small power angles (typical grid-connected operation), use this simplified expression:
""")

            # ===== Part 2: Formulas + deep analysis (locked in trial) =====
            if is_trial:
                trial_lock("Unlock with paid access: reactive/active power complete physical analysis and dq decoupling principle")
            else:
                st.markdown(f"""
$$Q \\approx \\frac{{V_s \\cdot (V_s - V_{{on}})}}{{X_L}}$$

So: **Raise V_s → over-excite → generate reactive → raises grid voltage; Lower V_s → under-excite → absorb reactive → lowers grid voltage.**

One often-overlooked detail: **when power angle δ is non-zero, phase also contributes some reactive power**.
When generating active power (δ > 0), cos δ < 1, so the converter must supply additional reactive power to sustain the inductor's magnetic field.
The larger the power angle (heavier loading), the less this coupled reactive component can be ignored.

> **Engineering takeaway**: For routine reactive control, adjust amplitude. Under heavy load, account for this coupling term.

---

#### 2. Active Power P is Determined Entirely by the Sign of Phase Difference

$$P \\propto V_s \\cdot V_{{on}} \\cdot \\sin\\delta / X_L$$

**The sign of P (generate or absorb) is determined solely by the sign of δ — amplitude cannot change this:**

- V_s Leading V_on (δ > 0) → converter sends active power to grid (generation/inversion)
- V_s Lagging V_on (δ < 0) → converter draws active power from grid (charging/rectification)

The role of V_s amplitude: once direction is fixed, it determines how much power is transferred. Double V_s → double P at same δ.

> **Engineering takeaway**: To control how much active power is generated, adjust phase angle. To deliver more power at the same angle, adjust amplitude.

---

#### 3. Why Does the Controller Need "Decoupling"?

The above shows that P and Q each have primary control knobs, but phase and amplitude are not fully independent — at large power angles, δ also affects Q.
**dq-axis decoupling control** compensates for this coupling term in software, so P and Q can be controlled as two independent knobs:

| Control Channel | Target | Primary Action |
|:---:|:---:|:---|
| **q-axis current** | Active P | Control phase angle δ of V_s relative to V_on |
| **d-axis current** | Reactive Q | Control V_s amplitude (modulation depth / excitation) |

Without decoupling, adjusting P disturbs Q and vice versa, making the controller difficult to stabilize.

> Reference: Kundur, *Power System Stability and Control*, McGraw-Hill 1994, §10.2;
> IEEE Std 1547-2018 Annex B
""")


    # ─────────────────────────────────────────────────────────────────
    # Tab 2: Voltage triangle (load convention, reuses Tab1 normalized geometry)
    # ─────────────────────────────────────────────────────────────────
    with tab2:
        c_vec1, c_vec2 = st.columns([5, 2])
        with c_vec1:
            # ── Key fix: P_actual placed at true voltage ratio ──────────────────
            # In normalized coords, R=1 represents modulation limit limit_v
            # Actual V_s = req_v_phase, so actual P-point distance from center = req_v_phase/limit_v * R
            # Over-modulation: ratio > 1 → P_actual outside circle; normal: inside
            ratio    = req_v_phase / limit_v        # true ratio, may be > 1
            mod_ok   = ratio <= 1.0
            P_actual = CTR + ratio * R * np.array([np.cos(theta_P), np.sin(theta_P)])

            # i_s endpoint follows P_actual (load convention, from P_actual outward)
            I_load_end_t2 = P_actual - i_len * I_gen_dir

            # Unified viewport (includes ABCD endpoints + P_actual + i_s endpoint)
            all_t2 = np.vstack([O, CTR, P_actual, I_load_end_t2,
                                CTR+[R,0], CTR+[-R,0], CTR+[0,R], CTR+[0,-R]])
            xs2, ys2 = all_t2[:,0], all_t2[:,1]
            cx2  = (xs2.min()+xs2.max())/2
            cy2  = (ys2.min()+ys2.max())/2
            half2 = max(xs2.max()-xs2.min(), ys2.max()-ys2.min())/2 * 1.38
            xlo2, xhi2 = cx2-half2, cx2+half2
            ylo2, yhi2 = cy2-half2, cy2+half2

            fig_t2 = go.Figure()

            def _arr2(fig, x0, y0, x1, y1, color, width=2.5):
                fig.add_annotation(
                    x=x1, y=y1, ax=x0, ay=y0,
                    xref='x', yref='y', axref='x', ayref='y',
                    showarrow=True, arrowhead=3, arrowsize=1.5,
                    arrowwidth=width, arrowcolor=color, text=""
                )
                fig.add_trace(go.Scatter(
                    x=[x0, x1], y=[y0, y1], mode='lines',
                    line=dict(color=color, width=width),
                    hoverinfo='skip', showlegend=False
                ))

            # PWM limit circle (red dashed, centered at CTR, R=1)
            ct2 = np.linspace(0, 2*np.pi, 300)
            fig_t2.add_trace(go.Scatter(
                x=CTR[0]+R*np.cos(ct2), y=CTR[1]+R*np.sin(ct2),
                mode='lines',
                line=dict(color='#cc1111', width=1.8, dash='dash'),
                name=(f'{mod_label} Limit' if is_trial else f'{mod_label} Limit  {limit_v:.0f} Vrms (Vdc={u_dc:.0f}V)'),
                hoverinfo='skip', showlegend=True
            ))

            # If over-modulation, show overflow segment from circle to P_actual
            if not mod_ok:
                P_on_circle = CTR + R * np.array([np.cos(theta_P), np.sin(theta_P)])
                fig_t2.add_trace(go.Scatter(
                    x=[P_on_circle[0], P_actual[0]],
                    y=[P_on_circle[1], P_actual[1]],
                    mode='lines',
                    line=dict(color='#cc1111', width=2.5, dash='dot'),
                    name=('Exceeds limit' if is_trial else f'Exceeds by {req_v_phase-limit_v:.0f} V'),
                    hoverinfo='skip', showlegend=True
                ))

            # ABCD ticks
            tk2 = 0.045 * R
            for (px, py, tx, ty, lbl, anc, dx, dy) in [
                (CTR[0]-R, CTR[1],    CTR[0]-R-0.13, CTR[1],        'A', 'right',  tk2, 0),
                (CTR[0],   CTR[1]-R,  CTR[0],        CTR[1]-R-0.15, 'B', 'center', 0,   tk2),
                (CTR[0]+R, CTR[1],    CTR[0]+R+0.13, CTR[1],        'C', 'left',   tk2, 0),
                (CTR[0],   CTR[1]+R,  CTR[0],        CTR[1]+R+0.15, 'D', 'center', 0,   tk2),
            ]:
                fig_t2.add_shape(type='line',
                    x0=px-dx, y0=py-dy, x1=px+dx, y1=py+dy,
                    line=dict(color='#555', width=1.5))
                fig_t2.add_annotation(x=tx, y=ty, text=f"<b>{lbl}</b>",
                    showarrow=False, font=dict(size=13, color='#333'), xanchor=anc)

            # V_on: O → CTR (green)
            _arr2(fig_t2, O[0], O[1], CTR[0], CTR[1], '#1a7a1a', width=2.8)
            fig_t2.add_annotation(
                x=(O[0]+CTR[0])/2, y=(O[1]+CTR[1])/2 - 0.13,
                text=(f"<b><i>V</i><sub>on</sub></b>" if is_trial else f"<b><i>V</i><sub>on</sub></b> <span style='font-size:11px'>{u_grid_rms:.0f}V</span>"),
                showarrow=False, font=dict(size=13, color='#1a7a1a'))

            # V_s: O → P_actual (dark gray, actual voltage, may exceed limit circle)
            _arr2(fig_t2, O[0], O[1], P_actual[0], P_actual[1], '#555555', width=1.8)
            vs2_dir = P_actual - np.array([O[0], O[1]])
            vs2_n   = np.array([-vs2_dir[1], vs2_dir[0]]) / (np.linalg.norm(vs2_dir)+1e-9) * 0.13
            vs2_mid = (np.array([O[0], O[1]]) + P_actual) / 2
            fig_t2.add_annotation(
                x=vs2_mid[0]+vs2_n[0], y=vs2_mid[1]+vs2_n[1],
                text=(f"<b><i>V</i><sub>s</sub></b>" if is_trial else f"<b><i>V</i><sub>s</sub></b> <span style='font-size:11px'>{req_v_phase:.0f}V</span>"),
                showarrow=False, font=dict(size=13, color='#555555'))

            # V_L_load (load convention): P_actual → CTR (black)
            vl_mag_rms = np.linalg.norm([vec_vl.real, vec_vl.imag])
            _arr2(fig_t2, P_actual[0], P_actual[1], CTR[0], CTR[1], '#111111', width=2.2)
            vl2_d   = CTR - P_actual
            vl2_n   = np.array([-vl2_d[1], vl2_d[0]]) / (np.linalg.norm(vl2_d)+1e-9) * 0.13
            vl2_mid = (P_actual + CTR) / 2
            fig_t2.add_annotation(
                x=vl2_mid[0]+vl2_n[0], y=vl2_mid[1]+vl2_n[1],
                text=(f"<b><i>V</i><sub>L</sub></b>" if is_trial else f"<b><i>V</i><sub>L</sub></b> <span style='font-size:11px'>{vl_mag_rms:.0f}V</span>"),
                showarrow=False, font=dict(size=13, color='#111111'))

            # i_s (load convention): P_actual → I_load_end_t2 (orange)
            _arr2(fig_t2, P_actual[0], P_actual[1],
                  I_load_end_t2[0], I_load_end_t2[1], '#cc6600', width=2.5)
            is2_off = (-I_gen_dir) * 0.11
            fig_t2.add_annotation(
                x=I_load_end_t2[0]+is2_off[0], y=I_load_end_t2[1]+is2_off[1],
                text=(f"<b><i>i</i><sub>s</sub></b>" if is_trial else f"<b><i>i</i><sub>s</sub></b> <span style='font-size:11px'>{I_mag:.0f}A</span>"),
                showarrow=False, font=dict(size=13, color='#cc6600'))

            # P_actual point marker
            p_marker_color = '#cc1111' if not mod_ok else '#222'
            fig_t2.add_trace(go.Scatter(
                x=[P_actual[0]], y=[P_actual[1]], mode='markers',
                marker=dict(size=10, color=p_marker_color,
                            symbol='x' if not mod_ok else 'circle'),
                hoverinfo='skip', showlegend=False
            ))
            if not mod_ok:
                fig_t2.add_annotation(
                    x=P_actual[0]+0.07, y=P_actual[1]+0.07,
                    text=("<b>Point P (over-limit)</b>" if is_trial else f"<b>Point P ({req_v_phase:.0f}V, over-limit)</b>"),
                    showarrow=False, font=dict(size=10, color='#cc1111'),
                    bgcolor='rgba(255,230,230,0.92)', bordercolor='#cc1111', borderwidth=1)

            # Legend trace
            for (c, lbl) in [
                ('#1a7a1a', 'V_on (Grid Voltage)' if is_trial else f'V_on = {u_grid_rms:.0f} Vrms'),
                ('#555555', 'V_s (Converter Output)' if is_trial else f'V_s = {req_v_phase:.0f} Vrms (O→P actual)'),
                ('#111111', 'V_L (Inductor)' if is_trial else f'V_L = {vl_mag_rms:.0f} Vrms (P→center)'),
                ('#cc6600', 'i_s (Load Convention, Grid→VSC)' if is_trial else f'i_s = {I_mag:.0f} A (Grid→VSC)'),
            ]:
                fig_t2.add_trace(go.Scatter(
                    x=[None], y=[None], mode='lines',
                    line=dict(color=c, width=3), name=lbl, showlegend=True
                ))

            # Top status label
            fig_t2.add_annotation(
                x=0.5, y=1.0, xref='paper', yref='paper',
                text=(f"<b>{state_txt}</b>  |  🔒 Modulation status requires paid access"
                      if is_trial else
                      (f"<b>{state_txt}</b>  |  "
                       f"{'✅ Modulation Margin +' + f'{limit_v-req_v_phase:.0f} V' if mod_ok else '❌ Over-modulation exceeds by ' + f'{req_v_phase-limit_v:.0f} V'}  "
                       f"({mod_label} limit={limit_v:.0f}V)")),
                showarrow=False, xanchor='center', yanchor='top',
                font=dict(size=10, color='#198754' if mod_ok else '#dc3545'),
                bgcolor="rgba(255,255,255,0.93)",
                bordercolor='#198754' if mod_ok else '#dc3545', borderwidth=1)

            fig_t2.update_layout(
                xaxis=dict(range=[xlo2, xhi2], visible=False, showgrid=False, zeroline=False),
                yaxis=dict(range=[ylo2, yhi2], visible=False, showgrid=False, zeroline=False,
                           scaleanchor='x', scaleratio=1),
                height=520, margin=dict(l=5, r=5, t=40, b=5),
                plot_bgcolor='rgba(245,255,250,1)', paper_bgcolor='rgba(255,255,255,0)',
                showlegend=True,
                legend=dict(
                    x=0.01, y=0.01, xanchor='left', yanchor='bottom',
                    bgcolor='rgba(255,255,255,0.92)', bordercolor='#ccc', borderwidth=1,
                    font=dict(size=11), itemsizing='constant', tracegroupgap=1
                ),
                title=dict(
                    text=("<b>🔌 Voltage Triangle (Load Convention)</b>"
                          f"<span style='font-size:11px;color:#666;'>  {mod_label}</span>"
                          if is_trial else
                          "<b>🔌 Voltage Triangle (Load Convention)</b>"
                          f"<span style='font-size:11px;color:#666;'>"
                          f"  L={l_total*1000:.2f}mH | {mod_label}</span>"),
                    font=dict(size=13), x=0.02, xanchor='left')
            )
            st.plotly_chart(fig_t2, use_container_width=True)

        with c_vec2:
            st.markdown("#### 📐 Vector Status")
            if is_trial:
                trial_lock("Unlock with paid access: modulation status diagnosis, voltage margin, and inductor parameter details")
            else:
                if not mod_ok:
                    st.error("❌ **Over-modulation**")
                    st.markdown(f"V_s = **{req_v_phase:.1f} Vrms**")
                    st.markdown(f"Limit = **{limit_v:.1f} Vrms**")
                    st.markdown(f"Exceeds by **{req_v_phase-limit_v:.1f} V**")
                    st.markdown(f"Ratio **{ratio:.3f}×** (reduce required)")
                else:
                    st.success("✅ **Modulation Normal**")
                    st.markdown(f"V_s = **{req_v_phase:.1f} Vrms**")
                    st.markdown(f"Margin **+{limit_v-req_v_phase:.1f} V**")
                    st.markdown(f"Utilization **{ratio*100:.1f}%**")
                st.divider()
                st.markdown("**Modulation Limit Calculation:**")
                st.caption(f"Vdc = {u_dc:.0f} V")
                if "SPWM" in modulation and "SVPWM" not in modulation:
                    st.caption(f"SPWM = Vdc ÷ (2√2)")
                    st.caption(f"= {u_dc:.0f} ÷ {2*np.sqrt(2):.4f}")
                    st.caption(f"= **{limit_v:.2f} Vrms**")
                elif "Over-modulation" not in modulation:
                    st.caption(f"SVPWM = Vdc ÷ √6")
                    st.caption(f"= {u_dc:.0f} ÷ {np.sqrt(6):.4f}")
                    st.caption(f"= **{limit_v:.2f} Vrms**")
                    st.caption("15.47% higher than SPWM")
                else:
                    st.caption(f"Over-modulation estimate = **{limit_v:.2f} Vrms**")
                st.divider()
                st.caption(f"L1={l_filter_mh:.3f}mH  XL1={xl1:.4f}Ω")
                st.caption(f"L2={l2_mh:.3f}mH  XL2={xl2:.4f}Ω")
                st.caption(f"L_total={l_total*1000:.3f}mH  XL={xl:.4f}Ω")

    # ─────────────────────────────────────────────────────────────────
    # Tab 3: PQ Operating Circle (load convention, one chart with two circles, view centered on current circle)
    # Current limit circle: centered at origin, radius = Vg·|I| = s_rated (VA)
    # Voltage limit circle: center (0, +Vg²/XL), radius = limit_v·Vg/XL (usually much larger than current circle)
    # View centered on current circle (operating point readable); voltage circle drawn only within view
    # ─────────────────────────────────────────────────────────────────
    # ─────────────────────────────────────────────────────────────────
    # ─────────────────────────────────────────────────────────────────
    # Tab 3: PQ Operating Circles (load convention, current + voltage circles on one chart)
    # Rules:
    #   View range centered on current circle (operating point clearly readable)
    #   Voltage circle:
    #     If lower edge (Qc-Rv) is within view → draw normally
    #     If completely outside view → show text reminder, give value on right
    #   Both circles always on the same chart, rendered once
    # ─────────────────────────────────────────────────────────────────
    with tab3:
        # ── Complete all calculations before any st rendering ────────────────────────
        # PQ circles use three-phase total power coordinates
        P_t3      = P_load_3ph
        Q_t3      = Q_load_3ph
        p_t3_kw   = P_t3 / 1000
        q_t3_kvar = Q_t3 / 1000
        op_dist   = np.sqrt(P_t3**2 + Q_t3**2)

        # Voltage limit circle (load convention, rigorous derivation)
        # P² + (Q - Vg²/XL)² ≤ (limit_v·Vg/XL)²
        beyond_v = False
        Qc_load  = 0.0
        Rv       = 0.0
        if xl > 1e-6:
            Qc_load  = 3 * (u_grid_rms**2) / xl   # 3-phase voltage circle center
            Rv       = 3 * (limit_v * u_grid_rms) / xl  # 3-phase voltage circle radius
            dist_v   = np.sqrt(P_t3**2 + (Q_t3 - Qc_load)**2)
            beyond_v = dist_v > Rv

        # Current limit circle (read from right-column number_input, use default first)
        # ⚠️ number_input is in right column — cannot render outside (would duplicate chart)
        # Use session_state or default value here
        i_max_default = float(3 * u_grid_rms * I_mag * 1.2)  # 3-phase VA estimate
        i_rated_max   = st.session_state.get('m4_i_max', i_max_default)
        i_max_is_est  = abs(i_rated_max - i_max_default) < 1.0
        i_limit_r     = 3 * u_grid_rms * i_rated_max  # 3-phase total VA
        beyond_i      = op_dist > i_limit_r

        # View range: centered on current circle; expand to include voltage circle lower edge if within reasonable range
        q_lo_v  = (Qc_load - Rv) if xl > 1e-6 else 0.0
        # : y_ax >= i_limit_r × 1.3
        y_ax = max(i_limit_r * 1.35, abs(P_t3)*1.4, abs(Q_t3)*1.4, 1000.0)
        # 3, 
        v_circle_visible = xl > 1e-6 and abs(q_lo_v) < 3.0 * i_limit_r
        if v_circle_visible:
            y_ax = max(y_ax, abs(q_lo_v) * 1.2)

        # ── ()────────────────────────────────────────
        theta_t = np.linspace(0, 2*np.pi, 400)
        fig_pq  = go.Figure()

        # ① ()
        fig_pq.add_trace(go.Scatter(
            x=i_limit_r * np.cos(theta_t),
            y=i_limit_r * np.sin(theta_t),
            fill='toself', fillcolor='rgba(255,165,0,0.07)',
            name=('Current limit circle (est.)' if is_trial else
                  (f'Current limit circle (est.)  I≈{i_rated_max:.0f} A  |S|≤{i_limit_r/1000:.1f} kVA'
                   if i_max_is_est else
                   f'Current limit circle  I_max={i_rated_max:.0f} A  |S|≤{i_limit_r/1000:.1f} kVA')),
            line=dict(color='#e08000', width=2,
                      dash='dot' if i_max_is_est else 'solid')
        ))

        # ② (): (Plotlyclip)
        if xl > 1e-6:
            if v_circle_visible:
                fig_pq.add_trace(go.Scatter(
                    x=Rv * np.cos(theta_t),
                    y=Qc_load + Rv * np.sin(theta_t),
                    fill='toself', fillcolor='rgba(204,17,17,0.06)',
                    name=(f'{mod_label} voltage limit circle' if is_trial else
                          f'{mod_label} voltage limit circle  center Q={Qc_load/1000:.0f} kvar  R={Rv/1000:.0f} kvar'),
                    line=dict(color='#cc1111', width=2, dash='dash')
                ))
            else:
                # , 
                v_status = "⚠️ exceeds" if beyond_v else "✅ Margin"
                fig_pq.add_annotation(
                    x=0, y=y_ax * 0.87, xanchor='center',
                    text=(f"<b>{mod_label} voltage limit circle (exceeds plot range)</b><br>🔒 Detailed values require paid access"
                          if is_trial else
                          (f"<b>{mod_label} voltage limit circle ({v_status})</b><br>"
                           f"Center Q={Qc_load/1000:.0f} kvar, radius={Rv/1000:.0f} kvar (exceeds plot range)<br>"
                           f"|V_s|={req_v_phase:.0f} V, limit={limit_v:.0f} V, "
                           f"{'exceeds by' if beyond_v else 'margin'} {abs(req_v_phase-limit_v):.0f} V")),
                    showarrow=False,
                    font=dict(color='#cc1111' if beyond_v else '#198754', size=10),
                    bgcolor='rgba(255,230,230,0.90)' if beyond_v else 'rgba(225,245,238,0.90)',
                    bordercolor='#cc1111' if beyond_v else '#198754', borderwidth=1,
                    align='center'
                )

        # ③ Operating point
        op_color = '#dc3545' if (beyond_i or beyond_v) else '#0d6efd'
        fig_pq.add_trace(go.Scatter(
            x=[P_t3], y=[Q_t3], mode='markers',
            marker=dict(size=16, color=op_color, symbol='x',
                        line=dict(width=3, color=op_color)),
                        name='Current Operating Point' if is_trial else f'current P={p_t3_kw:+.1f}kW Q={q_t3_kvar:+.1f}kvar',
                        hoverinfo='none' if is_trial else 'all',
                        hovertemplate=None if is_trial else
                           (f'<b>Current operating point (Load convention)</b><br>'
                           f'P = {p_t3_kw:+.3f} kW<br>'
                           f'Q = {q_t3_kvar:+.3f} kvar<br>'
                           f'|S| = {op_dist/1000:.2f} kVA<br>'
                           f'PF = {pf_val:.4f}<extra></extra>')
        ))
        fig_pq.add_annotation(
            x=P_t3, y=Q_t3,
            text="  Current Point 🔒" if is_trial else f"  P={p_t3_kw:+.0f}kW<br>  Q={q_t3_kvar:+.0f}kvar",
            showarrow=False, xanchor='left', yanchor='middle',
            font=dict(size=10, color=op_color),
            bgcolor='rgba(255,255,255,0.88)')

        # ABCD 
        lbl_r = y_ax * 0.65
        for (ax_, ay_, albl, aclr) in [
            (-lbl_r*0.55,  lbl_r*0.7,  'D generates active (P<0)',  '#666'),
            (-lbl_r*0.55, -lbl_r*0.7,  'B absorbs active (P>0)',   '#666'),
            ( lbl_r*0.55,  lbl_r*0.7,  'A absorbs reactive (Q>0)',  '#1155cc'),
            ( lbl_r*0.55, -lbl_r*0.7,  'C generates reactive (Q<0)',  '#cc1111'),
        ]:
            fig_pq.add_annotation(
                x=ax_, y=ay_, text=f"<b>{albl}</b>",
                showarrow=False, font=dict(size=10, color=aclr),
                bgcolor='rgba(255,255,255,0.75)', align='center')

        fig_pq.update_layout(
            title=dict(
            text=("P-Q Diagram (Load convention, 3-phase power)<br>"
                      "<span style='font-size:11px;color:#666;'>"
                      "P>0 = 3-phase absorbs active (Quadrant B); Q>0 = 3-phase absorbs reactive / under-excited (A); Q<0 = 3-phase generates reactive / over-excited (C)"
                      "</span>"),
                font=dict(size=13)
            ),
            xaxis_title="3-phase active P (W)  —  P>0: VSC absorbs active from grid",
            yaxis_title="3-phase reactive Q (var)  —  Q>0: absorbs reactive / under-excited",
            xaxis=dict(range=[-y_ax, y_ax],
                       showgrid=True, gridcolor='#eee',
                       zeroline=True, zerolinecolor='#999', zerolinewidth=1.5),
            yaxis=dict(range=[-y_ax, y_ax],
                       scaleanchor="x", scaleratio=1,
                       showgrid=True, gridcolor='#eee',
                       zeroline=True, zerolinecolor='#999', zerolinewidth=1.5),
            height=580, margin=dict(l=60, r=20, t=80, b=50),
            legend=dict(x=0.01, y=0.99, xanchor='left', yanchor='top',
                        bgcolor='rgba(255,255,255,0.90)',
                        bordercolor='#ccc', borderwidth=1, font=dict(size=10))
        )

        # ── : () + () ─────────────────────────────────
        c_pq1, c_pq2 = st.columns([3, 1])
        with c_pq1:
            st.plotly_chart(fig_pq, use_container_width=True)
            if i_max_is_est:
                st.caption(
                    "💡 The current limit circle is drawn at **rated current × 1.2**. Enter the rated current in the right panel."
                    if is_trial else
                    f"💡 The current limit circle is drawn at **rated current × 1.2** ({I_mag:.0f} A × 1.2 ≈ {i_rated_max:.0f} A). "
                    "Enter the rated current in the right panel."
                )

        with c_pq2:
            st.markdown("#### 🧠 Diagnosis")

            # Current limit input (kept even in trial mode — needed for circle calculation)
            i_rated_max_input = st.number_input(
                "Hardware Rated 3-Phase Apparent Power S_max (VA)",
                value=i_max_default, min_value=1.0, step=1000.0,
                key="m4_i_max",
                help="Current limit circle radius = 3-phase apparent power upper limit (VA)\n"
                     "= 3 × V_phase × I_max\n"
                     "Leave default for estimated 1.2× current"
            )

            st.divider()

            if is_trial:
                trial_lock("Unlock with paid access: current/voltage margin diagnosis and operating condition analysis")
            else:
                # Current diagnosis
                if beyond_i:
                    st.error(
                        f"**🚨 Exceeds Current Limit**\n\n"
                        f"|S|={op_dist/1000:.1f}kVA > {i_limit_r/1000:.1f}kVA\n\n"
                        "**Suggestion:** Reduce operating current or use higher-rated hardware."
                    )
                else:
                    margin_i = (1 - op_dist/i_limit_r) * 100
                    st.success(f"**✅ Current Margin {margin_i:.1f}%**")
                    st.caption(
                        f"|S|={op_dist/1000:.1f}kVA\n\n"
                        f"Current limit {i_limit_r/1000:.1f}kVA\n\n"
                        f"Margin {(i_limit_r-op_dist)/1000:.1f}kVA"
                    )

                st.divider()

                # Voltage diagnosis
                if beyond_v:
                    need_vdc = u_dc * (req_v_phase / limit_v) * 1.05
                    st.warning(
                        f"**⚠️ Exceeds Voltage Limit**\n\n"
                        f"|V_s|={req_v_phase:.0f}V > {limit_v:.0f}V\n\n"
                        "**Suggestions:**\n\n"
                        f"1. Increase Vdc to ≥{need_vdc:.0f}V\n\n"
                        "2. Reduce inductor L\n\n"
                        "3. SPWM→SVPWM (+15.5%)"
                    )
                else:
                    if xl > 1e-6:
                        dist_v2  = np.sqrt(P_t3**2 + (Q_t3 - Qc_load)**2)
                        margin_v = (1 - dist_v2/Rv) * 100
                        st.success(f"**✅ Voltage Margin {margin_v:.1f}%**")
                        st.caption(
                            f"|V_s|={req_v_phase:.0f}V\n\n"
                            f"{mod_label}={limit_v:.0f}V\n\n"
                            f"Margin+{limit_v-req_v_phase:.0f}V"
                        )

                st.divider()
                st.markdown("**Current operating point (Load convention):**")
                st.write(f"P(3-phase)= **{p_t3_kw:+.3f} kW**")
                st.write(f"Q(3-phase)= **{q_t3_kvar:+.3f} kvar**")
                st.write(f"PF = **{pf_val:.4f}**")
                st.caption(phase_desc_phys)
                if xl > 1e-6:
                    st.divider()
                    st.caption(
                        "**Voltage circle description:**\n\n"
                        "The voltage circle center is on the Q-axis (physically symmetrical)\n\n"
                        f"Qc=Vg²/XL={u_grid_rms:.0f}²/{xl:.4f}={Qc_load/1000:.0f}kvar\n\n"
                        f"Rv={mod_label}×Vg/XL={Rv/1000:.0f}kvar"
                    )












# ==========================================
#  5: Modulation Index & Voltage Utilization
# ==========================================

elif selection == nav_options[4]:
    # 
    from utils.inverter_algo import get_voltage_limits, calculate_modulation_info, generate_modulation_waveforms

    st.header("📊 Modulation Index & Voltage Utilization")

    # -------------------------------------------------------------------------
    # PART 1: Result ()
    # -------------------------------------------------------------------------
    st.markdown("### 1️⃣ Resultdefinition")
    
    col_def1, col_def2 = st.columns([1.2, 1])
    with col_def1:
        st.markdown("#### A. linearmodulationlimit")
        df_limits = pd.DataFrame({
 "modulation": ["SPWM", "SVPWM"],
 "voltage ($V_{ph,peak}$)": ["$0.5 V_{dc}$", "$1/\\sqrt{3} V_{dc} \\approx 0.577 V_{dc}$"],
            "Line Voltage Peak ($V_{ll,peak}$)": ["$0.866 V_{dc}$", "$1.0 V_{dc}$"],
            "Utilization Improvement": ["reference (100%)", "**+15.47%**"]
        })
        st.table(df_limits) #  table 
    with col_def2:
        st.markdown("#### B. Core Concept Definitions")
        st.info("""
        **Modulation Index (m)**: 
        $$ m = \\frac{V_{ph,peak}}{V_{dc}/2} $$
        *   **SPWM Limit**: $m=1.0$ (phase voltage reaches $\\pm V_{dc}/2$)
        *   **SVPWM Limit**: $m=1.155$ (line voltage reaches $\\pm V_{dc}$)
        """)

    # -------------------------------------------------------------------------
    # PART 1.C:  ()
    # -------------------------------------------------------------------------
   
    st.markdown("---")
    st.markdown("### 🔍 C. Physical Model & Waveform Formation Mechanism")
    st.caption("Starting from the topology, fully understand the relationship between voltage coefficients and modulation waveforms.")

    tab_2l, tab_3l, tab_cmp = st.tabs(["⚡ 2-Level Principle", "🧱 3-Level Principle", "🌊 SPWM vs SVPWM Waveform Essentials"])

   # === Tab 1: 2-Level ===
    with tab_2l:
        c_t1, c_t2 = st.columns([1.1, 1]) 
        with c_t1:
            st.markdown("**1. Topology & Reference Point**")
            # --- [SVG]  H  () ---
            svg_2level = """
            <svg width="100%" height="340" viewBox="0 0 460 340" xmlns="http://www.w3.org/2000/svg">
              <defs>
                <marker id="arrow_u2" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth">
                  <path d="M0,0 L0,6 L9,3 z" fill="#6610f2" />
                </marker>
              </defs>

              <!-- 1. DC Bus Lines -->
              <line x1="40" y1="30" x2="390" y2="30" stroke="#dc3545" stroke-width="3"/>
              <text x="395" y="35" fill="#dc3545" font-weight="bold" font-size="13" font-family="sans-serif">+ Vdc/2</text>
              <line x1="40" y1="310" x2="390" y2="310" stroke="#0d6efd" stroke-width="3"/>
              <text x="395" y="315" fill="#0d6efd" font-weight="bold" font-size="13" font-family="sans-serif">- Vdc/2</text>

 <!-- 2. DC Capacitors — IECstandard, coordinate, labelright side -->
 <!-- bus→C1 -->
              <line x1="80" y1="30" x2="80" y2="100" stroke="black" stroke-width="2"/>
 <!-- C1 () -->
              <line x1="62" y1="100" x2="98" y2="100" stroke="black" stroke-width="3"/>
 <!-- C1 (, , IECcapacitorstandard) -->
              <path d="M62,112 Q80,120 98,112" stroke="black" stroke-width="3" fill="none"/>
 <!-- C1 label: capacitorright side, perpendicular(midpointy=106), baseline4px -->
              <text x="103" y="113" font-size="13" font-family="sans-serif" font-weight="bold" fill="#333">C1</text>
 <!-- C1→midpointO connection -->
              <line x1="80" y1="120" x2="80" y2="165" stroke="black" stroke-width="2"/>

              <!-- midpoint O(bridge armnode y=170) -->
              <circle cx="80" cy="170" r="5" fill="#6610f2"/>
 <!-- O(0V) label: , baseline y=168, y=156,  -->
              <text x="2" y="168" fill="#6610f2" font-weight="bold" font-size="12" font-family="sans-serif">O (0V)</text>
 <!-- midpoint -->
              <line x1="85" y1="170" x2="155" y2="170" stroke="#6610f2" stroke-width="1" stroke-dasharray="4"/>

 <!-- midpointO→C2 connection -->
              <line x1="80" y1="175" x2="80" y2="220" stroke="black" stroke-width="2"/>
 <!-- C2 () -->
              <line x1="62" y1="220" x2="98" y2="220" stroke="black" stroke-width="3"/>
 <!-- C2 () -->
              <path d="M62,232 Q80,240 98,232" stroke="black" stroke-width="3" fill="none"/>
 <!-- C2 label: right side, perpendicular y=226, baseline4px -->
              <text x="103" y="233" font-size="13" font-family="sans-serif" font-weight="bold" fill="#333">C2</text>
 <!-- C2→bus connection -->
              <line x1="80" y1="240" x2="80" y2="310" stroke="black" stroke-width="2"/>

              <!-- 3. Leg A -->
              <line x1="190" y1="30" x2="190" y2="80" stroke="black" stroke-width="2"/>
              <rect x="175" y="80" width="30" height="50" fill="#f8f9fa" stroke="black" stroke-width="2"/>
              <text x="181" y="111" font-weight="bold" font-size="14" font-family="sans-serif">S1</text>
              <line x1="190" y1="130" x2="190" y2="210" stroke="black" stroke-width="2"/>
              <circle cx="190" cy="170" r="5" fill="black"/>
 <!-- Alabel: nodeabove, baseline y=158, y=146, node(cy=170)24px -->
              <text x="200" y="158" font-weight="bold" font-size="14" font-family="sans-serif">A</text>
              <rect x="175" y="210" width="30" height="50" fill="#f8f9fa" stroke="black" stroke-width="2"/>
              <text x="181" y="241" font-weight="bold" font-size="14" font-family="sans-serif">S2</text>
              <line x1="190" y1="260" x2="190" y2="310" stroke="black" stroke-width="2"/>

              <!-- 4. Leg B -->
              <line x1="310" y1="30" x2="310" y2="80" stroke="black" stroke-width="2"/>
              <rect x="295" y="80" width="30" height="50" fill="#f8f9fa" stroke="black" stroke-width="2"/>
              <text x="301" y="111" font-weight="bold" font-size="14" font-family="sans-serif">S3</text>
              <line x1="310" y1="130" x2="310" y2="210" stroke="black" stroke-width="2"/>
              <circle cx="310" cy="170" r="5" fill="black"/>
 <!-- Blabel: nodeabove, Alabel -->
              <text x="320" y="158" font-weight="bold" font-size="14" font-family="sans-serif">B</text>
              <rect x="295" y="210" width="30" height="50" fill="#f8f9fa" stroke="black" stroke-width="2"/>
              <text x="301" y="241" font-weight="bold" font-size="14" font-family="sans-serif">S4</text>
              <line x1="310" y1="260" x2="310" y2="310" stroke="black" stroke-width="2"/>

              <!-- 5. u_AO arrow -->
              <line x1="185" y1="170" x2="130" y2="170" stroke="#6610f2" stroke-width="2" marker-end="url(#arrow_u2)"/>
 <!-- u_AOlabel: arrowabove, baseline y=158, arrow(y=170)12px -->
              <text x="132" y="158" fill="#6610f2" font-weight="bold" font-size="13" font-family="sans-serif">u_AO</text>

              <!-- 6. Load A-B -->
              <path d="M190,170 Q250,170 310,170" stroke="#198754" stroke-width="2" stroke-dasharray="5,5"/>
 <!-- Loadlabel: connectionbelow, baseline y=188, connection(y=170)18px -->
              <text x="222" y="188" fill="#198754" font-size="12" font-family="sans-serif">Load (A-B)</text>
            </svg>
            """
            st.markdown('\n'.join(l.strip() for l in svg_2level.splitlines() if l.strip()), unsafe_allow_html=True)
            
            # --- [UPDATED CONTENT]  ---
            with st.expander("🤔 Deep Dive: Switch State S, Duty Cycle d, and Modulation Wave m", expanded=True):
                # First section: basic concept (open)
                st.markdown(r"""
                Many engineers confuse these three concepts. Here is a layered breakdown from **micro** to **macro** level:

                **1. Micro Level: Switch State (S)**
                *   **Definition**: The physical state of the IGBT at any instant.
                *   **Values**: Only **0** (off) or **1** (on).
                *   **Hardware reality**: The switch only knows on/off — **no** half-conducting state exists.
                """)

                # Second section: duty cycle and modulation wave concept chain (locked)
                if is_trial:
                    trial_lock("Unlock with paid access: complete concept chain of duty cycle d and modulation wave m", compact=True)
                else:
                    st.markdown(r"""
                **2. Macro Level: Duty Cycle (d)**
                *   **Duty cycle definition**: Over one **switching period $T_{sw}$** (e.g. $100\mu s$), the ratio of time the upper switch is ON.
                *   **Formula**: $d = \frac{T_{on}}{T_{sw}}$, range $[0, 1]$.
                *   **Physical meaning**: Via **volt-second balance**, micro 0/1 switching is equivalent to outputting $d$ times the voltage.

                **3. Control Level: Modulation Wave (m)**
                *   **Modulation index $m$ definition**: The control algorithm's normalized "desired voltage".
                *   **Value range notes**:
                    - **$[-1, +1]$ (midpoint O reference)**: Most common for AC converters. $m=+1$ outputs $+V_{dc}/2$, $m=-1$ outputs $-V_{dc}/2$, symmetric positive/negative swing. Used throughout this tool.
                    - **$[0, +1]$ (negative bus N reference)**: Common in single-supply or digital PWM registers. $m=1$ is equivalent to duty cycle $d=1$. Identical in essence, just shifted by $+0.5$.
                *   **Relationship**: $m$ is the goal; $d$ is the hardware means.
                    """)

        with c_t2:
            st.markdown("**2. PWM Generation & Waveform Analysis**")
            
            # --- [UPDATED CONTENT]  ---
            st.info("💡 **From Mathematical Intent to Physical Voltage**")
            
            st.markdown(r"""
            **Step 1: Translation**  
            Map control command $m \in [-1, 1]$ to hardware execution time $d \in [0, 1]$.
            $$ d(t) = 0.5 + 0.5 \cdot m(t) $$
            *(Note: 0.5 is the DC offset — time cannot be negative)*

            **Step 2: Execution**  
            Hardware performs high-frequency chopping based on $d$. By **volt-second balance**, the **average output voltage** $\bar{u}_{AO}$ is:
            $$ \bar{u}_{AO} = d \cdot (\frac{V_{dc}}{2}) + (1-d) \cdot (-\frac{V_{dc}}{2}) $$
            Simplifying:
            $$ \bar{u}_{AO} = (d - 0.5) \cdot V_{dc} $$

            """)
            if is_trial:
                trial_lock("Unlock with paid access: Step 3 derivation conclusion (modulation wave directly controls output voltage) and reference point analysis", compact=True)
            else:
                st.markdown(r"""
            **Step 3: Result**
            Substituting Step 1 into Step 2:
            $$ \bar{u}_{AO} = (d - 0.5) \cdot V_{dc} $$
            Substituting $d = 0.5 + 0.5m$:
            $$ \bar{u}_{AO} = \bigl((0.5 + 0.5m) - 0.5\bigr) \cdot V_{dc} = 0.5m \cdot V_{dc} $$
            Giving the final relationship:
            $$ \bar{u}_{AO}(t) = m(t) \cdot \frac{V_{dc}}{2} $$

            > **Key insight**: **The modulation index $m$ determines the ratio of output voltage to bus voltage.**
            > *   $m=1.0 \to$ output $+0.5 V_{dc}$
            > *   $m=0.0 \to$ output $0 V$
            > *   $m=-0.5 \to$ output $-0.25 V_{dc}$
                """)
                st.warning(r"""
            **📝 Alternative View: Using Negative Bus N as Reference (0V)**

            The derivation uses **midpoint O** as reference for symmetric positive/negative swing. In practice, hardware uses **bus N** as 0V reference:

            *   **Voltage range**: Changes from $\pm 0.5 V_{dc}$ to all-positive $[0, V_{dc}]$.
            *   **Simplified formula**: $\bar{u}_{AN} = d \cdot V_{dc}$ — duty cycle directly scales bus voltage.
                $$ \bar{u}_{AN} = d \cdot V_{dc} $$
            *   **Relationship**: $u_{AN}$ is $u_{AO}$ shifted up by half the bus voltage.
                $$ \bar{u}_{AN} = \underbrace{\bar{u}_{AO}}_{\text{AC swing}} + \underbrace{0.5 V_{dc}}_{\text{DC offset}} $$
                """)
            
            # ... ( t_demo = ...) ...

            # --- 2-Level PWM Plot (, , ) ---
            t_demo = np.linspace(0, 0.02, 1000)
            freq_c = 1000
            carrier = 2 * np.abs(2 * (t_demo * freq_c - np.floor(t_demo * freq_c + 0.5))) - 1 
            mod_wave = 0.8 * np.sin(2 * np.pi * 50 * t_demo)
            pwm_out = np.where(mod_wave > carrier, 0.5, -0.5)
            
            # 3-Level:  secondary_y,  + 
            pwm_scaled_2l = pwm_out * 2.0   # scale: +0.5→+1.0, -0.5→-1.0 for display

            fig_pwm = go.Figure()
            fig_pwm.add_trace(go.Scatter(x=t_demo, y=mod_wave, name="Modulation Wave m(t)", line=dict(color='orange', width=2)))
            fig_pwm.add_trace(go.Scatter(x=t_demo, y=carrier, name="Carrier", line=dict(color='gray', width=0.5), opacity=0.3))
            fig_pwm.add_trace(go.Scatter(x=t_demo, y=pwm_scaled_2l, name="Output Voltage uAO", line=dict(color='#0d6efd', width=1.5), opacity=0.8))

            #  annotation
            for val_data, label in [(1.0, "+0.5 Vdc"), (0.0, "0"), (-1.0, "-0.5 Vdc")]:
                fig_pwm.add_annotation(
                    x=1.01, y=val_data, xref="paper", yref="y",
                    text=f"<b>{label}</b>", showarrow=False, xanchor="left",
                    font=dict(size=10, color="#0d6efd"),
                )

            fig_pwm.update_layout(
                height=300,
                margin=dict(l=0, r=65, t=10, b=0),
                legend=dict(orientation="h", y=1.1, x=0),
                hovermode="x unified",
                plot_bgcolor="white",
            )
            fig_pwm.update_yaxes(
                title_text="Modulation / Voltage",
                range=[-1.15, 1.15],
                tickvals=[-1.0, 0, 1.0],
                ticktext=["-0.5 Vdc / -1", "0", "+0.5 Vdc / +1"],
                showgrid=True, gridcolor='rgba(0,0,0,0.08)'
            )
            fig_pwm.update_xaxes(showgrid=True, gridcolor='rgba(0,0,0,0.08)')
            st.plotly_chart(fig_pwm, use_container_width=True)

    # === Tab 2: 3-Level ===
    with tab_3l:
        c_3l_1, c_3l_2 = st.columns([1.1, 1])
        with c_3l_1:
            st.markdown("**1. 3-Leveltopology (NPC)**")
            # --- SVG  ( SVG) ---
            svg_3level = """
            <svg width="100%" height="420" viewBox="0 0 400 420" xmlns="http://www.w3.org/2000/svg">
              <!-- 1. DC Bus Lines -->
              <line x1="50" y1="20" x2="350" y2="20" stroke="#dc3545" stroke-width="3"/>
              <text x="355" y="25" fill="#dc3545" font-weight="bold" font-size="13" font-family="sans-serif">+ Vdc/2</text>
              <line x1="50" y1="400" x2="350" y2="400" stroke="#0d6efd" stroke-width="3"/>
              <text x="355" y="405" fill="#0d6efd" font-weight="bold" font-size="13" font-family="sans-serif">- Vdc/2</text>

 <!-- 2. DC Capacitors — IECstandardcapacitor -->
 <!-- bus→C1 -->
              <line x1="100" y1="20" x2="100" y2="103" stroke="black" stroke-width="2"/>
 <!-- C1 () -->
              <line x1="82" y1="103" x2="118" y2="103" stroke="black" stroke-width="3"/>
 <!-- C1 (, ) -->
              <path d="M82,115 Q100,123 118,115" stroke="black" stroke-width="3" fill="none"/>
 <!-- C1 label — ,  -->
              <text x="122" y="117" font-size="13" font-weight="bold" font-family="sans-serif" fill="#333">C1</text>
 <!-- C1→midpoint -->
              <line x1="100" y1="123" x2="100" y2="210" stroke="black" stroke-width="2"/>

              <!-- midpoint O -->
              <circle cx="100" cy="210" r="5" fill="#6610f2"/>
              <text x="32" y="214" fill="#6610f2" font-weight="bold" font-size="12" font-family="sans-serif">O (0V)</text>

 <!-- midpoint→C2 -->
              <line x1="100" y1="215" x2="100" y2="277" stroke="black" stroke-width="2"/>
 <!-- C2 () -->
              <line x1="82" y1="277" x2="118" y2="277" stroke="black" stroke-width="3"/>
 <!-- C2 () -->
              <path d="M82,289 Q100,297 118,289" stroke="black" stroke-width="3" fill="none"/>
              <!-- C2 label -->
              <text x="122" y="291" font-size="13" font-weight="bold" font-family="sans-serif" fill="#333">C2</text>
 <!-- C2→bus -->
              <line x1="100" y1="297" x2="100" y2="400" stroke="black" stroke-width="2"/>

              <!-- 3. Switch Leg (S1-S4) -->
              <line x1="250" y1="20" x2="250" y2="60" stroke="black" stroke-width="2"/>

              <rect x="235" y="60" width="30" height="40" fill="#f8f9fa" stroke="black" stroke-width="2"/>
              <text x="242" y="85" font-size="13" font-weight="bold" font-family="sans-serif">S1</text>

              <line x1="250" y1="100" x2="250" y2="130" stroke="black" stroke-width="2"/>
              <circle cx="250" cy="130" r="3" fill="black"/>

              <rect x="235" y="130" width="30" height="40" fill="#f8f9fa" stroke="black" stroke-width="2"/>
              <text x="242" y="155" font-size="13" font-weight="bold" font-family="sans-serif">S2</text>

              <line x1="250" y1="170" x2="250" y2="250" stroke="black" stroke-width="2"/>
              <circle cx="250" cy="210" r="5" fill="black"/>
              <line x1="250" y1="210" x2="310" y2="210" stroke="black" stroke-width="2"/>
              <text x="315" y="215" font-weight="bold" font-size="13" font-family="sans-serif">A (Out)</text>

              <rect x="235" y="250" width="30" height="40" fill="#f8f9fa" stroke="black" stroke-width="2"/>
              <text x="242" y="275" font-size="13" font-weight="bold" font-family="sans-serif">S3</text>

              <line x1="250" y1="290" x2="250" y2="320" stroke="black" stroke-width="2"/>
              <circle cx="250" cy="290" r="3" fill="black"/>

              <rect x="235" y="320" width="30" height="40" fill="#f8f9fa" stroke="black" stroke-width="2"/>
              <text x="242" y="345" font-size="13" font-weight="bold" font-family="sans-serif">S4</text>

              <line x1="250" y1="360" x2="250" y2="400" stroke="black" stroke-width="2"/>

              <!-- 4. Clamping Diodes -->
              <line x1="100" y1="210" x2="190" y2="210" stroke="#6610f2" stroke-width="2"/>
              <line x1="190" y1="130" x2="190" y2="290" stroke="#6610f2" stroke-width="2"/>
              <!-- Upper Clamp -->
              <line x1="190" y1="130" x2="250" y2="130" stroke="#6610f2" stroke-width="2"/>
              <polygon points="215,137 215,123 225,130" fill="white" stroke="#6610f2" stroke-width="2"/>
              <line x1="225" y1="123" x2="225" y2="137" stroke="#6610f2" stroke-width="2"/>
              <text x="193" y="120" font-size="11" fill="#6610f2" font-weight="bold">Dc1</text>
              <!-- Lower Clamp -->
              <line x1="190" y1="290" x2="250" y2="290" stroke="#6610f2" stroke-width="2"/>
              <polygon points="225,297 225,283 215,290" fill="white" stroke="#6610f2" stroke-width="2"/>
              <line x1="215" y1="283" x2="215" y2="297" stroke="#6610f2" stroke-width="2"/>
              <text x="193" y="315" font-size="11" fill="#6610f2" font-weight="bold">Dc2</text>

              <!-- 5. Voltage Arrow u_AO -->
              <line x1="305" y1="235" x2="155" y2="235" stroke="#6610f2" stroke-width="2"/>
              <polygon points="155,235 164,232 164,238" fill="#6610f2"/>
              <text x="195" y="228" fill="#6610f2" font-weight="bold" font-size="14" font-family="sans-serif">u_AO</text>
            </svg>
            """
            st.markdown('\n'.join(l.strip() for l in svg_3level.splitlines() if l.strip()), unsafe_allow_html=True)
            
            # --- [NEW]  ---
            with st.expander("🧐 How does 3-Level work? Understanding P, O, N states", expanded=True):
                st.markdown(r"""
The 3-Level inverter uses the **midpoint O (0V)** as an additional switching state.

**1. Three Physical States**
*   **P (+)**: S1, S2 on. Output $+V_{dc}/2$.
*   **O (0)**: S2, S3 on. Output $0V$ (clamped to midpoint via diodes).
*   **N (-)**: S3, S4 on. Output $-V_{dc}/2$.
                """)

                if is_trial:
                    trial_lock("Unlock with paid access: 3-Level duty cycle partition control and voltage output derivation", compact=True)
                else:
                    st.markdown(r"""
**2. Duty Cycle Partition Definition (Local Duty Cycle)**
Unlike 2-Level (one duty cycle $d$), 3-Level uses **partitioned control** based on sign of $m$:
*   **When $m > 0$ (positive half-cycle)**:
    *   System switches between **P** and **O**.
    *   $d_P = m$ (fraction of time connected to positive bus).
*   **When $m < 0$ (negative half-cycle)**:
    *   System switches between **O** and **N**.
    *   $d_N = |m| = -m$ (fraction of time connected to negative bus).

**3. Average Voltage Output**
$$ \bar{u}_{AO}(t) = m(t) \cdot \frac{V_{dc}}{2} $$
*Conclusion: Despite more complex switching, the macroscopic output voltage still strictly follows the modulation wave $m$.*
                    """)
        
        with c_3l_2:
            st.markdown("**2. 3-Level PWM generate (PD modulation)**")
            st.info("💡 **Key takeaway**: Regardless of topology, the modulation index $m$ directly determines the output voltage.")
            
            # --- 3-Level PWM Plot Calculation ---
            # Fix: 10002001, 1000Hz(t=0.25ms)
            # ,  c_top/c_bot .
            # : N=1000dt=20μs, t=0.25ms, 0.25/0.02=12.5, 
            # (tri_max≈0.998)(-1.0), 1.
            # N=2001dt≈10μs, 0.25ms/0.01ms=25, .
            t_3l = np.linspace(0, 0.02, 2001)

            # 1.  m(t)
            m_wave_3l = 0.85 * np.sin(2 * np.pi * 50 * t_3l)

            # 2.  (PD: Phase Disposition)
            # Fix:  floor , 
            freq_c_3l = 1000
            _phase_3l = (t_3l * freq_c_3l) % 1.0  # phasenormalized [0, 1)
            # : 0→0.5+1, 0.5→1-1
            tri_base = np.where(_phase_3l < 0.5, 4 * _phase_3l - 1, 3 - 4 * _phase_3l)
            #  [0, 1]  C_top
            c_top = 0.5 * (tri_base + 1)
            #  [-1, 0]  C_bot
            c_bot = 0.5 * (tri_base - 1)

            # 3.  (PWM Logic)
            # m > c_top -> P state (+0.5 Vdc)
            # m < c_bot -> N state (-0.5 Vdc)
            #       -> O state (0 V)
            pwm_out_3l = np.where(m_wave_3l > c_top, 0.5,
                                  np.where(m_wave_3l < c_bot, -0.5, 0.0))
            # --- Plotting ---
            # ── "" ─────────────────────────────────────
            # : Plotly secondary_y  range , 
            #  uAO=+0.5  c_top=1.0.
            # : ,  pwm  2 : 
            #   +0.5 × 2 = +1.0  →  c_top 
            #   -0.5 × 2 = -1.0  →  c_bot 
            #    0.0 × 2 =  0.0  → 
            #  annotation , .
            pwm_scaled = pwm_out_3l * 2.0   # axis [-1, 1]

            fig_3l = go.Figure()

            # (, )
            fig_3l.add_trace(go.Scatter(
                x=t_3l*1000, y=c_top,
                name="1 (0~1)",
                line=dict(color='gray', width=0.8, dash='dot'),
                showlegend=True
            ))
            fig_3l.add_trace(go.Scatter(
                x=t_3l*1000, y=c_bot,
                name="2 (-1~0)",
                line=dict(color='gray', width=0.8, dash='dot'),
                showlegend=True
            ))
            fig_3l.add_trace(go.Scatter(
                x=t_3l*1000, y=m_wave_3l,
                name="Modulation Wave m(t)",
                line=dict(color='orange', width=2.5)
            ))

            # uAO: , +0.5(P)c_top=1.0
            fig_3l.add_trace(go.Scatter(
                x=t_3l*1000, y=pwm_scaled,
                name="uAO (3-Level)",
                line=dict(color='#198754', width=2),
                opacity=0.9
            ))

            # ( secondary_y axis)
            # annotation  x  paper (1.0=), y  data ()
            for val_data, label in [(1.0, "+0.5 (P)"), (0.0, "0 (O)"), (-1.0, "-0.5 (N)")]:
                fig_3l.add_annotation(
                    x=1.01, y=val_data,
                    xref="paper", yref="y",
                    text=f"<b>{label}</b>",
                    showarrow=False,
                    xanchor="left",
                    font=dict(size=11, color="#198754"),
                    bgcolor="rgba(255,255,255,0)"
                )

            # (, )
            for val_data in [1.0, 0.0, -1.0]:
                fig_3l.add_hline(
                    y=val_data,
                    line=dict(color='rgba(25,135,84,0.15)', width=1, dash='dot')
                )

            # 
            fig_3l.update_layout(
                height=370,
                margin=dict(l=0, r=70, t=35, b=0),
                legend=dict(orientation="h", y=1.07, x=0),
                hovermode="x unified",
                title="modulation (PD-PWM) principle",
                plot_bgcolor="white",
                paper_bgcolor="white",
            )
            fig_3l.update_xaxes(title_text="Time (ms)", showgrid=True, gridcolor='rgba(0,0,0,0.08)')
            fig_3l.update_yaxes(
                title_text="Modulation / Carrier   |   Voltage (×Vdc/2)",
                range=[-1.15, 1.15],
                tickvals=[-1.0, -0.5, 0, 0.5, 1.0],
                ticktext=["-1 / -0.5(N)", "-0.5", "0 / 0(O)", "0.5", "1 / +0.5(P)"],
                showgrid=True,
                gridcolor='rgba(0,0,0,0.08)'
            )

            st.plotly_chart(fig_3l, use_container_width=True)
            
            st.success(r"""
**🔬 Key Observation:**
1. **Positive half-cycle**: When $m > 0$, **Carrier 1** is active. Output switches between $+0.5 V_{dc}$ and $0V$. **Carrier 2** stays idle.
            """)
            if is_trial:
                trial_lock("Unlock with paid access: 3-Level PWM negative half-cycle behavior and harmonic advantage analysis", compact=True)
            else:
                st.success(r"""
2. **Negative half-cycle**: When $m < 0$, **Carrier 2** is active. Output switches between $0V$ and $-0.5 V_{dc}$.
3. **Result**: Compared to 2-Level (switching between $+V_{dc}$ and $-V_{dc}$), the 3-Level step size is halved ($V_{dc}/2$), and harmonic content is dramatically reduced.
                """)

 # === Tab 3: SPWM vs SVPWM  ===
    with tab_cmp:
        st.markdown("### 🐎 SVPWM Deep Dive:")
        
        # --- 1.  Expander ---
        with st.expander("📖 Understanding SVPWM: switches S, duty cycle d, modulation m, and voltage u", expanded=True):
            st.markdown(r"""
Common question: **How is SVPWM different from SPWM?** Comparison:

| Dimension | SPWM (Sinusoidal PWM) | SVPWM (Space Vector) | Change in Essence |
| :--- | :--- | :--- | :--- |
| **1. Switch State S** | 0 or 1 | 0 or 1 | **No change** (hardware only switches) |
| **2. Duty cycle d** | Sinusoidal variation | **Saddle-shaped variation** | Follows modulation wave |
| **3. Modulation wave m** | **Pure sinusoid** | **Saddle wave** (sine + zero-sequence injection) | **Key difference!** |
| **4. Phase voltage $u_{AO}$** | Sinusoidal | **Saddle wave** | Looks "distorted" |
| **5. Line voltage $u_{AB}$** | Sinusoidal | **Sinusoidal (larger amplitude)** | **Common-mode cancels — perfect result** |

            """)
            # Core logic explanation (locked for trial users)
            if is_trial:
                trial_lock("Unlock with paid access: SVPWM common-mode injection and the physical essence of improved voltage utilization")
            else:
                st.markdown(r"""
 **💡 Core logic:**
 SVPWM simultaneously injects a **"common-mode component"** into all three phase modulation waves (flattening peaks, raising troughs — forming a saddle shape).
 * **Phase voltage $u_{AO}$**: The waveform becomes saddle-shaped, so the fundamental can be made larger without hitting the $[-1, 1]$ rail.
 * **Line voltage $u_{AB}$**: Since $u_{AB} = u_{AO} - u_{BO}$, the common-mode component **cancels out**. The load sees a pure sinusoid, but voltage utilization is improved by **15.47%**.
                """)

        st.divider()

        # --- 2.  ---
        st.markdown("**📊 Simulation verification: Is the output actually the same?**")
        
        col_sv_1, col_sv_2 = st.columns([1, 3])
        
        with col_sv_1:
            st.info("Adjust the parameters below to see the effect on waveforms.")
            #  m ,  SVPWM 
            sim_m_sv = st.slider("Modulation Index m (SVPWM)", 0.1, 1.15, 1.15, 0.05, key='sv_slider_m')
            st.caption(f"When m > 1.0, SPWM enters over-modulation, but SVPWM has a safety margin up to 1.155.")

        with col_sv_2:
            # ---  ---
            deg = np.linspace(0, 360, 360)
            rad = np.radians(deg)
            
            # 1.  (SPWM Base)
            u_a_sin = sim_m_sv * np.sin(rad)
            u_b_sin = sim_m_sv * np.sin(rad - 2*np.pi/3)
            
            # 2. Zero-Sequence Component (SVPWM Injection)
            #  Min-Max : Offset = -0.5 * (Max + Min)
            u_c_sin = sim_m_sv * np.sin(rad + 2*np.pi/3)
            v_max = np.maximum(np.maximum(u_a_sin, u_b_sin), u_c_sin)
            v_min = np.minimum(np.minimum(u_a_sin, u_b_sin), u_c_sin)
            u_offset = -0.5 * (v_max + v_min)
            
            # 3.  (SVPWM Phase Voltage)
            u_a_sv = u_a_sin + u_offset
            u_b_sv = u_b_sin + u_offset
            
            # 4.  (Line Voltage)
            # u_ab = u_a - u_b. : Offset !
            u_ab_sv = u_a_sv - u_b_sv
            # ,  sqrt(3) , 
            u_ab_norm = u_ab_sv / np.sqrt(3)

            # ---  ---
            fig_sv = make_subplots(rows=2, cols=1, shared_xaxes=True,
                                   subplot_titles=("1. Phase voltage uAO (SVPWM saddle)", "2. Line voltage uAB (identical)"))

            # Subplot 1: Phase Voltages
            fig_sv.add_trace(go.Scatter(x=deg, y=u_a_sin, name='Sine reference', line=dict(color='gray', dash='dot', width=1)), row=1, col=1)
            fig_sv.add_trace(go.Scatter(x=deg, y=u_a_sv, name='A uAO', line=dict(color='#d63384', width=3)), row=1, col=1)
            fig_sv.add_trace(go.Scatter(x=deg, y=u_b_sv, name='B uBO', line=dict(color='#0d6efd', width=3)), row=1, col=1)
            
            # 
            fig_sv.add_hline(y=1.0, line_dash="dash", line_color="red", row=1, col=1)
            fig_sv.add_hline(y=-1.0, line_dash="dash", line_color="red", row=1, col=1)

            # Subplot 2: Line Voltage
            fig_sv.add_trace(go.Scatter(x=deg, y=u_ab_sv, name='Line voltage uAB (actual output)', line=dict(color='#198754', width=3)), row=2, col=1)
            
            fig_sv.update_layout(height=500, hovermode="x unified")
            fig_sv.update_yaxes(title_text="amplitude (x Vdc/2)", row=1, col=1)
            fig_sv.update_yaxes(title_text="amplitude (x Vdc/2)", row=2, col=1)
            fig_sv.update_xaxes(title_text="phase (°)", row=2, col=1)
            
            st.plotly_chart(fig_sv, use_container_width=True)

            if is_trial:
                trial_lock("Unlock with paid access: saddle wave and line voltage restoration — interactive verification conclusion")
            else:
                st.success(f"""
**Conclusion:**
1. **Phase voltage**: At $m={sim_m_sv}$, the pure sinusoidal reference **exceeds ±1** (over-modulation, clips). But SVPWM's saddle wave stays within ±1 boundary!
2. **Line voltage**: After subtraction ($u_{{AO}} - u_{{BO}}$), both give **identical amplitude and waveform**. This proves SVPWM's voltage utilization advantage is real physics, not a trick.
                """)
    # -------------------------------------------------------------------------
    # PART 2:  ()
    # -------------------------------------------------------------------------
    st.markdown("### 2️⃣ simulation: ""physical")
    # ... ( Part 2 ) ...
    
    # ,  Part 2 , 
    col_sim_ctrl, col_sim_plot = st.columns([1, 2])
    with col_sim_ctrl:
        sim_vdc = st.number_input("simulation Vdc (V)", value=600.0, step=50.0, key="sim_vdc")
        limit_line = sim_vdc / 2.0
        sim_mode = st.radio("modulation", ["SPWM", "SVPWM"], horizontal=True, key="sim_mode")
        sim_m = st.slider(f"Modulation Index m", 0.1, 1.25, 1.05 if sim_mode=="SVPWM" else 0.9, 0.01)

    # 
    wave_data = generate_modulation_waveforms(sim_vdc, sim_m)
    v_fund_peak = sim_m * (sim_vdc / 2.0)
    wave_fundamental = v_fund_peak * np.sin(np.radians(wave_data['deg']))

    fig = go.Figure()
    # 1. 
    fig.add_hrect(y0=limit_line, y1=limit_line*1.4, fillcolor="red", opacity=0.1, line_width=0)
    fig.add_hrect(y0=-limit_line*1.4, y1=-limit_line, fillcolor="red", opacity=0.1, line_width=0)
    fig.add_hline(y=limit_line, line_dash="dash", line_color="#dc3545", opacity=0.8)
    fig.add_hline(y=-limit_line, line_dash="dash", line_color="#dc3545", opacity=0.8)
    # 2.
    fig.add_trace(go.Scatter(x=wave_data['deg'], y=wave_fundamental, name='Desired fundamental output', line=dict(color='gray', dash='dash', width=2), opacity=0.6))
    # 3.
    if sim_mode == "SPWM":
        y_mod = wave_data['spwm']
        color_mod = "#0d6efd"
    else:
        y_mod = wave_data['svpwm']
        color_mod = "#198754"
    fig.add_trace(go.Scatter(x=wave_data['deg'], y=y_mod, name=f'Actual modulation wave ({sim_mode})', line=dict(color=color_mod, width=4)))
    

    # ── Modulation status warning ────────────────────────────────────────────
    if is_trial:
        trial_lock("Unlock with paid access: modulation status diagnosis and over-modulation risk analysis")
    else:
        if sim_mode == "SPWM":
            if sim_m <= 1.0:
                st.success(f"✅ **SPWM Linear Zone**: m = {sim_m:.2f}, safely within linear range (m ≤ 1.0). "
                           f"Output voltage = {sim_m * sim_vdc / 2:.1f} V, no distortion.")
            else:
                st.error(f"❌ **SPWM Over-modulation (clipping distortion)**: m = {sim_m:.2f} > 1.0. "
                         f"Required phase voltage {sim_m * sim_vdc / 2:.1f} V exceeds half-bus voltage {sim_vdc/2:.1f} V. "
                         f"Waveform is hard-clipped, output contains **severe harmonic distortion**. "
                         f"\n\n**Solutions**: ① Increase Vdc; ② Switch to SVPWM (safe up to m = 1.155).")
        else:  # SVPWM
            if sim_m <= 1.0:
                st.success(f"✅ **SVPWM Safe Zone (with margin)**: m = {sim_m:.2f}, well below SVPWM limit of 1.155. "
                           f"Saddle wave peak has not reached rail yet — voltage utilization not yet maximized, consider increasing m.")
            elif sim_m <= 1.1547:
                st.success(f"⚡ **SVPWM Near Linear Limit (golden zone)**: m = {sim_m:.2f}. "
                           f"Pure sinusoidal fundamental (grey dashed) already exceeds ±{sim_vdc/2:.0f} V, but after "
                           f"zero-sequence injection the saddle wave (green solid) peak stays within ±{sim_vdc/2:.0f} V safety boundary. "
                           f"Line voltage utilization at maximum — **this is SVPWM's core value**.")
            else:
                st.error(f"🚨 **SVPWM Over-modulation (deep over-modulation)**: m = {sim_m:.2f} > 1.155. "
                         f"Exceeds SVPWM linear limit; saddle wave peak breaks through ±{sim_vdc/2:.0f} V boundary. "
                         f"Output waveform enters **nonlinear over-modulation region** — harmonics surge, low-order harmonics (5th, 7th) increase significantly. "
                         f"\n\n**Consequences**: transformer core saturation risk, protection mis-operation, grid harmonic violations. "
                         f"\n**Solutions**: Increase Vdc immediately, or reduce modulation index target below 1.155.")

    fig.update_layout(title=dict(text=f"waveformView: {sim_mode} @ m={sim_m:.2f}"), xaxis_title="phase (°)", yaxis_title="Voltage (V)", height=450, margin=dict(l=20, r=20, t=40, b=20), hovermode="x unified")
    st.plotly_chart(fig, use_container_width=True)

    st.divider()

    # -------------------------------------------------------------------------
    # PART 3:  ()
    # -------------------------------------------------------------------------
    st.markdown("### 3️⃣ engineeringdesignverification")
    with st.expander("🛠️ openvoltageMargincalculate", expanded=True):
        col_calc1, col_calc2 = st.columns(2)
        with col_calc1:
            c_vdc = st.number_input("actualbus (V)", value=sim_vdc, key="calc_vdc")
            c_vac = st.number_input("grid-connectedvoltage (V)", value=380.0, step=10.0, key="calc_vac")
        with col_calc2:
            if is_trial:
                trial_lock("Unlock with paid access: modulation index calculation and voltage margin diagnosis")
            else:
                req_v_peak = c_vac * np.sqrt(2) / np.sqrt(3)
                curr_m = req_v_peak / (c_vdc / 2)
                st.markdown(f"**Required Phase Voltage Peak**: `{req_v_peak:.1f} V`")
                st.markdown(f"**Current Modulation Index**: `m = {curr_m:.3f}`")
                if curr_m > 1.1547: st.error("❌ Severely insufficient (increase Vdc)")
                elif curr_m > 1.0: st.warning("⚠️ Must use SVPWM")
                else: st.success("✅ Voltage sufficient")









# ==========================================
#  6: Ripple Current Analysis
# ==========================================

elif selection == nav_options[5]:
    st.header("🌊 Ripple Current Calculation & Analysis")
    st.caption("Module flow: 1. Design tool (Results) → 2. Principle derivation (Math) → 3. Simulation verification (Waveforms)")

    # -------------------------------------------------------------------------
    # PART 1:  (The Calculator) - 
    # -------------------------------------------------------------------------
    st.markdown("### 1️⃣ Core Formula & Calculation")
    st.info("💡 **Engineering result**: Select your topology and modulation scheme, read the coefficient **K**, then substitute into the ripple formula.")

    col_c1, col_c2, col_c3 = st.columns([1, 1.5, 1])
    
    with col_c1:
        st.markdown("**Parameter Settings**")
        c_topo = st.selectbox("topology", ["2-Level", "3-Level (NPC)"], key="calc_topo")
        c_mod = st.selectbox("modulation", ["SVPWM", "SPWM"], key="calc_mod")
        c_vdc = st.number_input("Vdc (V)", value=st.session_state['vec_udc'], step=10.0)
        c_fsw = st.number_input("fsw (Hz)", value=st.session_state['sys_fsw'], step=100.0)
        c_L = st.number_input("L1 (mH)", value=st.session_state['lcl_l1'], step=0.01, format="%.3f") * 1e-3

    with col_c2:
        st.markdown("**Calculation Formula**")
        
        # 
        if c_topo == "2-Level":
            # 2-Level SPWM/SVPWM  4 ()
            k_val = 4
            # :  r,  \Delta 
            formula_tex = r"\Delta I_{pp,max} = \frac{V_{dc}}{4 \cdot L \cdot f_{sw}}"
            desc = "2-Level standard formula (applies to both SPWM and SVPWM)"
        else:
            # 3-Level
            if c_mod == "SPWM":
                k_val = 8
                formula_tex = r"\Delta I_{pp,max} = \frac{V_{dc}}{8 \cdot L \cdot f_{sw}}"
                desc = "3-Level SPWM (equivalent to half-voltage 2-Level)"
            else:
                k_val = 12
                formula_tex = r"\Delta I_{pp,max} = \frac{V_{dc}}{12 \cdot L \cdot f_{sw}}"
                desc = "3-Level SVPWM (optimized switching, lower ripple)"

        st.latex(formula_tex)
        st.caption(f"📌 {desc}")

    with col_c3:
        st.markdown("**Calculation Result**")
        if c_L > 0 and c_fsw > 0:
            res_ipp = c_vdc / (k_val * c_L * c_fsw)
            res_ipk = res_ipp / 2
        else:
            res_ipp = 0
            res_ipk = 0

        if is_trial:
            trial_lock("Ripple current results require paid access")
        else:
            st.metric("Peak-to-Peak (ΔIpp)", f"{res_ipp:.2f} A")
            st.metric("Amplitude (0-to-Peak)", f"{res_ipk:.2f} A")

            #
            i_ref = st.session_state['vec_i'] * np.sqrt(2)
            if i_ref > 0:
                st.progress(min(res_ipp/i_ref, 1.0))
                st.caption(f"Ratio to rated peak: {res_ipp/i_ref:.1%}")

    st.divider()

    # -------------------------------------------------------------------------
    # PART 2:  (The Derivation) - ?
    # -------------------------------------------------------------------------
    st.markdown("### 2️⃣ Why are the coefficients 1/4 and 1/12? (Mathematical Derivation)")
    
    # Mathematical derivation expander
    with st.expander("📐 Mathematical Derivation — Why 1/4? (full derivation with microscopic breakdown)", expanded=False):
        # Steps 1-2: setup (open)
        st.markdown(r"""
        For the classic **2-Level** topology, here is how coefficient **1/4** is derived from first principles.

        #### Step 1: Microscopic Voltage Analysis
        Within one switching period $T_{sw}$, when the IGBT turns on and off, the voltage across inductor $L$ jumps:
        1.  **On-state ($T_{on} = D \cdot T_{sw}$)**:
            *   Converter output: $V_{dc}$
            *   Inductor voltage: $V_L = V_{dc} - V_{grid}$
            *   Current behaviour: **rises linearly**
        2.  **Off-state ($T_{off} = (1-D) \cdot T_{sw}$)**:
            *   Converter output: $0$
            *   Inductor voltage: $V_L = 0 - V_{grid} = -V_{grid}$
            *   Current behaviour: **falls linearly**

        #### Step 2: Build the Integral Equation
        Using inductor characteristic $\Delta I = \frac{1}{L} \int V_L \, dt$, the current rise during the on-state equals the peak-to-peak ripple:
        $$ \Delta I_{pp} = \frac{(V_{dc} - V_{grid}) \cdot (D \cdot T_{sw})}{L} $$

        In steady state, volt-second balance gives $V_{grid} \approx V_{out,avg} = D \cdot V_{dc}$.
        Substituting:
        $$ \Delta I_{pp} = \frac{(V_{dc} - D \cdot V_{dc}) \cdot D}{L \cdot f_{sw}} = \frac{V_{dc}}{L \cdot f_{sw}} \cdot \underbrace{D(1-D)}_{\text{parabolic factor}} $$
        """)

        # Steps 3-4: key conclusion (locked)
        if is_trial:
            trial_lock("Unlock with paid access: extremum derivation (why exactly 1/4) and 3-Level generalisation conclusion", compact=True)
        else:
            st.markdown(r"""
        #### Step 3: Find the Maximum (The Mathematics)
        We need to find when the ripple is largest.
        *   The function $f(D) = D(1-D)$ is a downward-opening parabola.
        *   At **$D = 0.5$**, it reaches its maximum: $0.5 \times (1-0.5) = 0.25 = \mathbf{1/4}$.

        #### Step 4: Conclusion
        $$ \Delta I_{pp, max} = \frac{V_{dc}}{L \cdot f_{sw}} \cdot \frac{1}{4} $$

        > **Generalisation to 3-Level SVPWM**:
        > 3-Level halves the voltage step to $V_{dc}/2$, and SVPWM optimises the switching sequence, reducing the effective integral area. Full derivation yields a maximum coefficient of **1/12**.
            """)

    st.divider()

    # -------------------------------------------------------------------------
    # PART 3:  (The Simulation) - 
    # -------------------------------------------------------------------------
    st.markdown("### 3️⃣ Simulation Lab")

    # ===  A:  (D=0.5 ) ===
    st.markdown("#### 🔬 Simulation A: Why does maximum ripple occur at D=0.5?")
    col_sim_a1, col_sim_a2 = st.columns([1, 2])
    
    with col_sim_a1:
        st.write("Drag the slider to change duty cycle **D** and observe how the **inductor current ripple amplitude** changes.")
        d_slider = st.slider("duty cycle D", 0.1, 0.9, 0.5, 0.1)
        # , 
        st.caption(f"current D(1-D) = {d_slider*(1-d_slider):.2f}")
        if abs(d_slider - 0.5) < 0.01:
            st.success("🔥 Maximum ripple confirmed at D = 0.5!")
    
    with col_sim_a2:
        # ---  1: Parameter,  () ---
        #  session_state Parameter, 
        real_vdc = st.session_state.get('vec_udc', 1100.0) # V
        real_fsw = st.session_state.get('sys_fsw', 2500.0) # Hz
        # :  L  Henry
        real_l = st.session_state.get('lcl_l1', 0.5) * 1e-3 
        
        #  Base = Vdc / (L * fsw)
        if real_l > 0 and real_fsw > 0:
            base_ripple = real_vdc / (real_l * real_fsw)
        else:
            base_ripple = 0
            
        # ---  2:  ---
        t_micro = np.linspace(0, 1, 500) 
        factor_val = d_slider * (1 - d_slider) # D(1-D)
        
        # 
        real_ipp = base_ripple * factor_val
        
        #  (,  real_ipp)
        # :  t < D,  = real_ipp
        i_micro = []
        for t in t_micro:
            if t <= d_slider:
                # :  -ipp/2 ,  = ipp / D
                val = -real_ipp/2 + (real_ipp / d_slider) * t
            else:
                # :  +ipp/2 ,  = -ipp / (1-D)
                val = real_ipp/2 - (real_ipp / (1-d_slider)) * (t - d_slider)
            i_micro.append(val)
        
        fig_micro = go.Figure()
        
        # 
        fig_micro.add_trace(go.Scatter(
            x=t_micro, 
            y=i_micro, 
            fill='tozeroy', 
            name='Inductor ripple current',
            line=dict(color='#0d6efd', width=2)
        ))
        
        # ---  3:  ---
        # , 
        label_text = ("🔒 Paid access" if is_trial else
                      f"<b>Peak-to-Peak ΔIpp = {real_ipp:.2f} A</b><br>(coefficient k={factor_val:.2f})")
        
        fig_micro.add_annotation(
            x=d_slider,        # arrow X
            y=real_ipp/2,      # arrow Y
            ax=0,              # 
            ay=-40,            # perpendicular ()
            text=label_text,
            showarrow=True,
            arrowhead=2,
            bgcolor="rgba(255,255,255,0.8)", # 
            font=dict(size=12, color="#d63384")
        )
        
        fig_micro.update_layout(
        title=dict(text="Single-period Ripple Current Waveform (current design parameters)", font=dict(size=14)),
        xaxis_title="Normalised time (t/Tsw)",
        yaxis_title="Ripple Current (A)",
            height=300,
            margin=dict(t=40, b=20, l=40, r=20),
            hovermode=False if is_trial else "x unified"
        )
        st.plotly_chart(fig_micro, use_container_width=True)

        if not is_trial:
            st.caption(f"💡 **Design insight**: With current parameters, if duty cycle runs at 0.5 (worst case), the inductor carries **{real_ipp:.2f} A** peak-to-peak high-frequency ripple.")


    # ===  B:  () ===
    st.markdown("#### 🔭 Simulation B: Ripple distribution over one power cycle (SPWM)")
    st.info("💡 **Tip**: Use the slider to explore how duty cycle $D$ affects ripple. Note: SVPWM follows the same physical principle as SPWM for ripple calculation.")
    
    col_sim_b1, col_sim_b2 = st.columns([1, 3])
    
    with col_sim_b1:
        # 
        p_vdc = st.number_input("simulation Vdc (V)", value=st.session_state.get('vec_udc', 1100.0), step=50.0)
        p_fsw = st.number_input("simulation fsw (Hz)", value=st.session_state.get('sys_fsw', 2500.0), step=100.0)
        default_l1 = st.session_state.get('lcl_l1', 0.5)
        p_l1_mh = st.number_input("simulation L1 (mH)", value=default_l1, step=0.05, format="%.3f")
        p_l1 = p_l1_mh * 1e-3

        m_index = st.slider("Modulation Index m", 0.1, 1.15, 0.9, 0.05)
        
        # 
        k_factor = p_vdc / (p_l1 * p_fsw) if p_l1 > 0 else 0
        max_ripple_theory_2l = k_factor / 4  # 1/4
        max_ripple_theory_3l = k_factor / 8  # 1/8 (SPWM)

        st.markdown("---")
        st.caption(f"🏁 **SPWM theoretical peak**")
        if is_trial:
            st.caption("2-Level Max: **🔒**")
            st.caption("3-Level Max: **🔒**")
        else:
            st.caption(f"2-Level Max: **{max_ripple_theory_2l:.1f} A**")
            st.caption(f"3-Level Max: **{max_ripple_theory_3l:.1f} A**")
        
        st.warning("Note: For SVPWM, the 3-Level topology reduces peak ripple to approximately 85% of equivalent SPWM.")

    with col_sim_b2:
        # 1. 
        theta = np.linspace(0, 2*np.pi, 400)
        theta_deg = theta * 180 / np.pi
        
        # 2.  D(t) - SPWM
        d_t = 0.5 + 0.5 * m_index * np.sin(theta)
        
        # 3.  - 2 Level
        rip_2l = k_factor * d_t * (1 - d_t)
        
        # 4.  - 3 Level
        rip_3l = []
        for d in d_t:
            if d < 0.5:
                d_local = d * 2 
                val = (k_factor / 2) * d_local * (1 - d_local)
            else:
                d_local = (d - 0.5) * 2
                val = (k_factor / 2) * d_local * (1 - d_local)
            rip_3l.append(val)
        rip_3l = np.array(rip_3l)

        # 5.  ()
        #  rip_3l 
        max_idx_3l = np.argmax(rip_3l)
        max_val_3l_sim = rip_3l[max_idx_3l]
        max_theta_3l = theta_deg[max_idx_3l]
        
        #  2-Level  (0180, )
        # ,  0 ()
        
        # 6. 
        fig_macro = make_subplots(
            rows=2, cols=1, 
            shared_xaxes=True,
            vertical_spacing=0.1,
            subplot_titles=("(Cause) Duty cycle D(t) — SPWM sinusoidal variation", "(Effect) Ripple current peak-to-peak distribution"),
            row_heights=[0.3, 0.7]
        )
        
        # 1
        fig_macro.add_trace(go.Scatter(x=theta_deg, y=d_t, name='D(t)', line=dict(color='gray', dash='dot')), row=1, col=1)
        fig_macro.add_hline(y=0.5, line_width=1, line_color='black', row=1, col=1)
        # 
        fig_macro.add_hline(y=0.25, line_width=1, line_dash='dot', line_color='#198754', opacity=0.3, row=1, col=1)
        fig_macro.add_hline(y=0.75, line_width=1, line_dash='dot', line_color='#198754', opacity=0.3, row=1, col=1)

        # 2
        fig_macro.add_trace(go.Scatter(
            x=theta_deg, y=rip_2l, name='2-Level SPWM', 
            line=dict(color='#dc3545', width=3),
            fill='tozeroy', fillcolor='rgba(220, 53, 69, 0.1)'
        ), row=2, col=1)
        
        fig_macro.add_trace(go.Scatter(
            x=theta_deg, y=rip_3l, name='3-Level SPWM', 
            line=dict(color='#198754', width=3),
            fill='tozeroy', fillcolor='rgba(25, 135, 84, 0.2)'
        ), row=2, col=1)
        
        fig_macro.update_layout(height=550, hovermode=False if is_trial else "x unified")
        fig_macro.update_yaxes(title_text="Ripple (A)", row=2, col=1)
        fig_macro.update_xaxes(title_text="phase (°)", row=2, col=1, tickvals=[0, 90, 180, 270, 360])
        
        # Note: 2-Level ()
        fig_macro.add_annotation(
            x=0, y=rip_2l[0], row=2, col=1,
            text="<b>2-Level worst case</b><br>zero-crossing (D=0.5)",
            showarrow=True, arrowhead=2, ax=60, ay=-10, bgcolor="rgba(255,255,255,0.9)", bordercolor="#dc3545"
        )
        
        # Note: 3-Level (!)
        #  max_theta_3l  max_val_3l_sim
        fig_macro.add_annotation(
            x=max_theta_3l, y=max_val_3l_sim, row=2, col=1,
            text="<b>3-Level worst case</b><br>D=0.25/0.75<br>(half-amplitude voltage)",
            showarrow=True, arrowhead=2, 
            ax=0, ay=-50, # arrow pointing upward
            bgcolor="rgba(255,255,255,0.9)", bordercolor="#198754", font=dict(color="#198754")
        )
        
        st.plotly_chart(fig_macro, use_container_width=True)
        
        # Analysis result
        # Point 1: open to all users
        st.success(f"""
        ### 💡 Expert Analysis: How to read this chart?

        **1. Physical root: duty cycle D determines everything**
        *   **Fundamental law**: Regardless of topology, inductor ripple is always proportional to D(1-D). When D is close to **0.5**, ripple is highest; when D is close to **0** or **1**, ripple is lowest.
        """)

        # Points 2-4: locked for trial users
        if is_trial:
            trial_lock("Unlock with paid access: 2-Level vs 3-Level comparison, counter-intuitive modulation index effects, and SVPWM ripple advantage quantification")
        else:
            st.success(f"""
        **2. 2-Level vs 3-Level: the worst case shifts**
        *   **Near zero-crossing (0°/180°)**: This is 2-Level's **darkest moment** — to output 0V, D must stay near 0.5, driving ripple to its cycle maximum. 3-Level has near-zero ripple here, a major advantage.
        *   **Near half-amplitude voltage (D ≈ 0.25/0.75)**: This is 3-Level's **hardest moment** (the twin peaks in the green curve). 3-Level avoids the 2-Level zero-crossing worst case, but trades it for a different peak.

        **3. Counter-intuitive modulation index effect**
        *   Increasing m pushes output voltage amplitude higher, sweeping D across a wider range (closer to 0 and 1).
        *   **Observation**: You will notice the 2-Level ripple **dips** near the peak (90°) as m increases.
        *   **Reason**: D moves away from the worst-case 0.5 region (toward 1.0), so at high m the ripple at the voltage peak is actually less than at the zero-crossing.

        **4. Why show SPWM here?**
        *   SPWM is used because its duty cycle D(t) is a clean sinusoid — the direct mapping between "D variation" and "ripple envelope" is the clearest path to understanding the physics.
        *   **Engineering conclusion**: In real products using **SVPWM**, injecting a third-harmonic "saddle wave" keeps D away from the high-ripple zone for longer. **Result: SVPWM peak ripple is typically 15%–20% lower than SPWM under the same conditions.**
            """)









# ==========================================
#  7:  (IEEE 519 / GB/T 14549)
# v2 : 
#   [Fix1] inv_harmonic_amplitude 3-Level/2-Level
#          3-Level: Vdc, , 
#   [Fix2] FFTGB/T 14549, IEEE 519/
#          GB/T%TDD, 
#   [Fix3] GB/T: 
#          (<25) = , LCL, 
#          (≥25) = LCL, 
#          IEEE 519, LCL
# ==========================================

elif selection == nav_options[6]:

    st.header("📜 Harmonic Standard Verification (IEEE 519 / GB/T 14549)")
    st.markdown("""
 > **Engineering purpose**: After the LCL filter, verify that the injected harmonic current at the PCC
 > satisfies grid standards. This module checks your current design parameters against the standard.
    >
 > 📚 **References**: ① IEEE Std 519-2022 | ② GB/T 14549-1993 | ③ Holmes & Lipo, *Pulse Width Modulation for Power Converters*, IEEE Press 2003
    """)

    # ── Parameter ──────────────────────────────────────────────────────
    with st.expander("⚙️ System Parameter Input (with transformer step-up decoupling)", expanded=True):
        st.markdown("##### Input System Parameters (PCC and Inverter Side Voltage)")
        col_p1, col_p2, col_p3, col_p4 = st.columns(4)

        with col_p1:
            st.markdown("**Grid Parameters (PCC)**")
            h7_ug_pcc = st.number_input("PCC voltage Ug (Vrms)", value=35000.0, step=1000.0, key='h7_ug_pcc',
        help="PCC rated voltage")
            h7_fg     = st.number_input("Grid Frequency f₀ (Hz)", value=float(st.session_state.get('sys_fg', 50.0)),
                                        step=1.0, key='h7_fg')
            h7_scc    = st.number_input("PCC Short-Circuit Capacity Ssc (MVA)", value=10.0, step=1.0, key='h7_scc',
            help="Grid PCC 3-phase short-circuit capacity. Formula: Ug²/Zg.")

        with col_p2:
            st.markdown("**Converter Parameters**")
            h7_ug_inv = st.number_input("Inverter Side Voltage (Vrms)", value=float(st.session_state.get('sys_ug', 690.0)),
                                        step=10.0, key='h7_ug_inv')
            h7_prated = st.number_input("System rated power P (kW)", value=500.0, step=10.0, key='h7_prated',
                                        help="3-phase system rated active power")
            h7_il = st.number_input("Rated current IL (Arms)", value=float(st.session_state.get('vec_i', 500.0)),
                                    step=10.0, key='h7_il',
                                    help="IEEE 519 TDD base current: rated inverter-side current (RMS)")
            h7_vdc = st.number_input("DC Bus Vdc (V)", value=float(st.session_state.get('vec_udc', 1100.0)),
                                     step=50.0, key='h7_vdc')

        with col_p3:
            st.markdown("**Switching Frequency & Topology**")
            h7_fsw  = st.number_input("Switching Frequency fsw (Hz)", value=float(st.session_state.get('sys_fsw', 2500.0)),
                                      step=100.0, key='h7_fsw')
            h7_topo = st.selectbox("Converter topology", ["2-Level (2-Level)", "3-Level (3-Level NPC/T-Type)"],
                                   index=0, key='h7_topo',
                                   help="3-Level topology has different harmonic characteristics than 2-Level")
            h7_mod = st.selectbox("Modulation strategy", ["SVPWM", "SPWM"], index=0, key='h7_mod')

        with col_p4:
            st.markdown("**LCL Filter Parameters**")
            h7_l1 = st.number_input("L1 (mH)", value=float(st.session_state.get('lcl_l1', 0.5)),
                                    step=0.05, format="%.3f", key='h7_l1')
            h7_l2 = st.number_input("L2 (mH)", value=float(st.session_state.get('lcl_l2', 0.25)),
                                    step=0.05, format="%.3f", key='h7_l2')
            h7_cf = st.number_input("Cf (μF)", value=float(st.session_state.get('lcl_c', 100.0)),
                                    step=1.0, key='h7_cf')

    # ──  ──────────────────────────────────────────────────────────────
    _L1  = h7_l1 * 1e-3
    _L2  = h7_l2 * 1e-3
    _Cf  = h7_cf * 1e-6
    _omega0   = 2 * np.pi * h7_fg
    _omega_sw = 2 * np.pi * h7_fsw
    _Prated   = h7_prated * 1e3
    _k_tr     = h7_ug_pcc / h7_ug_inv if h7_ug_inv > 0 else 1.0
    _is_3level = ("3-Level" in h7_topo)

    if _L1 > 0 and _L2 > 0 and _Cf > 0:
        _omega_res = np.sqrt((_L1 + _L2) / (_L1 * _L2 * _Cf))
        _f_res     = _omega_res / (2 * np.pi)
    else:
        _omega_res = 0
        _f_res     = 0

    _isc_pcc = h7_scc * 1e6 / (np.sqrt(3) * h7_ug_pcc) if h7_ug_pcc > 0 else 0
    _il_pcc  = h7_il / _k_tr if _k_tr > 0 else h7_il
    _scr     = _isc_pcc / _il_pcc if _il_pcc > 0 else 0

    # ──  ────────────────────────────────────────────────
    if h7_vdc > 0 and h7_ug_inv > 0:
        if h7_mod == "SPWM":
            _ma_auto = (2.0 * np.sqrt(2) / np.sqrt(3)) * (h7_ug_inv / h7_vdc)
            _ma_note = "SPWM linear"
        else:
            _ma_auto = (2.0 / np.sqrt(3)) * (h7_ug_inv / h7_vdc)
            _ma_note = "SVPWM (boost factor 2/√3)"
        _ma_auto = float(np.clip(_ma_auto, 0.0, 1.15))
    else:
        _ma_auto = 0.90
        _ma_note = ""

    _i1_hat_auto = float(np.sqrt(2) * h7_il)

# ── SCR (Tab)──────────────────────────────────────────────────
    # : IEEE Std 519-2022, Table 2(th Harmonic, %TDD)
    # 2014: SCR, 17~2350~1004.0%→4.0%; 
    #               2014.2022(Table 4).
    # th Harmonic:  =  × 25%(IEEE 519-2022 Table 2  b)
    if _scr < 20:
        scr_label, scr_color = "Weak Grid (SCR<20)", "#dc3545"
        tdd_limit  = 5.0
        # IEEE 519-2022 Table 2, Column 1 (Isc/IL < 20)
        i_h_limits = {(3,11):4.0, (11,17):2.0, (17,23):1.5, (23,35):0.6, (35,999):0.3}
    elif _scr < 50:
        scr_label, scr_color = "Weak Grid (SCR 20~50)", "#fd7e14"
        tdd_limit  = 8.0
        # IEEE 519-2022 Table 2, Column 2 (20 ≤ Isc/IL < 50)
        i_h_limits = {(3,11):7.0, (11,17):3.5, (17,23):2.5, (23,35):1.0, (35,999):0.5}
    elif _scr < 100:
        scr_label, scr_color = "Medium Grid (SCR 50~100)", "#ffc107"
        tdd_limit  = 12.0
        # IEEE 519-2022 Table 2, Column 3 (50 ≤ Isc/IL < 100)
        i_h_limits = {(3,11):10.0, (11,17):4.5, (17,23):4.0, (23,35):1.5, (35,999):0.7}
    elif _scr < 1000:
        scr_label, scr_color = "Strong Grid (SCR 100~1000)", "#198754"
        tdd_limit  = 15.0
        # IEEE 519-2022 Table 2, Column 4 (100 ≤ Isc/IL < 1000)
        i_h_limits = {(3,11):12.0, (11,17):5.5, (17,23):5.0, (23,35):2.0, (35,999):1.0}
    else:
        scr_label, scr_color = "Very Strong Grid (SCR≥1000)", "#0d6efd"
        tdd_limit  = 20.0
        # IEEE 519-2022 Table 2, Column 5 (Isc/IL ≥ 1000)
        i_h_limits = {(3,11):15.0, (11,17):7.0, (17,23):6.0, (23,35):2.5, (35,999):1.4}

    # GB/T 
    ug_kv = h7_ug_pcc / 1000
    if ug_kv <= 1.2:
        v_level_str, s_ref_gbt, v_idx_gbt = "0.38 kV (LV)", 10.0, 0
    elif ug_kv <= 8.0:
        v_level_str, s_ref_gbt, v_idx_gbt = "6 kV", 100.0, 1
    elif ug_kv <= 12.0:
        v_level_str, s_ref_gbt, v_idx_gbt = "10 kV", 100.0, 2
    elif ug_kv <= 40.0:
        v_level_str, s_ref_gbt, v_idx_gbt = "35 kV", 250.0, 3
    else:
        v_level_str, s_ref_gbt, v_idx_gbt = "110 kV", 750.0, 4

    s_user_gbt = _Prated / 1e6
    ratio_gbt  = s_user_gbt / s_ref_gbt

    # ── GB/T 14549-1993  ────────────────────────────────────────────
    # : GB/T 14549-1993" "2, 3
    # 2: Individual Harmonics(: A)
    #   0.38kV  = 10 MVA
    #   6kV/10kV  = 100 MVA
    #   35kV  = 250 MVA(110kV: 750 MVA)
    #  = 2 × ( / )  [GB/T 14549 §4.2]
    # : PCC  = Inverter Side /  k_tr,  gbt_user_A_full  k_tr
    # ⚠️ : GB/T 14549 22, 4, 6th Harmonic, 
    gbt_orders_full = list(range(2, 26))
    if v_idx_gbt == 0:
        # 0.38kV, 10MVA, GB/T 14549-1993 2
        gbt_ref_odd = {2:78,3:62,4:39,5:62,6:26,7:44,8:19,9:21,10:16,11:28,12:13,13:24,14:11,15:16,16:10,17:20,18:9,19:18,20:8,21:15,22:7,23:15,24:7,25:12}
    elif v_idx_gbt in (1, 2):
        # 6kV/10kV, 100MVA, GB/T 14549-1993 2, 
        gbt_ref_odd = {2:20,3:16,4:10,5:16,6:6.4,7:11,8:4.8,9:5.3,10:4.0,11:7.1,12:3.2,13:6.0,14:2.7,15:4.0,16:2.4,17:5.1,18:2.1,19:4.6,20:1.9,21:3.8,22:1.7,23:3.8,24:1.6,25:3.0}
    else:
        # 35kV, 250~750MVA, GB/T 14549-1993 2, 
        gbt_ref_odd = {2:9.6,3:7.6,4:4.8,5:7.6,6:3.0,7:5.3,8:2.3,9:2.5,10:1.9,11:3.4,12:1.5,13:2.8,14:1.3,15:1.9,16:1.1,17:2.4,18:1.0,19:2.2,20:0.9,21:1.8,22:0.8,23:1.8,24:0.7,25:1.4}

    gbt_ref_A_full  = [gbt_ref_odd.get(h, 0) for h in gbt_orders_full]
    # :  × (§4.2)× (PCC)
    gbt_user_A_full = [v * ratio_gbt * _k_tr for v in gbt_ref_A_full]

    # ════════════════════════════════════════════════════════════════════════════
    # Tab 
    # ════════════════════════════════════════════════════════════════════════════
    tab_ieee, tab_gbt, tab_fft, tab_verdict = st.tabs([
        "📋 IEEE 519-2022 Standard",
        "📋 GB/T 14549-1993 (China)",
        "📊 FFT Simulation Analysis",
        "✅ Verification Result"
    ])

    # ════════════════════════════════════════════════════════════════════════════
    # TAB 1: IEEE 519-2022
    # ════════════════════════════════════════════════════════════════════════════
    with tab_ieee:
        st.subheader("IEEE Std 519-2022 requirement")

        with st.expander("🗺️ System topology: Where are the PCC, transformer, and converter? (Click to expand)", expanded=True):
            st.markdown("##### Grid-Connected Topology (PCC / Transformer / Converter)")

            def render_grid_topology():
                fig = go.Figure()
                X = dict(grid=0.07, pcc=0.22, box=0.40, lcl=0.60, inv=0.78, dc=0.95, other=0.36)
                Y_MAIN = 0.52
                fig.add_shape(type="rect", x0=0.01, x1=0.18, y0=0.15, y1=0.88,
                              fillcolor="rgba(13,110,253,0.05)", line=dict(color="#0d6efd", width=1.5, dash="dot"))
                fig.add_shape(type="rect", x0=0.24, x1=0.99, y0=0.08, y1=0.92,
                              fillcolor="rgba(25,135,84,0.04)", line=dict(color="#198754", width=1, dash="dot"))
                # Fix: , marker, 
                # :  → PCC →  → LCL →  → 
                main_segs = [
                    (X["grid"], X["pcc"],  Y_MAIN),   # grid → PCC
                    (X["pcc"],  X["box"],  Y_MAIN),   # PCC → transformer
                    (X["box"],  X["lcl"],  Y_MAIN),   # transformer → LCL
                    (X["lcl"],  X["inv"],  Y_MAIN),   # LCL → converter
                    (X["inv"],  X["dc"],   Y_MAIN),   # converter → 
                ]
                for x0, x1, y in main_segs:
                    fig.add_shape(type="line", x0=x0, y0=y, x1=x1, y1=y,
                                  xref="paper", yref="paper", line=dict(color="#555", width=2.5))
                # PCC : , 
                fig.add_shape(type="line", x0=X["pcc"], y0=Y_MAIN, x1=X["pcc"], y1=0.32,
                              xref="paper", yref="paper", line=dict(color="#555", width=2.5))
                fig.add_shape(type="line", x0=X["pcc"], y0=0.32, x1=X["other"], y1=0.32,
                              xref="paper", yref="paper", line=dict(color="#555", width=2.5))
                node_data = [
                    (X["grid"],  Y_MAIN, "#e8f4fd","#0d6efd",32,"grid",        "110kV/35kV utility grid"),
                    (X["box"],   Y_MAIN, "#e8f5e9","#198754",30,"transformer",  "Step-up 10kV→690V"),
                    (X["lcl"],   Y_MAIN, "#f3e5f5","#7b2d8b",30,"LCL filter",  "This tool's design scope"),
                    (X["inv"],   Y_MAIN, "#fce4ec","#d63384",30,"converter",    "Harmonic source"),
                    (X["dc"],    Y_MAIN, "#fff8e1","#fd7e14",26,"DC Bus",       "DC Bus / PV / Battery"),
                    (X["other"], 0.32,   "#fff3cd","#fd7e14",22,"other loads",  "PCC connection node"),
                ]
                for xn,yn,cf,cl,sz,lbl,hov in node_data:
                    fig.add_trace(go.Scatter(x=[xn], y=[yn], mode="markers+text",
                                            marker=dict(symbol="square", size=sz, color=cf, line=dict(color=cl, width=2.5)),
                                            text=[f"<b>{lbl}</b>"], textposition="top center",
                                            textfont=dict(size=10, color=cl),
                                            hovertext=[hov], hoverinfo="text", showlegend=False))
                fig.add_annotation(x=X["pcc"], y=0.86, xref="paper", yref="paper",
                text="<b>PCC</b><br>standard", showarrow=True,
                                   arrowhead=2, arrowcolor="#dc3545", arrowwidth=2, ax=0, ay=30,
                                   font=dict(size=11, color="white"), bgcolor="#dc3545",
                                   bordercolor="#dc3545", borderwidth=2, borderpad=5)
                fig.update_layout(height=300, margin=dict(l=5,r=5,t=5,b=5),
                                  paper_bgcolor="#fafafa", plot_bgcolor="#fafafa",
                                  xaxis=dict(visible=False, range=[0,1]),
                                  yaxis=dict(visible=False, range=[0,1]))
                return fig

            st.plotly_chart(render_grid_topology(), use_container_width=True)
            st.info("💡 **Key point**: Harmonic current (%TDD) is calculated at the PCC after the transformer. This tool evaluates at the inverter side and applies the transformer ratio.")

        st.divider()
        col_ieee_l, col_ieee_r = st.columns([1, 1.2])

        with col_ieee_l:
            st.markdown("""
            #### 📖 Key Concepts: TDD and SCR

 IEEE 519 regulates **harmonic injection at the PCC**. The limit depends on the **Short Circuit Ratio (SCR)**:
 A stronger grid (higher SCR) allows more harmonic injection.
            """)
            st.latex(r"SCR = \frac{I_{SC,PCC}}{I_{L,PCC}}")
            st.latex(r"TDD = \frac{\sqrt{\sum_{h=2}^{\infty} I_h^2}}{I_L} \times 100\%")
            if is_trial:
                trial_lock("Unlock with paid access: key engineering difference between TDD and THD", compact=True)
            else:
                st.error("**TDD uses rated current $I_L$ as the base, not actual current! TDD differs from THD — always use rated current for design verification.**")

            st.markdown("---")
            st.markdown("#### 📐 Current System SCR Calculation")
            if is_trial:
                trial_lock("Unlock with paid access: current system SCR calculation result and grid strength classification")
            else:
                st.latex(rf"I_{{SC,PCC}} = \frac{{{h7_scc:.1f}\text{{ MVA}}}}{{\sqrt{{3}}\times{h7_ug_pcc:.0f}\text{{ V}}}} = {_isc_pcc:.0f}\text{{ A}}")
                st.latex(rf"I_{{L,PCC}} = \frac{{{h7_il:.0f}\text{{ A}}}}{{{_k_tr:.2f}}} = {_il_pcc:.0f}\text{{ A}}")
                st.latex(rf"SCR = {_scr:.1f}")
                st.markdown(f"""
                <div style="background:{scr_color}22; border-left:4px solid {scr_color};
                            padding:12px; border-radius:4px; margin:8px 0;">
                    <b>Grid strength: {scr_label}</b> &nbsp;|&nbsp; TDD limit: <b>{tdd_limit}%</b>
                </div>
                """, unsafe_allow_html=True)

        with col_ieee_r:
            st.markdown("#### 📊 IEEE 519-2022 Harmonic Limits (Table 2)")
            scr_levels = [10, 30, 75, 500, 1500]
            scr_labels = ['<20','20~50','50~100','100~1000','>1000']
            colors_scr = ['#dc3545','#fd7e14','#ffc107','#198754','#0d6efd']
            harm_orders_plot = [3,5,7,9,11,13,17,19,23,25,35,49]

            def get_ieee_limit(scr_idx, h):
                lim = [[4.,7.,10.,12.,15.],[2.,3.5,4.5,5.5,7.],[1.5,2.5,4.,5.,6.],[.6,1.,1.5,2.,2.5],[.3,.5,.7,1.,1.4]]
                row = 0 if h<11 else 1 if h<17 else 2 if h<23 else 3 if h<35 else 4
                return lim[row][scr_idx]

            fig_ieee = go.Figure()
            for si, (_, slabel, clr) in enumerate(zip(scr_levels, scr_labels, colors_scr)):
                w = 4 if abs(scr_levels[si]-_scr)==min(abs(s-_scr) for s in scr_levels) else 1.5
                fig_ieee.add_trace(go.Scatter(x=harm_orders_plot,
                                              y=[get_ieee_limit(si,h) for h in harm_orders_plot],
                                              mode='lines+markers', name=f'SCR {slabel}',
                                              line=dict(color=clr, width=w)))
            fig_ieee.update_layout(height=280, template="plotly_white", margin=dict(t=10,b=10),
                                   legend=dict(orientation="h", yanchor="bottom", y=-0.35, xanchor="center", x=0.5))
            st.plotly_chart(fig_ieee, use_container_width=True)

        st.markdown("#### 📋 IEEE 519-2022 Table 2")
        df_ieee = pd.DataFrame({
            'Harmonic Order Range': ['3~11','11~17','17~23','23~35','≥35'],
            'SCR<20 (5%)': ['4.0%','2.0%','1.5%','0.6%','0.3%'],
            'SCR 20~50 (8%)': ['7.0%','3.5%','2.5%','1.0%','0.5%'],
            'SCR 50~100 (12%)': ['10.0%','4.5%','4.0%','1.5%','0.7%'],
            'SCR 100~1000 (15%)': ['12.0%','5.5%','5.0%','2.0%','1.0%'],
            'SCR>1000 (20%)': ['15.0%','7.0%','6.0%','2.5%','1.4%'],
        })
        st.dataframe(df_ieee, width="stretch", hide_index=True)
        st.warning("⚠️ Even harmonic limit = odd harmonic limit × 25% (IEEE 519-2022 Table 2, Note b)")
        st.caption(
 "📚 Source: IEEE Std 519-2022, Table 2 'Harmonic Current Distortion in Percent of IL'."
 " Compared to IEEE 519-2014: updated SCR thresholds and values;"
 "2022 adds Table 4 (flicker: Pst 10min, Plt 2h) description."
 "Refer to IEEE Std 519-2022 Section 5, Table 2."
        )

    # ════════════════════════════════════════════════════════════════════════════
    # TAB 2: GB/T 14549-1993
    # ════════════════════════════════════════════════════════════════════════════
    with tab_gbt:
        st.subheader("GB/T 14549-1993 requirement")
        st.info(f"📌 **Note**: IEEE 519 harmonic limits apply at the PCC ({h7_ug_pcc/1000:.1f} kV). This tool applies transformer ratio N={_k_tr:.2f} (from {h7_ug_inv:.0f} V inverter side).")

        col_g1, col_g2 = st.columns([1, 1])
        with col_g1:
            st.markdown("#### 📖 GB/T limit calculation: voltage level + capacity ratio")
            st.latex(r"I_{h,user} = I_{h,ref} \times \frac{S_{user}}{S_{ref}}")
            gbt_tier_data = [("0.38 kV","10 MVA","LV"),("6 kV","100 MVA","MV"),
                             ("10 kV","100 MVA","MV (most common)"),("35 kV","250 MVA","MHV"),("110 kV","750 MVA","HV")]
        df_tiers = pd.DataFrame(gbt_tier_data, columns=["Voltage","Reference Short-Circuit Capacity Sref","Note"])
        def hl_tier(row): return ['background-color:#d4edda;font-weight:bold']*3 if row.name==v_idx_gbt else ['']*3
        st.dataframe(df_tiers.style.apply(hl_tier,axis=1), width="stretch", hide_index=True)
        if is_trial:
            trial_lock("Unlock with paid access: GB/T capacity scaling parameters and applicable voltage tier", compact=True)
        else:
            st.caption(f"✅ Current PCC voltage: {h7_ug_pcc:.0f} V ({ug_kv:.3f} kV)")
            st.markdown(f"- Voltage level: **{v_level_str}**, Reference capacity Sref = **{s_ref_gbt:.0f} MVA**")
            st.markdown(f"- User capacity: **{s_user_gbt:.3f} MVA**, ratio k = **{ratio_gbt:.4f}**, transformer N = **{_k_tr:.2f}**")

        with col_g2:
            st.markdown("#### 📊 IEEE 519 vs GB/T 14549 Key Differences")
            df_diff = pd.DataFrame({
                'Item': ['Published','Classification','Limit form','Even harmonics','Measurement window','Region'],
                'IEEE 519': ['2022','SCR (grid strength)','TDD %','Odd × 25%','10min/2h','International'],
                'GB/T 14549': ['1993','Voltage level','Direct ampere value','Separate value (~odd × 50%)','Not specified','China'],
            })
            st.dataframe(df_diff, width="stretch", hide_index=True)

        st.divider()
        st.markdown(f"#### 📋 GB/T 14549 Harmonic Limits (scaled to LV side {h7_ug_inv:.0f} V)")
        fig_gbt2 = make_subplots(rows=2, cols=1, shared_xaxes=True, vertical_spacing=0.12,
                                 subplot_titles=[f"① Reference limits ({v_level_str}, Sref={s_ref_gbt:.0f} MVA)",
                                                 f"② Adjusted limits (×k={ratio_gbt:.4f} ×N={_k_tr:.2f})"],
                                 row_heights=[0.5, 0.5])
        x_labels = [f'{h}' for h in gbt_orders_full]
        if is_trial:
            fig_gbt2.add_trace(go.Bar(x=x_labels, y=gbt_ref_A_full, name='Reference limits',
                                      marker_color='#0d6efd', hoverinfo='none'), row=1, col=1)
            fig_gbt2.add_trace(go.Bar(x=x_labels, y=gbt_user_A_full, name='Adjusted limits',
                                      marker_color='#198754', hoverinfo='none'), row=2, col=1)
            fig_gbt2.update_layout(height=520, template="plotly_white", hovermode=False,
                                   legend=dict(orientation="h", yanchor="bottom", y=-0.08, xanchor="center", x=0.5))
            fig_gbt2.update_yaxes(showticklabels=False, row=1, col=1)
            fig_gbt2.update_yaxes(showticklabels=False, row=2, col=1)
            st.plotly_chart(fig_gbt2, use_container_width=True)
            trial_lock("Unlock with paid access: GB/T 14549 harmonic limits (including transformer ratio scaling)", compact=True)
        else:
            fig_gbt2.add_trace(go.Bar(x=x_labels, y=gbt_ref_A_full, name='Reference limits',
                                      marker_color='#0d6efd', text=[f'{v:.1f}' for v in gbt_ref_A_full],
                                      textposition='outside', textfont=dict(size=9)), row=1, col=1)
            fig_gbt2.add_trace(go.Bar(x=x_labels, y=gbt_user_A_full, name='Adjusted limits',
                                      marker_color='#198754', text=[f'{v:.2f}' for v in gbt_user_A_full],
                                      textposition='outside', textfont=dict(size=9)), row=2, col=1)
            fig_gbt2.update_layout(height=520, template="plotly_white",
                                   legend=dict(orientation="h", yanchor="bottom", y=-0.08, xanchor="center", x=0.5))
            fig_gbt2.update_yaxes(title_text="Current (A)", row=1, col=1)
            fig_gbt2.update_yaxes(title_text="Current (A)", row=2, col=1)
            st.plotly_chart(fig_gbt2, use_container_width=True)
            st.caption("Source: GB/T 14549-1993, Table 2")

    # ════════════════════════════════════════════════════════════════════════════
    # TAB 3: FFT 
    # ════════════════════════════════════════════════════════════════════════════
    with tab_fft:
        st.subheader("📊 LCL Filter Current FFT Simulation")

        # ── [Fix1] 3-Level ────────────────────────────────────────────────
        with st.expander("📖 2-Level vs 3-Level Harmonic Characteristics (click to expand theoretical description)", expanded=False):
            st.markdown(f"""
            **Currently selected topology: `{h7_topo}`**

            | Characteristic | 2-Level | 3-Level NPC/T-Type |
            |---|---|---|
            | Phase voltage swing | $V_{{DC}}$ | $V_{{DC}}/2$ (per bridge arm) |
            | Dominant harmonic | $f_{{sw}}$ | $2f_{{sw}}$ (doubled) |
            | Even-order sideband (m=2,4…) | Present | **Theoretically cancelled** |
            | Switching frequency harmonic amplitude | Higher | **Lower** (same switching frequency) |

 Reference: Holmes & Lipo, *Pulse Width Modulation for Power Converters*, IEEE Press 2003, Chapter 5 (Multilevel Inverters);
 Nabae et al., IEEE Trans. Ind. Appl., Vol. IA-17, 1981(NPC ).

 **Engineering note**: 3-Level topology at the same switching frequency provides harmonic performance equivalent to 2-Level at **2× switching frequency**,
 allowing smaller LCL inductors while still meeting harmonic standards.
            """)

        st.markdown("""
 > **Simulation method**: Uses SPWM/SVPWM modulation theory (Bessel model) to calculate converter output current harmonic amplitudes,
 > then applies LCL filter transfer function attenuation to determine grid-injected harmonics.
        >
 > ⚠️ **Model accuracy**: ±10% error; low-order harmonics (<25th) depend on control algorithm — within model range.
        """)

        col_fft_ctrl, col_fft_plot = st.columns([1, 2.5])

        with col_fft_ctrl:
            st.markdown("**Simulation Control Parameters**")

            # 
            st.markdown(f"""
            <div style="background:#e8f4fd; border-left:4px solid #0d6efd; padding:10px 12px; border-radius:6px; margin-bottom:6px;">
              <b style="color:#0d6efd; font-size:13px;">⚡ Modulation Index mₐ (auto-calculated)</b><br>
 <span style="font-size:13px;">Modulation: {h7_mod} | Result: <b>{_ma_auto:.3f}</b></span>
            </div>
            """, unsafe_allow_html=True)
            with st.expander("📐 Calculation Formula"):
                if h7_mod == "SPWM":
                    st.latex(r"m_a = \frac{2\sqrt{2}}{\sqrt{3}} \cdot \frac{U_{LL}}{U_{DC}}")
                else:
                    st.latex(r"m_a = \frac{2}{\sqrt{3}} \cdot \frac{U_{LL}}{U_{DC}}")
                st.markdown(f"$U_{{LL}}$ = {h7_ug_inv:.0f} V, $U_{{DC}}$ = {h7_vdc:.0f} V → $m_a$ = **{_ma_auto:.4f}**")
                st.caption("Source: Holmes & Lipo 2003, Chapter 2")

                fft_expert = st.checkbox("🔧 Expert mode: manually set mₐ", value=False)
            if fft_expert:
                fft_ma = st.slider("Modulation index mₐ", 0.5, 1.15, float(_ma_auto), 0.01)
                st.warning(f"Manual override: {fft_ma:.3f} (auto-calculated: {_ma_auto:.3f})")
            else:
                fft_ma = _ma_auto

            st.markdown("---")

            # 
            st.markdown(f"""
            <div style="background:#e8f5e9; border-left:4px solid #198754; padding:10px 12px; border-radius:6px; margin-bottom:6px;">
              <b style="color:#198754; font-size:13px;">⚡ Fundamental Current Î₁ (auto-calculated)</b><br>
              <span style="font-size:13px;">√2 × IL = √2 × {h7_il:.1f} = <b>{_i1_hat_auto:.2f} A</b></span>
            </div>
            """, unsafe_allow_html=True)
            fft_ibase = _i1_hat_auto

            st.markdown("---")
            fft_n_max = st.slider("Analysis harmonic order N", 10, 100, 50, 5)
            fft_show_limit = st.checkbox("Show standard limits", value=True)
            fft_rd = st.number_input("damping Rd (Ω)", value=0.05, step=0.01, format="%.3f",
            help="LCL capacitor series damping resistance. Set 0 for undamped.")

        with col_fft_plot:
            # ── LCL  ─────────────────────────────────────────────────────
            def lcl_tf_gain(freq_hz, L1, L2, Cf, Rd=0.0):
                w = 2 * np.pi * freq_hz
                s = 1j * w
                denom = (s**3 * L1 * L2 * Cf
                         + s**2 * (L1 + L2) * Rd * Cf
                         + s  * (L1 + L2) + 0j)
                return abs(1.0 / denom) if abs(denom) > 1e-30 else 0.0

            # ── [Fix1] (2-Level/3-Level) ──────────────────
            # : 
            # 2-Level: i_peak(m,n) = 4*Vdc*J_n(m*π*ma/2) / (m*π²*ω*L1)
            # 3-Level: Vdc(Vdc/2); 
            #         (m=2,4...), ; 
            #         2-LevelVdc/2
            # : Holmes & Lipo 2003, Chapter 5, Eq.(5.31)~(5.35)
            def inv_harmonic_amplitude(freq_hz, Vdc, L1, fsw, ma, f0, is_3level):
                from scipy.special import jv as bessel_j
                h_order = freq_hz / f0
                omega   = 2 * np.pi * freq_hz

                # Switching Frequency
                is_sw_band  = False
                m_carrier   = 1
                n_order_sw  = 0
                for m in range(1, 5):
                    if abs(freq_hz - m * fsw) < 5 * f0:
                        is_sw_band = True
                        m_carrier  = m
                        n_order_sw = round((freq_hz - m * fsw) / f0)
                        break

                if is_sw_band:
                    # ── 3-Level:  ──
                    if is_3level and (m_carrier % 2 == 0):
                        return 1e-9   # 

                    # : 3-Level Vdc/2
                    Vdc_eff = Vdc / 2.0 if is_3level else Vdc

                    try:
                        J_val  = abs(bessel_j(n_order_sw, m_carrier * np.pi * ma / 2.0))
                        i_peak = (4.0 * Vdc_eff * J_val) / (m_carrier * np.pi**2 * omega * L1)
                    except Exception:
                        i_peak = 0.01
                else:
                    # th Harmonic: , 
                    # , 
                    i_base = fft_ibase
                    i_peak = i_base * 0.02 / ((max(h_order, 2) / 5.0) ** 1.5)

                return max(i_peak, 1e-9)

            # ──  ─────────────────────────────────────────────────
            f0  = h7_fg
            fsw = h7_fsw
            analysis_freqs = set()
            for h in range(2, fft_n_max + 1):
                analysis_freqs.add(h * f0)
            for m in range(1, 5):
                for n in range(-4, 5):
                    f_sb = m * fsw + n * f0
                    if f_sb > 0:
                        analysis_freqs.add(f_sb)
            analysis_freqs = sorted(analysis_freqs)

            # ──  ───────────────────────────────────────────────────────
            harm_freqs   = []
            harm_orders_fft = []
            i_inv_pct    = []
            i_grid_pct   = []
            i_grid_A     = []
            i_tdd_pct    = []

            for freq in analysis_freqs:
                if freq < f0 * 1.5:
                    continue
                h_num = freq / f0
                i_inv_peak  = inv_harmonic_amplitude(freq, h7_vdc, _L1, fsw, fft_ma, f0, _is_3level)
                omega_h     = 2 * np.pi * freq
                tf_val      = lcl_tf_gain(freq, _L1, _L2, _Cf, fft_rd)
                i_grid_peak = i_inv_peak * abs(1j * omega_h * _L1) * tf_val
                i_inv_rms   = i_inv_peak  / np.sqrt(2)
                i_grid_rms  = i_grid_peak / np.sqrt(2)

                harm_freqs.append(freq)
                harm_orders_fft.append(h_num)
                i_inv_pct.append(i_inv_rms / (fft_ibase / np.sqrt(2)) * 100)
                i_grid_pct.append(i_grid_rms / (fft_ibase / np.sqrt(2)) * 100)
                i_grid_A.append(i_grid_rms)
                i_tdd_pct.append(i_grid_rms / h7_il * 100)

            harm_freqs   = np.array(harm_freqs)
            i_grid_pct   = np.array(i_grid_pct)
            i_inv_pct    = np.array(i_inv_pct)
            i_tdd_pct    = np.array(i_tdd_pct)
            i_grid_A_arr = np.array(i_grid_A)
            # ── [v2]  FFT  session_state, 8 ─────────
            # Exceeds Limit(IEEE 519 LCL Exceeds Limit)
            _export_violations = []
            for _fi, _freq in enumerate(harm_freqs):
                _h_ord = _freq / h7_fg
                _i_g   = float(i_grid_A_arr[_fi])
                _tdd_h = float(i_tdd_pct[_fi])
                # LCL & Exceeds Limit
                _lcl_resp_freq = max(_f_res * 1.5, h7_fsw * 0.4) if _f_res > 0 else h7_fsw * 0.4
                if _freq >= _lcl_resp_freq:
                    if _h_ord < 11:   _lim = i_h_limits[(3,11)]
                    elif _h_ord < 17: _lim = i_h_limits[(11,17)]
                    elif _h_ord < 23: _lim = i_h_limits[(17,23)]
                    elif _h_ord < 35: _lim = i_h_limits[(23,35)]
                    else:             _lim = i_h_limits[(35,999)]
                    if _tdd_h > _lim:
                        _export_violations.append({
                            'Order': f"{_h_ord:.1f}",
                            'Frequency (Hz)': f"{_freq:.0f}",
                            'Grid Current (Arms)': f"{_i_g:.4f}",
                            'TDD (%)': f"{_tdd_h:.3f}",
                            'IEEE 519 Limit (%)': f"{_lim:.2f}",
                            'Exceeds by': f"{_tdd_h/_lim:.2f}x"
                        })

            st.session_state['h7_harm_export'] = {
                'freqs':      harm_freqs.tolist(),
                'i_grid_A':   i_grid_A_arr.tolist(),
                'i_tdd_pct':  i_tdd_pct.tolist(),
                'il':         h7_il,
                'fg':         h7_fg,
                'violations': _export_violations,
            }

            # 
            if _export_violations:
                st.warning(
                    f"⚠️ Detected **{len(_export_violations)}** IEEE 519 violation(s) in the LCL-responsible frequency range. "
                    f"These results are passed to Module 8 (Passive Harmonic Filter Design) → Verification tab."
                )
            else:
                st.success(
                    "✅ No IEEE 519 violations detected in the LCL-responsible frequency range. "
                    "Module 8 will use these FFT results for system parallel impedance analysis."
                )
            # ──  ─────────────────────────────────────────────────────────
            # ──  FFT  ──────────────────────────────────────────────────────
            fig_fft = make_subplots(rows=2, cols=1, shared_xaxes=True,
                                    vertical_spacing=0.12,
                                    subplot_titles=[
 f"Inverter side current harmonics (before/after LCL filter, % of IL) — {h7_topo}",
 "Current harmonics after LCL filter (% TDD, reference: IL)"
                                    ],
                                    row_heights=[0.38, 0.62])

            if is_trial:
                fig_fft.add_trace(go.Bar(x=harm_freqs, y=i_inv_pct,
                                         name='Inverter Side (before filter)',
                                         marker_color='rgba(220,53,69,0.6)',
                                         hoverinfo='none'),
                                  row=1, col=1)
                fig_fft.add_trace(go.Bar(x=harm_freqs, y=i_tdd_pct,
                                         name='Grid current (after LCL filter)',
                                         marker_color='rgba(25,135,84,0.7)',
                                         hoverinfo='none'),
                                  row=2, col=1)
                fig_fft.update_layout(
                    height=640, template="plotly_white", hovermode=False,
                    legend=dict(orientation="h", yanchor="bottom", y=-0.10, xanchor="center", x=0.5)
                )
                fig_fft.update_yaxes(showticklabels=False, row=1, col=1)
                fig_fft.update_yaxes(showticklabels=False, row=2, col=1)
                fig_fft.update_xaxes(showticklabels=False)
                st.plotly_chart(fig_fft, use_container_width=True)
                trial_lock("Unlock with paid access: LCL-filtered grid current FFT spectrum simulation (including individual harmonic amplitudes and standard limit comparison)", compact=True)
            else:
                fig_fft.add_trace(go.Bar(x=harm_freqs, y=i_inv_pct,
                                         name='Inverter Side (before filter)',
                                         marker_color='rgba(220,53,69,0.6)',
                                         hovertemplate='%{x:.0f} Hz<br>%{y:.2f}% fundamental<extra>before filter</extra>'),
                                  row=1, col=1)

                fig_fft.add_trace(go.Bar(x=harm_freqs, y=i_tdd_pct,
                                         name='Grid current (after LCL filter)',
                                         marker_color='rgba(25,135,84,0.7)',
                                         hovertemplate='%{x:.0f} Hz<br>%{y:.3f}% TDD<extra>after filter</extra>'),
                                  row=2, col=1)

                # ── [Fix2] Overlay dual standard limit lines ──────────────────────────────────────────
                # IEEE 519 limit line (in %TDD, same axis as grid current)
                # GB/T: read from gbt_user_A → %TDD = gbt_user_A / h7_il * 100
                # Limit lines only apply to lower subplot (row=2)
                if fft_show_limit and len(harm_freqs) > 0:
                    ieee_tdd_limit = []
                    gbt_tdd_limit  = []
                    for freq in harm_freqs:
                        h = freq / h7_fg
                        # IEEE 519
                        if h < 11:   ieee_tdd_limit.append(i_h_limits[(3,11)])
                        elif h < 17: ieee_tdd_limit.append(i_h_limits[(11,17)])
                        elif h < 23: ieee_tdd_limit.append(i_h_limits[(17,23)])
                        elif h < 35: ieee_tdd_limit.append(i_h_limits[(23,35)])
                        else:        ieee_tdd_limit.append(i_h_limits[(35,999)])

                        # GB/T: interpolate, convert to %TDD
                        h_int = int(round(h))
                        if h_int in gbt_orders_full:
                            idx_g = gbt_orders_full.index(h_int)
                            gbt_a = gbt_user_A_full[idx_g]
                        elif h_int > 25:
                            # above 25th: I_h = I_5_user * 5/h
                            idx5 = gbt_orders_full.index(5) if 5 in gbt_orders_full else 3
                            gbt_a = gbt_user_A_full[idx5] * 5.0 / h_int
                        else:
                            gbt_a = 0.0
                        gbt_tdd_limit.append(gbt_a / h7_il * 100 if h7_il > 0 else 0.0)

                    fig_fft.add_trace(go.Scatter(
                        x=harm_freqs, y=ieee_tdd_limit,
                        mode='lines', name=f'IEEE 519 (SCR={_scr:.0f})',
                        line=dict(color='#dc3545', width=2, dash='dash'),
                        hovertemplate='%{x:.0f} Hz<br>IEEE 519 limit: %{y:.2f}%<extra></extra>'
                    ), row=2, col=1)

                    fig_fft.add_trace(go.Scatter(
                        x=harm_freqs, y=gbt_tdd_limit,
                        mode='lines', name='GB/T 14549 (%TDD)',
                        line=dict(color='#fd7e14', width=2, dash='dot'),
                        hovertemplate='%{x:.0f} Hz<br>GB/T limit: %{y:.2f}%<extra></extra>'
                    ), row=2, col=1)

                # Resonance Frequency marker
                if _f_res > 0:
                    for row_n in [1, 2]:
                        fig_fft.add_vline(x=_f_res, line_dash="dot", line_color="purple",
                                          annotation_text=f"fres={_f_res:.0f}Hz",
                                          annotation_position="top right",
                                          row=row_n, col=1)

                fig_fft.update_layout(
                    height=640, template="plotly_white", hovermode="x unified",
                    legend=dict(orientation="h", yanchor="bottom", y=-0.10, xanchor="center", x=0.5)
                )
                fig_fft.update_xaxes(title_text="Frequency (Hz)", row=2, col=1)
                fig_fft.update_yaxes(title_text="% of fundamental", row=1, col=1)
                fig_fft.update_yaxes(title_text="% TDD reference", row=2, col=1)
                st.plotly_chart(fig_fft, use_container_width=True)

                if fft_show_limit:
                    st.caption("📌 **Legend**: ● = IEEE 519 limit; ■ = GB/T 14549 limit (%TDD). Red = Exceeds Limit.")

        # ── Bode  ──────────────────────────────────────────────────────────────
        st.markdown("---")
        st.markdown("#### 📉 LCL Filter Frequency Response (Bode Plot)")
        f_bode = np.logspace(1, 5, 1000)
        w_bode = 2 * np.pi * f_bode
        if _L1 > 0 and _L2 > 0 and _Cf > 0:
            sys_nd = signal.TransferFunction([1], [_L1*_L2*_Cf, 0, _L1+_L2, 0])
            _, mag_nd, _ = signal.bode(sys_nd, w_bode)
            num_d = [fft_rd, 1] if fft_rd > 0 else [1]
            den_d = [_L1*_L2*_Cf, (_L1+_L2)*fft_rd*_Cf if fft_rd>0 else 0, _L1+_L2, 0]
            sys_d = signal.TransferFunction(num_d, den_d)
            _, mag_d, _ = signal.bode(sys_d, w_bode)
        else:
            mag_nd = mag_d = np.zeros_like(f_bode)

        fig_bode = go.Figure()
        fig_bode.add_trace(go.Scatter(x=f_bode, y=mag_nd, name='Undamped LCL',
                                      line=dict(color='#0d6efd', width=2, dash='dot')))
        if fft_rd > 0:
            fig_bode.add_trace(go.Scatter(x=f_bode, y=mag_d, name=f'With damping (Rd={fft_rd}Ω)',
                                          line=dict(color='#198754', width=2.5)))
        fig_bode.add_vline(x=h7_fg, line_dash="dot", line_color="gray", annotation_text="f₀")
        if _f_res > 0:
            fig_bode.add_vline(x=_f_res, line_dash="dot", line_color="purple",
                               annotation_text=f"fres={_f_res:.0f}Hz")
            idx_sw = np.argmin(abs(f_bode - h7_fsw))
            fig_bode.add_annotation(x=h7_fsw, y=mag_nd[idx_sw],
                                    text=f"attenuation={mag_nd[idx_sw]:.1f}dB@fsw",
                                    showarrow=True, arrowhead=2, ax=60, ay=-30,
                                    bgcolor="rgba(255,255,255,0.9)", bordercolor="#fd7e14")
        fig_bode.add_vline(x=h7_fsw, line_dash="dash", line_color="orange",
                           annotation_text=f"fsw={h7_fsw:.0f}Hz")
        fig_bode.update_layout(xaxis_title="Frequency (Hz)", yaxis_title="Gain (dB)",
                               xaxis_type="log", height=350, template="plotly_white",
                               legend=dict(yanchor="top", y=0.99, xanchor="right", x=0.99))
        st.plotly_chart(fig_bode, use_container_width=True)

    # ════════════════════════════════════════════════════════════════════════════
    # TAB 4: Result
    # ════════════════════════════════════════════════════════════════════════════
    with tab_verdict:
        st.subheader("✅ Verification Result")

        # ──  ────────────────────────────────────────────────────────
        st.info("""
**📌 Engineering Note: Two Types of Harmonics and Their Responsibilities**

Harmonics fall into two fundamentally different categories with different mitigation methods:

| Harmonic type | Source | Mitigation method | LCL effective? |
|---|---|---|---|
| **Low-order harmonics (5th, 7th, 11th, 13th…)** | Non-ideal modulation, linear modulation limits | **Control algorithm** (PR control, repetitive control, harmonic compensation) | ❌ LCL gain≈1, no attenuation |
| **Switching frequency harmonics (≥ fsw)** | PWM switching action | **LCL filter** (−60 dB/dec attenuation) | ✅ LCL is responsible |

**Therefore, in the verification below: low-order harmonics are annotated as "control algorithm reference" and are NOT counted as LCL design failures.**
        """)

        # ──  ─────────────────────────────────────────────────────
        tdd_calc = float(np.sqrt(np.sum(i_tdd_pct**2))) if len(i_tdd_pct) > 0 else 0.0

        sw_freq  = h7_fsw
        idx_sw   = np.argmin(abs(harm_freqs - sw_freq)) if len(harm_freqs) > 0 else 0
        tdd_sw   = float(i_tdd_pct[idx_sw]) if len(i_tdd_pct) > 0 else 0.0

        sw_order = sw_freq / h7_fg
        if sw_order < 11:   limit_sw_pct = i_h_limits[(3,11)]
        elif sw_order < 17: limit_sw_pct = i_h_limits[(11,17)]
        elif sw_order < 23: limit_sw_pct = i_h_limits[(17,23)]
        elif sw_order < 35: limit_sw_pct = i_h_limits[(23,35)]
        else:               limit_sw_pct = i_h_limits[(35,999)]

        f_res_ok = (10 * h7_fg <= _f_res <= 0.5 * h7_fsw) if _f_res > 0 else False
        tdd_pass = tdd_calc <= tdd_limit
        sw_pass  = tdd_sw   <= limit_sw_pct

        # ── LCL  ─────────────────────────────────────────────────
        # : fres  LCL ≈1,  ≈ max(1.5×fres, 0.4×fsw)
        # : Liserre et al., IEEE Trans. Ind. Electron. 52(5), 2005
        if _f_res > 0:
            lcl_responsible_freq = max(_f_res * 1.5, h7_fsw * 0.4)
        else:
            lcl_responsible_freq = h7_fsw * 0.4
        lcl_responsible_order = lcl_responsible_freq / h7_fg

        # ════════════════════════════════════════════════════════════════════════
        # 1:  —  GB/T 14549 Result(2)
        # ════════════════════════════════════════════════════════════════════════
        st.markdown("### 🔍 Summary")
        def _card(icon, title, value, note, pass_flag=None):
            """render"""
            if pass_flag is True:
                val_color = "#198754"; bg = "#d4edda"; border = "#198754"
            elif pass_flag is False:
                val_color = "#dc3545"; bg = "#f8d7da"; border = "#dc3545"
            else:
                val_color = "#333333"; bg = "#f8f9fa"; border = "#adb5bd"
            return f"""
            <div style="background:{bg}; border:1px solid {border}; border-radius:8px;
                        padding:12px 14px; height:100%; box-sizing:border-box;">
              <div style="font-size:12px; color:#555; margin-bottom:4px;">{icon} {title}</div>
              <div style="font-size:22px; font-weight:700; color:{val_color}; line-height:1.2; margin-bottom:6px;">{value}</div>
              <div style="font-size:12px; color:#444; line-height:1.4; word-break:break-all;">{note}</div>
            </div>"""

        if is_trial:
            st.markdown("**IEEE 519-2022**")
            _tc1, _tc2, _tc3, _tc4 = st.columns(4)
            with _tc1: st.markdown(_card("📊", "IEEE TDD (Overall)", "🔒", "Unlock with paid access", None), unsafe_allow_html=True)
            with _tc2: st.markdown(_card("🔢", "Switching Frequency Harmonic", "🔒", "Unlock with paid access", None), unsafe_allow_html=True)
            with _tc3: st.markdown(_card("📐", "LCL Resonance Frequency", "🔒", "Unlock with paid access", None), unsafe_allow_html=True)
            with _tc4: st.markdown(_card("⚡", "Topology / Modulation", "🔒", "Unlock with paid access", None), unsafe_allow_html=True)
            st.markdown("<div style='margin-top:8px'></div>", unsafe_allow_html=True)
            st.markdown("**GB/T 14549-1993**")
            _gc1, _gc2, _gc3, _gc4 = st.columns(4)
            with _gc1: st.markdown(_card("📋", "GB/T LCL Responsible Band", "🔒", "Unlock with paid access", None), unsafe_allow_html=True)
            with _gc2: st.markdown(_card("⚙️", "Control Algorithm Zone (Reference)", "🔒", "Unlock with paid access", None), unsafe_allow_html=True)
            with _gc3: st.markdown(_card("📐", "LCL Effective Attenuation Start", "🔒", "Unlock with paid access", None), unsafe_allow_html=True)
            with _gc4: st.markdown(_card("🔄", "Transformer Scaling Parameters", "🔒", "Unlock with paid access", None), unsafe_allow_html=True)
            st.markdown("<div style='margin-top:8px'></div>", unsafe_allow_html=True)
            trial_lock("Unlock with paid access: IEEE 519 / GB/T 14549 key metrics summary and compliance determination", compact=True)
        else:
            st.markdown("**IEEE 519-2022**")
            col_v1, col_v2, col_v3, col_v4 = st.columns(4)
            with col_v1:
                st.markdown(_card(
                    "📊", "IEEE TDD (Overall)",
                    f"{tdd_calc:.2f}%",
                    f"{'✅ Pass' if tdd_pass else '❌ Exceeds Limit'} | Limit: {tdd_limit:.1f}% ({scr_label})",
                    tdd_pass), unsafe_allow_html=True)
            with col_v2:
                st.markdown(_card(
                    "🔢", f"Switching Frequency Harmonic ({sw_order:.0f}th order)",
                    f"{tdd_sw:.2f}% TDD",
                    f"{'✅ Pass' if sw_pass else '❌ Exceeds Limit'} | Limit: {limit_sw_pct:.1f}%",
                    sw_pass), unsafe_allow_html=True)
            with col_v3:
                fres_str = f"{_f_res:.0f} Hz" if _f_res > 0 else "N/A"
                st.markdown(_card(
                    "📐", "LCL Resonance Frequency",
                    fres_str,
                    f"{'✅ In range' if f_res_ok else '⚠️ Out of range'} | Required: {10*h7_fg:.0f}~{0.5*h7_fsw:.0f} Hz",
                    f_res_ok if _f_res > 0 else None), unsafe_allow_html=True)
            with col_v4:
                st.markdown(_card(
                    "⚡", "Topology / Modulation",
                    "3-Level" if _is_3level else "2-Level",
                    f"Modulation index mₐ = {fft_ma:.3f} | Switching frequency fsw = {h7_fsw:.0f} Hz",
                    None), unsafe_allow_html=True)
            st.markdown("<div style='margin-top:8px'></div>", unsafe_allow_html=True)

        #  GB/T LCL 
        # GB/T 25, 25 I_h = I_5_user * 5/h 
        #  harm_freqs  lcl_responsible_freq 
        gbt_high_check_freqs = [f for f in harm_freqs if f >= lcl_responsible_freq] if len(harm_freqs) > 0 else []
        gbt_lcl_violations   = 0
        gbt_lcl_total        = 0
        for f_h in gbt_high_check_freqs:
            h_int = int(round(f_h / h7_fg))
            if h_int < 2:
                continue
            idx_f = np.argmin(abs(harm_freqs - f_h))
            ig_h  = float(i_grid_A_arr[idx_f]) if len(i_grid_A_arr) > 0 else 0.0
            if h_int in gbt_orders_full:
                idx_g   = gbt_orders_full.index(h_int)
                lim_gbt = gbt_user_A_full[idx_g]
            elif h_int > 25:
                idx5    = gbt_orders_full.index(5) if 5 in gbt_orders_full else 3
                lim_gbt = gbt_user_A_full[idx5] * 5.0 / h_int
            else:
                continue
            if lim_gbt > 0:
                gbt_lcl_total += 1
                if ig_h > lim_gbt:
                    gbt_lcl_violations += 1

        gbt_lcl_pass    = (gbt_lcl_violations == 0) and (gbt_lcl_total > 0)
        gbt_no_lcl_data = (gbt_lcl_total == 0)

        # : th Harmonic()
        gbt_low_freqs      = [f for f in harm_freqs if f < lcl_responsible_freq] if len(harm_freqs) > 0 else []
        gbt_ctrl_needs_fix = 0
        for f_h in gbt_low_freqs:
            h_int = int(round(f_h / h7_fg))
            if h_int not in gbt_orders_full:
                continue
            idx_f = np.argmin(abs(harm_freqs - f_h))
            ig_h  = float(i_grid_A_arr[idx_f]) if len(i_grid_A_arr) > 0 else 0.0
            idx_g   = gbt_orders_full.index(h_int)
            lim_gbt = gbt_user_A_full[idx_g]
            if lim_gbt > 0 and ig_h > lim_gbt:
                gbt_ctrl_needs_fix += 1

        # ── 2: GB/T 14549 (, )────────────────────────
        if not is_trial:
            st.markdown("**GB/T 14549-1993**")
            col_g1, col_g2, col_g3, col_g4 = st.columns(4)
            with col_g1:
                if gbt_no_lcl_data:
                    g1_val = "No data"
                    g1_note = "LCL attenuation start point is above GB/T range — see explanation below"
                    g1_pass = None
                elif gbt_lcl_pass:
                    g1_val = "✅ All pass"
                    g1_note = f"LCL zone: {gbt_lcl_total} harmonics checked, 0 exceed limit"
                    g1_pass = True
                else:
                    g1_val = f"❌ {gbt_lcl_violations} exceed limit"
                    g1_note = f"LCL zone: {gbt_lcl_total} checked — adjust LCL parameters"
                    g1_pass = False
                st.markdown(_card("📋", "GB/T LCL Zone", g1_val, g1_note, g1_pass), unsafe_allow_html=True)
            with col_g2:
                if gbt_ctrl_needs_fix == 0:
                    g2_val = "✅ Within limit"
                    g2_note = "Low-order harmonics (control algorithm zone) — model reference only; not an LCL failure"
                    g2_pass = None
                else:
                    g2_val = f"⚠️ {gbt_ctrl_needs_fix} need compensation"
                    g2_note = "Low-order harmonics exceed estimate — use PR / repetitive control compensation; not an LCL design failure"
                    g2_pass = None
                st.markdown(_card("⚙️", "Control Algorithm Zone (reference)", g2_val, g2_note, g2_pass), unsafe_allow_html=True)
            with col_g3:
                st.markdown(_card(
                    "📐", "LCL Attenuation Start Point",
                    f"≥ {lcl_responsible_freq:.0f} Hz",
                    f"Above ~{lcl_responsible_order:.0f}th order LCL enters effective attenuation (ref: Liserre 2005)",
                    None), unsafe_allow_html=True)
            with col_g4:
                st.markdown(_card(
                    "🔄", "Transformer & Scaling",
                    f"N = {_k_tr:.2f}",
                    f"Capacity ratio k = {ratio_gbt:.4f} | Voltage level: {v_level_str} | Sref = {s_ref_gbt:.0f} MVA",
                    None), unsafe_allow_html=True)
            st.markdown("<div style='margin-top:8px'></div>", unsafe_allow_html=True)

        st.divider()

        # ════════════════════════════════════════════════════════════════════════
        # 2: IEEE 519  — th Harmonic
        # : IEEE 519-2022 Table 2 Note:  =  × 25%
        # ════════════════════════════════════════════════════════════════════════
        st.markdown("#### 📋 IEEE 519-2022 Harmonic Verification (Itemized)")

        # Build data for all harmonics: [3~49] odd + [2~50] even
        ieee_odd_orders  = [3, 5, 7, 9, 11, 13, 17, 19, 23, 25, 35, 49]
        ieee_even_orders = [2, 4, 6, 8, 10, 12, 14, 16, 18, 20, 22, 24, 26]
        ieee_all_orders  = sorted(set(ieee_odd_orders + ieee_even_orders))

        ieee_rows      = []
        violations_ieee = []

        for h in ieee_all_orders:
            f_check  = h * h7_fg
            is_even  = (h % 2 == 0)

            # Look up measured TDD from FFT results
            if len(harm_freqs) > 0:
                idx_c = np.argmin(abs(harm_freqs - f_check))
                tdd_h = float(i_tdd_pct[idx_c]) if abs(harm_freqs[idx_c] - f_check) < 5 * h7_fg else 0.0
            else:
                tdd_h = 0.0

            # Odd harmonic limit from SCR table
            if h < 11:   lim_odd = i_h_limits[(3,11)]
            elif h < 17: lim_odd = i_h_limits[(11,17)]
            elif h < 23: lim_odd = i_h_limits[(17,23)]
            elif h < 35: lim_odd = i_h_limits[(23,35)]
            else:        lim_odd = i_h_limits[(35,999)]

            # Even harmonic limit = odd limit × 25%
            lim_h = lim_odd * 0.25 if is_even else lim_odd

            is_lcl_range = (f_check >= lcl_responsible_freq)
            harmonic_type = "Even" if is_even else "Odd"

            if is_lcl_range:
                status = "✅ Pass" if tdd_h <= lim_h else "❌ Exceeds Limit"
            else:
                status = "⚙️ Control algo" if tdd_h <= lim_h else "⚠️ Control algo (needs compensation)"

            ieee_rows.append({
                'Order': h,
                'Odd/Even': harmonic_type,
                'Frequency (Hz)': f'{f_check:.0f}',
                'Measured TDD (%)': f'{tdd_h:.3f}',
                'Limit TDD (%)': f'{lim_h:.2f}',
                'Responsibility': 'LCL' if is_lcl_range else 'Control algo',
                'Result': status
            })
            if is_lcl_range and tdd_h > lim_h:
                violations_ieee.append({
                    'Frequency (Hz)': f'{f_check:.0f}',
                    'Order': f'{h} ({harmonic_type})',
                    'Measured TDD (%)': f'{tdd_h:.3f}',
                    'Limit (%)': f'{lim_h:.2f}',
                    'Exceeds by': f'{tdd_h/lim_h:.2f}x'
                })

        if is_trial:
            _skel_ieee = pd.DataFrame([
                {'Order': h, 'Odd/Even': 'Even' if h%2==0 else 'Odd',
                 'Frequency (Hz)': f'{int(h*h7_fg)}',
                 'Measured TDD (%)': '🔒', 'Limit TDD (%)': '🔒',
                 'Responsibility': '🔒', 'Result': '🔒'}
                for h in [3, 5, 7, 9, 11, 13, 17, 19, 23, 25]
            ])
            st.dataframe(_skel_ieee, width="stretch", hide_index=True)
            trial_lock("Unlock with paid access: IEEE 519-2022 itemized harmonic verification table (including pass/fail judgment)", compact=True)
        else:
            st.caption("Includes odd and even harmonics. Even harmonic limit = same-tier odd limit × 25% (IEEE 519-2022 Table 2 footnote). Blue rows = control algorithm responsibility (reference, not counted toward LCL assessment).")
            df_ieee_check = pd.DataFrame(ieee_rows)

            def color_ieee_row(row):
                c = row['Result']
                if 'Pass' in c:           return ['background-color:#d4edda'] * len(row)
                elif 'Exceeds Limit' in c: return ['background-color:#f8d7da'] * len(row)
                elif 'compensation' in c:  return ['background-color:#fff3cd'] * len(row)
                else:                      return ['background-color:#e8f4fd'] * len(row)

            st.dataframe(df_ieee_check.style.apply(color_ieee_row, axis=1),
                         width="stretch", hide_index=True)
            st.caption("🟢 = LCL pass | 🔴 = LCL exceeds limit | 🔵 = control algorithm (reference, not LCL) | 🟡 = control algorithm exceeds limit (needs compensation)")

            if violations_ieee:
                st.error(f"⚠️ IEEE 519 **LCL responsible range** has **{len(violations_ieee)}** harmonic(s) exceeding limits:")
                st.dataframe(pd.DataFrame(violations_ieee), width="stretch", hide_index=True)

        st.divider()

        # ════════════════════════════════════════════════════════════════════════
        # GB/T 14549 itemized verification (split by LCL responsible / control algo zones)
        # LCL zone: scan all harm_freqs >= lcl_responsible_freq
        # Control algo zone: low-frequency harmonics (2~25th), display as reference
        # ════════════════════════════════════════════════════════════════════════
        st.markdown("#### 📋 GB/T 14549-1993 Harmonic Verification (by Zone)")

        if is_trial:
            _skel_gbt = pd.DataFrame([
                {'Order': h, 'Odd/Even': 'Even' if h%2==0 else 'Odd',
                 'Frequency (Hz)': f'{int(h*h7_fg)}',
                 'Measured (Arms)': '🔒', 'Limit (A)': '🔒',
                 'Margin': '🔒', 'Result': '🔒'}
                for h in [5, 7, 11, 13, 17, 19, 23, 25]
            ])
            st.dataframe(_skel_gbt, width="stretch", hide_index=True)
            trial_lock("Unlock with paid access: GB/T 14549-1993 itemized harmonic verification (including scaled limits and pass margin)", compact=True)
        else:
            # ── LCL responsible zone: scan harm_freqs >= lcl_responsible_freq ──
            st.markdown(f"""
            <div style="background:#e8f5e9; border-left:4px solid #198754; padding:8px 12px; border-radius:4px; margin-bottom:8px;">
              <b style="color:#198754;">✅ LCL Responsible Zone (≥ {lcl_responsible_freq:.0f} Hz, above ~{lcl_responsible_order:.0f}th order)</b>
              <span style="font-size:12px; color:#444;"> — LCL enters effective attenuation region; harmonics here are formally assessed against LCL design</span>
            </div>
            """, unsafe_allow_html=True)

            gbt_high_rows = []
            if len(harm_freqs) > 0:
                high_freq_candidates = sorted(set([
                    int(round(f / h7_fg)) for f in harm_freqs if f >= lcl_responsible_freq
                ]))
                # Only consider harmonics >= 2nd order
                high_freq_candidates = [h for h in high_freq_candidates if h >= 2]

                for h_int in high_freq_candidates:
                    f_check = h_int * h7_fg
                    idx_f   = np.argmin(abs(harm_freqs - f_check))
                    ig_h    = float(i_grid_A_arr[idx_f]) if abs(harm_freqs[idx_f] - f_check) < 3 * h7_fg else 0.0

                    # Determine GB/T limit
                    if h_int in gbt_orders_full:
                        idx_g   = gbt_orders_full.index(h_int)
                        lim_gbt = gbt_user_A_full[idx_g]
                    elif h_int > 25:
                        # GB/T extrapolation: I_h = I_5_ref * 5/h, even harmonics ×50%
                        idx5     = gbt_orders_full.index(5) if 5 in gbt_orders_full else 3
                        i5_user  = gbt_user_A_full[idx5]
                        if h_int % 2 == 0:
                            lim_gbt = i5_user * 5.0 / h_int * 0.5
                        else:
                            lim_gbt = i5_user * 5.0 / h_int
                    else:
                        continue

                    if lim_gbt <= 0:
                        continue

                    is_even_h  = (h_int % 2 == 0)
                    status_gbt = "✅ Pass" if ig_h <= lim_gbt else "❌ Exceeds Limit"
                    margin     = (lim_gbt - ig_h) / lim_gbt * 100

                    gbt_high_rows.append({
                        'Order': h_int,
                        'Odd/Even': 'Even' if is_even_h else 'Odd',
                        'Frequency (Hz)': f'{f_check:.0f}',
                        'Measured (Arms)': f'{ig_h:.4f}',
                        'Limit (A)': f'{lim_gbt:.4f}',
                        'Margin': f'{margin:+.1f}%',
                        'Result': status_gbt
                    })

            if gbt_high_rows:
                df_gbt_high = pd.DataFrame(gbt_high_rows)
                def color_gbt_high(row):
                    if 'Pass' in row['Result']:
                        return ['background-color:#d4edda'] * len(row)
                    return ['background-color:#f8d7da'] * len(row)
                st.dataframe(df_gbt_high.style.apply(color_gbt_high, axis=1),
                             width="stretch", hide_index=True)
                n_gbt_pass = sum(1 for r in gbt_high_rows if 'Pass' in r['Result'])
                n_gbt_fail = len(gbt_high_rows) - n_gbt_pass
                if n_gbt_fail == 0:
                    st.success(f"✅ LCL responsible zone: **{len(gbt_high_rows)}** harmonic(s) checked, **all pass**.")
                else:
                    st.error(f"❌ LCL responsible zone: **{len(gbt_high_rows)}** checked, **{n_gbt_fail}** exceed limit — adjust LCL parameters.")
            else:
                st.markdown(f"""
                <div style="background:#fff8e1; border-left:4px solid #ffc107; padding:10px 14px; border-radius:4px;">
                  <b style="color:#856404;">📌 No GB/T LCL-zone data for current parameters — explained below</b><br><br>
                  <span style="font-size:13px;">
                  GB/T 14549 defines limits for harmonics up to the <b>25th order</b>. The LCL attenuation start point
                  is <b>{lcl_responsible_freq:.0f} Hz (≈{lcl_responsible_order:.0f}th order)</b>.<br><br>
                  Harmonics above the <b>{lcl_responsible_order:.0f}th order</b> (i.e., {int(lcl_responsible_order)+1}th, {int(lcl_responsible_order)+2}th…)
                  fall outside GB/T's 2–25th range; the FFT model includes LCL attenuation for these,
                  so no "exceeds limit" data is expected.<br><br>
                  <b>This is expected behavior — it confirms LCL effectively attenuates high-frequency harmonics.</b><br>
                  Low-frequency harmonics (2nd–{int(lcl_responsible_order)-1}th) are shown as "control algorithm" reference below.
                  </span>
                </div>
                """, unsafe_allow_html=True)

            # ── Control algorithm zone: below lcl_responsible_order (GB/T 2~25th) ──
            st.markdown(f"""
            <div style="background:#e8f4fd; border-left:4px solid #0d6efd; padding:8px 12px; border-radius:4px; margin:14px 0 8px 0;">
              <b style="color:#0d6efd;">⚙️ Control Algorithm Zone (&lt; {lcl_responsible_freq:.0f} Hz, below ~{lcl_responsible_order:.0f}th order)</b>
              <span style="font-size:12px; color:#555;"> — LCL has no physical attenuation here (gain≈1); use PR control / repetitive control / harmonic compensation</span>
            </div>
            """, unsafe_allow_html=True)

            gbt_low_orders = [h for h in gbt_orders_full if h * h7_fg < lcl_responsible_freq]
            if gbt_low_orders:
                gbt_low_rows = []
                for h in gbt_low_orders:
                    f_check = h * h7_fg
                    if len(harm_freqs) > 0:
                        idx_c = np.argmin(abs(harm_freqs - f_check))
                        ig_h  = float(i_grid_A_arr[idx_c]) if abs(harm_freqs[idx_c] - f_check) < 5 * h7_fg else 0.0
                    else:
                        ig_h = 0.0
                    idx_g   = gbt_orders_full.index(h)
                    lim_gbt = gbt_user_A_full[idx_g]
                    if lim_gbt <= 0:
                        continue
                    margin_pct = (lim_gbt - ig_h) / lim_gbt * 100
                    is_even_h  = (h % 2 == 0)
                    if ig_h <= lim_gbt:
                        status_ref = f"📊 Reference (margin {margin_pct:.0f}%)"
                    else:
                        status_ref = f"⚠️ Reference (exceeds by {(ig_h/lim_gbt-1)*100:.0f}%, control algo compensation needed)"
                    gbt_low_rows.append({
                        'Order': h,
                        'Odd/Even': 'Even' if is_even_h else 'Odd',
                        'Frequency (Hz)': f'{f_check:.0f}',
                        'Measured (Arms)': f'{ig_h:.3f}',
                        'Limit (A)': f'{lim_gbt:.3f}',
                        'Status (reference, not LCL)': status_ref
                    })

                df_gbt_low = pd.DataFrame(gbt_low_rows)
                def color_gbt_low(row):
                    s = row['Status (reference, not LCL)']
                    if 'exceeds' in s:
                        return ['background-color:#fff3cd'] * len(row)
                    return ['background-color:#e8f4fd'] * len(row)
                st.dataframe(df_gbt_low.style.apply(color_gbt_low, axis=1),
                             width="stretch", hide_index=True)

            # Legend
            st.markdown("""
            <div style="font-size:13px; margin-top:6px; padding:8px; background:#f8f9fa; border-radius:4px;">
              <b>Legend: </b>
              &nbsp;🟢 <span style="background:#d4edda; padding:1px 6px; border-radius:3px;">Green</span> = LCL zone, pass &nbsp;|&nbsp;
              🔴 <span style="background:#f8d7da; padding:1px 6px; border-radius:3px;">Red</span> = LCL zone, exceeds limit (adjust LCL parameters) &nbsp;|&nbsp;
              🔵 <span style="background:#e8f4fd; padding:1px 6px; border-radius:3px;">Blue</span> = Control algo zone, within limit (reference, not LCL assessment) &nbsp;|&nbsp;
              🟡 <span style="background:#fff3cd; padding:1px 6px; border-radius:3px;">Yellow</span> = Control algo zone, estimate exceeds limit (needs compensation — reference, not LCL design failure)
            </div>
            """, unsafe_allow_html=True)
            st.caption(f"Limits include capacity scaling k={ratio_gbt:.4f}, transformer ratio N={_k_tr:.2f} | Above 25th order: odd Iₕ=I₅×5/h, even ×50%")

        st.divider()

        # ── Engineering Comprehensive Conclusion ──────────────────────────────────────────────────────────
        st.markdown("### 🏁 Engineering Comprehensive Conclusion")
        all_pass = tdd_pass and sw_pass and (len(violations_ieee) == 0) and (gbt_no_lcl_data or gbt_lcl_pass)

        if is_trial:
            trial_lock("Unlock with paid access: engineering comprehensive conclusion and compliance determination")
        else:
            if all_pass and f_res_ok:
                st.success(f"""
                ## ✅ Current LCL design meets harmonic standard requirements

                **IEEE 519-2022**
                - TDD = {tdd_calc:.2f}% ≤ limit {tdd_limit:.1f}% (SCR={_scr:.1f}, {scr_label}) ✅
                - Switching frequency ({sw_order:.0f}th order) harmonic = {tdd_sw:.2f}% TDD ≤ {limit_sw_pct:.1f}% ✅

                **GB/T 14549-1993**
                - LCL responsible zone (≥{lcl_responsible_freq:.0f} Hz): {'All pass ✅' if gbt_lcl_pass else 'No data (LCL fully attenuates) ✅'}

                **LCL Parameters**
                - Resonance frequency {_f_res:.0f} Hz, satisfies {10*h7_fg:.0f}~{0.5*h7_fsw:.0f} Hz ✅
                - Topology: {h7_topo} | Modulation index: mₐ={fft_ma:.3f} ({_ma_note})

                Recommend MATLAB/Simulink time-domain simulation verification; for formal certification, measure at PCC with an analyzer meeting **IEC 61000-4-7**.
                """)
            else:
                issues = []
                if not tdd_pass:
                    issues.append(f"• **IEEE TDD exceeds limit**: {tdd_calc:.2f}% > {tdd_limit:.1f}% — increase L2 or Cf to improve high-frequency attenuation.")
                if not sw_pass:
                    issues.append(f"• **Switching frequency harmonic exceeds limit**: {tdd_sw:.2f}% > {limit_sw_pct:.1f}% — LCL attenuation insufficient at fsw.")
                if not f_res_ok:
                    issues.append(f"• **Resonance frequency out of range**: {_f_res:.0f} Hz should be in {10*h7_fg:.0f}~{0.5*h7_fsw:.0f} Hz.")
                if violations_ieee:
                    issues.append(f"• **IEEE 519 LCL zone: {len(violations_ieee)} harmonic(s) exceed limit**.")
                if not gbt_lcl_pass and not gbt_no_lcl_data:
                    issues.append(f"• **GB/T 14549 LCL zone: {gbt_lcl_violations} harmonic(s) exceed limit** — improve high-frequency attenuation.")
                st.error("## ❌ Current LCL design has issues\n\n" + "\n".join(issues))
                col_f1, col_f2 = st.columns(2)
                with col_f1:
                    st.info("**🔧 TDD / GB/T improvement directions**\n1. Increase L2 (improve high-frequency attenuation)\n2. Increase Cf (lower resonance frequency)\n3. Raise fsw (push harmonics to higher frequencies)")
                with col_f2:
                    st.warning("**🔧 Resonance frequency improvement directions**\n1. Adjust L1/L2 ratio (recommend L2≈0.3~0.5×L1)\n2. Ensure fres < fsw/2\n3. Damping resistor starting point: 1/(3·ωres·Cf)")

            # ── :  L1, L2, Cf Parameter ──────────────────────────
            # : 
            #   [1] Liserre M, Blaabjerg F, Hansen S. IEEE Trans Ind Electron. 2005;52(5):1202-1214.
            #       — ripple current: L_total = Vdc / (8 * fsw * ΔiL1),  ΔiL1 = 0.10~0.20 * I_rated_peak
            #       — : Cf ≤ 0.05 * C_base, C_base = 1/(ω0 * Z_base)
            #   [2] Beres R N, et al. IEEE Trans Power Electron. 2016;31(7):4954-4970.
            #       — L1:L2  r = 2~4,  r=2.5, 
            #       — fres : 10f0 ≤ fres ≤ 0.5fsw
            # ⚠️  ±10~15%, Parameter, .
            st.markdown("---")
            st.markdown("### 🔧 Optimization: Recommended LCL Parameters")
            st.caption("Based on Liserre 2005 & Beres 2016 methods — derives parameters that satisfy harmonic requirements. Use as engineering reference.")

            if st.button("⚡ Calculate Recommended Parameters", key="h7_opt_btn", type="primary"):
                # ── Step 1:  ────────────────────────────────────
                # Z_base = Ug_phase² / P_rated = (Ug_LL/√3)² / (P/3) = Ug_LL² / P  []
                if _Prated > 0 and h7_ug_inv > 0 and h7_fg > 0:
                    Z_base  = h7_ug_inv ** 2 / _Prated
                    C_base  = 1.0 / (_omega0 * Z_base)
                    I_rated_peak = np.sqrt(2) * h7_il

                    # ── Step 2: Cf(,  3%~5%C_base)──────────────────
                    #  4% : Resonance Frequency, ; 
                    Cf_opt_pct = 0.04
                    Cf_opt     = Cf_opt_pct * C_base  # unit F

                    # ── Step 3: L_total(, ΔiL1 = 15% )─────────
                    # L_total = Vdc / (8 * fsw * ΔiL)   [Holmes & Lipo 2003, Ch.2, SPWM]
                    delta_i = 0.15 * I_rated_peak
                    L_total = h7_vdc / (8.0 * h7_fsw * delta_i)  # unit H

                    # ── Step 4: L1, L2 (Beres 2016  r=2.5)──────────────────
                    r_opt = 2.5  # L1/L2 
                    L2_opt = L_total / (1.0 + r_opt)
                    L1_opt = L_total - L2_opt

                    # ── Step 5:  fres ────────────────────────────────────────────
                    if L1_opt > 0 and L2_opt > 0 and Cf_opt > 0:
                        fres_opt = np.sqrt((L1_opt + L2_opt) / (L1_opt * L2_opt * Cf_opt)) / (2 * np.pi)
                    else:
                        fres_opt = 0.0

                    fres_ok_opt = (10 * h7_fg <= fres_opt <= 0.5 * h7_fsw) if fres_opt > 0 else False

                    # ── Step 6:  fres ,  Cf ─────────────────────────────
                    # fres :  fsw/6(:  [fsw/10, fsw/2] , )
                    fres_target = h7_fsw / 6.0
                    if not fres_ok_opt and L1_opt > 0 and L2_opt > 0:
                        # Cf = (L1+L2) / [L1*L2*(2π*fres_target)²]
                        Cf_adj = (L1_opt + L2_opt) / (L1_opt * L2_opt * (2 * np.pi * fres_target) ** 2)
                        if Cf_adj <= 0.05 * C_base:  # satisfies reactive power constraint
                            Cf_opt = Cf_adj
                            fres_opt = fres_target
                            fres_ok_opt = True

                    # ── Show recommended parameters ────────────────────────────────────────────────
                    st.success("✅ Recommended parameters calculated (engineering estimate, ±15% error — verify with simulation)")

                    if is_trial:
                        trial_lock("Optimization results require paid access")
                    else:
                        col_o1, col_o2, col_o3, col_o4 = st.columns(4)
                        with col_o1:
                            st.markdown(_card("🔵", "Recommended L1",
                                f"{L1_opt*1e3:.3f} mH",
                                f"Current: {h7_l1:.3f} mH | L1/L2 = {r_opt:.1f} (Beres 2016)",
                                None), unsafe_allow_html=True)
                        with col_o2:
                            st.markdown(_card("🟣", "Recommended L2",
                                f"{L2_opt*1e3:.3f} mH",
                                f"Current: {h7_l2:.3f} mH | L_total = {L_total*1e3:.3f} mH",
                                None), unsafe_allow_html=True)
                        with col_o3:
                            st.markdown(_card("🟠", "Recommended Cf",
                                f"{Cf_opt*1e6:.2f} μF",
                                f"Current: {h7_cf:.2f} μF | {Cf_opt/C_base*100:.1f}% of C_base (≤5%)",
                                None), unsafe_allow_html=True)
                        with col_o4:
                            st.markdown(_card("📐", "Resulting fres",
                                f"{fres_opt:.0f} Hz",
                                f"{'✅ In range' if fres_ok_opt else '⚠️ Out of range'} | Required: {10*h7_fg:.0f}~{0.5*h7_fsw:.0f} Hz",
                                fres_ok_opt if fres_opt > 0 else None), unsafe_allow_html=True)

                    st.markdown("<div style='margin-top:10px'></div>", unsafe_allow_html=True)

                    # 
                    with st.expander("📐 Derivation of Recommended Parameters"):
                        st.markdown(f"""
**Step 1 — System Reference Parameters**
- Reference impedance: $Z_{{base}} = U_{{LL}}^2 / P_{{rated}} = {h7_ug_inv:.0f}^2 / {_Prated/1e3:.0f}\\text{{kW}} = {Z_base:.4f}\\,\\Omega$
- Reference capacitance: $C_{{base}} = 1/(\\omega_0 Z_{{base}}) = {C_base*1e6:.2f}\\,\\mu F$

**Step 2 — Capacitor Cf (reactive power constraint, target {Cf_opt_pct*100:.0f}% of C_base)**
- Basis: GB/T 14549 §4.2 limits reactive power to 5% of rated; Liserre 2005 recommends $C_f \\leq 5\\%\\, C_{{base}}$
- $C_f = {Cf_opt_pct*100:.0f}\\%\\times {C_base*1e6:.2f}\\,\\mu F = {Cf_opt*1e6:.2f}\\,\\mu F$

**Step 3 — Total inductance L_total (ripple current constraint, $\\Delta i_{{L1}} = 15\\%\\, \\hat{{I}}_{{rated}}$)**
- Basis: Holmes & Lipo 2003, Ch.2, SPWM ripple current formula
- $L_{{total}} = V_{{dc}} / (8 f_{{sw}} \\Delta i_L) = {h7_vdc:.0f} / (8 \\times {h7_fsw:.0f} \\times {delta_i:.2f}) = {L_total*1e3:.3f}\\,mH$

**Step 4 — Split L1 and L2 ($r = L1/L2 = {r_opt}$, Beres 2016)**
- $L_2 = L_{{total}} / (1+r) = {L2_opt*1e3:.3f}\\,mH$, $L_1 = {L1_opt*1e3:.3f}\\,mH$

**Step 5 — Resonance frequency verification**
- $f_{{res}} = \\frac{{1}}{{2\\pi}}\\sqrt{{\\frac{{L_1+L_2}}{{L_1 L_2 C_f}}}} = {fres_opt:.0f}\\,Hz$
- Required range: $[{10*h7_fg:.0f},\\,{0.5*h7_fsw:.0f}]\\,Hz$ → **{'✅ Satisfied' if fres_ok_opt else '⚠️ Out of range — Cf adjusted in Step 6'}**

**References**
1. Liserre M, Blaabjerg F, Hansen S. *Design and Control of an LCL-Filter-Based Three-Phase Active Rectifier.* IEEE Trans Ind Appl. 2005;41(5):1281–1291.
2. Beres R N, Wang X, Blaabjerg F, et al. *Optimal Design of High-Order Passive-Damped Filters for Grid-Connected Applications.* IEEE Trans Power Electron. 2016;31(4):2083–2098.
3. Holmes D G, Lipo T A. *Pulse Width Modulation for Power Converters.* IEEE Press, 2003.
                        """)

                        st.info("💡 **Next step**: Enter the recommended parameters at the top of the page in the 「LCL Filter Parameters」 section, then re-run the verification. For formal design, supplement with MATLAB/Simulink simulation.")
                else:
                    st.warning("⚠️ Please fill in the system parameters at the top of the page (rated power, converter voltage, DC bus voltage) before calculating.")

        st.divider()
        with st.expander("📌 Simulation Limitations & Engineering Verification Guide"):
            st.markdown(f"""
#### Model Limitations

1. **Low-order harmonics**: This model uses ideal linear modulation theory. Low-order harmonic (5th, 7th, 11th, 13th) errors can be ±50% — use Simulink time-domain simulation for precise results.
2. **3-Level model note**: The 3-Level harmonic model assumes ideal symmetry; actual delays and dead-time distortion may increase low-order content by 5~10%.
3. **Inductor saturation**: Recommended safety margin of 15~20% on inductor peak current rating.
4. **Background harmonics**: Grid background harmonics are not modelled (see IEEE 519-2022 Annex B).

#### Engineering Verification Checklist

| Step | Tool | Standard |
|---|---|---|
| Simulation verification | MATLAB/Simulink + PLECS | Error target <5% |
| Lab harmonic measurement | IEC 61000-4-7 analyzer | IEEE 519-2022 Sec.5 |
| Measurement point | **At PCC** (after transformer) | IEEE 519-2022 Sec.4 |
| Measurement window | Pst: 10 min; Plt: 2 h | IEEE 519-2022 Table 4 |
| Recommended analyzer | Fluke 435-II, HIOKI PW6001 | IEC 61000-4-30 |
            """)










# ==========================================
#  8: ()  v2
#
# v2 : 
#   [Fix1] Tab3 (, )
#          : Z_filt = abs(...), I_grid = I_src * Z_filt / (Z_filt + Zg)
#          : Z_filt = , Zg = , I_grid = I_src * |Zf|/|Zf+Zg|
#   [Add1] Tab2 Impedance Curve Zsys = Zf*Zg/(Zf+Zg)
#          , 
#          : Das J.C., Power System Harmonics, Wiley-IEEE, 2015, §8.5
#   [Add2] Tab3 "7Exceeds Limit" 
#           session_state['h7_harm_export'](7 FFT Tab )
#          : 7 FFT, 
#   [Add3] Tab2 (Ssc ±20% )
#           Ssc  Zsys Parallel Anti-Resonance
#          (3/5/7/11/13)
#          : Das J.C., 2015, §8.5; IEEE Std 1531-2003, §5.7
#   [Add4] Tab3 (/)
#          Individual Harmonics, 
#          
#          Note: (), , 
#
# : 7 FFT Tab : 
#   st.session_state['h7_harm_export'] = {
#       'freqs': harm_freqs.tolist(),
#       'i_grid_A': i_grid_A_arr.tolist(),
#       'i_tdd_pct': i_tdd_pct.tolist(),
#       'il': h7_il,
#       'fg': h7_fg,
#       'violations': violations_ieee,  # list of dicts
#   }
#
# : 
#   [1] IEEE Std 519-2022: Recommended Practice for Harmonic Control
#   [2] IEEE Std 1531-2003: Guide for Application and Specification of
#       Harmonic Filters
#   [3] Das J.C., Power System Harmonics and Passive Filter Designs,
#       Wiley-IEEE Press, 2015. §8.5()
#   [4] Dugan R.C. et al., Electrical Power Systems Quality, 3rd ed.,
#       McGraw-Hill, 2012. Chapter 8.
# ==========================================

elif selection == nav_options[7]:

    # ──  session_state Parameter ─────────────────────────────────────────
    _fg     = float(st.session_state.get('sys_fg',   50.0))
    _fsw    = float(st.session_state.get('sys_fsw',  2500.0))
    _L1_mH  = float(st.session_state.get('lcl_l1',  0.500))
    _L2_mH  = float(st.session_state.get('lcl_l2',  0.250))
    _Cf_uF  = float(st.session_state.get('lcl_c',   100.0))
    _Ug     = float(st.session_state.get('sys_ug',   690.0))
    _IL     = float(st.session_state.get('vec_i',    500.0))
    _Vdc    = float(st.session_state.get('vec_udc',  1100.0))
    _Prated = float(st.session_state.get('h7_prated',500.0)) * 1e3

    L1 = _L1_mH * 1e-3
    L2 = _L2_mH * 1e-3
    Cf = _Cf_uF  * 1e-6

    f_res_lcl = (1 / (2 * np.pi)) * np.sqrt((L1 + L2) / (L1 * L2 * Cf)) if (L1*L2*Cf) > 0 else 0.0

    # ══════════════════════════════════════════════════════════════════════════
    # Page header
    # ══════════════════════════════════════════════════════════════════════════
    st.header("🎯 Passive Harmonic Filter Design")
    st.markdown("""
> **Overview**: When Section 7 harmonic verification shows a specific harmonic order exceeds the IEEE 519 limit,
> and the LCL filter cannot suppress low-order harmonics, a **Single-Tuned Passive Filter** must be designed
> for the specific exceeding order so the system meets the grid-connection harmonic standard.
    """)
    st.divider()

    tab_why, tab_design, tab_verify = st.tabs([
        "📖 1. Why a Passive Filter?",
        "🔧 2. Filter Design + Impedance Characteristic (Anti-Resonance)",
        "✅ 3. Harmonic Verification (Meets IEEE 519?)"
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 1 — Why a passive filter is needed
    # ══════════════════════════════════════════════════════════════════════════
    with tab_why:
        st.markdown("### 🧩 LCL Filter Limitation: Why Can't It Suppress Low-Order Harmonics?")
        c1, c2 = st.columns(2, gap="large")
        with c1:
            st.success(f"""
**✅ LCL Strength: High-Frequency Attenuation**

LCL filter attenuates high-frequency harmonics at **-60 dB/dec** slope.

- Resonance Frequency $f_{{res}}$ = **{f_res_lcl:.1f} Hz** (≈ {f_res_lcl/_fg:.0f}× fundamental)
- Switching Frequency $f_{{sw}}$ = **{_fsw:.0f} Hz** ({_fsw/_fg:.0f}× fundamental)

Above $f_{{sw}}$ and its sidebands, LCL provides strong attenuation —
this is the core value of the LCL topology.
            """)
        with c2:
            st.error(f"""
**❌ LCL Limitation: Cannot Suppress Low-Order Harmonics**

Below $f_{{res}}$, the LCL transfer function gain **≈ 1**:

$$G_{{LCL}}(j\\omega) \\approx 1 \\quad (\\omega \\ll \\omega_{{res}})$$

**5th ({5*_fg:.0f} Hz), 7th ({7*_fg:.0f} Hz),
11th ({11*_fg:.0f} Hz), 13th ({13*_fg:.0f} Hz)**
— these low-order harmonics pass through the LCL with almost no attenuation.

A **passive filter** or **control algorithm** (PR/repetitive control) is required.
            """)

        st.divider()
        st.markdown("### 📐 Passive Filter Operating Principle")
        c3, c4 = st.columns([1.2, 1], gap="large")
        with c3:
            st.markdown(f"""
**Single-Tuned Passive Filter** connects $L_f$ and $C_f$ in series,
with near-zero impedance at the tuning frequency:

$$Z_{{filter}}(j\\omega) = R_f + j\\omega L_f + \\frac{{1}}{{j\\omega C_f}}$$

**Tuning condition (series resonance)**:

$$f_{{tuned}} = \\frac{{1}}{{2\\pi\\sqrt{{L_f C_f}}}} \\Rightarrow Z(f_{{tuned}}) = R_f \\approx 0$$

Harmonic current is **preferentially diverted into the filter** rather than the grid, significantly reducing that harmonic order.

> 📚 Dugan R.C. et al., *Electrical Power Systems Quality*, 3rd ed.,
> McGraw-Hill, 2012, §8.4.
            """)
        with c4:
            svg_tf = """
<svg width="100%" height="320" viewBox="0 0 440 320" xmlns="http://www.w3.org/2000/svg" font-family="Arial,sans-serif">
  <rect width="440" height="320" fill="#f8f9fa" rx="8"/>
  <!-- main line -->
  <line x1="30" y1="100" x2="80" y2="100" stroke="#333" stroke-width="2"/>
  <path d="M80,100 Q90,82 100,100 Q110,82 120,100 Q130,82 140,100 Q150,82 160,100" fill="none" stroke="#d63384" stroke-width="2.5"/>
  <text x="105" y="72" font-size="14" font-weight="bold" fill="#d63384">L₁</text>
  <line x1="160" y1="100" x2="200" y2="100" stroke="#333" stroke-width="2"/>
  <path d="M200,100 Q210,82 220,100 Q230,82 240,100 Q250,82 260,100" fill="none" stroke="#d63384" stroke-width="2.5"/>
  <text x="215" y="72" font-size="14" font-weight="bold" fill="#d63384">L₂</text>
  <line x1="260" y1="100" x2="290" y2="100" stroke="#333" stroke-width="2"/>
  <!-- filter connection point -->
  <circle cx="290" cy="100" r="3" fill="#333"/>
  <line x1="290" y1="100" x2="360" y2="100" stroke="#333" stroke-width="2"/>
  <!-- Grid circle (shifted right) -->
  <line x1="360" y1="100" x2="400" y2="100" stroke="#333" stroke-width="2"/>
  <line x1="400" y1="100" x2="400" y2="112" stroke="#333" stroke-width="2"/>
  <circle cx="400" cy="140" r="26" fill="none" stroke="#198754" stroke-width="2"/>
  <text x="385" y="144" font-size="12" font-weight="bold" fill="#198754">Grid</text>
  <line x1="400" y1="166" x2="400" y2="240" stroke="#333" stroke-width="2"/>
  <!-- LCL Cf branch -->
  <line x1="180" y1="100" x2="180" y2="135" stroke="#333" stroke-width="2"/>
  <line x1="165" y1="135" x2="195" y2="135" stroke="#333" stroke-width="2"/>
  <line x1="165" y1="144" x2="195" y2="144" stroke="#333" stroke-width="2"/>
  <line x1="180" y1="144" x2="180" y2="165" stroke="#333" stroke-width="2"/>
  <line x1="170" y1="165" x2="190" y2="165" stroke="#333" stroke-width="2"/>
  <text x="200" y="142" font-size="12" font-weight="bold" fill="#0d6efd">Cf</text>
  <!-- bottom return line -->
  <line x1="30" y1="240" x2="400" y2="240" stroke="#333" stroke-width="2"/>
  <!-- inverter source -->
  <rect x="12" y="78" width="18" height="160" fill="#fce4ec" stroke="#d63384" stroke-width="1.5" rx="2"/>
  <text x="12" y="170" font-size="9" fill="#d63384" writing-mode="tb">Vinv</text>
  <!-- ===== Passive filter branch (orange, at x=290) ===== -->
  <rect x="248" y="42" width="88" height="22" fill="#fff3cd" stroke="#fd7e14" rx="4"/>
  <text x="253" y="57" font-size="11" font-weight="bold" fill="#fd7e14">Passive Filter</text>
  <!-- Lf inductor -->
  <line x1="290" y1="100" x2="290" y2="120" stroke="#fd7e14" stroke-width="2.5"/>
  <path d="M278,120 Q283,108 290,120 Q297,108 302,120" fill="none" stroke="#fd7e14" stroke-width="2.5"/>
  <text x="305" y="120" font-size="11" font-weight="bold" fill="#fd7e14">Lf</text>
 <!-- Rf -->
  <line x1="290" y1="126" x2="290" y2="135" stroke="#fd7e14" stroke-width="2.5"/>
  <rect x="280" y="135" width="20" height="12" fill="none" stroke="#fd7e14" stroke-width="2" rx="1"/>
  <text x="305" y="146" font-size="11" font-weight="bold" fill="#fd7e14">Rf</text>
  <!-- Cf_f capacitor -->
  <line x1="290" y1="147" x2="290" y2="170" stroke="#fd7e14" stroke-width="2.5"/>
  <line x1="277" y1="170" x2="303" y2="170" stroke="#fd7e14" stroke-width="2.5"/>
  <line x1="277" y1="178" x2="303" y2="178" stroke="#fd7e14" stroke-width="2.5"/>
  <text x="305" y="178" font-size="11" font-weight="bold" fill="#fd7e14">Cf_f</text>
  <!-- filter to ground -->
  <line x1="290" y1="178" x2="290" y2="240" stroke="#fd7e14" stroke-width="2.5"/>
  <!-- harmonic current annotation -->
  <text x="255" y="218" font-size="10" fill="#fd7e14">← Ih harmonic current</text>
</svg>"""
            st.markdown(svg_tf, unsafe_allow_html=True)
            st.caption("Orange branch = passive filter (in parallel between LCL and grid)")

        st.divider()
        st.markdown("### 🔍 Passive Filter vs. Control Algorithm: Which to Choose?")
        st.markdown("""
| Comparison | Passive Filter (hardware) | Control Algorithm (PR/repetitive control) |
|------------|--------------------------|------------------------------------------|
| **Response speed** | Instantaneous (passive) | Limited by control bandwidth |
| **Coverage** | Fixed harmonic orders | Usually ≤ 25th order |
| **Cost** | Hardware required (LC components) | Software only, lower cost |
| **Tuning accuracy** | Affected by component parameter drift | Real-time adaptive adjustment |
| **Reactive compensation** | Yes (provides reactive compensation as by-product) | No |
| **Recommended when** | Higher harmonics exceed limit, or control bandwidth insufficient | Low-order harmonics (≤13th), dynamic harmonic compensation |

> 📚 IEEE Std 1531-2003, §4.1: In engineering practice, both are often used together —
> control algorithms handle low-order harmonics (≤13th), passive filters address higher-order harmonics.
        """)

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 2 — Filter design + impedance characteristic (anti-resonance)
    # ══════════════════════════════════════════════════════════════════════════
    with tab_design:
        st.markdown("### 🔧 Single-Tuned Passive Filter — Parameter Design")
        st.info("""
**Design workflow** (per IEEE Std 1531-2003, §5):
1. Confirm harmonic orders exceeding limit from Section 7 → 2. Determine filter rated reactive power → 3. Calculate Lf / Cf
→ 4. Select Quality Factor Q → 5. Verify tuning frequency and impedance characteristic → **6. Check system Parallel Anti-Resonance**
        """)

        # ── Step 1: Select target harmonic orders ─────────────────────────────────────────────────
        st.markdown("#### Step 1: Select Target Harmonic Orders")
        col_s1a, col_s1b = st.columns([1, 1])
        with col_s1a:
            n_filters = st.number_input("Number of filter branches to design",
                                         min_value=1, max_value=5, value=2, step=1,
                                         key="s8_nf")
        with col_s1b:
            st.markdown("""
**Typical exceeding harmonics (reference)**:
- 5th (250 Hz): primary product of dead-time effect
- 7th (350 Hz): secondary product of dead-time effect
- 11th (550 Hz): control non-linearity
- 13th (650 Hz): control non-linearity
            """)
            with st.expander("💡 Why Does Dead Time Generate Low-Order Harmonics?", expanded=False):
                st.markdown("""
**Dead Time** is a delay inserted between upper/lower bridge arm switching to prevent shoot-through short-circuit (typical 2~5 μs).

During dead time, the output voltage is **uncontrolled** — the actual output deviates from the ideal sinusoidal command.
This deviation is a **rectangular pulse sequence** whose Fourier decomposition contains mainly **low-order odd harmonics**:

$$\\Delta V_{dead} \\approx \\frac{4 \\cdot t_d \\cdot f_{sw} \\cdot V_{dc}}{\\pi} \\sum_{n=1,3,5,...} \\frac{1}{n} \\sin(n\\omega_1 t)$$

where $t_d$ is dead time and $f_{sw}$ is switching frequency.

Due to the $1/n$ attenuation characteristic, **the 5th harmonic is largest, 7th is next** — this is the root cause of 5th and 7th harmonic limit violations.

> 📚 Mohan N., *Power Electronics*, 6th ed., Wiley, 2014, §8.3;
> Teodorescu R. et al., *Grid Converters for PV and Wind*, Wiley, 2011, §10.2
                """)

        filter_targets = []
        col_inputs = st.columns(min(n_filters, 3))
        for i in range(n_filters):
            col = col_inputs[i % len(col_inputs)]
            with col:
                default_orders = [5, 7, 11, 13, 3]
                h_n = st.number_input(
                    f"Filter #{i+1}: Target Harmonic Order",
                    min_value=2, max_value=50,
                    value=default_orders[i] if i < len(default_orders) else 5 + i*2,
                    step=1, key=f"s8_h_{i}"
                )
                filter_targets.append(int(h_n))

        st.divider()

        # ── Step 2: Filter capacity ───────────────────────────────────────────────────
        st.markdown("#### Step 2: Determine Filter Rated Reactive Power $Q_f$")
        st.markdown("""
Engineering recommendation: $Q_f$ ≈ **5%~20%** of system rated capacity
(Das J.C., *Power System Harmonics*, Wiley-IEEE, 2015, §6.3).
        """)
        col_s2a, col_s2b, col_s2c = st.columns(3)
        with col_s2a:
            st.metric("System Rated Power", f"{_Prated/1e3:.0f} kW")
        with col_s2b:
            Qf_pct = st.slider("Filter reactive capacity as % of system rated power (%)",
                                1.0, 30.0, 10.0, 0.5, key="s8_qf")
        Qf_total = _Prated * Qf_pct / 100.0
        Qf_per_branch = Qf_total / n_filters
        with col_s2c:
            if is_trial:
                trial_lock("Unlock with paid access: filter rated reactive power calculation", compact=True)
            else:
                st.metric("Per-Branch Capacity Qf", f"{Qf_per_branch/1e3:.1f} kVar",
                          delta=f"Total capacity {Qf_total/1e3:.1f} kVar")

        st.divider()

        # ── Step 3: Parameter calculation ──────────────────────────────────────────────────────
        omega1 = 2 * np.pi * _fg
        Ug_phase = _Ug / np.sqrt(3)
        st.markdown("#### Step 3: Calculate Lf / Cf Parameters")
        if is_trial:
            trial_lock("Unlock with paid access: branch Lf/Cf parameter design results", compact=True)
        else:
            st.latex(r"""
C_f = \frac{Q_f}{U_g^2 \cdot \omega_1} \qquad
L_f = \frac{1}{(h_n \cdot \omega_1)^2 \cdot C_f}
            """)
            st.caption("📚 Das J.C., Power System Harmonics and Passive Filter Designs, Wiley-IEEE Press, 2015, §6.3")

        st.markdown("#### Step 4: Select Quality Factor Q")
        if is_trial:
            trial_lock("Unlock with paid access: Quality Factor Q selection guide and engineering recommendations", compact=True)
        else:
            st.markdown("""
**Quality Factor Q** determines filter "sharpness" (selectivity). Definition:

$$Q = \\frac{\\omega_n L_f}{R_f} = \\frac{1}{R_f} \\sqrt{\\frac{L_f}{C_f}}$$

Higher Q = sharper impedance notch at tuning point (better filtering), but narrower bandwidth (more sensitive to parameter tolerance).
            """)

            col_q1, col_q2 = st.columns([1, 1])
            with col_q1:
                st.markdown("""
| Q Range | Characteristic | Recommended Use |
|---------|---------------|-----------------|
| **20~30** | Wide bandwidth, robust | Parameter uncertainty, large temperature variation |
| **30~60** | Balanced (most common) | Standard grid-connected filter design |
| **60~100** | Narrow bandwidth, high precision | Sharp tuning target, high component accuracy |
| **>100** | Extremely narrow | Laboratory only; impractical in engineering |

> 📚 IEEE Std 1531-2003, §5.3; Das J.C., 2015, §6.4
                """)
            with col_q2:
                st.warning("""
**Engineering selection guide**:
- **Recommended Q = 30~60**: optimal balance for grid-connected filter design
- With capacitor parameter tolerance (±5%~±10%), choose **lower Q** (larger Rf) for wider bandwidth
- If anti-resonance risk detected (Step 6), choose **lower Q** to suppress peak amplitude
- Very high Q (>100) is impractical due to component aging and thermal drift — **avoid in engineering design**
                """)

        Q_factor = st.slider("Quality Factor Q", 20, 120, 50, 5, key="s8_Q",
                              help="Q = ωn·Lf/Rf. Industrial range: 30~80. Start at 50 and adjust based on anti-resonance analysis results.")

        filter_params = []
        for i, h_n in enumerate(filter_targets):
            omega_n  = h_n * omega1
            f_n      = h_n * _fg
            Cf_f     = Qf_per_branch / (Ug_phase**2 * omega1)
            Lf       = 1.0 / (omega_n**2 * Cf_f)
            Rf       = omega_n * Lf / Q_factor
            f_tuned  = 1 / (2 * np.pi * np.sqrt(Lf * Cf_f))
            Z_at_tuned = abs(Rf + 1j * (omega_n * Lf - 1.0 / (omega_n * Cf_f)))
            Z_at_fg    = abs(Rf + 1j * (omega1 * Lf  - 1.0 / (omega1 * Cf_f)))
            Iq_fund    = Ug_phase / Z_at_fg
            filter_params.append({
                'branch': i + 1, 'h_n': h_n, 'f_n': f_n,
                'Cf_f_uF': Cf_f * 1e6, 'Lf_mH': Lf * 1e3, 'Rf_Ohm': Rf,
                'Qf_kVar': Qf_per_branch / 1e3, 'f_tuned': f_tuned,
                'Z_at_tuned': Z_at_tuned, 'Z_at_fg': Z_at_fg, 'Iq_fund': Iq_fund,
                'Q_factor': Q_factor, 'Cf_f': Cf_f, 'Lf': Lf, 'Rf': Rf,
            })

        df_params = pd.DataFrame([{
            'Branch': f"#{p['branch']}", 'Target Order': p['h_n'],
            'Target Frequency (Hz)': f"{p['f_n']:.0f}",
            'Cf (μF)': f"{p['Cf_f_uF']:.2f}", 'Lf (mH)': f"{p['Lf_mH']:.3f}",
            'Rf (Ω)': f"{p['Rf_Ohm']:.4f}", 'Actual Tuned Frequency (Hz)': f"{p['f_tuned']:.1f}",
            'Impedance at Tuned Freq (Ω)': f"{p['Z_at_tuned']:.4f}",
            'Capacity Qf (kVar)': f"{p['Qf_kVar']:.1f}"
        } for p in filter_params])
        st.dataframe(df_params, width="stretch", hide_index=True)

        st.divider()

        # ── Step 5: Impedance characteristic + Parallel Anti-Resonance ───────────────────────
        st.markdown("#### Step 5: Impedance-Frequency Characteristic (System Parallel Anti-Resonance Analysis)")

        # Grid impedance parameters (for Zg calculation)
        with st.expander("⚙️ Grid Impedance Parameters (expand to configure — required for anti-resonance analysis)", expanded=True):
            col_zg1, col_zg2, col_zg3 = st.columns(3)
            with col_zg1:
                zg_ssc = st.number_input("PCC Short-Circuit Capacity Ssc (MVA)",
                                          value=float(st.session_state.get('h7_scc', 10.0)),
                                          min_value=0.1, step=1.0, key="s8z_ssc",
                                          help="Determines grid inductor Lg = Ug²/(2πf·Ssc)")
            with col_zg2:
                zg_xr = st.number_input("Grid X/R Ratio",
                                         value=5.0, min_value=0.1, max_value=50.0,
                                         step=0.5, key="s8z_xr",
                                         help="Strong grid X/R=5~10; use lower value for weak grid")
            with col_zg3:
                zg_ug = st.number_input("Grid Phase Voltage Ug (Vrms)",
                                         value=float(_Ug / np.sqrt(3)),
                                         step=10.0, key="s8z_ug")

        # Compute grid complex impedance parameters
        # Zg = Rg + jωLg, where Lg = Ug²/(ω₁·Ssc), Rg = Xg/XR_ratio
        Lg_grid  = zg_ug**2 / (omega1 * zg_ssc * 1e6) if zg_ssc > 0 else 1e-3
        Xg_fund  = omega1 * Lg_grid
        Rg_grid  = Xg_fund / zg_xr if zg_xr > 0 else Xg_fund / 5.0

        # Frequency scan range
        f_scan    = np.linspace(10, min(_fsw * 1.2, 5000), 8000)
        omega_scan = 2 * np.pi * f_scan

        # Total filter impedance (complex, all branches in parallel)
        Zf_total_cmplx = np.zeros(len(f_scan), dtype=complex)
        for p in filter_params:
            Zf_branch = (p['Rf'] + 1j * (omega_scan * p['Lf'] - 1.0 / (omega_scan * p['Cf_f'])))
            # Parallel combination (multi-branch)
            if np.all(Zf_total_cmplx == 0):
                Zf_total_cmplx = Zf_branch
            else:
                # Two-impedance parallel: Z_parallel = Z1*Z2 / (Z1+Z2)
                denom = Zf_total_cmplx + Zf_branch
                denom = np.where(np.abs(denom) < 1e-12, 1e-12, denom)
                Zf_total_cmplx = Zf_total_cmplx * Zf_branch / denom

        # Grid impedance (complex inductive)
        Zg_cmplx = Rg_grid + 1j * omega_scan * Lg_grid

        # [Add1] System parallel total impedance Zsys = Zf * Zg / (Zf + Zg)
        denom_sys = Zf_total_cmplx + Zg_cmplx
        denom_sys = np.where(np.abs(denom_sys) < 1e-15, 1e-15, denom_sys)
        Zsys_cmplx = Zf_total_cmplx * Zg_cmplx / denom_sys
        Zsys_abs   = np.abs(Zsys_cmplx)
        Zf_abs     = np.abs(Zf_total_cmplx)
        Zg_abs     = np.abs(Zg_cmplx)

        # Detect anti-resonance peaks (local maxima of Zsys)
        zsys_smooth = Zsys_abs.copy()
        # Identify peaks: local maxima wider than window
        peak_candidates = []
        window = 50
        for i in range(window, len(zsys_smooth) - window):
            if zsys_smooth[i] == zsys_smooth[i-window:i+window].max():
                if zsys_smooth[i] > Rg_grid * 3:  # peak significantly above grid resistance
                    peak_candidates.append((f_scan[i], zsys_smooth[i]))

        # Sort by amplitude, keep top 3
        peak_candidates.sort(key=lambda x: x[1], reverse=True)
        top_peaks = peak_candidates[:3]

        # Build impedance chart
        fig_z = go.Figure()
        colors_branch = ['#fd7e14', '#0d6efd', '#198754', '#dc3545', '#6f42c1']

        # Per-branch filter impedance
        for i, p in enumerate(filter_params):
            Z_arr = np.abs(p['Rf'] + 1j * (omega_scan * p['Lf'] - 1.0 / (omega_scan * p['Cf_f'])))
            fig_z.add_trace(go.Scatter(
                x=f_scan, y=Z_arr,
                name=f"Zf_{p['branch']}({p['h_n']})",
                line=dict(color=colors_branch[i % len(colors_branch)], width=2, dash='dot'),
                opacity=0.7
            ))
            fig_z.add_vline(x=p['f_tuned'], line_dash="dot",
                             line_color=colors_branch[i % len(colors_branch)],
                             annotation_text=f"{p['h_n']}")

        # Grid impedance reference line
        fig_z.add_trace(go.Scatter(
            x=f_scan, y=Zg_abs,
            name=f"Zg (grid impedance, Lg={Lg_grid*1000:.2f}mH, Rg={Rg_grid:.3f}Ω)",
            line=dict(color='#6c757d', width=1.5, dash='dash'),
            opacity=0.8
        ))

        # [Add1] System parallel total impedance Zsys (key trace)
        fig_z.add_trace(go.Scatter(
            x=f_scan, y=Zsys_abs,
            name="Zsys (system parallel impedance, seen by harmonics)",
            line=dict(color='#dc3545', width=3),
            fill='tozeroy', fillcolor='rgba(220,53,69,0.05)'
        ))

        # Mark fundamental frequency
        fig_z.add_vline(x=_fg, line_dash="dot", line_color="#6c757d",
                         annotation_text=f"Fund. {_fg:.0f}Hz")

        # Mark anti-resonance peaks on Zsys curve with prominent markers
        if top_peaks:
            peak_freqs = [f_pk for f_pk, _ in top_peaks]
            peak_zvals = [z_pk for _, z_pk in top_peaks]
            peak_labels = [f"AR{i+1}: {f_pk:.0f}Hz ({f_pk/_fg:.1f}×)" for i, (f_pk, _) in enumerate(top_peaks)]

            # Overlay prominent red triangle markers on Zsys curve
            fig_z.add_trace(go.Scatter(
                x=peak_freqs, y=peak_zvals,
                mode='markers+text',
                name='⚠ Anti-Resonance Peaks (see below)',
                marker=dict(symbol='triangle-up', size=14, color='#dc3545',
                             line=dict(width=2, color='white')),
                text=[f"AR{i+1}" for i in range(len(top_peaks))],
                textposition='top center',
                textfont=dict(size=11, color='#dc3545', family='Arial Black'),
                hovertext=[f"⚠ Anti-Resonance AR{i+1}<br>Frequency: {f_pk:.0f} Hz ({f_pk/_fg:.1f}× fund.)<br>|Zsys| = {z_pk:.3f} Ω"
                           for i, (f_pk, z_pk) in enumerate(top_peaks)],
                hoverinfo='text',
            ))

        fig_z.update_layout(
            title="Filter Impedance Characteristic (System Parallel Anti-Resonance Analysis)",
            xaxis_title="Frequency (Hz)",
            yaxis_title="|Z| (Ω)",
            yaxis_type="log",
            height=520,
            template="plotly_white",
            legend=dict(
                yanchor="top", y=-0.12, xanchor="center", x=0.5,
                orientation="h", font=dict(size=9),
                bgcolor="rgba(255,255,255,0.8)"
            ),
            margin=dict(b=100)
        )
        # Mode C: show chart to all users; disable hover for trial users
        for trace in fig_z.data:
            if is_trial:
                trace.hoverinfo = 'none'
        st.plotly_chart(fig_z, use_container_width=True)

        # ── Anti-resonance peak details and warning — locked for trial users ────────────────────────
        if is_trial:
            trial_lock("Unlock with paid access: anti-resonance peak analysis, system impedance warning, and chart legend", compact=True)
        else:
            if top_peaks:
                st.markdown("**📍 Anti-Resonance Peak Details** (marked by ▲ in chart above):")
                ar_cols = st.columns(min(len(top_peaks), 3))
                for idx_pk, (f_pk, z_pk) in enumerate(top_peaks):
                    with ar_cols[idx_pk % len(ar_cols)]:
                        h_pk = f_pk / _fg
                        risk = "🔴 High Risk" if h_pk < 10 else "🟡 Medium Risk"
                        st.error(f"""
**AR{idx_pk+1}** — {f_pk:.0f} Hz (harmonic order {h_pk:.1f}×)

|Zsys| = {z_pk:.3f} Ω　　{risk}
                        """)

            if top_peaks:
                st.error(f"""
### ⚠️ System Parallel Anti-Resonance Warning!

**Anti-resonance mechanism**: The filter capacitor forms a parallel LC resonance with the grid inductor $L_g$.
At the resonance frequency, the system impedance $Z_{{sys}}$ spikes sharply.
Harmonic current encounters high impedance and **flows entirely into the grid — amplified**, which may cause:
- 🔴 **PCC voltage harmonic exceeding limit** (overvoltage stress on capacitor)
- 🔴 **Filter capacitor overvoltage** (risk of capacitor failure in the field)

**Detected anti-resonance peaks**:

| Resonance Frequency | Harmonic Order | System Impedance | Risk |
|-----------|------------|------------|--------|
{chr(10).join([f"| **{f_pk:.0f} Hz** | {f_pk/_fg:.1f}× | {z_pk:.3f} Ω | {'🔴 High risk (close to low-order harmonics)' if f_pk/_fg < 10 else '🟡 Medium risk (monitor background harmonics)'} |" for f_pk, z_pk in top_peaks])}

**Engineering countermeasures** (Das J.C., Wiley-IEEE 2015, §8.5):

1. **Adjust grid short-circuit capacity Ssc** or filter capacity to shift resonance away from harmonic frequencies (3rd, 5th, etc.)
2. **Use a high-pass filter (HPF) branch** instead of low-order tuned branches — HPF avoids parallel resonance
3. **Reduce Q factor** (increase Rf) to widen bandwidth and lower resonance amplitude
4. **Add series damping resistor** $R_d$ in the filter branch to directly suppress resonance

> 📚 Das J.C., *Power System Harmonics and Passive Filter Designs*,
> Wiley-IEEE Press, 2015, §8.5 "Parallel Resonance"
                """)
            else:
                st.success("✅ With the current grid parameters, no Parallel Anti-Resonance peaks detected (|Zsys| remains smooth).")

            st.markdown("""
**Chart legend (3 curves)**:
- **Orange/colored dotted lines**: per-branch filter impedance $|Z_f(f)|$, drops to ≈0 at tuning frequency (series resonance)
- **Gray dashed line**: grid impedance $|Z_g(f)| = |R_g + j\\omega L_g|$, increases linearly with frequency
- **Red solid line**: system parallel impedance $|Z_{sys}(f)|$ — the actual impedance seen by harmonics.
  Peaks (anti-resonance) = danger signal; valleys = filter effectively diverting harmonic current
            """)

        # ── Step 5: Design engineering verification ─────────────────────────────────────────────────────
        st.divider()
        st.markdown("#### ⚠️ Design Engineering Verification")
        if is_trial:
            trial_lock("Filter design verification results require paid access")
        else:
            for p in filter_params:
                freq_err_pct = abs(p['f_tuned'] - p['f_n']) / p['f_n'] * 100
                dist_to_lcl  = abs(p['f_n'] - f_res_lcl)
                S_fund_kvar  = Ug_phase * p['Iq_fund'] * 3 / 1e3

                # Tuning Frequency(!)
                near_antires = any(abs(p['f_n'] - f_pk) < 50 for f_pk, _ in top_peaks)

                with st.expander(f"Filter Branch #{p['branch']} (H{p['h_n']}, {p['f_n']:.0f} Hz) Design Verification",
                                  expanded=True):
                    col_w1, col_w2, col_w3, col_w4 = st.columns(4)
                    with col_w1:
                        ok1 = freq_err_pct < 0.1
                        st.metric("Tuning Frequency Error", f"{freq_err_pct:.3f}%",
                                  delta="✅ OK" if ok1 else "⚠ Check",
                                  delta_color="normal" if ok1 else "inverse")
                    with col_w2:
                        ok2 = dist_to_lcl > 50
                        st.metric("Distance to LCL Resonance", f"{dist_to_lcl:.0f} Hz",
                                  delta="✅ Safe" if ok2 else f"⚠ Too close to LCL resonance at {f_res_lcl:.0f} Hz!",
                                  delta_color="normal" if ok2 else "inverse")
                    with col_w3:
                        ok3 = S_fund_kvar < _Prated * 0.15 / 1e3
                        st.metric("Injected Reactive Power", f"{S_fund_kvar:.1f} kVar",
                                  delta="✅ Within limit" if ok3 else "⚠ Excess reactive power",
                                  delta_color="normal" if ok3 else "inverse")
                    with col_w4:
                        if near_antires:
                            st.metric("Anti-Resonance Risk", "🔴 Warning",
                                      delta="Resonance peak within 50 Hz of tuning freq",
                                      delta_color="inverse")
                        else:
                            st.metric("Anti-Resonance Risk", "✅ Safe",
                                      delta="No nearby anti-resonance",
                                      delta_color="normal")

                st.caption(f"""
Tuned Frequency: {p['f_tuned']:.2f} Hz (target {p['f_n']:.0f} Hz) |
Impedance at tuning point: {p['Z_at_tuned']:.4f} Ω (≈ Rf = {p['Rf_Ohm']:.4f} Ω) |
Quality Factor Q = {p['Q_factor']}
                """)

        # ── [Add3] Anti-resonance sensitivity analysis (Robustness Test) ──────────────────────────
        # Grid short-circuit capacity Ssc varies with operating conditions.
        # Lg changes as Ssc changes, shifting the anti-resonance frequency.
        # Reference: Das J.C., Power System Harmonics, Wiley-IEEE, 2015, §8.5;
        #            IEEE Std 1531-2003, §5.7 "Resonance with System Impedance"
        # ─────────────────────────────────────────────────────────────────────
        st.divider()
        st.markdown("#### 🔬 Step 6: Anti-Resonance Sensitivity Analysis (Ssc Variation Robustness)")
        st.info("""
**Engineering background**: Grid short-circuit capacity $S_{sc}$ is not constant — it changes with substation
operating configurations, upstream line maintenance, and load variations. Changes in $S_{sc}$ directly alter
the grid inductance $L_g = U_g^2 / (\\omega_1 \\cdot S_{sc})$, which **shifts the anti-resonance peak frequency**.

This analysis shows how the anti-resonance peak frequency shifts as $S_{sc}$ varies, helping engineers assess
whether the design remains safe under worst-case grid conditions.

> 📚 Das J.C., *Power System Harmonics*, Wiley-IEEE, 2015, §8.5;
> IEEE Std 1531-2003, §5.7
        """)

        col_sens1, col_sens2 = st.columns([1, 1])
        with col_sens1:
            ssc_variation = st.slider(
                "Ssc Variation Range (%)",
                min_value=5, max_value=50, value=20, step=5,
                key="s8_ssc_var",
                help="Grid short-circuit capacity variation range. Strong grid: ±20%, Weak grid: ±40%"
            )
        with col_sens2:
            ssc_steps = st.slider(
                "Number of Ssc Sweep Steps",
                min_value=3, max_value=11, value=7, step=2,
                key="s8_ssc_steps",
                help="Number of Ssc sampling points across the variation range"
            )

        # Build Ssc sweep array
        ssc_ratios = np.linspace(1 - ssc_variation/100, 1 + ssc_variation/100, ssc_steps)
        ssc_values = zg_ssc * ssc_ratios  # MVA

        fig_sens = go.Figure()

        # Mark dangerous background harmonic frequencies (3/5/7/11/13)
        danger_harmonics = [3, 5, 7, 11, 13]
        for idx_d, h_d in enumerate(danger_harmonics):
            f_danger = h_d * _fg
            if f_danger < f_scan[-1]:
                # Vertical dashed line (low opacity)
                fig_sens.add_vline(
                    x=f_danger, line_dash="dot", line_color="rgba(108,117,125,0.3)",
                )
                # Stagger annotation y-offset to avoid overlap
                fig_sens.add_annotation(
                    x=f_danger, y=0, yref="y domain",
                    yshift=10 + (idx_d % 2) * 14,
                    text=f"{h_d}({f_danger:.0f}Hz)",
                    showarrow=False,
                    font=dict(size=8, color="#6c757d"),
                    bgcolor="rgba(255,255,255,0.7)",
                )

        # Color gradient: red (low Ssc/weak grid) → blue (high Ssc/strong grid)
        colorscale_sens = []
        for idx_s, ssc_ratio in enumerate(ssc_ratios):
            ratio_norm = (ssc_ratio - ssc_ratios[0]) / (ssc_ratios[-1] - ssc_ratios[0]) if len(ssc_ratios) > 1 else 0.5
            r_val = int(220 * (1 - ratio_norm) + 13 * ratio_norm)
            g_val = int(53 * (1 - ratio_norm) + 110 * ratio_norm)
            b_val = int(69 * (1 - ratio_norm) + 253 * ratio_norm)
            colorscale_sens.append(f"rgb({r_val},{g_val},{b_val})")

        # Track anti-resonance peak across Ssc sweep
        antires_track = []  # [(ssc_mva, f_peak, z_peak), ...]

        for idx_s, ssc_val in enumerate(ssc_values):
            # Compute grid impedance for this Ssc value
            Lg_s = zg_ug**2 / (omega1 * ssc_val * 1e6) if ssc_val > 0 else 1e-3
            Rg_s = (omega1 * Lg_s) / zg_xr if zg_xr > 0 else (omega1 * Lg_s) / 5.0
            Zg_s = Rg_s + 1j * omega_scan * Lg_s

            # Compute system parallel impedance
            denom_s = Zf_total_cmplx + Zg_s
            denom_s = np.where(np.abs(denom_s) < 1e-15, 1e-15, denom_s)
            Zsys_s = np.abs(Zf_total_cmplx * Zg_s / denom_s)

            is_nominal = abs(ssc_val - zg_ssc) < 0.01
            line_w = 3 if is_nominal else 1.5
            line_dash = 'solid' if is_nominal else 'dot'
            opacity_s = 1.0 if is_nominal else 0.6

            fig_sens.add_trace(go.Scatter(
                x=f_scan, y=Zsys_s,
                name=f"Ssc={ssc_val:.1f}MVA ({ssc_ratio*100-100:+.0f}%)" + (" ← nominal" if is_nominal else ""),
                line=dict(color=colorscale_sens[idx_s], width=line_w, dash=line_dash),
                opacity=opacity_s,
                showlegend=True
            ))

            # Find anti-resonance peaks for this Ssc value
            for i_pk in range(window, len(Zsys_s) - window):
                if Zsys_s[i_pk] == Zsys_s[i_pk-window:i_pk+window].max():
                    if Zsys_s[i_pk] > Rg_s * 3:
                        antires_track.append((ssc_val, f_scan[i_pk], Zsys_s[i_pk]))

        fig_sens.update_layout(
            title=f"Anti-Resonance Sensitivity Analysis: |Zsys| vs Frequency — Ssc {zg_ssc*(1-ssc_variation/100):.1f}~{zg_ssc*(1+ssc_variation/100):.1f} MVA",
            xaxis_title="Frequency (Hz)",
            yaxis_title="|Zsys| (Ω)",
            yaxis_type="log",
            height=550,
            template="plotly_white",
            legend=dict(
                yanchor="top", y=-0.15, xanchor="center", x=0.5,
                orientation="h", font=dict(size=9),
                bgcolor="rgba(255,255,255,0.8)"
            ),
            margin=dict(b=120)
        )
        if is_trial:
            fig_sens.update_layout(hovermode=False)
            fig_sens.update_yaxes(showticklabels=False)
            st.plotly_chart(fig_sens, use_container_width=True)
            trial_lock("Unlock with paid access: anti-resonance sensitivity analysis (including peak trajectory table and robustness conclusions)", compact=True)
        else:
            st.plotly_chart(fig_sens, use_container_width=True)

        # ── Anti-resonance peak tracking results ──────────────────────────────────────────────
        if not is_trial and antires_track:
            df_track = pd.DataFrame(antires_track, columns=['Ssc (MVA)', 'Anti-Resonance Frequency (Hz)', '|Zsys| (Ω)'])
            df_track['Harmonic Order'] = df_track['Anti-Resonance Frequency (Hz)'] / _fg
            df_track['Harmonic Order'] = df_track['Harmonic Order'].round(1)

            # Check if resonance peak sweeps through dangerous harmonic frequencies
            swept_danger = set()
            f_ar_min = df_track['Anti-Resonance Frequency (Hz)'].min()
            f_ar_max = df_track['Anti-Resonance Frequency (Hz)'].max()
            for h_d in danger_harmonics:
                f_d = h_d * _fg
                if f_ar_min <= f_d <= f_ar_max:
                    swept_danger.add(h_d)

            col_ar1, col_ar2, col_ar3 = st.columns(3)
            with col_ar1:
                st.metric("Anti-Resonance Frequency Range",
                           f"{f_ar_min:.0f} ~ {f_ar_max:.0f} Hz",
                           delta=f"Span: {f_ar_max - f_ar_min:.0f} Hz")
            with col_ar2:
                st.metric("Harmonic Order Range",
                           f"{f_ar_min/_fg:.1f} ~ {f_ar_max/_fg:.1f}×")
            with col_ar3:
                if swept_danger:
                    danger_str = ', '.join([f"{h}th" for h in sorted(swept_danger)])
                    st.metric("⚠ Swept Harmonic Orders", danger_str,
                               delta="Anti-resonance sweeps through these harmonic orders!",
                               delta_color="inverse")
                else:
                    st.metric("Swept Harmonic Orders", "✅ Safe",
                               delta="No dangerous harmonics in resonance sweep",
                               delta_color="normal")

            if swept_danger:
                st.error(f"""
### 🔴 Anti-Resonance Sweeps Through Dangerous Harmonic Frequencies!

Within $S_{{sc}}$ ±{ssc_variation}% variation, the anti-resonance frequency shifts between **{f_ar_min:.0f} Hz and {f_ar_max:.0f} Hz**,
passing through **{', '.join([str(h)+'th ('+str(int(h*_fg))+' Hz)' for h in sorted(swept_danger)])} harmonic frequencies**.

This means that under different grid operating conditions (transformer configurations, maintenance states, etc.),
the anti-resonance peak may coincide with background harmonic frequencies, amplifying those harmonics and
potentially causing **capacitor overvoltage** (Das J.C., 2015, §8.5).

**Recommended actions**:
1. Adjust filter capacity Qf to shift resonance away from {', '.join([str(h)+'th' for h in sorted(swept_danger)])} across the Ssc variation range
2. Reduce Q factor (add damping Rf) to lower resonance amplitude
3. Use high-pass filter (HPF) branch to replace the lowest-order tuned branch
4. Add series transformer impedance to narrow the Ssc variation range
                """)
            else:
                st.success(f"""
✅ **Robustness confirmed**: Within $S_{{sc}}$ ±{ssc_variation}% variation,
no anti-resonance peaks coincide with standard harmonic frequencies (3/5/7/11/13). Design is safe.
                """)

            with st.expander("📊 Anti-Resonance Sensitivity Data Table", expanded=False):
                st.dataframe(df_track.style.format({
                    'Ssc (MVA)': '{:.1f}',
                    'Anti-Resonance Frequency (Hz)': '{:.1f}',
                    '|Zsys| (Ω)': '{:.4f}',
                    'Harmonic Order': '{:.1f}'
                }), use_container_width=True, hide_index=True)
        elif not is_trial:
            st.success("No Parallel Anti-Resonance peaks detected across the Ssc variation range.")

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 3 — Harmonic verification (before vs after filter)
    # ══════════════════════════════════════════════════════════════════════════
    with tab_verify:
        st.markdown("### ✅ Passive Filter — Harmonic Verification")

        # ── [Fix1] Current divider using complex impedance ────────────────────────────────
        st.markdown("**Calculation principle** (impedance divider method):")
        st.latex(r"I_{grid,h} = I_{source,h} \cdot \frac{|Z_{filter,h}|}{|Z_{filter,h} + Z_{grid,h}|}")
        st.markdown("Note: $Z_{filter}$ and $Z_{grid}$ are **complex impedances** (with phase information), not scalar magnitudes. Formulas:")
        st.latex(r"""
Z_{filter,h} = R_f + j\left(\omega_h L_f - \frac{1}{\omega_h C_f}\right)
\qquad
Z_{grid,h} = R_g + j\omega_h L_g
        """)
        st.error("⚠️ Do NOT use scalar addition (`|Zf| + |Zg|`). This ignores phase and misses parallel resonance peaks!")

        st.divider()

        # ── Parameter ──────────────────────────────────────────────────────
        with st.expander("⚙️ System Parameters (pre-filled from Section 7 — expand to modify)", expanded=False):
            col_v1, col_v2, col_v3 = st.columns(3)
            with col_v1:
                v_ug = st.number_input("Grid voltage Ug (Vrms)", value=float(_Ug/np.sqrt(3)),
                                       step=10.0, key="s8v_ug")
                v_il = st.number_input("Rated current IL (A)", value=float(_IL),
                                       step=10.0, key="s8v_il")
                v_fg = st.number_input("Frequency (Hz)", value=float(_fg),
                                       step=1.0, key="s8v_fg")
            with col_v2:
                v_ssc  = st.number_input("PCC Short-Circuit Capacity Ssc (MVA)", value=10.0,
                                          step=1.0, key="s8v_ssc")
                v_fsw  = st.number_input("Switching frequency (Hz)", value=float(_fsw),
                                        step=100.0, key="s8v_fsw")
                v_xr = st.number_input("Grid X/R ratio", value=5.0, min_value=0.5,
                                       step=0.5, key="s8v_xr")
            with col_v3:
                v_scr  = st.number_input("Short Circuit Ratio (SCR)", value=20.0, step=1.0, key="s8v_scr")
                v_L1   = st.number_input("L1 (mH)", value=float(_L1_mH), step=0.05, key="s8v_l1")
                v_L2   = st.number_input("L2 (mH)", value=float(_L2_mH), step=0.05, key="s8v_l2")

        # IEEE 519 
        if v_scr < 20:
            v_i_h_limits = {(3,11):4.0,(11,17):2.0,(17,23):1.5,(23,35):0.6,(35,999):0.3}
            v_tdd_limit  = 5.0
        elif v_scr < 50:
            v_i_h_limits = {(3,11):7.0,(11,17):3.5,(17,23):2.5,(23,35):1.0,(35,999):0.5}
            v_tdd_limit  = 8.0
        elif v_scr < 100:
            v_i_h_limits = {(3,11):10.0,(11,17):4.5,(17,23):4.0,(23,35):1.5,(35,999):0.7}
            v_tdd_limit  = 12.0
        elif v_scr < 1000:
            v_i_h_limits = {(3,11):12.0,(11,17):5.5,(17,23):5.0,(23,35):2.0,(35,999):1.0}
            v_tdd_limit  = 15.0
        else:
            v_i_h_limits = {(3,11):15.0,(11,17):7.0,(17,23):6.0,(23,35):2.5,(35,999):1.4}
            v_tdd_limit  = 20.0

        def get_ieee_limit_v(h, limits):
            if h < 11:   return limits[(3,11)]
            elif h < 17: return limits[(11,17)]
            elif h < 23: return limits[(17,23)]
            elif h < 35: return limits[(23,35)]
            else:        return limits[(35,999)]

        # ── [Add2] Import exceeding harmonics from Section 7 ────────────────────────────────────
        st.markdown("#### 🔢 Harmonic Current Input")

        h7_export = st.session_state.get('h7_harm_export', None)

        col_imp1, col_imp2 = st.columns([2, 3])
        with col_imp1:
            if h7_export is not None:
                st.success(f"✅ Section 7 FFT data loaded ({len(h7_export.get('freqs',[]))} frequency points)")
                do_import = st.button("🔄 Import Section 7 Exceeding Harmonic Currents",
                                       key="s8_import_b7", type="primary")
            else:
                st.warning("⚠ No Section 7 FFT data detected.")
                st.caption('Go to "7. Harmonic Standard Verification" → "📊 FFT Simulation Analysis" tab, '
                           'run the simulation and click the "Save results to Section 8" button, then return here.')
                # Pre-populate with typical PWM inverter harmonic profile so Section 8 works without Section 7
                if not st.session_state.get('s8_imported_harm'):
                    _i1 = float(v_il) if v_il > 0 else 1.0
                    # Typical 3-phase PWM inverter harmonic current ratios (relative to rated current)
                    _pwm_pct = {
                        1: 1.000, 2: 0.005, 3: 0.000, 4: 0.004, 5: 0.080,
                        6: 0.000, 7: 0.050, 8: 0.003, 9: 0.000, 10: 0.003,
                        11: 0.030, 12: 0.000, 13: 0.025, 14: 0.002, 15: 0.000,
                        16: 0.002, 17: 0.015, 18: 0.000, 19: 0.012, 20: 0.001,
                    }
                    st.session_state['s8_imported_harm'] = {
                        h: round(_i1 * pct, 4) for h, pct in _pwm_pct.items()
                    }
                    st.caption("📌 Default PWM inverter harmonic profile loaded (typical values; run Section 7 for actual data).")
                do_import = False

        with col_imp2:
            if h7_export is not None:
                # 7Exceeds Limit
                violations = h7_export.get('violations', [])
                if violations:
                    st.dataframe(pd.DataFrame(violations), width="stretch",
                                  hide_index=True, height=130)
                else:
                    st.info("Loaded from Section 7: harmonics within limits (filter may already satisfy standard)")

        # Combine filter target orders with standard reference orders
        all_orders = sorted(set(filter_targets + [5, 7, 11, 13]))

        # Import harmonic currents from Section 7 session_state
        if do_import and h7_export is not None:
            freqs_b7  = np.array(h7_export['freqs'])
            i_grid_b7 = np.array(h7_export['i_grid_A'])
            fg_b7     = h7_export.get('fg', _fg)

            imported = {}
            for h_n in all_orders:
                f_target = h_n * fg_b7
                idx = np.argmin(np.abs(freqs_b7 - f_target))
                if abs(freqs_b7[idx] - f_target) < fg_b7 * 2:
                    imported[h_n] = float(i_grid_b7[idx])
                else:
                    imported[h_n] = 0.0
            st.session_state['s8_imported_harm'] = imported
            st.success(f"✅ Imported {len(imported)} harmonic order current values from Section 7 FFT data!")

        imported_vals = st.session_state.get('s8_imported_harm', {})

        st.caption("Source: Section 7 FFT analysis results or field-measured values. Unit: A (RMS). "
                   "Leave as default if no data is available.")

        # Build harmonic input / display columns
        harm_input_data = []
        n_cols = min(len(all_orders), 4)
        cols_harm = st.columns(n_cols)

        for i, h_n in enumerate(all_orders):
            col = cols_harm[i % n_cols]
            with col:
                # Use imported value if available, otherwise estimate
                if h_n in imported_vals and imported_vals[h_n] > 0:
                    i_default = imported_vals[h_n]
                    label_suffix = " [Section 7]"
                else:
                    i_default = max(0.1, v_il * 0.05 / (h_n / 5.0))
                    label_suffix = ""

                i_src = st.number_input(
                    f"{h_n}th Harmonic ({h_n*v_fg:.0f} Hz) Source Current{label_suffix} (A)",
                    min_value=0.0, value=round(i_default, 3), step=0.01,
                    key=f"s8_isrc_{h_n}",
                    help=f"{h_n}th harmonic current RMS (after LCL filter, before passive filter; from Section 7 or field measurement)"
                )
                harm_input_data.append({
                    'h': h_n,
                    'f': h_n * v_fg,
                    'I_source': i_src,
                    'has_filter': h_n in filter_targets
                })

        st.divider()

        # ── [Fix1] Harmonic comparison using complex impedance divider ──────────────────────────────────────
        st.markdown("#### 📊 Before/After Filter Harmonic Comparison")

        # Compute grid impedance using Tab parameters (v_ssc, v_xr)
        omega1_v = 2 * np.pi * v_fg
        Lg_v  = (v_ug**2) / (omega1_v * v_ssc * 1e6) if v_ssc > 0 else 1e-3
        Xg_v  = omega1_v * Lg_v
        Rg_v  = Xg_v / v_xr if v_xr > 0 else Xg_v / 5.0

        results = []
        for d in harm_input_data:
            h_n    = d['h']
            I_src  = d['I_source']
            f_h    = d['f']
            omega_h = 2 * np.pi * f_h

            I_grid_no = I_src

            if d['has_filter']:
                fp_list = [p for p in filter_params if p['h_n'] == h_n]
                if fp_list:
                    fp = fp_list[0]
                    # [Fix1] Use complex impedance for current divider
                    Zf_cmplx = fp['Rf'] + 1j * (omega_h * fp['Lf'] - 1.0 / (omega_h * fp['Cf_f']))
                    Zg_cmplx_h = Rg_v + 1j * omega_h * Lg_v
                    denom_cmplx = Zf_cmplx + Zg_cmplx_h
                    if abs(denom_cmplx) > 1e-12:
                        I_grid_with = I_src * abs(Zf_cmplx) / abs(denom_cmplx)
                    else:
                        I_grid_with = I_src  # near resonance — fallback
                else:
                    I_grid_with = I_src
            else:
                I_grid_with = I_src

            TDD_no   = I_grid_no   / v_il * 100
            TDD_with = I_grid_with / v_il * 100
            limit_pct = get_ieee_limit_v(h_n, v_i_h_limits)

            results.append({
                'h': h_n, 'f': f_h,
                'I_source': I_src,
                'I_no_filter': I_grid_no,
                'I_with_filter': I_grid_with,
                'TDD_no': TDD_no, 'TDD_with': TDD_with,
                'limit': limit_pct,
                'pass_no':   TDD_no   <= limit_pct,
                'pass_with': TDD_with <= limit_pct,
                'has_filter': d['has_filter']
            })

        # Build comparison bar chart
        fig_compare = go.Figure()
        r_orders = [f"{r['h']}th\n{r['f']:.0f}Hz" for r in results]
        fig_compare.add_trace(go.Bar(
            x=r_orders, y=[r['TDD_no'] for r in results],
            name='Without passive filter',
            marker_color='rgba(220,53,69,0.7)',
            text=[f"{v:.2f}%" for v in [r['TDD_no'] for r in results]],
            textposition='outside'
        ))
        fig_compare.add_trace(go.Bar(
            x=r_orders, y=[r['TDD_with'] for r in results],
            name='With passive filter',
            marker_color='rgba(25,135,84,0.8)',
            text=[f"{v:.2f}%" for v in [r['TDD_with'] for r in results]],
            textposition='outside'
        ))
        fig_compare.add_trace(go.Scatter(
            x=r_orders, y=[r['limit'] for r in results],
            name='IEEE 519 Limit',
            mode='markers+lines',
            marker=dict(symbol='diamond', size=12, color='#dc3545'),
            line=dict(color='#dc3545', dash='dash', width=2)
        ))
        fig_compare.update_layout(
            title=f"Per-harmonic TDD Comparison (IEEE 519, SCR={v_scr:.0f})",
            xaxis_title="Harmonic Order",
            yaxis_title="TDD (%)",
            barmode='group', height=420,
            template="plotly_white",
            legend=dict(yanchor="top", y=0.98, xanchor="right", x=0.98)
        )
        if is_trial:
            fig_compare_trial = go.Figure(fig_compare)
            fig_compare_trial.update_traces(hoverinfo='none', hovertemplate=None, text=None)
            fig_compare_trial.update_yaxes(showticklabels=False)
            st.plotly_chart(fig_compare_trial, use_container_width=True)
            trial_lock("Unlock with paid access: per-harmonic TDD comparison chart (including standard limit comparison)", compact=True)
        else:
            st.plotly_chart(fig_compare, use_container_width=True)

        # Itemized verification table and total TDD
        # Compute total TDD unconditionally (needed for waveform chart below)
        tdd_total_no   = np.sqrt(sum(r['TDD_no']**2   for r in results))
        tdd_total_with = np.sqrt(sum(r['TDD_with']**2 for r in results))

        st.markdown("#### 📋 Itemized Verification Result")
        if is_trial:
            trial_lock("Unlock with paid access: itemized verification result table (per-harmonic pass/fail status, total TDD statistics)")
        else:
            df_verify = pd.DataFrame([{
                'Harmonic Order': f"{r['h']}th ({r['f']:.0f} Hz)",
                'Source Current (A)': f"{r['I_source']:.3f}",
                'TDD without filter': f"{r['TDD_no']:.2f}%",
                'TDD with filter': f"{r['TDD_with']:.2f}%",
                'IEEE 519 Limit': f"{r['limit']:.1f}%",
                'Result (no filter)': "✅ Pass" if r['pass_no'] else "❌ Exceeds Limit",
                'Result (with filter)': "✅ Pass" if r['pass_with'] else "❌ Still exceeds",
                'Filter branch': "✅ Designed" if r['has_filter'] else "⚪ None (reference)"
            } for r in results])

            def color_verdict(row):
                if 'Still exceeds' in str(row['Result (with filter)']):
                    return ['background-color:#f8d7da'] * len(row)
                if '✅ Pass' in str(row['Result (with filter)']):
                    return ['background-color:#d4edda'] * len(row)
                return [''] * len(row)

            st.dataframe(df_verify.style.apply(color_verdict, axis=1),
                          width="stretch", hide_index=True)

            col_t1, col_t2, col_t3 = st.columns(3)
            col_t1.metric("Total TDD (no filter)", f"{tdd_total_no:.2f}%",
                          delta="Exceeds limit" if tdd_total_no > v_tdd_limit else "Within limit",
                          delta_color="inverse" if tdd_total_no > v_tdd_limit else "normal")
            col_t2.metric("Total TDD (with filter)", f"{tdd_total_with:.2f}%",
                          delta=f"{'✅ Pass' if tdd_total_with<=v_tdd_limit else '❌ Still exceeds'} (limit {v_tdd_limit}%)",
                          delta_color="normal" if tdd_total_with <= v_tdd_limit else "inverse")
            col_t3.metric("TDD Improvement", f"{tdd_total_no-tdd_total_with:.2f}%",
                          delta=f"Reduced {(1-tdd_total_with/tdd_total_no)*100:.0f}% ↓" if tdd_total_no > 0 else "—",
                          delta_color="normal")

            still_fail = [r for r in results if not r['pass_with']]
            if still_fail:
                fail_str = ', '.join([f"{r['h']}th ({int(r['f'])} Hz)" for r in still_fail])
                st.error(f"""
**❌ The following harmonic orders still exceed the limit after adding the passive filter:**

{fail_str}

**Suggested actions**:
1. Increase filter reactive power Qf for that order (improve absorption capacity)
2. Lower quality factor Q (widen filter bandwidth)
3. Verify harmonic source estimates (recommend field measurement)
4. Check for anti-resonance interference (see Tab2 Parallel Anti-Resonance analysis)
                """)
            else:
                st.success("**✅ All targeted harmonic orders meet IEEE 519 requirements after adding the passive filter!**")

        # ── [Add4] Current waveform synthesis comparison (before/after filter) ──────────────────────
        # Based on individual harmonic amplitudes, reconstruct time-domain waveform via inverse Fourier synthesis.
        # Visually demonstrates waveform distortion improvement after adding the passive filter.
        # Note: Assumes zero initial phase (ideal synthesis), for qualitative illustration only.
        # ──────────────────────────────────────────────────────────────────────
        st.divider()
        st.markdown("#### 🌊 Grid Current Waveform Comparison (without filter vs with filter)")
        st.caption("⚠️ Note: Assumes individual harmonics are in phase (worst case). Waveform is for illustration; use actual simulation for precise results.")

        # Display parameter
        n_cycles_show = st.slider("Number of cycles to display", 1, 5, 2, key="s8_wf_cycles")
        T_fund = 1.0 / v_fg
        t_wf = np.linspace(0, n_cycles_show * T_fund, 2000)
        omega1_wf = 2 * np.pi * v_fg

        # Start with fundamental component + add individual harmonics
        i_wf_no_filter   = v_il * np.sin(omega1_wf * t_wf)   # fundamental
        i_wf_with_filter = v_il * np.sin(omega1_wf * t_wf)   # fundamental

        for r in results:
            h_n = r['h']
            omega_h = 2 * np.pi * r['f']
            # Without filter: add all harmonics as-is
            i_wf_no_filter   += r['I_no_filter']   * np.sin(omega_h * t_wf)
            # With filter: add attenuated harmonics
            i_wf_with_filter += r['I_with_filter'] * np.sin(omega_h * t_wf)

        # Normalize x-axis to fundamental cycles
        t_norm = t_wf / T_fund

        fig_wf = go.Figure()

        # Ideal sinusoid reference
        fig_wf.add_trace(go.Scatter(
            x=t_norm,
            y=v_il * np.sin(omega1_wf * t_wf),
            name='Ideal sinusoid (fundamental only)',
            line=dict(color='rgba(108,117,125,0.3)', width=1, dash='dash'),
        ))

        # Without filter (red — distorted)
        fig_wf.add_trace(go.Scatter(
            x=t_norm,
            y=i_wf_no_filter,
            name=f'Without filter (THD≈{tdd_total_no:.1f}%)',
            line=dict(color='#dc3545', width=2),
            fill='tonexty',
            fillcolor='rgba(220,53,69,0.08)',
        ))

        # With filter (green — improved)
        fig_wf.add_trace(go.Scatter(
            x=t_norm,
            y=i_wf_with_filter,
            name=f'With filter (THD≈{tdd_total_with:.1f}%)',
            line=dict(color='#198754', width=2.5),
        ))

        fig_wf.update_layout(
            title="Grid Current Waveform Synthesis Comparison (fundamental + harmonics)",
            xaxis_title="Time (fundamental cycles)",
            yaxis_title="Current i(t) (A)",
            height=400,
            template="plotly_white",
            legend=dict(yanchor="top", y=0.98, xanchor="right", x=0.98),
            hovermode="x unified"
        )

        # Add cycle boundary lines
        for cyc in range(1, n_cycles_show + 1):
            fig_wf.add_vline(x=cyc, line_dash="dot",
                              line_color="rgba(108,117,125,0.2)")

        if is_trial:
            fig_wf_trial = go.Figure(fig_wf)
            fig_wf_trial.update_traces(hoverinfo='none', hovertemplate=None)
            fig_wf_trial.update_yaxes(showticklabels=False)
            st.plotly_chart(fig_wf_trial, use_container_width=True)
            trial_lock("Unlock with paid access: current waveform synthesis comparison chart (before vs after filter, time-domain waveform)", compact=True)
        else:
            st.plotly_chart(fig_wf, use_container_width=True)

            # Waveform peak distortion summary
            peak_no   = np.max(np.abs(i_wf_no_filter))
            peak_with = np.max(np.abs(i_wf_with_filter))
            peak_ideal = v_il
            crest_err_no   = abs(peak_no - peak_ideal) / peak_ideal * 100
            crest_err_with = abs(peak_with - peak_ideal) / peak_ideal * 100
            col_wf1, col_wf2 = st.columns(2)
            with col_wf1:
                st.metric("Peak distortion (no filter)",
                           f"{crest_err_no:.1f}%",
                           delta=f"Peak {peak_no:.1f} A vs ideal {peak_ideal:.1f} A")
            with col_wf2:
                st.metric("Peak distortion (with filter)",
                           f"{crest_err_with:.1f}%",
                           delta=f"Improved by {crest_err_no - crest_err_with:.1f}% ↓",
                           delta_color="normal")

        st.markdown("""
> **Reading guide**: Gray dashed line = ideal sinusoid (pure fundamental); red = distorted waveform without filter;
> green = improved waveform after adding the passive filter. The larger the shaded area between red and gray,
> the more severe the distortion. The closer the green curve is to gray, the better the filter performance.
        """)



















# ==========================================
#  9:  v3
#
# v3 (): 
#   [M1] TAB1 /, 
#        (Air Gap/Turns///)
#   [M2]  TAB2 
#   [M3] : "", ; 
#        ; 
#        4: ① ② ③ ④
#   [M4] : , 
#   [M5] TAB5  OCP ; 
#        ; 
#         Excel
#
# (): 
#   [1] IEC 60076-6:2007 — Power Transformers Part 6: Reactors
#   [2] GB/T 1094.6-2017 —  6: 
#   [3] NB/T 10503-2021 — 
#   [4] McLyman C.W.T., "Transformer and Inductor Design Handbook", 4th ed., CRC Press
#   [5] Kazimierczuk M.K., "High-Frequency Magnetic Components", Wiley 2014
#   [6] Dowell P.L., IEE Proc. 1966, Vol.113, No.8 — Copper Loss
# ==========================================

elif selection == nav_options[8]:

    import io

    # ──  session_state Parameter ─────────────────────────────────────
    _L1_mH   = float(st.session_state.get('lcl_l1',  0.5))
    _L2_mH   = float(st.session_state.get('lcl_l2',  0.25))
    _Cf_uF   = float(st.session_state.get('lcl_c',   100.0))
    _IL_rms  = float(st.session_state.get('vec_i',   500.0))
    _Vg_line = float(st.session_state.get('sys_ug',  690.0))
    _Vdc     = float(st.session_state.get('vec_udc', 1100.0))
    _fsw     = float(st.session_state.get('sys_fsw', 2500.0))
    _fg      = float(st.session_state.get('sys_fg',  50.0))
    _L1_H    = _L1_mH * 1e-3
    _I_peak  = _IL_rms * np.sqrt(2)

    # ── Page Header ─────────────────────────────────────────────────────────────────
    st.markdown("## 🔩 Module 9: Reactor Engineering Design")
    st.caption(
        "Known LCL parameters → this module completes the full chain of reactor design from **physical understanding** to **engineering implementation**. "
        "Covers core selection, winding design, loss calculation, scheme comparison, and supplier inquiry sheet generation."
    )

    c_r1, c_r2, c_r3, c_r4, c_r5 = st.columns(5)
    with c_r1: st.metric("L1 (Converter Side)", f"{_L1_mH:.3f} mH")
    with c_r2: st.metric("L2 (Grid Side)",       f"{_L2_mH:.3f} mH")
    with c_r3: st.metric("Rated Current",         f"{_IL_rms:.0f} Arms")
    with c_r4: st.metric("DC Bus Voltage",        f"{_Vdc:.0f} V")
    with c_r5: st.metric("Grid Voltage",          f"{_Vg_line:.0f} V")
    if is_trial:
        trial_lock("Design parameter results require paid access")
    st.divider()

    # [M2] Tabs
    tab_intro, tab_design, tab_test = st.tabs([
        "📚 1. Design Fundamentals",
        "🛠️ 2. Design Calculation",
        "📥 3. Test Requirements & Supplier Scheme"
    ])


    # ══════════════════════════════════════════════════════════════════════════
    # TAB 1 —   [M1]
    # ══════════════════════════════════════════════════════════════════════════
    with tab_intro:
        st.markdown("### 🧲 Dry-Core Reactor: Qualitative Rules Every Engineer Must Know")
        st.info(
            "💡 **Feynman Introduction**: A reactor acts like a spring for current — when current tries to change rapidly, "
            "it generates a back-EMF that resists the change, protecting IGBTs from current surges. "
            "The following are the most important qualitative rules for design — build intuition first, then review formulas."
        )

        # ══════════════════════════════════════════
        # Part 1: Three-Phase Core Engineering Structure
        # ══════════════════════════════════════════
        st.markdown("#### 🏗️ Three-Phase Core Reactor: Engineering Structure Overview")
        st.caption("Front view + side view comparison, to understand the spatial relationship between lamination width, stacking thickness, and core cross-sectional area")

        col_front, col_side, col_dim_rules = st.columns([5, 4, 4])

        with col_front:
            st.markdown('<div style="text-align:center;font-size:12px;color:#555;margin-bottom:4px">▼ Front View (Y-Z Plane)</div>', unsafe_allow_html=True)
            svg_3phase = """
<svg width="100%" height="370" viewBox="0 0 420 370" xmlns="http://www.w3.org/2000/svg">
<!-- Coordinate axis annotation -->
  <text x="10" y="16" font-size="9" fill="#888" font-family="sans-serif">Y →</text>
  <text x="2"  y="40" font-size="9" fill="#888" font-family="sans-serif" transform="rotate(-90,8,40)">Z ↑</text>

  <!-- Upper yoke -->
  <rect x="25" y="28" width="370" height="44" fill="#7070a8" rx="3"/>
  <g stroke="#333" stroke-width="0.5" opacity="0.35">
    <line x1="25" y1="38" x2="395" y2="38"/>
    <line x1="25" y1="48" x2="395" y2="48"/>
    <line x1="25" y1="58" x2="395" y2="58"/>
    <line x1="25" y1="66" x2="395" y2="66"/>
  </g>

  <!-- Lower yoke -->
  <rect x="25" y="298" width="370" height="44" fill="#7070a8" rx="3"/>
  <g stroke="#333" stroke-width="0.5" opacity="0.35">
    <line x1="25" y1="308" x2="395" y2="308"/>
    <line x1="25" y1="318" x2="395" y2="318"/>
    <line x1="25" y1="328" x2="395" y2="328"/>
    <line x1="25" y1="338" x2="395" y2="338"/>
  </g>

  <!-- Phase A core limb x=55 w=44 -->
  <rect x="55" y="72" width="44" height="74" fill="#7878b0" rx="2"/>
  <g stroke="#333" stroke-width="0.4" opacity="0.3">
    <line x1="60" y1="82"  x2="94" y2="82"/><line x1="60" y1="92"  x2="94" y2="92"/>
    <line x1="60" y1="102" x2="94" y2="102"/><line x1="60" y1="112" x2="94" y2="112"/>
    <line x1="60" y1="122" x2="94" y2="122"/><line x1="60" y1="132" x2="94" y2="132"/>
  </g>
  <rect x="55" y="146" width="44" height="16" fill="#f5c518" stroke="#b8860b" stroke-width="1.2"/>
  <rect x="55" y="162" width="44" height="136" fill="#7878b0" rx="2"/>
  <g stroke="#333" stroke-width="0.4" opacity="0.3">
    <line x1="60" y1="172" x2="94" y2="172"/><line x1="60" y1="182" x2="94" y2="182"/>
    <line x1="60" y1="192" x2="94" y2="192"/><line x1="60" y1="202" x2="94" y2="202"/>
    <line x1="60" y1="212" x2="94" y2="212"/><line x1="60" y1="222" x2="94" y2="222"/>
    <line x1="60" y1="232" x2="94" y2="232"/><line x1="60" y1="242" x2="94" y2="242"/>
    <line x1="60" y1="252" x2="94" y2="252"/><line x1="60" y1="262" x2="94" y2="262"/>
    <line x1="60" y1="272" x2="94" y2="272"/><line x1="60" y1="282" x2="94" y2="282"/>
  </g>
  <rect x="28"  y="88" width="20" height="188" fill="#c88000" rx="3" opacity="0.92"/>
  <rect x="106" y="88" width="20" height="188" fill="#c88000" rx="3" opacity="0.92"/>
  <text x="77" y="20" font-size="11" fill="#b87800" font-family="sans-serif" font-weight="bold" text-anchor="middle">Phase A</text>

  <!-- Phase B core limb x=188 -->
  <rect x="188" y="72" width="44" height="74" fill="#7878b0" rx="2"/>
  <g stroke="#333" stroke-width="0.4" opacity="0.3">
    <line x1="193" y1="82"  x2="227" y2="82"/><line x1="193" y1="92"  x2="227" y2="92"/>
    <line x1="193" y1="102" x2="227" y2="102"/><line x1="193" y1="112" x2="227" y2="112"/>
    <line x1="193" y1="122" x2="227" y2="122"/><line x1="193" y1="132" x2="227" y2="132"/>
  </g>
  <rect x="188" y="146" width="44" height="16" fill="#f5c518" stroke="#b8860b" stroke-width="1.2"/>
  <rect x="188" y="162" width="44" height="136" fill="#7878b0" rx="2"/>
  <g stroke="#333" stroke-width="0.4" opacity="0.3">
    <line x1="193" y1="172" x2="227" y2="172"/><line x1="193" y1="182" x2="227" y2="182"/>
    <line x1="193" y1="192" x2="227" y2="192"/><line x1="193" y1="202" x2="227" y2="202"/>
    <line x1="193" y1="212" x2="227" y2="212"/><line x1="193" y1="222" x2="227" y2="222"/>
    <line x1="193" y1="232" x2="227" y2="232"/><line x1="193" y1="242" x2="227" y2="242"/>
    <line x1="193" y1="252" x2="227" y2="252"/><line x1="193" y1="262" x2="227" y2="262"/>
    <line x1="193" y1="272" x2="227" y2="272"/><line x1="193" y1="282" x2="227" y2="282"/>
  </g>
  <rect x="161" y="88" width="20" height="188" fill="#2070c0" rx="3" opacity="0.92"/>
  <rect x="239" y="88" width="20" height="188" fill="#2070c0" rx="3" opacity="0.92"/>
  <text x="210" y="20" font-size="11" fill="#1a6faa" font-family="sans-serif" font-weight="bold" text-anchor="middle">Phase B</text>

  <!-- Phase C core limb x=321 -->
  <rect x="321" y="72" width="44" height="74" fill="#7878b0" rx="2"/>
  <g stroke="#333" stroke-width="0.4" opacity="0.3">
    <line x1="326" y1="82"  x2="360" y2="82"/><line x1="326" y1="92"  x2="360" y2="92"/>
    <line x1="326" y1="102" x2="360" y2="102"/><line x1="326" y1="112" x2="360" y2="112"/>
    <line x1="326" y1="122" x2="360" y2="122"/><line x1="326" y1="132" x2="360" y2="132"/>
  </g>
  <rect x="321" y="146" width="44" height="16" fill="#f5c518" stroke="#b8860b" stroke-width="1.2"/>
  <rect x="321" y="162" width="44" height="136" fill="#7878b0" rx="2"/>
  <g stroke="#333" stroke-width="0.4" opacity="0.3">
    <line x1="326" y1="172" x2="360" y2="172"/><line x1="326" y1="182" x2="360" y2="182"/>
    <line x1="326" y1="192" x2="360" y2="192"/><line x1="326" y1="202" x2="360" y2="202"/>
    <line x1="326" y1="212" x2="360" y2="212"/><line x1="326" y1="222" x2="360" y2="222"/>
    <line x1="326" y1="232" x2="360" y2="232"/><line x1="326" y1="242" x2="360" y2="242"/>
    <line x1="326" y1="252" x2="360" y2="252"/><line x1="326" y1="262" x2="360" y2="262"/>
    <line x1="326" y1="272" x2="360" y2="272"/><line x1="326" y1="282" x2="360" y2="282"/>
  </g>
  <rect x="294" y="88" width="20" height="188" fill="#c03030" rx="3" opacity="0.92"/>
  <rect x="372" y="88" width="20" height="188" fill="#c03030" rx="3" opacity="0.92"/>
  <text x="343" y="20" font-size="11" fill="#aa2222" font-family="sans-serif" font-weight="bold" text-anchor="middle">Phase C</text>

  <!-- Lamination width annotation (Phase A top) -->
  <line x1="55"  y1="62" x2="99"  y2="62" stroke="#333" stroke-width="1" stroke-dasharray="3,2"/>
  <line x1="55"  y1="57" x2="55"  y2="67" stroke="#333" stroke-width="1.2"/>
  <line x1="99"  y1="57" x2="99"  y2="67" stroke="#333" stroke-width="1.2"/>
  <text x="77"  y="58" font-size="9" fill="#333" font-family="sans-serif" text-anchor="middle">Width w (Y-axis)</text>

  <!-- Air Gap annotation -->
  <line x1="232" y1="154" x2="258" y2="148" stroke="#b8860b" stroke-width="1"/>
  <text x="260" y="147" font-size="9" fill="#b8860b" font-family="sans-serif" font-weight="bold">Air Gap δ</text>

  <!-- Yoke cross-section green box -->
  <rect x="128" y="28" width="76" height="44" fill="none" stroke="#27ae60" stroke-width="1.5" stroke-dasharray="4,2"/>
  <text x="165" y="22" font-size="8" fill="#27ae60" font-family="sans-serif" text-anchor="middle">Yoke Area ≥ Ae</text>

  <!-- Core cross-section red box -->
  <rect x="55" y="200" width="44" height="32" fill="none" stroke="#e74c3c" stroke-width="1.5" stroke-dasharray="4,2"/>
  <line x1="99" y1="216" x2="120" y2="216" stroke="#e74c3c" stroke-width="1"/>
  <text x="122" y="213" font-size="8" fill="#e74c3c" font-family="sans-serif">Ae=w×d×kfe</text>

  <!-- Z-axis height annotation (left side) -->
  <line x1="12" y1="72"  x2="12" y2="298" stroke="#555" stroke-width="1" stroke-dasharray="3,2"/>
  <line x1="7"  y1="72"  x2="17" y2="72"  stroke="#555" stroke-width="1.2"/>
  <line x1="7"  y1="298" x2="17" y2="298" stroke="#555" stroke-width="1.2"/>
  <text x="5" y="190" font-size="8" fill="#555" font-family="sans-serif" text-anchor="middle" transform="rotate(-90,5,190)">Window Height lw (Z-axis)</text>
</svg>"""
            st.markdown(svg_3phase, unsafe_allow_html=True)

        with col_side:
            st.markdown('<div style="text-align:center;font-size:12px;color:#555;margin-bottom:4px">▼ Side View (X-Z Plane, Phase A Core Limb)</div>', unsafe_allow_html=True)
            svg_side = """
<svg width="100%" height="370" viewBox="0 0 240 370" xmlns="http://www.w3.org/2000/svg">
<!-- Coordinate axes -->
  <text x="8"  y="16" font-size="9" fill="#888" font-family="sans-serif">X →</text>
  <text x="0"  y="40" font-size="9" fill="#888" font-family="sans-serif" transform="rotate(-90,6,40)">Z ↑</text>

  <!-- Upper yoke (side view: X-Z cross-section, stack depth d = X-direction width) -->
  <rect x="30" y="28" width="180" height="44" fill="#7070a8" rx="3"/>
  <!-- Lamination texture: vertical, as viewed from X direction -->
  <g stroke="#fff" stroke-width="0.6" opacity="0.25">
    <line x1="50"  y1="28" x2="50"  y2="72"/>
    <line x1="70"  y1="28" x2="70"  y2="72"/>
    <line x1="90"  y1="28" x2="90"  y2="72"/>
    <line x1="110" y1="28" x2="110" y2="72"/>
    <line x1="130" y1="28" x2="130" y2="72"/>
    <line x1="150" y1="28" x2="150" y2="72"/>
    <line x1="170" y1="28" x2="170" y2="72"/>
    <line x1="190" y1="28" x2="190" y2="72"/>
  </g>

  <!-- Lower yoke -->
  <rect x="30" y="298" width="180" height="44" fill="#7070a8" rx="3"/>
  <g stroke="#fff" stroke-width="0.6" opacity="0.25">
    <line x1="50"  y1="298" x2="50"  y2="342"/>
    <line x1="70"  y1="298" x2="70"  y2="342"/>
    <line x1="90"  y1="298" x2="90"  y2="342"/>
    <line x1="110" y1="298" x2="110" y2="342"/>
    <line x1="130" y1="298" x2="130" y2="342"/>
    <line x1="150" y1="298" x2="150" y2="342"/>
    <line x1="170" y1="298" x2="170" y2="342"/>
    <line x1="190" y1="298" x2="190" y2="342"/>
  </g>

  <!-- Core limb (side view: X-width = stack depth d, Z = window height) -->
  <rect x="78" y="72" width="84" height="74" fill="#7878b0" rx="2"/>
  <!-- Vertical lamination texture (viewed from X direction, silicon steel sheets stand vertically) -->
  <g stroke="#fff" stroke-width="0.6" opacity="0.2">
    <line x1="92"  y1="72" x2="92"  y2="146"/>
    <line x1="106" y1="72" x2="106" y2="146"/>
    <line x1="120" y1="72" x2="120" y2="146"/>
    <line x1="134" y1="72" x2="134" y2="146"/>
    <line x1="148" y1="72" x2="148" y2="146"/>
  </g>
  <!-- Air Gap (visible in side view, full X-width) -->
  <rect x="78" y="146" width="84" height="16" fill="#f5c518" stroke="#b8860b" stroke-width="1.2"/>
  <!-- Lower core limb segment -->
  <rect x="78" y="162" width="84" height="136" fill="#7878b0" rx="2"/>
  <g stroke="#fff" stroke-width="0.6" opacity="0.2">
    <line x1="92"  y1="162" x2="92"  y2="298"/>
    <line x1="106" y1="162" x2="106" y2="298"/>
    <line x1="120" y1="162" x2="120" y2="298"/>
    <line x1="134" y1="162" x2="134" y2="298"/>
    <line x1="148" y1="162" x2="148" y2="298"/>
  </g>

  <!-- Winding (side view: viewed from X direction, winding is a rectangular cross-section around core) -->
  <!-- Front winding (near observer side) -->
  <rect x="36"  y="88" width="36" height="188" fill="#c88000" rx="3" opacity="0.88"/>
  <!-- Rear winding (far observer side) -->
  <rect x="168" y="88" width="36" height="188" fill="#c88000" rx="3" opacity="0.88"/>

  <!-- Stack depth d annotation (X-axis direction, core limb width) -->
  <line x1="78"  y1="360" x2="162" y2="360" stroke="#e74c3c" stroke-width="1.2"/>
  <line x1="78"  y1="355" x2="78"  y2="365" stroke="#e74c3c" stroke-width="1.5"/>
  <line x1="162" y1="355" x2="162" y2="365" stroke="#e74c3c" stroke-width="1.5"/>
  <text x="120" y="356" font-size="9" fill="#e74c3c" font-family="sans-serif" text-anchor="middle">Stack depth d (X-axis)</text>

  <!-- Air Gap annotation -->
  <line x1="162" y1="154" x2="188" y2="148" stroke="#b8860b" stroke-width="1"/>
  <text x="190" y="147" font-size="9" fill="#b8860b" font-family="sans-serif" font-weight="bold">Air Gap</text>

  <!-- Z-axis height annotation -->
  <line x1="14" y1="72"  x2="14" y2="298" stroke="#555" stroke-width="1" stroke-dasharray="3,2"/>
  <line x1="9"  y1="72"  x2="19" y2="72"  stroke="#555" stroke-width="1.2"/>
  <line x1="9"  y1="298" x2="19" y2="298" stroke="#555" stroke-width="1.2"/>
  <text x="6" y="190" font-size="8" fill="#555" font-family="sans-serif" text-anchor="middle" transform="rotate(-90,6,190)">Window Height lw (Z-axis)</text>

  <!-- Cross-section description box -->
  <rect x="78" y="200" width="84" height="30" fill="none" stroke="#e74c3c" stroke-width="1.5" stroke-dasharray="4,2"/>
  <text x="120" y="218" font-size="8" fill="#e74c3c" font-family="sans-serif" text-anchor="middle">Ae = w×d×kfe</text>
</svg>"""
            st.markdown(svg_side, unsafe_allow_html=True)

        with col_dim_rules:
            st.markdown("**Core Dimensions and Cross-Sectional Area — Basic Knowledge**")
            st.markdown("""
**Three Key Dimensions**
> - **Lamination width w** (Y-axis): visible in front view, punching width of silicon steel sheets, constrained by tooling
> - **Stack depth d** (X-axis): **visible in side view**, total thickness of laminated stack in depth direction, continuously adjustable
> - **Window height lw** (Z-axis): height between upper and lower yokes, determines winding space

**Core Cross-Sectional Area**
> Ae = w × d × kfe
> kfe ≈ 0.95 (lamination fill factor)
>
> **Larger Ae → lower B = L·I/(N·Ae) → less likely to saturate**
> But heavier core means higher cost

**Stack Depth Adjustment Rule**
> Lamination width w typically follows standard series (fixed tooling); stack depth d is the main adjustment variable:
> - Need larger cross-section → increase stack depth d
> - Stack depth is generally 0.8~1.5× lamination width; oversized ratio worsens heat dissipation

**Yoke Cross-Sectional Area Rule**
> Yoke carries the same flux as the core limb
> **Engineering convention: yoke cross-section ≥ core limb cross-section**
> - Standard design uses 1.0× for sufficient margin
> - Lightweight design uses 0.85×; must carefully verify yoke flux density is not exceeded

**Lamination Direction and Core Loss**
> Silicon steel has highest permeability along the rolling direction (Z-axis).
> At yoke corners, flux must turn 90°, which is where core loss concentrates — design with slightly thicker yokes.
""")

        st.divider()

        # ══════════════════════════════════════════
        # :  + 
        # ══════════════════════════════════════════
        col_svg, col_rules = st.columns([1, 1])
        with col_svg:
            st.markdown("**Single-Phase Core Reactor Structure Diagram (Five Key Elements)**")
            svg_iron = """
<svg width="100%" height="300" viewBox="0 0 320 300" xmlns="http://www.w3.org/2000/svg">
<rect x="45"  y="40"  width="32" height="220" fill="#7070a8" rx="3"/>
  <rect x="243" y="40"  width="32" height="220" fill="#7070a8" rx="3"/>
  <rect x="45"  y="40"  width="230" height="30" fill="#7070a8" rx="3"/>
  <rect x="45"  y="230" width="230" height="30" fill="#7070a8" rx="3"/>
  <rect x="144" y="70"  width="32" height="78"  fill="#7070a8" rx="2"/>
  <rect x="144" y="148" width="32" height="16"  fill="#f5c518" stroke="#b8860b" stroke-width="1.5" rx="1"/>
  <rect x="144" y="164" width="32" height="66"  fill="#7070a8" rx="2"/>
  <g stroke="#555" stroke-width="0.4" opacity="0.5">
    <line x1="45"  y1="52"  x2="275" y2="52"/>
    <line x1="45"  y1="60"  x2="275" y2="60"/>
    <line x1="45"  y1="240" x2="275" y2="240"/>
    <line x1="45"  y1="248" x2="275" y2="248"/>
  </g>
  <ellipse cx="160" cy="100" rx="28" ry="9" fill="none" stroke="#d4a000" stroke-width="7" opacity="0.9"/>
  <ellipse cx="160" cy="114" rx="28" ry="9" fill="none" stroke="#d4a000" stroke-width="7" opacity="0.9"/>
  <ellipse cx="160" cy="128" rx="28" ry="9" fill="none" stroke="#d4a000" stroke-width="7" opacity="0.9"/>
  <ellipse cx="160" cy="180" rx="28" ry="9" fill="none" stroke="#d4a000" stroke-width="7" opacity="0.9"/>
  <ellipse cx="160" cy="194" rx="28" ry="9" fill="none" stroke="#d4a000" stroke-width="7" opacity="0.9"/>
  <ellipse cx="160" cy="208" rx="28" ry="9" fill="none" stroke="#d4a000" stroke-width="7" opacity="0.9"/>
  <line x1="164" y1="156" x2="210" y2="150" stroke="#b8860b" stroke-width="1.2"/>
  <text x="213" y="154" font-size="11" fill="#b8860b" font-family="sans-serif" font-weight="bold">Air Gap δ</text>
  <line x1="77"  y1="150" x2="35"  y2="140" stroke="#4a4a6a" stroke-width="1.2"/>
  <text x="2"   y="138" font-size="10" fill="#7c7cac" font-family="sans-serif">Core</text>
  <line x1="188" y1="128" x2="228" y2="115" stroke="#b8860b" stroke-width="1.2"/>
  <text x="231" y="118" font-size="10" fill="#b8860b" font-family="sans-serif">Winding N</text>
  <line x1="160" y1="92"  x2="160" y2="25"  stroke="#8B0000" stroke-width="3"/>
  <circle cx="160" cy="23"  r="5" fill="#8B0000"/>
  <text x="168" y="20"  font-size="11" fill="#8B0000" font-family="sans-serif">Terminal A</text>
  <line x1="160" y1="215" x2="160" y2="278" stroke="#8B0000" stroke-width="3"/>
  <circle cx="160" cy="280" r="5" fill="#8B0000"/>
  <text x="168" y="285" font-size="11" fill="#8B0000" font-family="sans-serif">Terminal B</text>
  <path d="M 100 80 Q 45 80 45 150 Q 45 220 100 220" stroke="#0d6efd" stroke-width="1.5"
        fill="none" stroke-dasharray="5,3" opacity="0.6"/>
  <polygon points="100,215 95,225 105,225" fill="#0d6efd" opacity="0.5"/>
  <text x="10" y="160" font-size="9" fill="#0d6efd" font-family="sans-serif" opacity="0.8">Flux Φ</text>
</svg>"""
            st.markdown(svg_iron, unsafe_allow_html=True)

        with col_rules:
            st.markdown("**⚡ Five Qualitative Rules — Quick Reference**")
            st.markdown("""
**① Larger air gap → lower inductance**
> Larger air gap increases total magnetic reluctance → inductance L = N²/R_m drops for the same turns count.
> But a larger air gap makes the core less likely to saturate (stronger DC bias capability).
> **Trade-off**: air gap size = balance between meeting inductance target vs. avoiding saturation.

**② More turns → higher inductance, but also higher flux density**
> L ∝ N²; more turns raises inductance;
> but for the same current, B = μ₀NI/δ — more turns means higher flux density, easier to saturate.
> **Conclusion**: increasing turns must be paired with a larger air gap, otherwise saturation occurs.

""")
            if is_trial:
                trial_lock("Unlock with paid access: magnetic saturation prevention, lamination frequency selection, and copper/core loss trade-off rules", compact=True)
            else:
                st.markdown("""
**③ How to prevent magnetic saturation**
> Magnetic saturation = B_peak > B_sat → inductance drops → ripple current out of control → IGBT damage.
> Three mitigation methods (can be combined):
> - Increase air gap (most direct)
> - Reduce turns (also reduces inductance — requires redesign)
> - Use a higher B_sat core material (e.g., switch from powder core to silicon steel)

**④ Higher switching frequency → thinner laminations preferred**
> Core loss ∝ f^1.6 — higher frequency increases core loss faster. Thinner laminations shorten eddy current paths, reducing core loss.
> fsw ≤ 1 kHz: 0.35 mm silicon steel satisfies requirement;
> fsw = 2~5 kHz: suggest 0.23 mm thin silicon steel or Sendust powder core;
> fsw > 10 kHz: consider Sendust powder core or amorphous alloy.

**⑤ Copper loss vs. core loss trade-off**
> More turns → higher winding resistance → higher copper loss;
> Fewer turns → inductance drops → need to increase air gap → core saturates more easily.
> Engineering target: design for copper loss ≈ core loss (loss balance), to achieve minimum total loss.
""")

        st.divider()

        # ══════════════════════════════════════════
        # : 
        # ══════════════════════════════════════════
        st.markdown("#### 📋 Dry Core Material Selection Quick Reference")
        mat_data = {
            "Core Material": [
                "Non-oriented silicon steel 0.35mm (35WW300)",
                "Grain-oriented silicon steel 0.35mm (B27R095)",
                "Grain-oriented silicon steel 0.23mm (B23R085)",
                "Sendust powder core (Fe-Si-Al)",
                "Amorphous alloy (1K101)",
                "Nanocrystalline",
            ],
            "B_sat (T)": ["1.7", "2.0", "1.8", "1.05", "1.56", "1.20"],
            "Applicable Frequency": [
                "≤ 2kHz",
                "≤ 1kHz",
                "1~3kHz",
                "1~200kHz",
                "50Hz~20kHz",
                "1~100kHz",
            ],
            "Core Loss Characteristics": [
                "Lowest cost, higher core loss at high frequency, common engineering trade-off",
                "High flux density, low loss at 50Hz, large core loss at high frequency",
                "Core loss ~60% lower than 0.35mm grain-oriented silicon steel",
                "Distributed air gap, very low high-frequency loss, strong DC bias capability",
                "Amorphous structure, low loss across wide frequency range",
                "Extremely low core loss, highest cost",
            ],
            "Cost Reference": [
                "Lowest (baseline 1×)",
                "Low (1.1×)",
                "Low (1.2×)",
                "Medium (2.0×)",
                "High (2.5×)",
                "Highest (3.5×)",
            ],
        }
        st.dataframe(pd.DataFrame(mat_data), use_container_width=True, hide_index=True)

        st.divider()

        # ══════════════════════════════════════════
        # : 
        # ══════════════════════════════════════════
        st.markdown("#### 🗒️ Insulation Paper — The Reactor's \"Firewall\"")
        st.caption("Insulation paper is one of the most overlooked yet most critical materials in a reactor, directly determining insulation class and service life.")

        col_ins1, col_ins2 = st.columns(2)
        with col_ins1:
            st.markdown("**Insulation Paper Usage Locations**")
            ins_pos_data = {
                "Usage Location": [
                    "Core outer wrapping insulation",
                    "Inter-layer insulation (between foil layers)",
                    "Winding outer insulation",
                    "End insulation (air gap side)",
                    "Terminal block insulation",
                ],
                "Function": [
                    "Main insulation between core and winding, prevents turn-to-core short circuit",
                    "Insulates adjacent aluminum/copper foil layers, withstands inter-layer voltage",
                    "Main insulation between winding and enclosure / between phases",
                    "Winding end near air gap, prevents fringing field edge effect causing insulation damage",
                    "Creepage insulation between busbar terminals and core/enclosure",
                ],
                "Typical Thickness": [
                    "0.13~0.25mm",
                    "0.10~0.18mm",
                    "0.25~0.50mm",
                    "0.18~0.50mm (thickened)",
                    "1~3mm rigid board",
                ],
            }
            if is_trial:
                trial_lock("Unlock with paid access: insulation paper usage location detail table", compact=True)
            else:
                st.dataframe(pd.DataFrame(ins_pos_data), use_container_width=True, hide_index=True)

        with col_ins2:
            st.markdown("**Common insulation paper material types**")
            ins_mat_data = {
                "Material Type": [
                    "Nomex 410 (aramid paper)",
                    "NMN (Nomex-Mylar-Nomex)",
                    "NHN (Nomex-Kapton-Nomex)",
                    "DMD (polyester film-polyester fiber)",
                    "AMA / NMA (mica composite)",
                    "PI film / Kapton (polyimide)",
                ],
                "Insulation Class": [
                    "H (180°C)",
                    "H (180°C)*",
                    "C (>220°C)",
                    "B/F (130~155°C)",
                    "H (180°C)",
                    "C (>220°C)",
                ],
                "Flame Rating": [
                    "UL94 V-0",
                    "UL94 V-0",
                    "UL94 V-0",
                    "UL94 V-1",
                    "UL94 V-0",
                    "UL94 V-0",
                ],
                "Description & Reference": [
                    "DuPont Nomex 410; IEC 60626-2; preferred H-class main insulation",
                    "* H-class rating requires UL 1446 insulation system certification (high-temp PET or special bonding); standard NMN is F-class (155°C); IEC 60626-2",
                    "Middle PI film layer is C-class, overall used as C-class (>220°C); high-temperature applications",
                    "Low cost, B~F class, not corona-resistant; IEC 60626-1",
                    "Contains mica layer, resistant to partial discharge/arc, high-voltage applications; IEC 60371",
                    "DuPont Kapton, high-temp and chemical resistant, high cost; IEC 60626-2",
                ],
            }
            if is_trial:
                trial_lock("Unlock with paid access: common insulation paper material types comparison (including insulation class and selection basis)", compact=True)
            else:
                st.dataframe(pd.DataFrame(ins_mat_data), use_container_width=True, hide_index=True)
                st.caption(
                    "Insulation class per: IEC 60085 (electrical insulation thermal classification) and material datasheets. "
                    "* NMN rated H-class (180°C) requires high-temperature grade PET or special bonding process "
                    "per UL 1446 insulation system certification. Standard NMN is F-class (155°C). "
                    "Always verify against UL/IEC certification report for actual selection."
                )

        st.info(
            "**Insulation Selection Recommendation for Renewable Energy Converter Reactors** (reference IEC 60076-6, GB/T 1094.6): \n\n"
            "- Rated temperature ≤ 155°C (F-class): inter-layer insulation use **NMN 0.15mm**, core wrapping use **DMD 0.20mm**\n"
            "- Rated temperature ≤ 180°C (H-class): inter-layer and main insulation use **Nomex 410, 0.13~0.18mm**, thicken at air gap ends\n"
            "- High altitude (>2000m) or high humidity: increase main insulation thickness by ~20%, specify varnish/potting process in TDS\n"
            "- **Air gap side end**: winding end near air gap should be thickened to 0.3~0.5mm to reduce eddy-current hot spots from fringing field"
        )

        st.divider()

        # ══════════════════════════════════════════
        # : 
        # ══════════════════════════════════════════
        st.markdown("#### 🔌 Winding Conductor Materials — Selection of Copper Foil, Aluminum Foil, and Wire")
        st.caption("Winding material directly affects copper loss, weight, cost, and manufacturing process — the second most important design decision after the core.")

        winding_data = {
            "Conductor Type": [
                "Copper Foil (Cu Foil)",
                "Aluminum Foil (Al Foil)",
                "Solid Copper Wire (Solid Cu Wire)",
                "Copper Rectangular Wire (Cu Rectangular Wire)",
                "Litz Wire",
            ],
            "Resistivity (Ω·m)": [
                "1.72×10⁻⁸",
                "2.83×10⁻⁸",
                "1.72×10⁻⁸",
                "1.72×10⁻⁸",
                "1.72×10⁻⁸ (multi-strand)",
            ],
            "Density (g/cm³)": ["8.9", "2.7", "8.9", "8.9", "8.9"],
            "Advantages": [
                "High fill factor (>85%), low copper loss, large heat dissipation area, neat winding",
                "Weight only ~1/3 of copper, cost ~1/4 of copper, preferred for large-power lightweight designs",
                "Mature process, easy soldering, suitable for small cross-section / many turns",
                "High fill factor, suitable for high-current low-turns applications, good mechanical strength",
                "Multi-strand stranded wire, effectively suppresses skin effect, very low high-frequency copper loss",
            ],
            "Disadvantages": [
                "Heavy, high cost (about 4× aluminum)",
                "Same resistance requires ~1.7× cross-section; aluminum surface oxidizes easily requiring protection; welding process more demanding",
                "Single wire diameter limited by skin depth (≤2.6mm when fsw>1kHz), large high-frequency loss",
                "Difficult to bend, stress concentration at corners",
                "High cost, low fill factor (~65%), advantage only at high frequency",
            ],
            "Applicable Scenarios": [
                "Mainstream selection for renewable energy converter reactors (both L1 and L2)",
                "Large-power lightweight applications, cost-sensitive projects",
                "Auxiliary power transformers, low-power filter inductors (I<100A)",
                "High-power DC inductors, bus reactors (I>500A, low turns)",
                "High-frequency reactors with fsw>5kHz requiring low high-frequency loss",
            ],
        }
        st.dataframe(pd.DataFrame(winding_data), use_container_width=True, hide_index=True)

        if is_trial:
            trial_lock("Unlock with paid access: winding conductor selection recommendation (copper foil / aluminum foil trade-off and skin effect analysis)", compact=True)
        else:
            col_cu_al1, col_cu_al2 = st.columns(2)
            with col_cu_al1:
                st.success(
                    "**✅ Copper foil vs. aluminum foil: how to choose?**\n\n"
                    "- Power < 500 kW or space-sensitive → prefer copper foil (higher fill factor, lower copper loss)\n"
                    "- Power > 1 MW or weight reduction needed → consider aluminum foil (≈40% lighter for same inductance)\n"
                    "- Aluminum foil terminals can use **aluminum terminal bars** directly or connect via **copper-aluminum transition plates**;\n"
                    "  both are acceptable — choose based on downstream busbar material; key: **contact surface must be oxidation-protected**\n"
                    "- Aluminum foil welding requires inert gas protection (TIG) — more complex than copper\n"
                    "- Reference: GB/T 5585 (*Copper and Aluminum Busbars Standard*); IEC 62271 series"
                )
            with col_cu_al2:
                st.warning(
                    "**⚠ Skin effect impact on conductor selection**\n\n"
                    "- At fsw = 2500 Hz, copper skin depth ≈ 1.33 mm\n"
                    "- Recommend copper foil thickness ≤ 2×δ_skin = 2.66 mm (for effective high-frequency utilization)\n"
                    "- Aluminum skin depth ≈ 1.70 mm (higher resistivity) — slightly more relaxed than copper\n"
                    "- Foil thickness beyond the skin depth barely reduces copper loss — only adds weight and cost"
                )

        # ══════════════════════════════════════════
        # ══════════════════════════════════════════
        # : 
        # ══════════════════════════════════════════
        st.markdown("#### 💡 Design Notes")
        col_e1, col_e2 = st.columns(2)
        with col_e1:
            st.warning(
                "**📌 Note 1: Core material selection must consider both frequency and cost**\n\n"
                "From a core loss perspective, higher switching frequency calls for thinner laminations or powder core materials. "
                "However, in actual engineering, cost pressure often leads to non-oriented silicon steel (e.g. 35WW300), "
                "which is a reasonable engineering trade-off — the price is higher core loss and temperature rise, "
                "so design must verify that temperature rise still stays within the insulation class limit with adequate thermal margin."
            )
        if is_trial:
            trial_lock("Unlock with paid access: key electromagnetic design notes (magnetic saturation / air gap / assembly points)", compact=True)
        else:
            with col_e1:
                st.warning(
                    "**📌 Note 2: Small air gap easily causes magnetic saturation**\n\n"
                    "Air gap too small → peak current drives B past B_sat → inductance drops abruptly → ripple current out of control.\n"
                    "During testing this typically shows as normal inductance at low current but sudden ripple spike at high current.\n\n"
                    "Recommend designing with B_peak ≤ 0.75×B_sat to maintain adequate margin."
                )
            with col_e2:
                st.warning(
                    "**📌 Note 3: Saturation risk must be evaluated alongside protection capability**\n\n"
                    "When saturation occurs, the abrupt inductance drop causes current to rise rapidly. "
                    "Whether the inverter over-current protection (OCP) can respond in time directly determines whether the IGBT survives. "
                    "Design must consider both the core's saturation margin and the inverter's protection response speed — "
                    "only together do they truly protect the power device."
                )
                st.warning(
                    "**📌 Note 4: Air gap fringing increases local winding loss**\n\n"
                    "Magnetic flux at the air gap does not travel straight up-and-down — it bulges outward (fringing effect). "
                    "These fringing field lines, if they pass through winding conductors near the air gap, "
                    "induce additional eddy currents in the conductors, creating local hot spots and potentially damaging insulation. "
                    "When designing, keep winding turns away from the air gap, or use an air gap correction formula to compensate."
                )

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 2 —   [M3]
    # ══════════════════════════════════════════════════════════════════════════
    with tab_design:
        st.markdown("### 🛠️ Dry-Core Reactor Core Design Calculation (Three-Phase Dry Core)")
        st.markdown("#### Basic Design Input Parameters")

        col_di1, col_di2, col_di3 = st.columns(3)
        with col_di1:
            d9_target = st.selectbox("Design Target Inductor", ["L1 (Converter Side)", "L2 (Grid Side)"], key="d9_target")
            d9_L_mH = st.number_input(
                "Target Inductance L (mH)",
                value=_L1_mH if "L1" in d9_target else _L2_mH,
                min_value=0.01, step=0.01, format="%.3f", key="d9_L"
            )
            d9_I_rms = st.number_input(
                "Rated Current Irms (Arms)",
                value=_IL_rms, min_value=1.0, step=10.0, key="d9_I",
                help="Line rated RMS current, continuous operating condition"
            )
            d9_Ipeak = st.number_input(
                "Hardware OCP Peak Ipeak (Apk)",
                value=round(_IL_rms * np.sqrt(2) * 2.0, 1),
                min_value=1.0, step=10.0, key="d9_Ipk",
                help="Maximum peak current when IGBT hardware over-current protection triggers (Apk). Core air gap must not saturate at this current. Typically 2×Irms×√2."
            )
        with col_di2:
            _L_for_rip = (_L1_mH if "L1" in st.session_state.get("d9_target","L1 (Converter Side)") else _L2_mH) * 1e-3
            _Irip_default = round(_Vdc / (4.0 * max(_L_for_rip, 1e-5) * _fsw), 1)
            d9_Iripple = st.number_input(
                "Ripple Current Peak-to-Peak ΔI (App)",
                value=min(_Irip_default, d9_I_rms * 0.3),
                min_value=0.1, step=1.0, key="d9_Irip",
                help="Reactor ripple current peak-to-peak value (Apk-pk). Topology coefficient varies (2-Level: 1/4, 3-Level NPC: ~1/8, SVPWM 3-Level: ~1/12). Use simulation or Module 6 actual result — do not rely on auto-calculation."
            )
            d9_fsw = st.number_input("Switching Frequency fsw (Hz)", value=_fsw, min_value=100.0, step=100.0, key="d9_fsw")
            d9_fg = st.number_input("Fundamental Frequency fg (Hz)", value=_fg, min_value=1.0, step=1.0, key="d9_fg")
            d9_Vline = st.number_input("System Line Voltage Vrms (V)", value=_Vg_line, step=10.0, key="d9_Vline")
        with col_di3:
            d9_ins_class = st.selectbox(
                "Insulation Class",
                ["H (180°C, ≤125K)", "F (155°C, ≤100K)", "B (130°C, ≤80K)"],
                key="d9_ins"
            )
            d9_ins_paper = st.selectbox(
                "Insulation Paper",
                ["Nomex 410, 0.13mm (H-class)",
                 "Nomex 410, 0.18mm (H-class, thickened)",
                 "NMN, 0.15mm (H-class, UL1446 certification required)",
                 "NHN, 0.13mm (C-class, >220°C)",
                 "AMA (mica composite, H-class, partial discharge resistant)",
                 "DMD, 0.20mm (B/F-class, low cost)"],
                key="d9_ins_paper",
                help="NMN rated H-class requires UL 1446 insulation system certification. Reference: IEC 60085; IEC 60626-2"
            )
            d9_conductor = st.selectbox(
                "Winding Conductor",
                ["Copper Foil (Cu Foil)", "Aluminum Foil (Al Foil)",
                 "Solid Copper Wire (Solid Cu)", "Copper Rectangular Wire (Cu Rect)", "Litz Wire"],
                key="d9_cond"
            )
            d9_cooling = st.selectbox(
                "Cooling Method",
                ["AN (Natural Air Cooling)", "AF (Forced Air Cooling)", "WF (Water Cooling)"],
                key="d9_cooling"
            )
            if "AF" in st.session_state.get("d9_cooling", "AF (Forced Air Cooling)"):
                d9_wind_v = st.number_input(
                    "Air Velocity v (m/s)",
                    value=3.0, min_value=0.5, max_value=15.0, step=0.5,
                    key="d9_wind_v",
                    help="Forced air cooling wind speed, directly affects convective heat transfer coefficient. 1m/s≈weak, 3m/s≈moderate, 5m/s≈strong. Reference: Incropera §7.2"
                )
            else:
                d9_wind_v = 0.0

        Irip9  = d9_Iripple   # Use user input directly; do not auto-calculate

        # Derived constants
        mu_0   = 4 * np.pi * 1e-7
        d9_L_H = d9_L_mH * 1e-3
        rho9 = 1.72e-8 if "Copper" in d9_conductor else 2.83e-8
        dens9 = 8900 if "Copper" in d9_conductor else 2700
        is_foil9 = "Foil" in d9_conductor
        J9_map = {"AN (Natural Air Cooling)": 3.0, "AF (Forced Air Cooling)": 4.5, "WF (Water Cooling)": 6.5}
        J9     = J9_map.get(d9_cooling, 3.0)
        ins9_map = {
            "Nomex 410, 0.13mm (H-class)": 0.13,
            "Nomex 410, 0.18mm (H-class, thickened)": 0.18,
            "NMN, 0.15mm (H-class, UL1446 certification required)": 0.15,
            "NHN, 0.13mm (C-class, >220°C)": 0.13,
            "AMA (mica composite, H-class, partial discharge resistant)": 0.20,
            "DMD, 0.20mm (B/F-class, low cost)": 0.20,
        }
        ins9_t = ins9_map.get(d9_ins_paper, 0.15)
        T9_hot_map = {"H (180°C, ≤125K)": 155, "F (155°C, ≤100K)": 130, "B (130°C, ≤80K)": 100}
        T9_lim_map = {"H (180°C, ≤125K)": 125, "F (155°C, ≤100K)": 100, "B (130°C, ≤80K)": 80}
        T9_hot = T9_hot_map.get(d9_ins_class, 155)
        T9_lim = T9_lim_map.get(d9_ins_class, 125)

        # ── Core material database
        if "L1" in d9_target:
            mat9_opts = {
                "Sendust Powder Core (B_sat=1.05T, low high-freq loss)":       {"Bsat":1.05,"Cm":0.008,"al":1.5,"be":1.8,"dens":6.8, "cost":2.0,"kBr":0.70},
                "Non-oriented Silicon Steel 35WW300 (B_sat=1.7T, lowest cost)": {"Bsat":1.70,"Cm":0.055,"al":1.35,"be":1.9,"dens":7.65,"cost":1.0,"kBr":0.75},
                "Grain-oriented Silicon Steel 0.23mm B23R085 (B_sat=1.8T, thin-lamination low loss)": {"Bsat":1.80,"Cm":0.020,"al":1.4,"be":1.9,"dens":7.65,"cost":1.2,"kBr":0.75},
                "Amorphous Alloy 1K101 (B_sat=1.56T, wide-band low loss)":     {"Bsat":1.56,"Cm":0.012,"al":1.5,"be":1.8,"dens":7.1, "cost":2.5,"kBr":0.72},
            }
            mat9_def = "Sendust Powder Core (B_sat=1.05T, low high-freq loss)"
        else:
            mat9_opts = {
                "Grain-oriented Silicon Steel 0.35mm B27R095 (B_sat=2.0T, L2 standard)": {"Bsat":2.00,"Cm":0.060,"al":1.3,"be":1.9,"dens":7.65,"cost":1.1,"kBr":0.75},
                "Non-oriented Silicon Steel 35WW300 (B_sat=1.7T, lowest cost)":           {"Bsat":1.70,"Cm":0.055,"al":1.35,"be":1.9,"dens":7.65,"cost":1.0,"kBr":0.75},
                "Amorphous Alloy 1K101 (B_sat=1.56T, wide-band low loss)":                {"Bsat":1.56,"Cm":0.012,"al":1.5,"be":1.8,"dens":7.1, "cost":2.5,"kBr":0.72},
            }
            mat9_def = "Grain-oriented Silicon Steel 0.35mm B27R095 (B_sat=2.0T, L2 standard)"

        st.divider()

        # ── All sub-tabs share calculations (executed before st.tabs() to avoid cross-tab NameError)
        _mk9   = st.session_state.get("d9_cmat", mat9_def)
        if _mk9 not in mat9_opts: _mk9 = mat9_def
        mat9   = mat9_opts[_mk9]
        Bsat9  = mat9["Bsat"]; kBr9 = mat9["kBr"]
        kB9    = float(st.session_state.get("d9_kB", kBr9))
        Bop9   = Bsat9 * kB9

        dsk9   = np.sqrt(rho9 / (np.pi * d9_fsw * mu_0)) * 1000   #  mm @fsw
        k_sf9  = 0.96

        # ── : Tab①
        w9     = float(st.session_state.get("d9_core_w",   65.0))
        d9s    = float(st.session_state.get("d9_core_d",  215.0))
        lw9_cm = float(st.session_state.get("d9_lw_cm",   28.5))
        Ae9    = w9 * d9s * k_sf9 / 100.0
        Ae9_m2 = Ae9 * 1e-4
        lw9_m  = lw9_cm * 1e-2

        # ── Turns: designed for rated peak current (Irms×√2), OCP current only for short-time verification
        # Rated peak = Irms×√2 + ΔI/2
        I9_rated_pk = d9_I_rms * np.sqrt(2) + Irip9 / 2.0
        N9     = max(1, round(d9_L_H * I9_rated_pk / (Bop9 * Ae9_m2)))
        # OCP peak verification (allow short-time saturation, since IGBT Desat protection blocks in μs)
        B9_rated_pk = d9_L_H * I9_rated_pk / (N9 * Ae9_m2) if (N9 * Ae9_m2) > 0 else 0.0
        B9_ocp_pk   = d9_L_H * d9_Ipeak     / (N9 * Ae9_m2) if (N9 * Ae9_m2) > 0 else 0.0

        dlt9_b   = mu_0 * N9**2 * Ae9_m2 / d9_L_H if d9_L_H > 0 else 1e-3
        dlt9_bmm = dlt9_b * 1000
        if dlt9_b > 0 and Ae9_m2 > 0 and lw9_m > dlt9_b:
            F9       = 1.0 + (dlt9_b / np.sqrt(Ae9_m2)) * np.log(2 * lw9_m / dlt9_b)
            dlt9_c   = dlt9_b / F9; dlt9_cmm = dlt9_c * 1000; frg9 = (F9-1)*100
        else:
            F9=1.0; dlt9_c=dlt9_b; dlt9_cmm=dlt9_bmm; frg9=0.0
        Bpk9 = B9_rated_pk  # for display (rated operating condition)

        # ── Foil thickness selection: standard series, skin depth constraint
        # Standard thickness series (common supplier specifications)
        foil_t_std_cu = [0.5, 0.8, 1.0, 1.2, 1.5, 2.0, 2.5, 3.0]   # copper foil mm
        foil_t_std_al = [0.5, 0.8, 1.0, 1.2, 1.4, 1.5, 2.0, 2.5, 3.0]  # aluminum foil mm
        foil_t_sel_key = "d9_foil_t_sel"
        _foil_t_std = foil_t_std_cu if "Copper" in d9_conductor else foil_t_std_al
        # Skin depth constraint: t ≤ 2×δ_skin
        _foil_t_max = 2 * dsk9
        _foil_t_rec = next((t for t in reversed(_foil_t_std) if t <= _foil_t_max), _foil_t_std[0])
        foil9_t = float(st.session_state.get(foil_t_sel_key, _foil_t_rec))

        # ── Foil width and window utilization (geometric relationship correction)
        # Coordinate system (consistent with TAB1 diagram):
        #   Z-axis = window height direction lw → foil width w_foil along Z-axis, constrained by lw
        #   X-axis = stack depth direction → turns stacked along X-axis, occupying winding total thickness N×(t+t_ins)
        # End insulation clearance: H-class ≥5mm/side, F-class ≥3mm/side (IEC 60076-6 §6.4)
        ins_edge9 = 10.0 if "H" in d9_ins_class else 6.0  # Z-axis both sides total mm
        foil9_w_max = lw9_cm * 10 - ins_edge9              # Z-axis maximum foil width
        # Foil width = window height - end clearance, aligned to 5mm steps
        # Target Z-axis utilization 80~90% (leaving sufficient end insulation clearance)
        # Design directly to target value
        _foil_w_target_pct = 0.85  # target utilization 85%
        foil9_w = round(lw9_cm * 10 * _foil_w_target_pct / 5) * 5
        foil9_w = max(20.0, min(foil9_w, foil9_w_max))  # physical
        # ()
        S9        = foil9_t * foil9_w if is_foil9 else d9_I_rms / J9
        J9_actual = d9_I_rms / S9 if S9 > 0 else J9
        deq9      = 2 * np.sqrt(S9 / np.pi)
        h9pt      = foil9_t + ins9_t   # Zaxis mm

        # Window utilization has two dimensions:
        # Z-axis utilization (foil width direction): foil width / lw, target 80~90%
        wu9_z   = foil9_w / (lw9_cm * 10) * 100 if lw9_cm > 0 else 0.0
        # X-axis stacking (turns direction): N×(t+t_ins) is total winding thickness, for reference only, not constrained by lw
        wh9     = N9 * h9pt  # total winding thickness mm (X-axis direction)
        wu9_pct = wu9_z      # main window utilization uses Z-axis (foil width vs. window height)

        MLT9   = max(0.15, 2*(w9+d9s)/1000 + np.pi*(foil9_t+ins9_t)*N9/1000)
        ltot9  = N9 * MLT9

        kt9    = 1 + 0.00393*(T9_hot - 20)
        Rdc9   = rho9 * ltot9 / (S9*1e-6) if S9 > 0 else 0.0
        Rh9    = Rdc9 * kt9
        xi9    = foil9_t / (2*dsk9) if dsk9 > 0 else 1.0
        kac9   = max(1.0, xi9/2.0)
        Irip9r = (Irip9/2)/np.sqrt(3)
        Pcuf9  = Rh9 * d9_I_rms**2
        Pcuhf9 = kac9 * Rh9 * Irip9r**2
        Pcu9   = Pcuf9 + Pcuhf9

        Cm9=mat9["Cm"]; a9=mat9["al"]; b9=mat9["be"]
        Ve9=Ae9*lw9_cm*4; Ve9m=Ve9*1e-6
        # ── Core Loss (Steinmetz) correction notes:
        # Fundamental core loss: B_m is the actual flux amplitude at rated 50Hz condition (not design max safety flux density B_op)
        #   B_m = L×Irms×√2 / (N×Ae)  Reference: Steinmetz Trans.AIEE 1892; IEC 60404-2
        # Ripple core loss: Steinmetz B_m refers to AC flux amplitude (half of peak-to-peak ΔB/2)
        #   Using full peak-to-peak ΔB directly overestimates by ~2^1.9 ≈ 3.7× due to β≈1.9 exponent
        B9_m   = d9_L_H * d9_I_rms * np.sqrt(2) / (N9 * Ae9_m2) if (N9 * Ae9_m2) > 0 else 0.0
        dB9    = d9_L_H * Irip9 / (N9 * Ae9_m2) if (N9 * Ae9_m2) > 0 else 0.0  # peak-to-peak value
        dB9_amp = dB9 / 2.0  # amplitude (Steinmetz requires amplitude input, not peak-to-peak)
        Pfef9  = Cm9 * (d9_fg  / 1000.0) ** a9 * B9_m  ** b9 * Ve9m * 1e6
        Pfesw9 = Cm9 * (d9_fsw / 1000.0) ** a9 * dB9_amp ** b9 * Ve9m * 1e6 if dB9_amp > 0 else 0.0
        Pfe9   = Pfef9 + Pfesw9
        Ptot9  = Pcu9 + Pfe9
        cufr9  = Pcu9 / max(Pfe9, 0.01)

        # Heat dissipation area: three-phase three-limb, per-phase coil lateral area ≈ π×(w+d)×lw (equivalent rectangular cross-section perimeter × height)
        # Reference: Incropera "Fundamentals of Heat and Mass Transfer" 7th ed. §3.3
        Asurf9_per_phase = np.pi * ((w9+d9s)/1000) * (lw9_cm/100)  # single-phase m²
        Asurf9 = Asurf9_per_phase * 3  # three-phase total
        # Convective heat transfer coefficient (Incropera "Fundamentals of Heat and Mass Transfer" 7th §7.2)
        # AF cooling: Nu method, characteristic length = coil height, air Pr=0.71, ν=1.6e-5 m²/s, k=0.028 W/(m·K)
        _v9   = float(st.session_state.get("d9_wind_v", 3.0)) if "AF" in d9_cooling else 0.0
        _L9   = max(lw9_m, 0.1)   # feature = 
        _nu9  = 1.6e-5; _k9 = 0.028; _Pr9 = 0.71
        if "WF" in d9_cooling:
            hconv9 = 300   # water cooling estimate (engineering experience)
        elif "AF" in d9_cooling and _v9 > 0:
            _Re9  = _v9 * _L9 / _nu9
            _Nu9  = (0.037 * _Re9**0.8 - 871) * _Pr9**(1/3) if _Re9 >= 5e5 else 0.664 * _Re9**0.5 * _Pr9**(1/3)
            hconv9 = max(10, _Nu9 * _k9 / _L9)
        else:
            hconv9 = 10    # AN natural cooling
        dT9    = Ptot9 / (hconv9 * max(Asurf9, 0.01))
        cool_ok9 = dT9 <= T9_lim
        cool_up9 = {"AN (Natural Air Cooling)":"AF (Forced Air Cooling)","AF (Forced Air Cooling)":"WF (Water Cooling)"}
        if not cool_ok9 and d9_cooling in cool_up9:
            hup9  = {"AF (Forced Air Cooling)":25,"WF (Water Cooling)":60}[cool_up9[d9_cooling]]
            dTup9 = Ptot9 / (hup9 * max(Asurf9, 0.01))
            upnote9 = f"upgrade to {cool_up9[d9_cooling]} → Temperature Rise={dTup9:.0f}K ({'✅ satisfied' if dTup9<=T9_lim else 'still needs optimization'})"
        else:
            dTup9=dT9; upnote9=""

        S9_term  = d9_I_rms / 1.5
        std9s    = [10,16,20,25,30,40,50,60,80,100,120,150,200,250,300,400]
        S9_std   = next((s for s in std9s if s >= S9_term), 400)
        clr9_map = [(d9_Vline<=660,12),(d9_Vline<=1000,16),(d9_Vline<=3000,25),(d9_Vline<=6000,40)]
        clr9     = next((v for c,v in clr9_map if c), 60)

        # ── :  d9_xxx  TAB4 
        d_L_mH       = d9_L_mH
        d_I_rms      = d9_I_rms
        d_Ipeak      = d9_Ipeak
        d_Vline      = d9_Vline
        d_fg         = d9_fg
        d_ins_class  = d9_ins_class
        d_conductor  = d9_conductor
        d_cooling    = d9_cooling
        d_target     = d9_target
        topo_choice  = "Three-phase (default)"
        N_final      = N9
        S_wire_mm2   = S9
        foil_thick_mm = foil9_t
        MLT_m        = MLT9
        l_wire_total  = ltot9
        core_width_mm = w9
        core_stack_mm = d9s
        lw_cm        = lw9_cm
        delta_corr_mm = dlt9_cmm
        R_hot        = Rh9
        k_ac         = kac9
        P_Cu_total   = Pcu9
        P_Fe_total   = Pfe9
        P_total      = Ptot9
        term_rec_type = ("Crimp Copper Terminal" if d9_I_rms<=400 else
                         "Copper Bus-bar (Cu Bus-bar) tin-plated" if d9_I_rms<=2000 else "Aluminum Bus-bar (Al Bus-bar) nickel-plated")
        J_term_rec    = (2.0 if d9_I_rms<=400 else 1.5 if d9_I_rms<=2000 else 1.2)
        term_rec_note = ""
        S_std         = S9_std
        clr_mm        = clr9
        # Legacy variable compatibility (for TAB4 Excel generation)
        _Iripple_pp      = Irip9          # Ripple Current Peak-to-Peak
        delta_skin_sw_mm = dsk9           # skin depth
        h_conv           = hconv9         # convective heat transfer coefficient
        A_surface_m2     = Asurf9         # heat dissipation area (three-phase total)
        delta_T          = dT9            # estimated temperature rise
        T_max_ins        = T9_lim         # insulation class temperature rise limit
        J_rated          = J9             # rated current density recommended value
        # Extra variables for Excel sheet
        topo_coeff_note  = "Three-phase (read topology from Module 6)"
        core_material    = _mk9           # core material name
        safety_factor    = kB9            # flux density safety factor
        delta_basic_mm   = dlt9_bmm       # theoretical air gap (before correction)
        foil_width_mm    = foil9_w        # foil width
        window_height_used = wh9          # winding occupied thickness (X-axis)
        t_ins            = ins9_t         # inter-layer insulation paper thickness
        # Core parameter aliases
        B_sat            = Bsat9
        B_op             = Bop9
        B_actual         = B9_rated_pk
        Ae_cm2           = Ae9
        k_sf             = k_sf9
        F_fringe         = F9
        # Loss parameter aliases
        Cm               = Cm9
        alpha_s          = a9
        beta_s           = b9
        Ve_cm3           = Ve9
        P_Cu             = Pcuf9
        P_Cu_hf          = Pcuhf9
        P_Fe_fund        = Pfef9
        P_Fe_sw          = Pfesw9
        # Winding/terminal parameter aliases
        J_rated_adj      = J9_actual
        S_term_mm2       = S9_term

        # ── Five sub-tabs
        t9_1, t9_2, t9_3, t9_4, t9_5 = st.tabs([
            "🧲 ① Core Design",
            "🔌 ② Winding Design",
            "🔥 ③ Loss Calculation",
            "🔩 ④ Terminal Design",
            "📊 ⑤ Design Safe Operating Area"
        ])

        if is_trial:
            with t9_1:
                st.markdown("#### 🧲 Core Design")
                trial_lock("Unlock with paid access: core turns, air gap, and flux density design calculation full process")
            with t9_2:
                st.markdown("#### 🔌 Winding Design")
                trial_lock("Unlock with paid access: winding layers, foil width, and current density calculation full process")
            with t9_3:
                st.markdown("#### 🔥 Loss Calculation")
                trial_lock("Unlock with paid access: Dowell copper loss + Steinmetz core loss calculation and balance analysis")
            with t9_4:
                st.markdown("#### 🔩 Terminal Design")
                trial_lock("Unlock with paid access: terminal cross-section and temperature rise design")
            with t9_5:
                st.markdown("#### 📊 Design Safe Operating Area")
                # Build safe operating area chart (unconditional execution)
                Ae9r = np.linspace(Ae9*0.35, Ae9*2.5, 45)
                N9r  = np.arange(max(1, N9-8), N9+13, 1)
                AEg,NNg = np.meshgrid(Ae9r, N9r)
                AEgm = AEg * 1e-4
                Bg   = d9_L_H * d9_I_rms * np.sqrt(2) / (np.maximum(NNg,1) * np.maximum(AEgm,1e-6))
                lg   = NNg * MLT9
                Rg   = rho9 * lg / (S9*1e-6) * kt9
                Pcug = Rg * d9_I_rms**2 * kac9
                Veg  = AEgm * lw9_m * 4
                Pfeg = Cm9*(d9_fg/1000)**a9*np.minimum(Bg,Bsat9*0.99)**b9*Veg*1e6
                Ptg  = Pcug + Pfeg
                wtcg = Veg * mat9["dens"] * 1e3
                wtug = S9*1e-6 * lg * dens9
                wtg  = wtcg + wtug
                from plotly.subplots import make_subplots as _msp_t9
                fig9c = _msp_t9(
                    rows=1, cols=3,
                    subplot_titles=["① Peak Flux Density B_peak (T)","② Total Loss P_total (W)","③ Estimated Total Weight (kg)"],
                    horizontal_spacing=0.12
                )
                for col_i, (Z, csc) in enumerate([(Bg,"RdYlGn_r"),(Ptg,"YlOrRd"),(wtg,"Blues")], 1):
                    fig9c.add_trace(go.Contour(
                        z=Z, x=Ae9r, y=N9r,
                        colorscale=csc, showscale=True,
                        contours=dict(showlabels=True, labelfont=dict(size=11, color="white")),
                        line=dict(width=0.5),
                        colorbar=dict(len=0.8, thickness=14,
                                      x=0.33*col_i-0.005,
                                      tickfont=dict(size=10))
                    ), row=1, col=col_i)
                    fig9c.add_trace(go.Scatter(
                        x=[Ae9], y=[N9], mode="markers",
                        marker=dict(symbol="star", size=18, color="red",
                                    line=dict(width=2, color="white")),
                        showlegend=False
                    ), row=1, col=col_i)
                Bsat_x, Bsat_y = [], []
                for Av in Ae9r:
                    Nv = d9_L_H * d9_Ipeak / (Bsat9 * Av * 1e-4)
                    if N9r[0] <= Nv <= N9r[-1]:
                        Bsat_x.append(Av); Bsat_y.append(Nv)
                if Bsat_x:
                    fig9c.add_trace(go.Scatter(
                        x=Bsat_x, y=Bsat_y, mode="lines",
                        line=dict(color="red", dash="dash", width=3),
                        name=f"Saturation Boundary B_sat={Bsat9}T"
                    ), row=1, col=1)
                fig9c.update_layout(
                    height=500, template="plotly_white",
                    margin=dict(t=70, b=100, l=70, r=20),
                )
                # Trial version: hide axis labels and hover, show chart, add lock
                fig9c_trial = go.Figure(fig9c)
                fig9c_trial.update_traces(hoverinfo='none', hovertemplate=None)
                fig9c_trial.update_xaxes(showticklabels=False)
                fig9c_trial.update_yaxes(showticklabels=False)
                st.plotly_chart(fig9c_trial, use_container_width=True)
                trial_lock("Unlock with paid access: design scheme safe operating area visualization (flux density / loss / weight three-dimensional contour chart)", compact=True)
        else:

            # ══ ① Core Design ══════════════════════════════════════════════════════
            with t9_1:
                st.markdown("#### 🧲 Core Design")
                col_m1, col_m2 = st.columns(2)
                with col_m1:
                    if "L1" in d9_target:
                        st.warning(f"L1 converter-side reactor, fsw={d9_fsw:.0f}Hz — prioritize thin laminations or powder core to control core loss.")
                    else:
                        st.success("L2 grid-side reactor has very little high-frequency content — grain-oriented silicon steel 0.35mm is most cost-effective.")
                    st.selectbox("Core Material", list(mat9_opts.keys()),
                                 index=list(mat9_opts.keys()).index(_mk9), key="d9_cmat")
                    kBr_show = mat9_opts[st.session_state.get("d9_cmat", mat9_def)]["kBr"]
                    st.number_input(
                        f"Flux Density Safety Factor k_B (engineering recommendation: {kBr_show})",
                        value=kBr_show, min_value=0.50, max_value=0.90,
                        step=0.05, format="%.2f", key="d9_kB",
                        help="B_op=k_B×B_sat. Powder core: 0.70, silicon steel: 0.75, with DC bias: 0.65~0.70. McLyman §5.2"
                    )
                    st.info(f"B_op = {Bop9:.3f} T ({Bsat9}T × {kB9})")
                with col_m2:
                    st.markdown("**Core Dimensions**")
                    # ── Auxiliary input: installation dimensions → recommended core specifications
                    with st.expander("📐 Back-calculate recommended core specs from installation dimensions (optional)", expanded=False):
                        st.caption("Enter maximum reactor installation space; tool estimates recommended core specifications. Final parameters still require supplier confirmation.")
                        _c1, _c2, _c3 = st.columns(3)
                        with _c1:
                            _H_mm = st.number_input("Maximum Installation Height H (mm)", value=650.0, step=10.0, key="d9_H_mm",
                                                     help="Total height (Z-axis), including base insulator ~150mm")
                        with _c2:
                            _W_mm = st.number_input("Maximum Installation Width W (mm)", value=380.0, step=10.0, key="d9_W_mm",
                                                     help="Single-phase width (X-axis), including winding and insulation margin")
                        with _c3:
                            _D_mm = st.number_input("Installation Depth D (mm)", value=280.0, step=10.0, key="d9_D_mm",
                                                     help="Core depth direction (Y-axis), including winding lead-out margins on both sides")
                        # Estimation logic (engineering experience formulas)
                        # Lamination width standard series
                        _w_std = [40,50,60,70,80,90,100,110,120,130,140,150,160]
                        # Lamination width ≈ depth D / (2 + 1.5) empirical (two windings + core simplification)
                        _w_rec = max(40, min(160, round((_D_mm * 0.3) / 10) * 10))
                        _w_rec = min(_w_std, key=lambda x: abs(x - _w_rec))  # nearest standard value
                        # Yoke height ≈ lamination width
                        _yoke_h = _w_rec
                        # Window height lw = installation height - base - 2×yoke height
                        _lw_rec = max(100, _H_mm - 150 - 2*_yoke_h)
                        # Stack depth d = installation width - 2×(single-side winding thickness estimate) - 2×insulation margin
                        # Single-side winding thickness ≈ N×(t+t_ins) + 20mm insulation, N≈15 turns×1.55mm≈25mm
                        _winding_thick = max(20, N9 * (1.4+0.15) + 20)
                        _d_rec = max(50, _W_mm - 2*_winding_thick - 20)
                        _d_rec = round(_d_rec / 10) * 10
                        _rec_lines = [
                            "Recommended core specifications (for reference, confirm tooling with supplier):",
                            f"Lamination width w ≈ **{_w_rec} mm** (standard series)  |  Stack depth d ≈ **{_d_rec} mm**  |  Window height lw ≈ **{_lw_rec/10:.1f} cm**",
                            f"Effective cross-section Ae ≈ {_w_rec}×{_d_rec}×0.96 = {_w_rec*_d_rec*0.96/100:.1f} cm²",
                        ]
                        st.success("  \n".join(_rec_lines))
                        st.caption("Formula: lw ≈ H - base 150mm - 2×yoke height (≈ lamination width); d ≈ W - 2×winding thickness - insulation margin. IEC 60076-6 §6.3")

                    st.caption("Core limb lamination width/stack depth is determined by supplier tooling specifications — enter actual values or refer to recommendations above.")
                    col_cd1, col_cd2, col_cd3 = st.columns(3)
                    with col_cd1:
                        st.number_input("Lamination Width w (Y-axis, mm)", value=65.0, min_value=20.0, step=5.0, key="d9_core_w",
                                        help="Core limb cross-section width (Y-axis). Determined by punching tooling, typical range 60-150mm.")
                    with col_cd2:
                        st.number_input("Stack Depth d (X-axis, mm)", value=215.0, min_value=20.0, step=5.0, key="d9_core_d",
                                        help="Total lamination stack thickness (X-axis). Main adjustment variable for cross-section, typical range 100-300mm.")
                    with col_cd3:
                        st.number_input("Window Height lw (Z-axis, cm)", value=28.5, min_value=2.0, step=0.5, key="d9_lw_cm",
                                        help="Clear height between upper and lower yokes (Z-axis), determines foil width and winding space.")
                    st.info(
                        f"Effective cross-section Ae = w × d × k_sf = {w9:.0f} × {d9s:.0f} × {k_sf9} = **{Ae9:.1f} cm²**  "
                        f"|  Window height lw = {lw9_cm:.1f} cm ({lw9_cm*10:.0f} mm)"
                    )

                    # Result summary (displayed inside column)
                    N9_show = d9_L_H * I9_rated_pk / (Bop9 * Ae9_m2) if (Bop9*Ae9_m2)>0 else 1.0

                c1,c2,c3,c4,c5 = st.columns(5)
                c1.metric("Turns N",             f"{N9} turns")
                c2.metric("Ae",                   f"{Ae9:.1f} cm²")
                c3.metric("Theoretical Air Gap",  f"{dlt9_bmm:.2f} mm")
                c4.metric("Corrected Air Gap",    f"{dlt9_cmm:.2f} mm")
                c5.metric("Rated Flux Density",   f"{B9_rated_pk:.3f} T",
                           delta=f"Margin {(Bop9-B9_rated_pk)/Bop9*100:.1f}%",
                           delta_color="normal" if B9_rated_pk<=Bop9 else "inverse")

                if dlt9_cmm < 0.5:
                    st.warning("⚠ Air gap <0.5mm — recommend confirming machining capability with supplier (high grinding precision required).")
                if B9_rated_pk > Bsat9:
                    st.error("❌ Rated flux density exceeds B_sat — design infeasible! Increase Ae or reduce k_B.")
                elif B9_rated_pk > Bop9:
                    st.warning(f"⚠ Rated flux density {B9_rated_pk:.3f}T > B_op={Bop9:.3f}T — recommend reducing k_B or increasing Ae.")
                else:
                    st.success(f"✅ Rated flux density {B9_rated_pk:.3f}T ≤ B_op={Bop9:.3f}T — design safe. OCP peak {B9_ocp_pk:.3f}T allows short-time saturation (μs-level).")

                # Full-width calculation process
                with st.expander("📐 Core Calculation Full Process", expanded=True):
                    N9_show = d9_L_H * I9_rated_pk / (Bop9 * Ae9_m2) if (Bop9*Ae9_m2)>0 else 1.0
                    _calc_core = [
                        "**① Turns Calculation** (Faraday's Law; McLyman CRC Press 4th ed. §5.2; IEC 60076-6 §5)",
                        f"I_rated_pk = Irms×√2 + ΔI/2 = {d9_I_rms:.0f}×1.414 + {Irip9:.1f}/2 = **{I9_rated_pk:.1f} Apk** (rated continuous peak; OCP peak {d9_Ipeak:.0f}Apk for short-time verification only)",
                        f"N = L × I_rated_pk / (B_op × Ae) = {d9_L_mH:.3f}mH × {I9_rated_pk:.1f}Apk / ({Bop9:.3f}T × {Ae9:.1f}cm²) = **{N9_show:.1f} → rounded N = {N9} turns**",
                        "",
                        "**② Theoretical Air Gap** (McLyman CRC Press 4th ed. §14.4; GB/T 1094.6-2017 Appendix B)",
                        f"δ_basic = μ₀ × N² × Ae / L = 4π×10⁻⁷ × {N9}² × {Ae9:.1f}cm² / {d9_L_mH:.3f}mH = **{dlt9_bmm:.3f} mm**",
                        "",
                        "**③ McLyman Fringing Effect Correction** (Kazimierczuk M.K., \"High-Frequency Magnetic Components\", Wiley 2014 §3.4)",
                        "  Flux lines at the air gap bulge outward, effective cross-section is larger than core cross-section — without correction inductance is overestimated",
                        f"  F = 1 + (δ/√Ae) × ln(2lw/δ) = 1 + ({dlt9_bmm:.3f}mm/√{Ae9:.1f}cm) × ln(2×{lw9_cm*10:.0f}mm/{dlt9_bmm:.3f}mm) = **{F9:.4f}** (effective Ae expanded by {frg9:.1f}%)",
                        f"  δ_corr = δ_basic / F = {dlt9_bmm:.3f} / {F9:.4f} = **{dlt9_cmm:.3f} mm** ← physical slot dimension to specify to supplier",
                        "",
                        "**④ Flux Density Verification** (IEC 60076-6:2007 §3.1; GB/T 1094.6-2017 §7.1)",
                        f"  Rated condition (continuous): B_rated = L×I_rated_pk/(N×Ae) = {d9_L_mH:.3f}mH×{I9_rated_pk:.1f}Apk/({N9}×{Ae9:.1f}cm²) = **{B9_rated_pk:.4f} T**",
                        f"  → B_op = B_sat×k_B = {Bsat9}T×{kB9} = {Bop9:.3f}T, margin = {(Bop9-B9_rated_pk)/Bop9*100:.1f}%  {'✅ meets continuous operation requirement' if B9_rated_pk<=Bop9 else '⚠ exceeds B_op, adjustment needed'}",
                        f"  OCP protection (short-time μs-level): B_ocp = L×I_ocp/(N×Ae) = **{B9_ocp_pk:.4f} T**",
                        f"  → Short-time saturation permitted; IGBT Desat protection blocks within 2~10μs (IEC 60076-6 §3.1.3 filter reactor clause)",
                    ]
                    st.markdown("  \n".join(_calc_core))

            # ══ ② Winding Design ══════════════════════════════════════════════════════
            with t9_2:
                st.markdown("#### 🔌 Winding Design")

                # ── Foil thickness selection (standard series, skin depth constraint)
                _foil_std_list = foil_t_std_al if "Aluminum" in d9_conductor else foil_t_std_cu
                st.markdown(f"**① Foil Thickness Selection** (skin depth δ={dsk9:.3f}mm @{d9_fsw:.0f}Hz, recommend t ≤ 2δ={2*dsk9:.3f}mm)")
                st.caption(f"The following are supplier standard specification series; specifications exceeding 2δ increase high-frequency copper loss. Reference: Dowell IEE Proc.1966 Vol.113 No.8")

                st.selectbox(
                    "Select Foil Thickness (mm)",
                    options=[str(t) for t in _foil_std_list],
                    index=_foil_std_list.index(foil9_t) if foil9_t in _foil_std_list else 0,
                    key="d9_foil_t_sel",
                    format_func=lambda x: f"{x}mm {'← recommended (meets skin depth constraint)' if float(x)<=2*dsk9 else '⚠ exceeds 2δ, higher high-freq copper loss'}",
                    help="Aluminum/copper foil standard thickness series, determined by supplier tooling. t≤2δ_skin satisfies skin effect constraint."
                )

                st.divider()
                # ── Winding design results summary
                st.markdown("**② Winding Design Results Summary**")

                # Explicitly display key conclusions
                res_col1, res_col2, res_col3, res_col4, res_col5 = st.columns(5)
                res_col1.metric("Turns N", f"{N9} turns",
                                 help=f"N = L×Ipk_rated/(B_op×Ae), rated peak={I9_rated_pk:.0f}Apk")
                res_col2.metric("Foil Thickness t", f"{foil9_t:.1f} mm")
                res_col3.metric("Foil Width w", f"{foil9_w:.0f} mm",
                                 help=f"foil width = min(lw - end clearance, S/t), end clearance {ins_edge9:.0f}mm")
                res_col4.metric("Actual Cross-section S", f"{S9:.1f} mm²",
                                 help=f"S = t×w = {foil9_t:.1f}×{foil9_w:.0f}")
                res_col5.metric("Actual Current Density J", f"{J9_actual:.2f} A/mm²")

                st.markdown(f"""
| Parameter | Value | Description |
|------|-----|------|
| **Turns N** | **{N9} turns** | L×I_rated_pk/(B_op×Ae); I_rated_pk={I9_rated_pk:.0f}Apk |
| **Foil Thickness t** | **{foil9_t:.1f} mm** | User selection (≤2δ={2*dsk9:.2f}mm, skin depth constraint) |
| **Foil Width w** | **{foil9_w:.0f} mm** | Target 85%×lw={lw9_cm*10:.0f}mm, end clearance {(lw9_cm*10-foil9_w)/2:.1f}mm each side |
| **Conductor Cross-section S** | **{S9:.1f} mm²** | t×w = {foil9_t:.1f}×{foil9_w:.0f} |
| Actual Current Density J | {J9_actual:.2f} A/mm² | I/S ({d9_cooling} reference limit: {J9:.1f}A/mm²) |
| Inter-layer Insulation Paper | {ins9_t} mm | {d9_ins_paper.split(", ")[0]} |
| Per-turn Thickness (X-axis) | {h9pt:.2f} mm | Foil thickness + insulation paper thickness |
| Total Winding Thickness (X-axis) | {wh9:.1f} mm | N×{h9pt:.2f}mm = {N9} turns stacked |
| Z-axis Window Utilization | {wu9_z:.0f}% | {foil9_w:.0f}/{lw9_cm*10:.0f}mm ({"✅" if wu9_z<=92 else "⚠ high"}) |
| Mean Turn Length MLT | {MLT9*1000:.0f} mm | 2(w+d)+π(t+t_ins)N, McLyman §9.1 |
| Total Wire Length | {ltot9:.2f} m | N×MLT |
""")

                if wu9_z < 70:
                    st.info(f"ℹ Window utilization: {wu9_z:.0f}%. Design uses 85% of lw as target — adjust lw if needed.")
                if J9_actual > J9 * 1.5:
                    st.error(f"❌ Current density {J9_actual:.2f} A/mm² exceeds {d9_cooling} limit {J9} A/mm²×1.5 — excessive temperature rise expected.")
                elif J9_actual > J9 * 1.1:
                    st.warning(f"⚠ Current density J={J9_actual:.2f} A/mm² slightly above {d9_cooling} recommendation of {J9} A/mm² — monitor temperature rise.")

                st.divider()
                st.markdown("**③ Inter-layer Insulation Paper Notes**")
                # Inter-layer voltage calculation (per-turn induced voltage)
                _V_layer = d9_L_H * (d9_I_rms * np.sqrt(2)) * d9_fsw * 2 / N9 if N9 > 0 else 0.0
                # Inter-layer peak voltage ≈ L × ΔI_pp × fsw / N (simplified Faraday estimate)
                # Inter-layer voltage: V = L×di/dt, triangular wave di/dt = ΔI/(1/(2fsw)) = 2×ΔI×fsw
                # Reference: Kazimierczuk "High-Frequency Magnetic Components" §1.2
                _V_layer2 = d9_L_H * Irip9 / N9 * d9_fsw * 2.0 if N9 > 0 else 0.0
                _ins_rows = {
                    "H-class (180°C)": ["Nomex 410 (aramid paper) 0.13mm", "NHN (Nomex-Kapton-Nomex) 0.13mm", "NMN 0.15mm (requires UL1446 H-class certification)", "AMA/NMA (mica composite) 0.20mm"],
                    "F-class (155°C)": ["NMN 0.15mm (standard F-class)", "DMD 0.20mm (low cost)", "Nomex 410 0.13mm (downgraded use)"],
                    "B-class (130°C)": ["DMD 0.20mm", "DMDM 0.25mm"],
                }
                _cur_grade = next((k for k in _ins_rows if k.split(" (")[0].lower().replace("-class","") in d9_ins_class.lower()), "H-class (180°C)")
                _grade_options = _ins_rows[_cur_grade]
                _ins_tip_lines = [
                    f"**Current Selection**: {d9_ins_paper} (thickness {ins9_t}mm)",
                    "",
                    f"**Inter-layer Voltage Estimate** (determines minimum insulation thickness):",
                    f"- Inter-layer induced voltage = L×(2ΔI×fsw)/N = {d9_L_H*1000:.3f}mH × 2×{Irip9:.1f}App × {d9_fsw:.0f}Hz / {N9} = **{_V_layer2:.1f} V** (V=L×di/dt, triangular wave di/dt=2ΔI·fsw)",
                    f"- {ins9_t}mm insulation paper withstands ≥ 500V (Nomex/NMN), far exceeding {_V_layer2:.1f}V ✅",
                    f"- Insulation paper selection principle: ① material type determined by insulation class (H-class: select Nomex/NMN/AMA etc.);"
                    f"  ② thickness determined by safety electrical clearance requirements (IEC 60664-1); inter-layer voltage {_V_layer2:.1f}V is well below any spec paper withstand voltage (>500V),"
                    f"  minimum mechanical grade of 0.10-0.20mm is sufficient",
                    "",
                    f"**{_cur_grade} available insulation papers** (IEC 60085; IEC 60626-2):",
                ] + [f"- {opt}" for opt in _grade_options] + [
                    "",
                    f"**End Insulation Clearance** ({ins_edge9/2:.0f}mm each side):",
                    f"- Foil width {foil9_w:.0f}mm (target 85%×lw={lw9_cm*10:.0f}mm), end clearance {(lw9_cm*10-foil9_w)/2:.1f}mm each side ✅",
                ]
                st.info("\n".join(_ins_tip_lines))

                with st.expander("📐 Winding Calculation Full Process", expanded=True):
                    _t_ok = "✅" if foil9_t <= 2*dsk9 else "⚠ exceeds 2δ limit"
                    _calc_lines = [
                        "**① Turns** (rated current peak design; OCP peak for short-time verification only)",
                        f"I_rated_pk = Irms×√2 + ΔI/2 = {d9_I_rms:.0f}×1.414 + {Irip9:.1f}/2 = **{I9_rated_pk:.1f} Apk**",
                        f"N = L×I_rated_pk / (B_op×Ae) = {d9_L_mH:.3f}mH×{I9_rated_pk:.1f}Apk / ({Bop9:.3f}T×{Ae9:.1f}cm²) = **{N9} turns**",
                        "",
                        "**② Skin Depth Constraint** (Dowell IEE Proc.1966 Vol.113 No.8)",
                        f"δ_skin = √(ρ/(π·f·μ₀)) = **{dsk9:.3f}mm** @{d9_fsw:.0f}Hz, 2δ = {2*dsk9:.3f}mm",
                        f"t_foil = {foil9_t:.1f}mm ≤ 2δ = {2*dsk9:.3f}mm  {_t_ok}",
                        "",
                        "**③ Foil Width** (Z-axis direction, constrained by core window height lw; IEC 60076-6 §6.4)",
                        f"Foil width = lw × 85% = {lw9_cm*10:.0f} × 0.85 = {lw9_cm*10*0.85:.0f}mm → use **{foil9_w:.0f}mm** (5mm step alignment, target 85% utilization)",
                        f"Actual cross-section S = t×w = {foil9_t:.1f}×{foil9_w:.0f} = **{S9:.0f} mm²**",
                        f"Actual current density J = I/S = {d9_I_rms:.0f}/{S9:.0f} = **{J9_actual:.3f} A/mm²**",
                        "",
                        "**④ Mean Turn Length MLT** (McLyman §9.1)",
                        f"MLT ≈ 2(w+d) + π(t+t_ins)×N = 2({w9:.0f}+{d9s:.0f}) + π×{foil9_t+ins9_t:.2f}×{N9} ≈ **{MLT9*1000:.0f} mm/turn**",
                        f"Total wire length = {N9}×{MLT9*1000:.0f}mm = **{ltot9:.2f} m**",
                        "",
                        "**⑤ Window Utilization**",
                        f"Z-axis (foil width direction): {foil9_w:.0f}/{lw9_cm*10:.0f}mm = **{wu9_z:.0f}%** (target 80~90%, foil width entirely determined by lw)",
                        f"X-axis (turns stacking direction): N×(t+t_ins) = {N9}×{h9pt:.2f} = **{wh9:.1f}mm** total winding thickness",
                    ]
                    st.markdown("  \n".join(_calc_lines))

            # ══ ③ Loss Calculation ══════════════════════════════════════════════════════
            with t9_3:
                st.markdown("#### 🔥 Loss Calculation")

                if cufr9 > 3.0:
                    st.warning(
     f"⚠ Copper Loss ({Pcu9:.0f} W) / Core Loss ({Pfe9:.0f} W) = {cufr9:.1f} — Copper Loss significantly exceeds Core Loss.\n\n"
     f"**Optimization Suggestion**: Increase Ae (reduces Turns, lowers Copper Loss) or reduce current density J (increase conductor cross-section).\n"
     f"Target: Copper Loss ≈ Core Loss (balanced loss design). Ref: McLyman §1.3"
                    )
                elif cufr9 < 0.33:
                    st.warning(
     f"⚠ Core Loss ({Pfe9:.0f} W) / Copper Loss ({Pcu9:.0f} W) = {1/cufr9:.1f} — Core Loss significantly exceeds Copper Loss.\n\n"
     f"**Optimization Suggestion**: Reduce Ae (increases Turns), to lower Core Loss."
                    )
                else:
                    st.success(f"✅ Copper Loss ({Pcu9:.0f} W) ≈ Core Loss ({Pfe9:.0f} W) — loss distribution is balanced and optimized.")

                    with st.expander("📐 Copper Loss Calculation (Dowell Model)", expanded=True):
                        st.markdown(f"""
    **DC Resistance** (IEC 60028; ρ_Cu=1.72×10⁻⁸ Ω·m, ρ_Al=2.83×10⁻⁸ Ω·m)
    $$R_{{dc}} = \\rho l/S = {rho9:.2e}\\times{ltot9:.2f}\\,\\text{{m}} / ({S9:.2f}\\times10^{{-6}}\\,\\text{{m}}^2) = {Rdc9*1000:.4f}\\,\\text{{mΩ}}$$

    **Hot Resistance** (temperature coefficient 0.00393/°C; IEC 60028)
    $$R_{{hot}}({T9_hot}°C) = R_{{dc}}[1+0.00393({T9_hot}-20)] = {Rh9*1000:.4f}\\,\\text{{mΩ}}$$

    **Fundamental Copper Loss**
    $$P_{{Cu,f}} = R_{{hot}}\\cdot I_{{rms}}^2 = {Pcuf9:.1f}\\,\\text{{W}}$$

    **High-Frequency Copper Loss** (Dowell IEE Proc.1966 Vol.113 No.8)
    $$\\xi={foil9_t if is_foil9 else deq9:.2f}/(2\\times{dsk9:.3f})={xi9:.3f},\\;k_{{ac}}={kac9:.3f},\\;\\Delta P_{{Cu}}={Pcuhf9:.1f}\\,\\text{{W}}$$
    $$\\boxed{{P_{{Cu}} = {Pcuf9:.1f}+{Pcuhf9:.1f} = {Pcu9:.1f}\\,\\text{{W}}}}$$
    """)

                    with st.expander("📐 Core Loss Calculation (Steinmetz Equation)", expanded=True):
                        _fe_lines = [
                            f"**Steinmetz Equation** (Steinmetz Trans.AIEE 1892; IEC 60404-2; manufacturer datasheet)",
                            f"P_Fe = Cm × f^α × B_m^β × Ve  — where Cm={Cm9}, α={a9}, β={b9}",
                            f"**Note**: B_m is flux density amplitude (peak value). Two components: rated operating B_op (from fundamental current), and ripple ΔB (from switching harmonic current).",
                            "",
                            f"Core Volume Ve = Ae×lw×4 = {Ae9:.1f}×{lw9_cm:.1f}×4 = **{Ve9:.0f} cm³** (three-phase, 4 limbs)",
                            "",
                            f"**Fundamental Core Loss** (B_m = L×Irms×√2/(N×Ae) = rated RMS actual amplitude): ",
                        f"B_m = {d9_L_mH:.3f}mH × {d9_I_rms:.0f}A × 1.414 / ({N9} × {Ae9:.1f}cm²) = **{B9_m:.4f} T**",
                        f"P_Fe,f = {Cm9} × ({d9_fg:.0f}/1000)^{a9} × {B9_m:.4f}^{b9} × {Ve9:.0f}cm³ = **{Pfef9:.1f} W**",
                        "",
                        f"**Switching Harmonic Core Loss** (B_m = ΔB/2, switching frequency component; β={b9} → harmonic amplification factor {2**b9:.1f}×): ",
                        f"ΔB_pp = L×ΔI/(N×Ae) = {dB9:.5f} T (peak-to-peak), B_m = ΔB_pp/2 = **{dB9_amp:.5f} T**",
                        f"P_Fe,sw = {Cm9} × ({d9_fsw:.0f}/1000)^{a9} × {dB9_amp:.5f}^{b9} × {Ve9:.0f}cm³ = **{Pfesw9:.1f} W**",
                        "",
                        f"**Total Core Loss P_Fe = {Pfef9:.1f} + {Pfesw9:.1f} = {Pfe9:.1f} W**",
                        ]
                        st.markdown("  \n".join(_fe_lines))

                if is_trial:
                    trial_lock("Loss calculation results require paid access")
                else:
                    c1,c2,c3,c4,c5 = st.columns(5)
                    c1.metric("Copper Loss P_Cu",   f"{Pcu9:.1f} W")
                    c2.metric("Core Loss P_Fe",     f"{Pfe9:.1f} W")
                    c3.metric("Total Loss",          f"{Ptot9:.1f} W")
                    c4.metric("Copper/Core Loss Ratio", f"{cufr9:.2f}",
                              delta="✅ Balanced" if 0.33<=cufr9<=3 else "⚠ Needs Optimization",
                              delta_color="normal" if 0.33<=cufr9<=3 else "off")
                    c5.metric("Temperature Rise", f"{dT9:.0f} K",
                              delta=f"Limit {T9_lim} K {'✅' if cool_ok9 else '⚠'}",
                              delta_color="normal" if cool_ok9 else "inverse")

                if not cool_ok9:
                    st.warning(
     f"⚠ Temperature Rise {dT9:.0f} K exceeds {d9_ins_class} limit {T9_lim} K.\n\n"
     f"**Cooling Upgrade Suggestion**: {upnote9}\n\n"
     f"Other options: ① Reduce current density J (increase conductor cross-section); ② Reduce core loss (optimize Ae and switching frequency); ③ Improve ventilation."
                    )

                fig9p = go.Figure(go.Pie(
                    labels=["Cu Loss (Fundamental)", "Cu Loss (HF Harmonic)", "Core Loss (Fundamental)", "Core Loss (Switching)"],
                    values=[max(Pcuf9,0.01),max(Pcuhf9,0.01),max(Pfef9,0.01),max(Pfesw9,0.01)],
                    hole=0.38,
                    marker_colors=["#0d6efd","#85b7eb","#dc3545","#f08080"],
                    textinfo="label+percent", textfont_size=12
                ))
                fig9p.update_layout(title="Loss Distribution (Copper Loss vs. Core Loss)", height=260, template="plotly_white",
                                    margin=dict(t=40,b=10))
                st.plotly_chart(fig9p, use_container_width=True)

            # ══ ④ Terminal Design ══════════════════════════════════════════════════════
            with t9_4:
                st.markdown("#### 🔩 Terminal Design")
                if is_trial:
                    trial_lock("Unlock with paid access: terminal cross-section and temperature rise design")
                else:
                    col_ta, col_tb = st.columns([1, 2])
                    with col_ta:
                        term9_mat = st.radio(
                            "Terminal Busbar Material",
                            ["Copper Busbar (recommended)", "Aluminum Busbar (cost-optimized)"],
                            key="d9_term_mat"
                        )
                    is_al9 = "Aluminum" in term9_mat
                    J9t = 1.2 if is_al9 else (2.0 if d9_I_rms<=400 else 1.5 if d9_I_rms<=2000 else 1.2)
                    note9t = ("resistivity rho approx 1.64x copper, heavier; requires Al/Cu bimetallic connector; GB/T 5585"
                              if is_al9 else
                              ("<=400A: 2.0 A/mm²" if d9_I_rms<=400 else
                               "400~2000A: 1.5 A/mm²" if d9_I_rms<=2000 else ">2000A: 1.2 A/mm²"))
                    S9t     = d9_I_rms / J9t
                    S9t_std = next((s for s in std9s if s >= S9t), 400)

                    with col_tb:
                        st.success(
                            f"**Recommended**: {term9_mat}\n\n"
                            f"- J = {J9t} A/mm² ({note9t})\n"
                            f"- Required cross-section = {S9t:.1f} mm² → Standard size: **{S9t_std} mm²** (GB/T 5585)"
                        )

                    col_tc, col_td = st.columns(2)
                    with col_tc:
                        st.markdown("**Surface Plating Options**")
                        st.dataframe(pd.DataFrame({
                            "Plating": ["Tin (Sn)", "Silver (Ag)", "Nickel (Ni)"],
                            "Thickness": ["5~15 μm", "3~10 μm", "5~20 μm"],
                            "Max Temp.": ["150°C", "200°C", "300°C"],
                            "Application": ["Standard, most common", "High-current low-loss", "High-temperature / corrosion-resistant"],
                        }), use_container_width=True, hide_index=True)
                        st.caption("GB/T 5585 (*Busbars Standard*)")
                    with col_td:
                        st.markdown(f"**Connection Specifications** (GB/T 14048.2; IEC 60664-1)")
                        st.markdown(f"""
    | Parameter | Requirement |
    |------|------|
    | Bolt Torque (M10) | 25~30 N·m |
    | Contact Surface Roughness | Ra ≤ 3.2 μm |
    | Conductor Overlap | Flat contact, no bending |
    | Stress Relief | Flexible section recommended |
    | Phase-to-Phase Clearance | ≥{clr9} mm (at {d9_Vline:.0f} V) |
    | Creepage Distance | ≥{20 if d9_Vline<=1000 else 40} mm |
    | Terminal Marking | Phase label (L1/L2/L3), PE ground |
    """)

            # ══ ⑤ Safe Operating Area ════════════════════════════════════════════════
            with t9_5:
                st.markdown("#### 📊 Design Safe Operating Area")
                st.info(
                    "**How to Read This Chart**\n\n"
                    "- **X-axis** = Core Effective Cross-section Ae (cm²); **Y-axis** = Turns N\n"
                    "- **🌟 Red Star** = Current Design Point\n"
                    "- **① Flux Density**: color represents B_peak. Stay **below** B_sat boundary line (Safe Zone).\n"
                    "- **② Total Loss**: color represents total loss magnitude. Darker → higher Copper Loss region; redder → higher Core Loss.\n"
                    "- **③ Weight**: color represents estimated total weight. Use with ①② to find optimal design point."
                )

                Ae9r = np.linspace(Ae9*0.35, Ae9*2.5, 45)
                N9r  = np.arange(max(1, N9-8), N9+13, 1)
                AEg,NNg = np.meshgrid(Ae9r, N9r)
                AEgm = AEg * 1e-4
                # Flux density: B_m = L*Irms*sqrt(2)/(N*Ae)
                # OCP peak not used in contour — OCP is short-time only (us level)
                Bg   = d9_L_H * d9_I_rms * np.sqrt(2) / (np.maximum(NNg,1) * np.maximum(AEgm,1e-6))
                lg   = NNg * MLT9
                Rg   = rho9 * lg / (S9*1e-6) * kt9
                Pcug = Rg * d9_I_rms**2 * kac9
                Veg  = AEgm * lw9_m * 4
                Pfeg = Cm9*(d9_fg/1000)**a9*np.minimum(Bg,Bsat9*0.99)**b9*Veg*1e6
                Ptg  = Pcug + Pfeg
                wtcg = Veg * mat9["dens"] * 1e3
                wtug = S9*1e-6 * lg * dens9
                wtg  = wtcg + wtug

                from plotly.subplots import make_subplots
                fig9c = make_subplots(
                    rows=1, cols=3,
                    subplot_titles=["① Peak Flux Density B_peak (T)",
                                     "② Total Loss P_total (W)",
                                     "③ Estimated Total Weight (kg)"],
                    horizontal_spacing=0.12
                )

                for col_i, (Z, csc) in enumerate([(Bg,"RdYlGn_r"),(Ptg,"YlOrRd"),(wtg,"Blues")], 1):
                    fig9c.add_trace(go.Contour(
                        z=Z, x=Ae9r, y=N9r,
                        colorscale=csc, showscale=True,
                        contours=dict(showlabels=True, labelfont=dict(size=11, color="white")),
                        line=dict(width=0.5),
                        colorbar=dict(len=0.8, thickness=14,
                                      x=0.33*col_i-0.005,
                                      tickfont=dict(size=10))
                    ), row=1, col=col_i)
                    # Current design point star marker
                    fig9c.add_trace(go.Scatter(
                        x=[Ae9], y=[N9], mode="markers",
                        marker=dict(symbol="star", size=18, color="red",
                                    line=dict(width=2, color="white")),
                        showlegend=False
                    ), row=1, col=col_i)

                # B_sat boundary line (OCP peak defines saturation boundary)
                Bsat_x, Bsat_y = [], []
                for Av in Ae9r:
                    Nv = d9_L_H * d9_Ipeak / (Bsat9 * Av * 1e-4)
                    if N9r[0] <= Nv <= N9r[-1]:
                        Bsat_x.append(Av); Bsat_y.append(Nv)
                if Bsat_x:
                    fig9c.add_trace(go.Scatter(
                        x=Bsat_x, y=Bsat_y, mode="lines",
                        line=dict(color="red", dash="dash", width=3),
                        name=f"B_sat={Bsat9}T boundary (design must stay below)"
                    ), row=1, col=1)

                fig9c.update_xaxes(title_text="Core Ae (cm²)", title_font=dict(size=14), tickfont=dict(size=12))
                fig9c.update_yaxes(title_text="Turns N",         title_font=dict(size=14), tickfont=dict(size=12))
                fig9c.update_layout(
                    height=500, template="plotly_white",
                    legend=dict(font=dict(size=12), orientation="h", y=-0.18, x=0),
                    margin=dict(t=70, b=100, l=70, r=20),
                    title=dict(
                        text=f"Current Design: Ae={Ae9:.1f}cm² N={N9} B_peak={Bpk9:.3f}T Total Loss={Ptot9:.0f}W Temp. Rise≈{dT9:.0f}K",
                        font=dict(size=13)
                    )
                )
                st.plotly_chart(fig9c, use_container_width=True)

                # Design Recommendations
                st.markdown("**📋 Design Results & Optimization Recommendations**")
                col_r1, col_r2, col_r3 = st.columns(3)
                with col_r1:
                    mg = (Bop9-Bpk9)/Bop9*100
                    if Bpk9 > Bsat9*0.85:
                        st.error(f"**Flux Density**: Margin only {mg:.0f}%, dangerously close to B_sat. Suggestion: increase Ae (more cross-section) to reduce B_peak below B_sat.")
                    elif mg > 40:
                        st.info(f"**Flux Density**: Margin {mg:.0f}%, design is oversized. Consider reducing Ae to lower cost and weight, then verify losses.")
                    else:
                        st.success(f"**Flux Density**: Margin {mg:.0f}%, in Safe Zone. Design is optimal.")
                with col_r2:
                    if not cool_ok9:
                        st.error(f"**Loss / Temperature Rise**: {dT9:.0f} K exceeds limit {T9_lim} K. {upnote9}")
                    elif cufr9 > 2.5:
                        st.warning(f"**Loss Balance**: Copper Loss is {cufr9:.1f}× Core Loss, imbalanced. Consider increasing Turns N to reduce Copper Loss.")
                    elif cufr9 < 0.4:
                        st.warning(f"**Loss Balance**: Core Loss is {1/cufr9:.1f}× Copper Loss, imbalanced. Consider reducing Ae to lower Core Loss.")
                    else:
                        st.success(f"**Loss Balance**: Ratio = {cufr9:.2f}, well balanced. Total Loss = {Ptot9:.0f} W, Temperature Rise = {dT9:.0f} K.")
                with col_r3:
                    wt9_est = (Ae9_m2 * lw9_m * 4 * mat9["dens"] * 1e3
                               + S9*1e-6 * ltot9 * dens9)
                    st.info(f"**Estimated Weight**: {wt9_est:.1f} kg (core + winding). For reference only. Temperature Rise satisfies insulation class requirement.")




    with tab_test:
        st.markdown("### 📥 Test Requirements, Supplier Scheme & Output Documents")

        st.info("""
**Test Category Overview** (GB/T 1094.6 Clause 7 & NB/T 10503 Clause 6):
- 🟢 **Routine Test**: Tests performed on every unit before delivery
- 🔵 **Type Test**: Tests performed once to validate design type
- 🟡 **Special Test**: Optional tests performed on agreement with customer
        """)

        # Protection current calculation based on design parameters
        st.markdown("#### ⚡ Protection Current Reference Values (Auto-Calculated from Design Parameters)")
        I_sc_igbt     = d_I_rms * 2.0          # IGBT Desat protection: 2x rated
        I_sc_sw       = d_I_rms * 3.0          # Software overcurrent protection: 3x rated
        I_sc_breaker  = d_I_rms * 5.0          # Circuit breaker protection: 5x rated (IEC 60076-6 §3.1.3)
        I_sc_igbt_pk  = I_sc_igbt  * np.sqrt(2)
        I_sc_sw_pk    = I_sc_sw    * np.sqrt(2)
        I_sc_brk_pk   = I_sc_breaker * np.sqrt(2)

        if is_trial:
            trial_lock("Protection current results require paid access")
        else:
            col_sc1, col_sc2, col_sc3 = st.columns(3)
            with col_sc1:
                st.metric("IGBT Desat Protection Current",
                           f"{I_sc_igbt:.0f} A / {I_sc_igbt_pk:.0f} Apk",
                           help="2× rated current — hardware overcurrent; responds in < 10μs, IGBT shuts off")
                st.caption("Response time: < 10 μs")
            with col_sc2:
                st.metric("Software Overcurrent Protection Current",
                           f"{I_sc_sw:.0f} A / {I_sc_sw_pk:.0f} Apk",
                           help="3× rated current — software overcurrent; responds in 1~10ms, reduces PWM duty cycle")
                st.caption("Response time: 1~10 ms")
            with col_sc3:
                st.metric("Circuit Breaker Protection Current",
                           f"{I_sc_breaker:.0f} A / {I_sc_brk_pk:.0f} Apk",
                           help="5× rated current — breaker trip; responds in 50~100ms, isolates grid-side fault")
                st.caption("Response time: 50~100 ms · IEC 60076-6 §3.1.3")

        st.caption(
            "Note: The above are LCL filter reactor withstand current thresholds based on inverter protection hierarchy. "
            "Grid short-circuit current (typically 25× rated, lasts ~3s) is fundamentally different — "
            "filter reactors do not independently interrupt fault current; the inverter protection system handles this."
        )

        st.divider()

        # Test Item Table
        test_items = [
            {"Category": "Routine Test", "Test Item": "Inductance Measurement",
             "Method": "Measure at rated current using LCR meter, frequency 50/60 Hz",
             "qualifiedstandard": "Deviation ≤ ±3% (GB/T 1094.6) or ±5% (NB/T 10503)",
             "referencestandard": "GB/T 1094.6 Cl.7.1"},
            {"Category": "Routine Test", "Test Item": "DC Winding Resistance",
             "Method": "Four-wire method, corrected to 75°C",
             "qualifiedstandard": "Deviation ≤ ±2% (compared to design value)",
             "referencestandard": "GB/T 1094.6 Cl.7.2"},
            {"Category": "Routine Test", "Test Item": "Power Frequency Withstand Voltage (1 min)",
             "Method": "Apply 2×Un+1000V AC, hold 60s",
             "qualifiedstandard": f"No flashover or breakdown (test voltage {max(2000, 2*_Vg_line+1000):.0f} V, 60s)",
             "referencestandard": "IEC 60076-6 Cl.11.2"},
            {"Category": "Routine Test", "Test Item": "Insulation Resistance",
             "Method": "500V/1000V megohmmeter, winding to ground",
             "qualifiedstandard": "≥1000 MΩ (new equipment)",
             "referencestandard": "IEC 60060-1"},
            {"Category": "Type Test", "Test Item": "Temperature Rise Test",
             "Method": "Operate at rated current until thermal equilibrium (Δθ ≤ 1K/h), measure winding temperature",
             "qualifiedstandard": f"Temperature rise ≤ {'125K (H-class)' if 'H' in d_ins_class else '100K (F-class)' if 'F' in d_ins_class else '80K (B-class)'}",
             "referencestandard": "GB/T 1094.6 Cl.8.1"},
            {"Category": "Type Test", "Test Item": "Short-Circuit Withstand Current (filter reactor)",
             "Method": f"Apply {I_sc_brk_pk:.0f} Apk peak, duration 50~100ms (circuit breaker protection level)",
             "qualifiedstandard": "Inductance deviation ≤ ±3% after test, no deformation, no insulation damage",
             "referencestandard": "IEC 60076-6 §3.1.3 (filter reactor clause)"},
            {"Category": "Special Test", "Test Item": "Lightning Impulse Voltage (LI)",
             "Method": "Standard waveform 1.2/50μs, apply to winding-to-ground",
             "qualifiedstandard": "No waveform collapse (flashover)",
             "referencestandard": "IEC 60060-1 & IEC 60076-6 Cl.11.4"},
            {"Category": "Special Test", "Test Item": "Partial Discharge Measurement",
             "Method": "IEC 60270, apply 1.1×Un and measure pC level",
             "qualifiedstandard": "Partial discharge ≤ 10 pC (H-class insulation)",
             "referencestandard": "IEC 60270"},
            {"Category": "Special Test", "Test Item": "Noise Level",
             "Method": "Operate at rated current, measure at 1m distance, A-weighting",
             "qualifiedstandard": "≤70 dB(A) (standard), ≤65 dB(A) (customer-specified quiet requirement)",
             "referencestandard": "GB/T 1094.6 Cl.10 & IEC 60076-10"},
            {"Category": "Special Test", "Test Item": "Harmonic Current Loss Verification",
             "Method": "Apply actual current waveform with switching frequency harmonics, measure total loss",
             "qualifiedstandard": "Actual loss within expected range (harmonic extra loss 5~20%)",
             "referencestandard": "NB/T 10503-2021 Appendix B"},
        ]

        df_test = pd.DataFrame(test_items)

        def color_test(row):
            if row["Category"] == "Routine Test":
                return ["background-color:#d4edda"] * len(row)
            if row["Category"] == "Type Test":
                return ["background-color:#cce5ff"] * len(row)
            return ["background-color:#fff3cd"] * len(row)

        st.dataframe(df_test.style.apply(color_test, axis=1),
                     use_container_width=True, hide_index=True)
        st.caption("🟢 = Routine Test (every unit) | 🔵 = Type Test (design validation) | 🟡 = Special Test (by agreement)")

        st.divider()

        # Delivery Documentation Requirements
        st.markdown("#### 📦 Delivery Documentation Requirements (NB/T 10503-2021 Clause 8)")
        st.markdown(f"""
1. **Test Report**: All routine test data, type test reports (if applicable), signed and stamped
2. **Nameplate Data**: Inductance {d_L_mH:.3f} mH, rated current {d_I_rms:.0f} A, insulation class {d_ins_class.split("(")[0].strip()}, protection class IP23
3. **Installation Drawing**: Outline dimensions, mounting hole positions, electrical clearance diagram (→ ≥ {clr_mm} mm), maintenance access space
4. **Factory Test Records** (originals or certified copies)
5. **Quality Documents**: Material certificates, process inspection records, test equipment calibration records
        """)

        st.divider()

        # ── Excel Supplier Design Scheme Export ──────────────────────────────────
        st.markdown("#### 📥 Generate Supplier Design Scheme (Excel)")
        st.info(
            "Click the button below to compile all calculation results from this module into a standard Excel supplier design scheme, "
            "covering: design parameters, core design, winding design, loss calculation, terminal design, and test requirements (6 sheets). "
            "This file can be sent directly to the magnetic component supplier for quoting and manufacturing."
        )

        if is_trial:
            trial_lock("Design sheet generation requires paid access")
        elif st.button("📊 Generate Supplier Design Scheme (Excel)", key="r_gen_tds", type="primary"):

            # ── Sheet 1: Design Parameters
            tds_basic = {
                "Parameter": ["Target Inductance", "Inductance Tolerance", "Rated Current Irms", "Design Peak Current Ipeak",
                              "Ripple Current Peak-to-Peak ΔI", "Line Voltage", "Grid Frequency", "Switching Frequency",
                              "Converter Topology", "Insulation Class", "Conductor Material", "Cooling Method", "Mounting"],
                "Value":    [f"{d_L_mH:.3f} mH", "±3% (NB/T 10503 allows ±5%)",
                             f"{d_I_rms:.0f} Arms", f"{d_Ipeak:.1f} A",
                             f"{_Iripple_pp:.1f} A",
                             f"{d_Vline:.0f} Vrms (line voltage)", f"{d_fg:.0f} Hz",
                             f"{_fsw:.0f} Hz", topo_choice, d_ins_class,
                             d_conductor, d_cooling, "Vertical, single-phase side-by-side"],
                "Notes": ["Calculated from LCL filter design section", "GB/T 1094.6 Cl.7.1",
                          "", "Irms×√2 + ΔI/2",
                          f"Vdc/(k·L·fsw), {topo_coeff_note}",
                          "Three-phase system", "", "PWM carrier frequency",
                          "Affects ripple current coefficient", f"Max temperature rise ≤{T_max_ins}K",
                          "", "", ""],
            }

            # ── Sheet 2: Core Design
            tds_core = {
                "Parameter": ["Core Material", "Saturation Flux Density B_sat", "Design Operating Flux Density B_op",
                              "Safety Coefficient k_B", "Lamination Width w", "Stack Depth d",
                              "Stacking Factor k_sf", "Effective Cross-section Ae",
                              "Window Height lw",
                              "Theoretical Air Gap (before correction)", "McLyman Fringing Factor F",
                              "Recommended Physical Air Gap (corrected)",
                              "Rated Peak Flux Density B_peak", "Flux Density Safety Margin"],
                "Value":    [core_material.split("(")[0],
                             f"{B_sat} T", f"{B_op:.3f} T",
                             f"{safety_factor:.2f}",
                             f"{core_width_mm:.0f} mm", f"{core_stack_mm:.0f} mm",
                             f"{k_sf}", f"{Ae_cm2:.1f} cm²",
                             f"{lw_cm:.1f} cm",
                             f"{delta_basic_mm:.3f} mm",
                             f"{F_fringe:.4f}",
                             f"{delta_corr_mm:.3f} mm",
                             f"{B_actual:.4f} T",
                             f"{(B_op - B_actual)/B_op*100:.1f}%"],
                "Formula / Notes": ["", "", "B_op = k_B × B_sat",
                                     "Recommended 0.70~0.80, DC reactor use lower value",
                                     "", "", "Stacking factor typically 0.96",
                                     "Ae = w × d × k_sf",
                                     "McLyman correction parameter reference",
                                     "δ_basic = μ₀·N²·Ae/L, McLyman §14.4",
                                     "F = 1+(δ/√Ae)·ln(2lw/δ), Kazimierczuk §3.4",
                                     "δ_corr = δ_basic/F, specify this to supplier as physical slot dimension",
                                     "B_peak = L·Ipeak/(N·Ae), rated continuous",
                                     "Margin = (B_op - B_peak)/B_op × 100%"],
            }

            # ── Sheet 3: Winding Design
            tds_wire = {
                "Parameter": ["Conductor Material", "Turns N",
                              "Design Current Density J", "Required Cross-section S",
                              "Foil Thickness t", "Foil Width w",
                              "Skin Depth δ_skin (@fsw)",
                              "Mean Turn Length MLT (single-phase)",
                              "Hot Resistance R_hot", "Dowell AC Correction Factor k_ac",
                              "Fundamental Copper Loss P_Cu (fg)", "HF Copper Loss ΔP_Cu",
                              "Total Copper Loss P_Cu_total",
                              "Inter-layer Insulation Paper", "Total Winding Thickness"],
                "Value": [d_conductor, f"{N_final} turns",
                                     f"{J_rated_adj:.1f} A/mm²",
                                     f"{S_wire_mm2:.2f} mm²",
                                     f"{foil_thick_mm:.1f} mm",
                                     f"{foil_width_mm:.1f} mm",
                                     f"{delta_skin_sw_mm:.3f} mm",
                                     f"{MLT_m*100:.1f} cm",
                                     f"{R_hot*1000:.3f} mΩ",
                                     f"{k_ac:.3f}",
                                     f"{P_Cu:.1f} W",
                                     f"{P_Cu_hf:.1f} W",
                                     f"{P_Cu_total:.1f} W",
                                     "Nomex 410 0.13mm / NMN 0.15mm (H-class)",
                                     f"{window_height_used:.1f} mm"],
                "Formula / Notes": ["",
                                     "N = √(L·δ_corr/(μ₀·Ae)), McLyman §5.2",
                                     "AN:3.0 / AF:4.5 / WF:6.5 A/mm², IEC 60076-6 §6.4",
                                     "S = Irms/J",
                                     f"Recommended t ≤ 2×δ_skin = {2*delta_skin_sw_mm:.2f}mm",
                                     "w = S/t (constrained by lw)",
                                     "δ = √(ρ/(π·f·μ₀)), Dowell 1966",
                                     "MLT ≈ 2(w+d)_core + π·t·N (foil winding)",
                                     "R_hot = R_dc·[1+0.00393·(T-20)]",
                                     "k_ac = max(1, ξ/2), ξ=d/(2δ), Dowell 1966",
                                     "P_Cu = R_hot·Irms²",
                                     "ΔP_Cu = k_ac·R_hot·I_ripple_rms²",
                                     "",
                                     "H-class: Nomex/NMN, IEC 60076-6 recommended",
                                     "N×(t + t_ins)"],
            }

            # ── Sheet 4: Loss & Thermal Design
            tds_loss = {
                "Parameter": ["Core Volume Ve", "Core Loss P_Fe (fundamental fg)", "Core Loss P_Fe (switching fsw)",
                              "Total Core Loss P_Fe", "Total Copper Loss P_Cu",
                              "Total Loss P_total",
                              "Convection Heat Transfer Coefficient h (cooling type)", "Effective Radiating Surface Area A",
                              "Steady-State Temperature Rise ΔT", "Insulation Class Temperature Rise Limit",
                              "Temperature Rise Evaluation"],
                "Value":    [f"{Ve_cm3:.0f} cm³",
                             f"{P_Fe_fund:.1f} W",
                             f"{P_Fe_sw:.1f} W",
                             f"{P_Fe_total:.1f} W",
                             f"{P_Cu_total:.1f} W",
                             f"{P_total:.1f} W",
                             f"{h_conv} W/(m²·K)",
                             f"{A_surface_m2:.3f} m²",
                             f"{delta_T:.0f} K",
                             f"{T_max_ins} K",
                             "✅ Temperature Rise OK" if delta_T <= T_max_ins else "⚠ Temperature Rise Exceeds Limit, Optimization Needed"],
                "Formula / Notes": ["Ve ≈ Ae × lw × 4 (three-phase, 4 limbs)",
                                     f"P_Fe = Cm·(fg/1kHz)^α·B_op^β·Ve, Steinmetz; Cm={Cm},α={alpha_s},β={beta_s}",
                                     "P_Fe(sw) = Cm·(fsw/1kHz)^α·ΔB^β·Ve, Steinmetz",
                                     "", "", "",
                                     "AN:10 / AF:25 / WF:60 W/(m²·K), engineering estimate",
                                     "π·D·H (cylindrical surface estimate)",
                                     "ΔT = P_total/(h·A)",
                                     f"{d_ins_class.split('(')[0]}, GB/T 1094.6",
                                     ""],
            }

            # ── Sheet 5: Test Requirements
            tds_test = {
                "Category":          [x["Category"]          for x in test_items],
                "Test Item":         [x["Test Item"]          for x in test_items],
                "Method":            [x["Method"]             for x in test_items],
                "qualifiedstandard": [x["qualifiedstandard"]  for x in test_items],
                "referencestandard": [x["referencestandard"]  for x in test_items],
            }

            # ── Sheet 6: Terminal Design & Protection Parameters
            tds_terminal = {
                "Parameter": ["Recommended Busbar Material", "Required Cross-section S", "Standard Size (GB/T 5585)",
                              "Design Current Density J", "Surface Plating",
                              "Phase-to-Phase Clearance", "Bolt Torque (M10)",
                              "IGBT Desat Protection Current (peak)", "IGBT Desat Response Time",
                              "Software Overcurrent Protection Current (peak)", "Software OCP Response Time",
                              "Circuit Breaker Protection Current (peak)", "Circuit Breaker Response Time"],
                "Value":    [term_rec_type,
                             f"{S_term_mm2:.1f} mm²",
                             f"{S_std} mm²",
                             f"{J_term_rec} A/mm²",
                             "Tin (Sn, standard) / Silver (Ag, low-loss) / Nickel (Ni, high-temp)",
                             f"≥ {clr_mm} mm",
                             "25~30 N·m",
                             f"{I_sc_igbt_pk:.0f} Apk",
                             "< 10 μs",
                             f"{I_sc_sw_pk:.0f} Apk",
                             "1~10 ms",
                             f"{I_sc_brk_pk:.0f} Apk",
                             "50~100 ms"],
                "Notes": [term_rec_note,
                          "S = Irms/J",
                          "GB/T 5585",
                          f"Based on {d_cooling} cooling method",
                          "GB/T 5585",
                          f"Line voltage {_Vg_line:.0f}V, GB/T 1094.6 & IEC 60664-1",
                          "GB/T 14048.2",
                          "2× rated current",
                          "IGBT Desat hardware protection shuts off",
                          "3× rated current",
                          "PWM duty cycle software reduction",
                          "5× rated current, IEC 60076-6 §3.1.3",
                          "Grid-side circuit breaker trips"],
            }

            # ── Generate Excel File
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                pd.DataFrame(tds_basic).to_excel(writer, sheet_name="Design Parameters", index=False)
                pd.DataFrame(tds_core).to_excel(writer, sheet_name="Core Design", index=False)
                pd.DataFrame(tds_wire).to_excel(writer, sheet_name="Winding Design", index=False)
                pd.DataFrame(tds_loss).to_excel(writer, sheet_name="Loss & Thermal", index=False)
                pd.DataFrame(tds_test).to_excel(writer, sheet_name="Test Requirements", index=False)
                pd.DataFrame(tds_terminal).to_excel(writer, sheet_name="Terminal & Protection", index=False)

                # Apply formatting: column width + header color + alternating row fill
                from openpyxl.styles import Font, PatternFill, Alignment
                header_colors = {
                    "Design Parameters":   "1F4E79",
                    "Core Design":         "145A32",
                    "Winding Design":      "6B2C00",
                    "Loss & Thermal":      "7D3C98",
                    "Test Requirements":   "1A5276",
                    "Terminal & Protection": "784212",
                }
                for sheet_name, ws in writer.sheets.items():
                    # Auto column width
                    for col in ws.columns:
                        max_len = max(
                            (len(str(cell.value or "")) for cell in col), default=0
                        ) + 4
                        ws.column_dimensions[col[0].column_letter].width = min(max_len, 55)
                    # Header row formatting
                    hc = header_colors.get(sheet_name, "1F4E79")
                    for cell in ws[1]:
                        cell.font      = Font(bold=True, color="FFFFFF", size=11)
                        cell.fill      = PatternFill(fill_type="solid", fgColor=hc)
                        cell.alignment = Alignment(horizontal="center", vertical="center",
                                                    wrap_text=True)
                    ws.row_dimensions[1].height = 22
                    # Alternating row fill
                    from openpyxl.styles import PatternFill as PF
                    alt_fill = PF(fill_type="solid", fgColor="F2F2F2")
                    for row in ws.iter_rows(min_row=2):
                        if row[0].row % 2 == 0:
                            for cell in row:
                                if cell.fill.fgColor.rgb == "00000000":
                                    cell.fill = alt_fill

            output.seek(0)
            st.download_button(
                label="📥 Download Supplier Design Scheme (.xlsx, 6 Sheets)",
                data=output.getvalue(),
                file_name=(
                    f"SupplierScheme_{d_target.split('(')[0]}_"
                    f"{d_L_mH:.3f}mH_{d_I_rms:.0f}A.xlsx"
                ),
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="r_dl_tds"
            )
            st.success(
                "✅ Supplier Design Scheme generated! Contains 6 sheets:\n"
                "Design Parameters | Core Design | Winding Design | Loss & Thermal | Test Requirements | Terminal & Protection"
            )












# ==========================================
#  10: Capacitor Selection v5
#
# v4 (v3, ): 
#   [Fix-A] MKP vs : 
#           - "99%/LCL"(, MW
#             ); /
#           - : (Metalized Film), 
#             ; Film/Foil
#           - IEC 60831-1(), 
#             IEC 61071
#   [Fix-B] SVG: 
#           - : (Schoopage), 
#             (Inserted Tab)
#           - : (Staggered Margin), 
#             (Non-inductive Winding)
#           :  vs (, )
#   [Fix-C] →ESR(v3): 
#           - "→ESR"( R=ρL/S, )
#           - : PP→→ESR; 
#             →ESR, ()
#           - EPCOS Heavy Edge / Wave cut
#
# v3 (Opt1~3): , MKT, , SVG
# v2 : 
#   [Fix1]  3.5  7(IEC 61071 & TDK )
#   [Fix2]  dv/dt (IEC 61071 §5.1, )
#   [Fix3] ripple current Lg ( → )
#   [Fix4]  Y/Δ (Rated Voltage)
#   [Add1] AI (Claude API Parameter, )
#   [Add2] Miner ()
#   [Add3] (EPCOS/WIMA/Vishay)
#
# : 
#   [1] IEC 61071:2017 — Capacitors for power electronics(LCL)
#   [2] IEC 61071:2017 §5.1 dv/dt, §6.5~6.6 
#   [3] TDK EPCOS, "Film Capacitors General Technical Information", Rev.2023, §2~§4
#       → // n=7~9(MKP)
#   [4] IEC 60300-3-1:2003 — Dependability, Miner 
#   [5] Liserre M. et al., IEEE Trans. Ind. Appl., 2005
#   [6] TDK EPCOS B32776 / B32778  — , , Heavy Edge
#   Note: IEC 60831-1 , LCLCapacitor Selection
# ==========================================

elif selection == nav_options[9]:

    # ──  session_state Parameter ─────────────────────────────────────────
    _fg    = float(st.session_state.get('sys_fg',   50.0))
    _fsw   = float(st.session_state.get('sys_fsw', 2500.0))
    _Ug    = float(st.session_state.get('sys_ug',  690.0))
    _IL    = float(st.session_state.get('vec_i',   500.0))
    _Vdc   = float(st.session_state.get('vec_udc', 1100.0))
    _L1_mH = float(st.session_state.get('lcl_l1',  0.500))
    _L2_mH = float(st.session_state.get('lcl_l2',  0.250))
    _Cf_uF = float(st.session_state.get('lcl_c',   100.0))

    L1_base = _L1_mH * 1e-3
    L2_base = _L2_mH * 1e-3
    Cf_base = _Cf_uF * 1e-6
    omega1  = 2 * np.pi * _fg
    omega_sw = 2 * np.pi * _fsw
    Ug_line  = _Ug                        # voltage
    Ug_phase = _Ug / np.sqrt(3)           # voltage(Y)

    # ══════════════════════════════════════════════════════════════════════════
    # 
    # ══════════════════════════════════════════════════════════════════════════
    st.header("🔋 Filter Capacitor Selection & Design")
    st.markdown("""
> **Section engineering target**: LCL filter capacitor type selection, rated parameters, dv/dt verification,
> Ripple Current Analysis, lifetime & reliability prediction (Miner's Rule), AI datasheet parser & recommendation.
    """)
    st.divider()

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 
    # ══════════════════════════════════════════════════════════════════════════
    tab_type, tab_spec, tab_ripple, tab_life, tab_upload = st.tabs([
 "📚 1. Capacitor Type Selection",
 "📐 2. Rated Parameters & dv/dt",
        "〰️ 3. Ripple Current Analysis",
 "🌡️ 4. Temperature · Lifetime · Mission Profile",
 "📄 5. AI Datasheet Parser & Recommendation",
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 1 —   [Opt1] [Opt2] [Opt3]
    # ══════════════════════════════════════════════════════════════════════════
    with tab_type:
        st.markdown("### 📚 Power Capacitor Type Selection")

        # ── [Opt1] : MKT, / ────────────
        st.markdown("""
> 💡 **Guide**: LCL filter capacitors must withstand AC voltage + PWM ripple current + dv/dt simultaneously.
> The "LCL Filter Suitability" column shows engineering results; use it to narrow down your selection quickly.
        """)

        df_types = pd.DataFrame({
 "Type": [
                "Dry Metallized Film (MKP)",
 "Oil-Immersed Film",
                "Aluminum Electrolytic",
                "Ceramic Capacitor (X)",
                "CBB Capacitor (PP, Domestic)",
            ],
 "Dielectric / Structure": [
                "PP film + metallized Al, dry",
                "Film/Foil + oil-immersed, wet",
                "Aluminum oxide / wet electrolyte",
                "Ceramic (X7R/X5R)",
                "PP film (domestic), dry",
            ],
 "ESR Level": [
                "Low (1~5 mΩ)",
                "Very Low (1~3 mΩ)",
                "High (100~500 mΩ)",
                "Medium (mΩ range)",
                "Medium (mΩ range)",
            ],
 "dv/dt Capability": [
                "⭐⭐⭐⭐⭐ (kV/μs class)",
                "⭐⭐⭐⭐⭐ (high power)",
                "⭐ (very poor)",
                "⭐⭐⭐⭐",
                "⭐⭐⭐⭐",
            ],
 "Self-Healing": [
                "✅ Yes (metallized film)",
                "✅ Yes (metallized type)",
                "❌ No",
                "❌ No (ceramic cracks)",
                "✅ Yes (metallized film)",
            ],
 "Capacity Density": [
                "Medium (μF/dm³)",
                "High (large capacity/size)",
                "Very High (mF class)",
                "High (compact)",
                "Medium (compact)",
            ],
 "Voltage Lifetime Exponent n": [
                "7~9 (IEC 61071)",
                "7~9 (same as MKP)",
                "3~5",
                "—",
                "7~9",
            ],
 "LCL Filter Suitability": [
                "✅ First choice (standard)",
                "✅ High-power / offshore",
                "❌ Prohibited (ESR, dv/dt not rated)",
                "⚠ Small-signal filter only (low capacity)",
                "⚠ Limited capacity range",
            ],
 "Typical Applications": [
                "PV/wind inverter LCL filter, STATCOM reactive compensation",
                "Offshore wind SVG, MW converter, high-power UPS",
                "Bulk energy storage, low-frequency filter (not for PWM ESR)",
                "EMC X/Y capacitor, decoupling",
                "Low-power domestic inverter",
            ],
        })
        st.dataframe(df_types, width="stretch", hide_index=True)
        st.caption(
            "Source: IEC 61071:2017 Table 1 (power capacitors); "
            "TDK EPCOS Film Capacitors General Technical Information Rev.2023"
        )

        # ── [Fix-A]  vs () ──────────────────────
        st.divider()
        st.markdown("#### 🔍 Why Use MKP Film Capacitors? (Dry vs. Oil-Immersed)")
        st.info("""
💡 **Note**: "Dry MKP **(Metalized Film)**" and "Oil-immersed **(Metal Foil)**" are two different structures.
All LCL filter capacitors use Film (film electrode); the Foil (aluminum foil electrode) structure has **no self-healing capability**.
        """)
        if is_trial:
            trial_lock("Unlock with paid access: detailed comparison of dry MKP vs oil-immersed capacitors (application scenarios / pros & cons / selection criteria)", compact=True)
        else:
            col_dry, col_oil = st.columns(2, gap="large")
            with col_dry:
                st.markdown("""
**🧪 Dry metallized film capacitor (MKP)**

**Self-healing mechanism**: When the metallized aluminum layer (20–50 nm ultra-thin) is locally punctured,
the arc energy instantly vaporizes the aluminum at that point, automatically isolating the fault,
and the capacitor continues to operate (Self-Healing, IEC 61071 §3.3.1).

**Application scenarios (by power/voltage class)**:
- Tens of kW to hundreds of kW string/central PV inverters: ✅ **dry-type preferred**
- 690 V / 1140 V medium-voltage wind converters (centralized): ⚠ **check manufacturer custom offerings**
- Low-voltage commercial/industrial storage, UPS, STATCOM: ✅ dry-type mainstream

**Advantages**: compact, lightweight, low cost, easy modular installation

**Limitations**: rated voltage usually ≤ 1200 Vrms AC;
for high inrush applications, verify dv/dt and rated ripple current

> 📚 IEC 61071:2017 §3.3.1 (Self-healing definition and test);
> TDK EPCOS General Technical Information Rev.2023, §3
                """)
            with col_oil:
                st.markdown("""
**🛢 Oil-immersed film capacitor**

**Key advantage**: oil impregnation improves three performance areas:
① Higher dielectric strength (oil fills micro-voids, delays breakdown);
② Better heat dissipation (oil carries heat away from the film);
③ Suppresses partial discharge, extends lifetime under AC stress.

**Self-healing note**:
- Metallized film structure (MKP type): ✅ self-healing, even when oil-immersed
- Film/Foil (aluminum foil) structure (common in high-current models): ❌ no self-healing, short-circuit failure

**Application scenarios**:
- 1.5 MW to 15 MW offshore/onshore wind converters: ✅ **oil-immersed preferred**
- SVG / STATCOM (6 kV medium voltage): ✅ mainstream
- High-power SVC, reactive compensation banks: ✅ recommended

**Safety Valve feature**:
Large oil-immersed capacitors include an internal safety valve —
when internal pressure builds up from a fault,
the valve ruptures to release pressure, preventing explosion.
This is a safety feature, not a failure mode;
confirm operation requirements with manufacturer during procurement.

> 📚 IEC 61071:2017 §3.1 (power capacitor general requirements, including oil-immersed);
> IEC 61071:2017 §5.3 (safety requirements)
                """)

        # ══════════════════════════════════════════════════════════════════════
        # [Fix-B] ()
        # ══════════════════════════════════════════════════════════════════════
        st.divider()
        st.markdown("### 🔬 Filter Capacitor Engineering Parameters")
        st.markdown("""
> This section explains the physical meaning behind capacitor "selection parameters".
> Understanding **(MKP) film capacitor** internal structure and material properties,
> helps make informed comparisons between different capacitor types.
        """)

        # ══════════════════════════════════════════════════════════════════════
        # SVG: ( + ), All
        # ══════════════════════════════════════════════════════════════════════
        # SVG: (0~340), (370~860), 
        # All, 
        film_cap_svg = """
<svg viewBox="0 0 860 260" xmlns="http://www.w3.org/2000/svg" font-family="Arial, sans-serif">
  <rect width="860" height="260" fill="#f8f9fa" rx="10"/>

 <!-- ══════ left side 0~340:  ══════ -->
  <text x="170" y="20" text-anchor="middle" font-size="12" font-weight="bold" fill="#343a40">Capacitor Appearance</text>

 <!-- -->
  <rect x="50" y="32" width="180" height="130" fill="#dee2e6" stroke="#adb5bd" stroke-width="1.5"/>
 <!-- left end(, ) -->
  <ellipse cx="50"  cy="97" rx="20" ry="65" fill="#adb5bd" stroke="#868e96" stroke-width="1.5"/>
 <!-- right end(, ) -->
  <ellipse cx="230" cy="97" rx="20" ry="65" fill="#ffd43b" stroke="#f08c00" stroke-width="2"/>

 <!-- -->
  <rect x="22" y="83" width="16" height="28" fill="#868e96" rx="2"/>
 <!-- -->
  <rect x="238" y="83" width="16" height="28" fill="#c9a227" rx="2"/>

  <!-- left endannotation(left side, exceedsx=0) -->
  <text x="26" y="176" font-size="10" fill="#495057" text-anchor="middle">Busbar (Neg)</text>

  <!-- right endannotation(right side, exceedsx=340) -->
  <text x="246" y="176" font-size="10" fill="#7a5c00" text-anchor="middle">Schoopage (Pos)</text>

 <!-- Note: , below -->
  <line x1="230" y1="162" x2="230" y2="190" stroke="#f08c00" stroke-width="1.2" stroke-dasharray="3,2"/>
  <text x="230" y="204" text-anchor="middle" font-size="10.5" fill="#e67700" font-weight="bold">Schoopage End Face</text>
  <text x="230" y="217" text-anchor="middle" font-size="9.5"  fill="#555">Zinc/Al alloy spray, heavy current path</text>
  <text x="230" y="230" text-anchor="middle" font-size="9.5"  fill="#c92a2a" font-weight="bold">⚠ dv/dt thermal bottleneck</text>

 <!-- annotation -->
  <text x="140" y="248" text-anchor="middle" font-size="9" fill="#868e96">Note: Internal safety valve for pressure protection</text>

 <!-- ══════ right side 370~860: () ══════ -->
  <text x="615" y="20" text-anchor="middle" font-size="12" font-weight="bold" fill="#343a40">Internal Layer Structure (Magnified)</text>

 <!-- (, ) -->
  <rect x="380" y="32" width="14" height="175" fill="#adb5bd" stroke="#868e96" stroke-width="1.2"/>
  <text x="375" y="120" text-anchor="middle" font-size="9.5" fill="#495057"
        transform="rotate(-90,375,120)">Neg busbar</text>

 <!-- (, ) -->
  <rect x="842" y="32" width="14" height="175" fill="#ffd43b" stroke="#f08c00" stroke-width="1.2"/>
  <text x="861" y="120" text-anchor="middle" font-size="9.5" fill="#e67700"
        transform="rotate(90,861,120)">Pos busbar</text>

 <!-- ① PP (, ) -->
  <rect x="394" y="36" width="448" height="34" fill="#a5d8ff" stroke="#339af0" stroke-width="1.2"/>
  <text x="618" y="54"  text-anchor="middle" font-size="12" fill="#1864ab" font-weight="bold">① PP Film (Dielectric) — insulating layer</text>
  <text x="618" y="66"  text-anchor="middle" font-size="10" fill="#1864ab">Left-end aligned, right edge exposed</text>

 <!-- ② (, , ) -->
  <rect x="420" y="72" width="436" height="24" fill="#d4a017" stroke="#a67c00" stroke-width="1.2"/>
 <!-- right side(right end) -->
  <rect x="842" y="70" width="14" height="28" fill="#e67700" opacity="0.9"/>
  <text x="618" y="88"  text-anchor="middle" font-size="12" fill="#5c3a00" font-weight="bold">② Metallized Al (Pos) — electrode, right end→Schoopage</text>

 <!-- ③ PP (, ) -->
  <rect x="394" y="98" width="448" height="34" fill="#a5d8ff" stroke="#339af0" stroke-width="1.2"/>
  <text x="618" y="116" text-anchor="middle" font-size="12" fill="#1864ab" font-weight="bold">③ PP Film (Dielectric) — same as ①</text>
  <text x="618" y="128" text-anchor="middle" font-size="10" fill="#1864ab">Sandwiched between ② and ④</text>

 <!-- ④ (, , ) -->
  <rect x="394" y="134" width="436" height="24" fill="#d4a017" stroke="#a67c00" stroke-width="1.2"/>
 <!-- left side(left end) -->
  <rect x="394" y="132" width="14" height="28" fill="#7a7a7a" opacity="0.9"/>
  <text x="618" y="150" text-anchor="middle" font-size="12" fill="#5c3a00" font-weight="bold">④ Metallized Al (Neg) — electrode, left end→Schoopage</text>

 <!-- tip -->
  <text x="618" y="175" text-anchor="middle" font-size="10" fill="#868e96">↑ PP films  ↑</text>

 <!-- (, , right side) -->
  <rect x="394" y="184" width="12" height="12" fill="#a5d8ff" stroke="#339af0" stroke-width="1"/>
  <text x="410"  y="195" font-size="10" fill="#1864ab">= PP Film (Dielectric)</text>
  <rect x="555" y="184" width="12" height="12" fill="#d4a017" stroke="#a67c00" stroke-width="1"/>
  <text x="571"  y="195" font-size="10" fill="#5c3a00">= Metallized Al (electrode)</text>
  <rect x="718" y="184" width="12" height="12" fill="#e67700" opacity="0.9"/>
  <text x="734"  y="195" font-size="10" fill="#e67700">Orange = Schoopage end contact</text>

 <!-- principle(right sidebottom, exceeds860) -->
  <text x="618" y="218" text-anchor="middle" font-size="10" fill="#555">
  Pos → right end busbar;  Neg → left end busbar
  </text>
  <text x="618" y="232" text-anchor="middle" font-size="10" fill="#555">
  Staggered winding, parallel paths → low ESR, non-inductive ESL≈0
  </text>

 <!-- -->
  <line x1="355" y1="18" x2="355" y2="245" stroke="#dee2e6" stroke-width="1" stroke-dasharray="4,3"/>

</svg>
        """

        st.components.v1.html(
            f"""<div style="background:#f8f9fa;border-radius:12px;padding:6px;">{film_cap_svg}</div>""",
            height=278
        )
        st.caption(
            "Source: IEC 61071:2017 §3.3.1 (self-healing definition), §5.1 (dv/dt annotation); "
            "TDK EPCOS General Technical Information Rev.2023 §2-§4; "
            "EPCOS B32776 datasheet (Non-inductive Winding / Heavy Edge process)"
        )

        # ══════════════════════════════════════════════════════════════════════
        # ①: 
        # ══════════════════════════════════════════════════════════════════════
        st.markdown("#### 🧱 Core Material Concepts")
        col_c1, col_c2, col_c3 = st.columns(3, gap="medium")
        with col_c1:
            st.info("""
**① PP Film (Dielectric)**

**Physical nature**: Polypropylene (PP) insulating film, thickness 3~20 μm.

**Function**: The capacitor dielectric, responsible for insulating and storing electric field energy.
Its thickness determines voltage rating, and its surface accepts the metallized electrode — it is the physical foundation of the capacitor.

**Key relationships**:
- Thinner PP → lower voltage rating (breakdown field ~600 V/μm)
- Thinner PP → more turns per unit length → more parallel paths → lower ESR

> 📚 TDK EPCOS Rev.2023 §2.1
            """)
        if is_trial:
            trial_lock("Unlock with paid access: deep analysis of core capacitor material concepts (metallized aluminum layer and metallized film)", compact=True)
        else:
            with col_c2:
                st.info("""
**② Metallized aluminum layer (gold layer)**

**Physical nature**: A layer of aluminum atoms vacuum-deposited onto the PP film surface; thickness only 20–50 nm (about 1/5000th of a human hair).

**Function**: The capacitor electrode, responsible for storing charge.
Positive and negative electrodes sandwich the PP dielectric film to form the complete capacitor structure.

**Thickness significance**:
- Thicker aluminum → larger cross-section S → lower ESR → can handle higher ripple current
- Thinner aluminum → higher ESR, but **lower energy needed for self-healing → easier self-healing**

⚠ Note: aluminum layer thickness affects ESR through a completely different mechanism than PP film thickness — do not confuse the two.

> 📚 TDK EPCOS Rev.2023 §3.2
                """)
            with col_c3:
                st.info("""
**③ Metallized Film**

**Physical nature**: The composite thin film produced after step ②'s "vacuum-deposit aluminum onto PP film" — the PP insulating film and aluminum electrode are **integrated into one**.

**Difference from aluminum foil (Foil) structure**:
- **Metallized film**: aluminum is deposited on PP; electrode is ultra-thin; PP and aluminum are inseparable
  → ✅ Self-healing capable
- **Aluminum foil structure**: separate aluminum foil sheets (~5–10 μm thick) interleaved alternately with PP film
  → ❌ No self-healing (foil too thick; breakdown arc cannot vaporize and isolate it; results in permanent short-circuit)

Datasheet recognition: `Metalized Film` or `Self-healing` = metallized film structure;
`Film/Foil` = aluminum foil structure.

> 📚 IEC 61071:2017 §3.3.1
                """)

        # ══════════════════════════════════════════════════════════════════════
        # ②: 
        # ══════════════════════════════════════════════════════════════════════
        st.markdown("#### 🔑 Two Core Winding Processes")
        col_p1, col_p2 = st.columns(2, gap="medium")
        with col_p1:
            st.success("""
**Process 1: Staggered Margin Winding**

The two metallized layers are wound with opposite margins — this is the **root of non-inductive structure**:
- Positive layer → PP film extends to the left; right end is exposed → right end Schoopage contact
- Negative layer → PP film extends to the right; left end is exposed → left end Schoopage contact

**Result**:
All positive electrodes connect to the right Schoopage; all negative electrodes connect to the left Schoopage.
Thousands of layers wound in parallel, dramatically reducing ESR.

This is the core reason why LCL filter capacitors achieve ultra-low ESR.

> 📚 EPCOS B32776 datasheet; TDK EPCOS Rev.2023 §2.3
            """)
        if is_trial:
            trial_lock("Unlock with paid access: non-inductive winding process principle and its core connection to LCL high-frequency filtering", compact=True)
        else:
            with col_p2:
                st.success("""
**Process 2: Non-inductive Winding**

Because the positive and negative layers are **wound in opposite directions**, the current in each turn flows in opposite directions:
- Positive layer current: flows right to left
- Negative layer current: flows left to right

The magnetic fields generated by the two currents are in opposite directions, **flux mutually cancels**, resulting in an extremely low equivalent parasitic inductance ESL (only a few nH).

**Engineering significance**:
ESL presents high-frequency impedance to ripple current: $X_L = 2\\pi f \\cdot ESL$.
Ultra-low ESL → extremely low high-frequency impedance → capacitor remains effective at switching frequency (kHz to tens of kHz).

This is one of the core reasons why MKP film capacitors work in LCL high-frequency filtering while ordinary capacitors do not.

> 📚 TDK EPCOS General Technical Information Rev.2023 §2.3
                """)

        # ══════════════════════════════════════════════════════════════════════
        # Parameter(HTML, SVG)
        # ══════════════════════════════════════════════════════════════════════
        st.markdown("#### 📊 Engineering Parameters & Physical Meaning")
        st.caption("Parameters in datasheets have specific physical origins — understanding them leads to better selection decisions")

        params_html = """
<style>
.param-grid { display:grid; grid-template-columns:repeat(4,1fr); gap:14px; font-family:Arial,sans-serif; }
.param-card { border-radius:10px; overflow:hidden; border:2px solid; background:#fff; }
.param-head { padding:10px 14px 8px; font-size:14px; font-weight:bold; text-align:center; }
.param-body { padding:10px 14px; font-size:12.5px; line-height:1.7; color:#343a40; }
.param-body .fail { color:#e03131; font-weight:bold; font-size:13px; }
.param-body .warn { color:#c92a2a; font-weight:bold; }
.param-body .tip  { color:#666; font-size:11.5px; margin-top:6px; }
.param-body .act  { margin-top:8px; font-weight:bold; font-size:12px; padding:5px 8px;
                    border-radius:5px; text-align:center; }
.divider { border-top:1px solid #e9ecef; margin:8px 0; }
</style>
<div class="param-grid">

  <div class="param-card" style="border-color:#4dabf7;">
    <div class="param-head" style="background:#e7f5ff;color:#1864ab;">ESR(series)</div>
    <div class="param-body">
      Physical origin: electrode resistance + contact resistance + lead resistance
      <div class="divider"></div>
      <span class="fail">Engineering impact: thermal dissipation P = I² × ESR</span><br>
      MKP: 1~5 mΩ (film type)<br>
      Electrolytic: 100~500 mΩ (prohibited)
      <div class="divider"></div>
      Ways to lower ESR:<br>
      • Thin PP film → more turns → more parallel paths ↑<br>
      • Thick aluminum layer → cross-section S↑<br>
      <span class="warn">⚠ ESR increases with frequency</span><br>
      <span class="tip">Formula: R=ρL/S, lower S → higher ESR</span>
      <div class="act" style="background:#e7f5ff;color:#1864ab;">→ For selection: check ESR @ fsw (mΩ)<br>Use Tab 5 AI parser</div>
    </div>
  </div>

  <div class="param-card" style="border-color:#f08c00;">
    <div class="param-head" style="background:#fff4e6;color:#e67700;">dv/dt Rating</div>
    <div class="param-body">
      PWM switching → rapid voltage rise<br>→ displacement current through capacitor
      <div class="divider"></div>
      <span class="fail">Engineering impact: Schoopage overheating</span><br>
      I_peak = C × dv/dt<br>
      High current → heat → Schoopage delamination
      <div class="divider"></div>
      Typical rated values:<br>
      EPCOS B32776: 50 V/μs<br>
      WIMA MKP10: 100 V/μs<br>
      WIMA FKP1: 200 V/μs<br>
      <span class="warn">⚠ IEC 61071 §5.1 requirement — must be annotated</span><br>
      <span class="tip">This parameter must appear in the datasheet</span>
      <div class="act" style="background:#fff4e6;color:#e67700;">→ Verify in Tab 2 dv/dt calculation tool</div>
    </div>
  </div>

  <div class="param-card" style="border-color:#40c057;">
    <div class="param-head" style="background:#ebfbee;color:#2f9e44;">Self-Healing</div>
    <div class="param-body">
      Exclusive to metallized film structure<br>(Film/Foil type does NOT have this)
      <div class="divider"></div>
      <span class="fail">Protection mechanism: arc vaporizes aluminum</span><br>
      Localized breakdown → arc → aluminum vaporized<br>
      → fault zone isolated<br>
      → capacitor continues operating (loss &lt; 1%)
      <div class="divider"></div>
      Self-healing improves reliability<br>
      <span class="warn">⚠ Self-healing slightly raises ESR</span><br>
      Each event → local aluminum consumed → ESR increases slightly<br>
      <span class="tip">Trade-off: more self-healing events = higher ESR over time</span>
      <div class="act" style="background:#ebfbee;color:#2f9e44;">→ Identification: look for<br>Metalized Film / Self-healing in datasheet</div>
    </div>
  </div>

  <div class="param-card" style="border-color:#cc5de8;">
    <div class="param-head" style="background:#f8f0fc;color:#9c36b5;">Lifetime</div>
    <div class="param-body">
      Determined by hotspot temperature (Arrhenius)<br>
      Th = Tamb + P·Rth + ΔT
      <div class="divider"></div>
      <span class="fail">Rule: every +10°C halves lifetime</span><br>
      L = L₀ × 2^((Tmax−Th)/10)<br>
      　× (Un/Uc)^7 (voltage acceleration)
      <div class="divider"></div>
      Key engineering levers:<br>
      ① Reduce hotspot temp (improve thermal path)<br>
      ② Derate voltage (use higher Rated Voltage)<br>
      67% derating (1.5× rated voltage):<br>
      <strong style="color:#9c36b5;">Lifetime multiplied by 4.5×</strong><br>
      50% derating (2× rated voltage):<br>
      <strong style="color:#9c36b5;">Lifetime multiplied by 16×</strong>
      <div class="act" style="background:#f8f0fc;color:#9c36b5;">→ Use Tab 4 lifetime calculation tool<br>(supports Miner's Rule)</div>
    </div>
  </div>

</div>
        """
        st.components.v1.html(params_html, height=460)
        st.caption(
            "Source: IEC 61071:2017 §3.3.1, §5.1, §6.5-6.6; "
            "TDK EPCOS General Technical Information Rev.2023 §2-§4; "
            "EPCOS B32776 datasheet (Heavy Edge / Wave Cut process)"
        )


        # ──  vs  ─────────────────────────────────────────
        with st.expander("🔬 Dry MKP vs Oil-Immersed: Internal Structure Comparison (expand)", expanded=False):
            col_ds1, col_ds2 = st.columns(2, gap="large")
            with col_ds1:
                st.markdown("""
**🧪 Dry MKP — Internal structure highlights**

| Structural feature | Description |
|---------|------|
| Core filling | No filling; core gaps are **air or nitrogen** (inert gas encapsulation) |
| Housing | Metal can (aluminum/tin) or plastic housing, epoxy-sealed |
| Safety valve | **None** (small size; limited internal pressure) |
| Heat dissipation path | Via housing radiation and convection; relatively high thermal resistance |
                """)
                if is_trial:
                    trial_lock("Unlock with paid access: dry MKP electrode structure, mounting method, and reference details")
                else:
                    st.markdown("""
| Electrode structure | **Metallized aluminum layer** (Metalized Al): aluminum atoms vacuum-deposited onto PP film surface; thickness only 20–50 nm; electrode and insulating film integrated into one |
| Mounting | Stud/lead terminals; can be mounted in any orientation |

> 📚 TDK EPCOS General Technical Information Rev.2023, §2.3
                    """)
            with col_ds2:
                st.markdown("""
**🛢 Oil-immersed film capacitor — Internal structure highlights**

| Structural feature | Description |
|---------|------|
| Core filling | **Mineral oil or synthetic ester**, vacuum-impregnation injection |
| Housing | Metal tank (steel/aluminum), welded and sealed |
| Safety valve | ✅ **Equipped with pressure relief valve (Safety Valve)** |
| Heat dissipation path | Oil conducts heat + metal housing; thermal resistance significantly lower than dry type |
                """)
                if is_trial:
                    trial_lock("Unlock with paid access: oil-immersed capacitor electrode structure, safety valve notes, and self-healing judgment method")
                else:
                    st.markdown("""
| Electrode structure | **Metallized film** (Metalized) or **aluminum foil** (Foil), depending on model |
| Mounting | Usually has orientation restriction (prevent oil leakage) — check installation manual |

**Safety valve engineering notes**:
- Valve typically located at top; opening pressure approx. **3–6 bar**
- **Never point toward personnel or electrical equipment** during installation
- Some models have the safety valve linked to a circuit breaker for fail-safe operation on fault

**Self-healing determination** (important):
- Datasheet states **"Metalized Film"** or **"Self-healing"**: ✅ self-healing capable
- Datasheet states **"Film/Foil"** or **"Non-self-healing"**: ❌ no self-healing

> 📚 IEC 61071:2017 §5.3 (safety requirements);
> EPCOS MKV / MKK series datasheet (oil-immersed LCL filter capacitors)
                    """)

        with st.expander("📖 Structural Parameter Engineering Details (expand)", expanded=True):
            st.markdown(f"""
#### 1️⃣ PP Film Thickness → Rated Voltage & ESR

The PP (Polypropylene) insulating film is **3~20 μm** thick, determining both voltage withstand and turn count.
PP breakdown field strength is approximately **600 V/μm** (peak), significantly higher than PET films (~400 V/μm).

**PP thickness and ESR**: Thinner PP → more turns per unit length →
more layers in parallel → more parallel current paths → **lower ESR** (more parallel resistances = lower combined resistance).

> 📚 TDK EPCOS General Technical Information Rev.2023, §2.1; 
> Harper C.A., *Electronic Packaging and Interconnection Handbook*, 4th ed., §12

#### 2️⃣ Metallized Aluminum Layer Thickness (20~50 nm) → ESR Trade-off

> ⚠ **Engineering point (physical)**: PP thickness and aluminum thickness affect ESR through **completely different mechanisms** — do not confuse them.

From formula $R = \\rho L / S$, where $S$ is the aluminum layer cross-section (proportional to thickness),
a thicker aluminum layer → larger $S$ → lower ESR per layer → lower total ESR.

**Summary — thickness vs. parameters**:

| Dimension | ESR | Self-Healing | Rated Ripple Current |
|------|------------|------------|----------------|
| PP film **thinner** | ↓ (more parallel turns) | unchanged | unchanged |
| Al layer **thicker** | ↓ **lower** (↑ cross-section) | ↑ **harder** (more energy to vaporize) | ↓ slightly lower |
| Al layer **thinner** | ↑ higher (↓ cross-section) | ↓ **easier** (less energy)| ↑ |

            """)
            if is_trial:
                trial_lock("Unlock with paid access: Heavy Edge/Wave Cut process details, Schoopage dv/dt mechanism, thermal resistance lifetime model (second half of structural parameter engineering guide)")
            else:
                st.markdown(f"""
**Engineering solution for high-current models**: EPCOS B32776 series uses **Heavy Edge (edge thickening)**
or **Wave Cut (wave cutting)** process — locally thickening the metallized layer at the edges;
the Schoopage contact zone has extremely low ESR, while the wound center maintains moderate thickness to preserve self-healing capability — achieving both.

> 📚 TDK EPCOS General Technical Information Rev.2023, §3.2 (metallized layer design trade-off);
> EPCOS B32776 datasheet (Heavy Edge / Wave Cut process);
> IEC 61071:2017 §3.3.1 (Self-healing definition)

#### 3️⃣ Schoopage (end-spray metallization) → dv/dt weak point

Capacitor end faces are thermal-spray coated with zinc or aluminum alloy (the "end-spray" or Schoopage process, from French).
LCL filter capacitor busbars or bolts are **welded/crimped to the entire Schoopage end face** (large-area face contact),
with hundreds of metallized electrode edge strips paralleled through the Schoopage — achieving ultra-low ESL and ESR.

> ⚠ Note: The end-face Schoopage is the **thermal and current bottleneck of the capacitor** (not the film layers —
> current flows through the Schoopage, which is tested per IEC 61071 ripple current test standard).

PWM switching dv/dt causes displacement current:

$$I_{{peak}} = C_f \\times \\frac{{dv}}{{dt}}$$

Larger capacitance → higher dv/dt current → higher Schoopage temperature → accelerated aging.
Always verify system dv/dt in Tab2, and confirm it does not exceed **rated dv/dt**.

> 📚 IEC 61071:2017 §5.1 (dv/dt annotation); EPCOS B32776 datasheet (Schoopage)

#### 4️⃣ Thermal resistance Rth → Hotspot temperature → Lifetime

Capacitor generates heat internally ("hotspot"), which must be conducted outward through thermal resistance,
creating a hotspot temperature above ambient, typically **3–10°C** higher.

Lifetime model (Arrhenius factor + voltage acceleration, IEC 61071 §6.5–6.6):

$$L = L_0 \\times 2^{{\\frac{{T_{{max}}-T_h}}{{10}}}} \\times \\left(\\frac{{U_{{rated}}}}{{U_{{op}}}}\\right)^7$$

**Every 10°C increase halves lifetime**; **every 10% voltage derating multiplies lifetime by 7^(voltage step)**.
Tab4 of this section provides a lifetime calculation tool with the physical basis above.

> 📚 TDK EPCOS General Technical Information Rev.2023, §4; IEC 61071:2017 §6.5–6.6
                """)

        # ── [Opt2] "",  ──────────────────────────
        st.divider()
        st.markdown("#### ⚡ Why Can Only Film Capacitors Be Used in LCL Filters? (Quantitative Justification)")

        # Build ESR comparison figure before column split to ensure fig_esr is always defined
        f_pl = np.logspace(1, 5, 400)
        esr_mkp  = 3e-3 * (1 + (f_pl / 1e5) ** 0.3)
        esr_elec = 80e-3 * (50 / np.maximum(f_pl, 50)) ** 0.5 + 20e-3
        fig_esr = go.Figure()
        fig_esr.add_trace(go.Scatter(x=f_pl, y=esr_mkp*1000, name='MKP Film (preferred)',
                                     line=dict(color='#198754', width=2.5)))
        fig_esr.add_trace(go.Scatter(x=f_pl, y=esr_elec*1000, name='Aluminum Electrolytic (prohibited)',
                                     line=dict(color='#dc3545', width=2.5, dash='dash')))
        fig_esr.add_vline(x=_fsw, line_dash='dot', line_color='#fd7e14',
                          annotation_text=f'Current fsw={_fsw:.0f} Hz')
        fig_esr.update_layout(
            title='ESR Frequency Characteristic Comparison',
            xaxis_title='Frequency (Hz)',
            yaxis_title='ESR (mΩ)',
            xaxis_type='log',
            yaxis_type='log',
            height=320,
            template='plotly_white'
        )

        col_t1, col_t2 = st.columns(2, gap="large")
        with col_t1:
            st.markdown(f"""
There are two physical constraints that rule out non-film capacitors:

**Constraint ①: ESR → thermal dissipation**

$$P_{{loss}} = I_{{ripple}}^2 \\times ESR$$

 {_fsw:.0f} Hz Switching Frequency:

| Capacitor Type | ESR (typical) | Thermal dissipation at 50 A | Result |
|---------|------------|------------|------|
| **MKP Film** | 3 mΩ | **7.5 W** ✅ | Acceptable |
| Aluminum Electrolytic | 300 mΩ | **750 W** ❌ | Destroyed immediately |

**constraint②:  dv/dt rated → **

IEC 61071 §5.1 requirementpowercapacitor**annotation dv/dt Rating**,
designstandard(IEC 60384)Parameter,
 PWM .

> 📚 IEC 61071:2017 §3.1 & §5.1; IEC 60384(standard, dv/dtrequirement)
            """)

        if is_trial:
            trial_lock("Unlock with paid access: ESR frequency characteristic comparison chart (film vs. aluminum electrolytic, quantitative justification for film-only use)", compact=True)
        else:
            with col_t2:
                st.plotly_chart(fig_esr, use_container_width=True)
                st.caption("Data source: EPCOS B32xxx MKP datasheet; Nichicon aluminum electrolytic typical parameters")

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 2 — Parameter & dv/dt (v2, )
    # ══════════════════════════════════════════════════════════════════════════
    with tab_spec:
        st.markdown("### 📐 Rated Parameters & dv/dt Verification")

        # ── [Fix4] Y/Δ  ─────────────────────────────────────────────────
        st.markdown("#### ⚡ Connection Type: Capacitor Selection")
        st.warning("""
**Engineering note**: The LCL filter capacitor physical connection type (Y or Δ) determines the
actual capacitor voltage and current, which directly affects the Rated Voltage selection — wrong selection causes component destruction!
        """)
        col_conn1, col_conn2 = st.columns([1, 1.5])
        with col_conn1:
            conn_type = st.radio(
 "Capacitor connection type",
                ["Y (star, common)", "Δ (delta, less common)"],
                key="s10_conn",
                help="3-phasepowerconverterY; designΔ"
            )
            is_delta = "Δ" in conn_type

        with col_conn2:
            if is_delta:
                Uc_actual = Ug_line
                Ic_factor = 1.0 / np.sqrt(3)
                st.info(f"""
**Δ connection**:
- Capacitor voltage = **line voltage = {Ug_line:.0f} Vrms** (√3 times higher than Y — be careful!)
- Capacitor current = line current / √3 = **{_IL/np.sqrt(3):.1f} A**
- For the same filter effect, Δ capacitance = Y value × **1/3**

⚠ If you use a capacitor rated for Y voltage ({Ug_phase:.0f}V) in Δ connection, the actual voltage is {Ug_line:.0f}V — **voltage overstress, do not use!**
                """)
            else:
                Uc_actual = Ug_phase
                Ic_factor = 1.0
                st.info(f"""
**Y connection**:
- Capacitor voltage = **phase voltage = {Ug_phase:.1f} Vrms**
- Capacitor current ≈ line current
- Neutral point connected to system neutral; most common configuration

> In 3-phase systems with neutral (TN/TT grounding), the capacitor neutral carries Zero-Sequence Component.
> Recommend using individually rated capacitors with adequate Margin for neutral current.
                """)

        st.divider()

        # ── Parameter ────────────────────────────────────────────────────────
        with st.expander("⚙️ System Parameters (overrides session values, local to this section)", expanded=True):
            # : Ug / IL / fsw / Cf
            sc1, sc2, sc3, sc4 = st.columns(4)
            with sc1:
                s_ug = st.number_input("Grid voltage Ug (Vrms)", value=float(_Ug), step=10.0, key="s10_ug")
            with sc2:
                s_il = st.number_input("Rated current IL (Arms)", value=float(_IL), step=10.0, key="s10_il")
            with sc3:
                s_fsw = st.number_input("Switching Frequency fsw (Hz)",     value=float(_fsw),   step=100.0, key="s10_fsw")
            with sc4:
                s_Cf  = st.number_input("Designed Cf (μF)",          value=float(_Cf_uF), step=1.0,   key="s10_cf")

            # : fg / Vdc / L1 / L2
            sc5, sc6, sc7, sc8 = st.columns(4)
            with sc5:
                s_fg = st.number_input("Grid frequency fg (Hz)", value=float(_fg), step=1.0, key="s10_fg")
            with sc6:
                s_vdc = st.number_input("DC bus voltage Vdc (V)", value=float(_Vdc), step=50.0, key="s10_vdc")
            with sc7:
                s_L1  = st.number_input("Inverter-side L1 (mH)", value=float(_L1_mH), step=0.05, format="%.3f", key="s10_l1")
            with sc8:
                s_L2  = st.number_input("Grid-side L2 (mH)",     value=float(_L2_mH), step=0.05, format="%.3f", key="s10_l2")

            # : Lg()|  | ()| ()
            sc9, sc10, sc11, sc12 = st.columns([1.2, 0.8, 1, 1])
            with sc9:
                s_Lg  = st.number_input("Grid leakage inductance Lg (mH)", value=0.1, step=0.01, format="%.3f", key="s10_lg",
                help="Transformer leakage inductance or grid impedance inductor, typically 10%~30% of L2")
            with sc10:
                st.empty()
            with sc11:
                s_topology = st.selectbox(
                    "convertertopology",
 ["2-Level(2L-VSC)", "3-Level NPC(3L-NPC)", "3-Level T (3L-TNPC)"],
                    key="s10_topo",
                    help="Same switching frequency with different topology → different ΔI ripple coefficient"
                )
            with sc12:
                s_modulation = st.selectbox(
 "modulation",
                    ["SPWM (Sinusoidal PWM)", "SVPWM(PWM)"],
                    key="s10_mod",
                    help="SVPWM utilizes 15.5% more DC bus than SPWM, affecting Ripple Coefficient"
                )
                st.caption("↑ Topology & modulation determine ripple coefficient (2-Level SPWM/SVPWM=4; 3-Level SPWM=8; 3-Level SVPWM=12), used in Tab 3 analysis")

            # 
            # : Holmes & Lipo, "PWM for Power Converters", IEEE Press, 2003
            if "3-Level" in s_topology and "SVPWM" in s_modulation:
                ripple_denom = 12
            elif "3-Level" in s_topology:
                ripple_denom = 8
            else:
                ripple_denom = 4

        # 
        s_omega1  = 2 * np.pi * s_fg
        s_omegasw = 2 * np.pi * s_fsw
        sL1 = s_L1 * 1e-3
        sL2 = s_L2 * 1e-3
        sLg = s_Lg * 1e-3
        sCf = s_Cf * 1e-6

        # 
        if is_delta:
            s_Uc = s_ug
            s_Ic_line = s_il / np.sqrt(3)
        else:
            s_Uc = s_ug / np.sqrt(3)
            s_Ic_line = s_il

        st.divider()
        st.markdown("### 🔑 Rated Parameters")

        # 1. 
        Uc_margin  = s_Uc * 1.1
        Uc_select  = s_Uc * 1.5
        Uc_std     = int(np.ceil(Uc_select / 50) * 50)

        # 2. ripple current(Lg, /)
        Ic_fund = s_Uc * s_omega1 * sCf
        delta_I_pp = s_vdc / (ripple_denom * sL1 * s_fsw) if sL1 > 0 else 0.0
        ZL2_sw  = s_omegasw * (sL2 + sLg)
        ZCf_sw  = 1.0 / (s_omegasw * sCf) if sCf > 0 else 1e9
        denom_sw = abs(ZL2_sw - ZCf_sw)
        Ic_sw_pk = (delta_I_pp / 2) * ZL2_sw / denom_sw if denom_sw > 1e-9 else 0.0
        Ic_sw_rms = Ic_sw_pk / np.sqrt(2)
        Ic_total  = np.sqrt(Ic_fund**2 + Ic_sw_rms**2)
        Ic_rated_min = Ic_total * 1.2

        # 3. [Fix2] dv/dt 
        dvdt_actual = Ic_sw_pk / sCf / 1e6 if sCf > 0 else 0.0
        dvdt_conservative = s_vdc / (sL1 * sCf * s_omegasw) / 1e6 if (sL1 > 0 and sCf > 0) else 0.0
        dvdt_design = max(dvdt_actual, dvdt_conservative)

        if is_trial:
            trial_lock("Capacitor design results require paid access")
        else:
            col_p1, col_p2, col_p3 = st.columns(3)
            with col_p1:
                st.markdown("#### ① Rated Voltage")
                if not is_trial:
                    st.metric("Actual voltage (connection correction)", f"{s_Uc:.1f} Vrms")
                    st.metric("recommendedRated Voltage(1.5×)", f"{Uc_std} Vrms AC",
                              delta="IEC 60831: 67% derating recommended", delta_color="off")
                if is_delta:
                    st.error(f"⚠ Delta connection! Rated Voltage must handle **line voltage {s_ug:.0f}V** — much higher than phase voltage!")
                else:
                    st.success(f"Y connection, phase voltage {s_Uc:.1f}V — select capacitor with rated voltage ≥ {Uc_std}V AC")

            with col_p2:
                st.markdown("#### ② Ripple Current (with Lg correction)")
                with st.expander("📐 Calculation details (expand)", expanded=True):
                    st.markdown(f"""
    **Step 1 — Fundamental current** (capacitor voltage × capacitive admittance): 

    $$I_{{fund}} = U_c \\times \\omega_1 \\times C_f = {s_Uc:.1f} \\times 2\\pi \\times {s_fg:.0f} \\times {s_Cf:.1f}\\times10^{{-6}} = {Ic_fund:.2f}\\text{{ A}}$$

    **Step 2 — Inverter output ripple current**: 

    $$\\Delta I_{{pp}} = \\frac{{V_{{dc}}}}{{{ripple_denom} \\cdot L_1 \\cdot f_{{sw}}}} = \\frac{{{s_vdc:.0f}}}{{{ripple_denom} \\times {sL1*1e3:.2f}\\text{{mH}} \\times {s_fsw:.0f}}} = {delta_I_pp:.2f}\\text{{ A(peak-to-peak)}}$$

    > Ripple coefficient = **{ripple_denom}** (determined by topology/modulation, see Tab 3 Ripple Current Analysis)

    **Step 3 — Cf current at Switching Frequency** (with grid leakage inductance Lg correction): 

    $$Z_{{L2+Lg}} = \\omega_{{sw}}(L_2 + L_g) = 2\\pi \\times {s_fsw:.0f} \\times {(sL2+sLg)*1e3:.2f}\\text{{mH}} = {ZL2_sw:.2f}\\text{{ Ω}}$$

    $$Z_{{Cf}} = \\frac{{1}}{{\\omega_{{sw}} C_f}} = \\frac{{1}}{{2\\pi \\times {s_fsw:.0f} \\times {s_Cf:.1f}\\times10^{{-6}}}} = {ZCf_sw:.2f}\\text{{ Ω}}$$

    $$I_{{sw,pk}} = \\frac{{\\Delta I_{{pp}}/2 \\times Z_{{L2+Lg}}}}{{|Z_{{L2+Lg}} - Z_{{Cf}}|}} = {Ic_sw_pk:.2f}\\text{{ A(peak)}}$$

    $$I_{{sw,rms}} = \\frac{{I_{{sw,pk}}}}{{\\sqrt{{2}}}} = {Ic_sw_rms:.2f}\\text{{ Arms}}$$

    **Step 4 — Total ripple current (RSS)**: 

    $$I_{{total}} = \\sqrt{{I_{{fund}}^2 + I_{{sw,rms}}^2}} = \\sqrt{{{Ic_fund:.2f}^2 + {Ic_sw_rms:.2f}^2}} = {Ic_total:.2f}\\text{{ Arms}}$$

    > 📚 Liserre M. et al., IEEE Trans. Ind. Appl., 2005; IEC 61071 §6.5
                    """)
                    if not is_trial:
                        st.metric("Fundamental current", f"{Ic_fund:.2f} Arms")
                        st.metric("Switching freq. current (Lg corrected)", f"{Ic_sw_rms:.2f} Arms")
                        st.metric("Total ripple current (RSS)", f"{Ic_total:.2f} Arms")
                if not is_trial:
                    st.warning(f"Rated Ripple Current ≥ **{Ic_rated_min:.1f} A**(1.2×Margin, IEC 61071 §6.5)")

            with col_p3:
                st.markdown("#### ③ Reactive Power")
                Qc = s_Uc**2 * s_omega1 * sCf
                with st.expander("📐 Calculation details (expand)", expanded=True):
                    st.markdown(f"""
    **Single-phase design capacitance** (from LCL design section): 

    $$C_f = {s_Cf:.1f}\\text{{ μF}}$$( ±5%,  {s_Cf*0.95:.1f}~{s_Cf*1.05:.1f} μF)

    **Single-phase reactive power**: 

    $$Q_{{1\\phi}} = U_c^2 \\times \\omega_1 \\times C_f$$

    $$= {s_Uc:.1f}^2 \\times 2\\pi \\times {s_fg:.0f} \\times {s_Cf:.1f}\\times10^{{-6}}$$

    $$= {Qc:.1f}\\text{{ Var}} = {Qc/1e3:.3f}\\text{{ kVar}}$$

    **3-phase total reactive power**: 

    $$Q_{{3\\phi}} = 3 \\times Q_{{1\\phi}} = 3 \\times {Qc/1e3:.3f} = {Qc*3/1e3:.3f}\\text{{ kVar}}$$

    > Engineering note: The LCL filter injects a small capacitive reactive power into the grid;
    > confirm the system reactive power capacity (typically ≤ 5% of rated active power).
                    """)
                    if not is_trial:
                        st.metric("Design capacitance", f"{s_Cf:.1f} μF ±5%")
                        st.metric("Single-phase reactive power", f"{Qc/1e3:.3f} kVar")
                        st.metric("3-phase reactive power", f"{Qc*3/1e3:.3f} kVar")

            st.divider()

            # ── [Fix2] dv/dt  ──────────────────────────────────────────────────
            st.markdown("### ⚡ dv/dt (IEC 61071 §5.1 requirement)")
            st.error("""
    **Warning**: If actual dv/dt exceeds the capacitor's rated dv/dt, the Schoopage (end-face contact) will overheat,
    causing capacitor failure. IEC 61071 §5.1 requires power capacitors to meet the dv/dt Rating ≥ actual dv/dt.
            """)

            col_dv1, col_dv2, col_dv3 = st.columns(3)
            with col_dv1:
                st.markdown("**Calculation Principle**")
                st.markdown(f"""
    **How does dv/dt relate to ripple current?**

    During switching, the inverter output ripple current $I_{{sw,pk}}$ flows through $L_2$ (and $L_g$) into the $C_f$ impedance.
    The $C_f$ current charges the capacitor voltage:

    $$\\frac{{dv_{{Cf}}}}{{dt}} = \\frac{{i_{{Cf}}(t)}}{{C_f}}$$

     When $i_{{Cf}} = I_{{sw,pk}}$ (worst case): 

    $$\\left.\\frac{{dv}}{{dt}}\\right|_{{max}} = \\frac{{I_{{sw,pk}}}}{{C_f}} = \\frac{{{Ic_sw_pk:.2f}\\text{{ A}}}}{{{sCf*1e6:.1f}\\text{{ μF}}}} = {dvdt_actual:.2f}\\text{{ V/μs}}$$

    > 📚 IEC 61071:2017 §5.1
                """)
                if dvdt_conservative > dvdt_actual:
                    st.caption(f"($V_{{dc}}$ $L_1·C_f$ )= {dvdt_conservative:.2f} V/μs,  **{dvdt_design:.2f} V/μs**")

            with col_dv2:
                dvdt_rated = st.number_input(
     " dv/dt Rating (V/μs)",
                    min_value=0.0, value=50.0, step=1.0, key="s10_dvdt",
                    help="Capacitor rated max dv/dt or du/dt from datasheet"
                )
                dvdt_ok = dvdt_rated >= dvdt_design
                if dvdt_ok:
                    st.success(f"✅ Safe: rated {dvdt_rated:.1f} V/μs ≥ actual {dvdt_design:.2f} V/μs")
                else:
                    st.error(f"❌ Overstress: actual {dvdt_design:.2f} V/μs > rated {dvdt_rated:.1f} V/μs — select a capacitor with higher dv/dt rating!")

            with col_dv3:
                st.markdown("**Common MKP capacitor dv/dt reference**")
                st.markdown("""
    | | dv/dt Rating |
    |------|------------|
    | EPCOS B32373 | 40 V/μs |
    | EPCOS B32776 | 50 V/μs |
    | WIMA MKP 10  | 100 V/μs |
    | Vishay MKP1848 | 50 V/μs |
    | CBB65 | 25 V/μs |

    > Select a capacitor whose rated dv/dt ≥ the calculated value above;
    > ensure the Cf design capacitance is consistent with the dv/dt calculation.
                """)

            st.divider()

            # ──  ──────────────────────────────────────────────────────────
            st.markdown("### ⬇️ Voltage Derating Design")
            st.markdown(f"""
    Per **TDK EPCOS Film Capacitors General Technical Information Rev.2023, §4**,
    the MKP capacitor voltage–lifetime model is:

    $$L = L_0 \\times \\left(\\frac{{U_{{rated}}}}{{U_{{op}}}}\\right)^n, \\quad n = 7 \\sim 9 \\text{{(MKP, IEC 61071)}}$$

    Current operating voltage = **{s_Uc:.1f} Vrms** ({conn_type[:1]} connection)

    | Rated Voltage | Derating Ratio | Lifetime Factor (n=7) | Reliability |
    |---------|---------|--------------|------|
    | {s_Uc*1.1:.0f} Vrms (min.) | 91% | 1.0× | ⚠ Not recommended |
    | {s_Uc*1.25:.0f} Vrms | 80% | 1.7× | ✅ Acceptable |
    | **{s_Uc*1.5:.0f} Vrms** | **67%** | **4.5×** | **✅✅ Recommended** |
    | {s_Uc*2.0:.0f} Vrms | 50% | 16× | ✅✅✅ Conservative |

    > Note: Lifetime factor = $(U_{{rated}}/U_{{op}})^7$, normalized to the minimum rated voltage as reference.
    > 📚 TDK EPCOS General Technical Information §4; IEC 61071:2017 §6.5
            """)

            st.divider()

            # ══════════════════════════════════════════════════════════════════════
            # ()
            # ══════════════════════════════════════════════════════════════════════
            st.markdown("### 🔢 Capacitor Parallel Configuration Calculator")
            st.info("""
    **Engineering context**: In medium/large power converters (kW–MW class), the LCL filter capacitor may require
    hundreds of μF and 100+ Arms ripple current — far beyond a single capacitor; multiple units in parallel are used.
    Parallel configuration must satisfy all constraints simultaneously: capacitance, ripple current, dv/dt, and rated voltage.
            """)

            with st.expander("⚙️ Parallel Configuration Calculator (expand)", expanded=True):
                pa_col1, pa_col2 = st.columns([1, 1.5])

                with pa_col1:
                    st.markdown("**① System Requirements** (from above calculations, adjustable)")
                    pa_Cf_total = st.number_input(
                        "Required Cf (μF, single-phase)",
                        value=float(s_Cf), step=5.0, key="pa_cf",
                        help="LCL design requirement per phase"
                    )
                    pa_Ir_total = st.number_input(
                        "Required Ripple Current (Arms, 1.2× margin)",
                        value=float(Ic_rated_min), step=1.0, key="pa_ir",
                        help="Rated Ripple Current from above calculation"
                    )
                    pa_dvdt_req = st.number_input(
                        "Required dv/dt (V/μs)",
                        value=float(max(dvdt_design, 1.0)), step=1.0, key="pa_dvdt",
                        help="dv/dt design value from above dv/dt calculation"
                    )
                    pa_Un_req = st.number_input(
                        "Required Rated Voltage (Vrms)",
                        value=float(Uc_std), step=50.0, key="pa_un",
                        help="1.5× derating recommended Rated Voltage"
                    )

                    st.markdown("**② Per-unit Capacitor Parameters** (from recommended series)")
                    pa_Cn = st.number_input("Unit capacitance Cn (μF)", value=50.0, step=5.0, key="pa_cn")
                    pa_Ir1 = st.number_input("Unit Rated Ripple Current (A)", value=30.0, step=1.0, key="pa_ir1")
                    pa_dvdt1 = st.number_input("Unit dv/dt rated (V/μs)", value=50.0, step=5.0, key="pa_dvdt1")
                    pa_Un1 = st.number_input("Unit Rated Voltage (Vrms)", value=float(Uc_std), step=50.0, key="pa_un1")
                    pa_ESR1 = st.number_input("Unit ESR @ fsw (mΩ)", value=4.0, step=0.5, key="pa_esr1")

                with pa_col2:
                    st.markdown("**③ Parallel Configuration Result**")

                    # 
                    n_Cf  = int(np.ceil(pa_Cf_total / pa_Cn))   if pa_Cn > 0   else 999
                    n_Ir  = int(np.ceil(pa_Ir_total / pa_Ir1))  if pa_Ir1 > 0  else 999
                    n_dvdt = 1 if pa_dvdt1 >= pa_dvdt_req else 999  # dv/dtparallel, satisfied
                    n_Un   = 1 if pa_Un1 >= pa_Un_req else 999

                    # dv/dtRated Voltage, 
                    n_min = max(n_Cf, n_Ir)  # capacitance and ripple current both drive parallel count

                    st.markdown(f"""
    | Constraint | Calculation | Count | Status |
    |---------|---------|---------|------|
    | Capacitance met | {pa_Cf_total:.0f} μF ÷ {pa_Cn:.0f} μF/unit | **{n_Cf} units** | {'✅' if n_Cf <= 20 else '⚠ Too many'} |
    | Ripple current met | {pa_Ir_total:.1f} A ÷ {pa_Ir1:.0f} A/unit | **{n_Ir} units** | {'✅' if n_Ir <= 20 else '⚠ Too many'} |
    | dv/dt met | Rated {pa_dvdt1:.0f} {'≥' if pa_dvdt1>=pa_dvdt_req else '<'} {pa_dvdt_req:.1f} V/μs | **1 unit sufficient** | {'✅' if pa_dvdt1>=pa_dvdt_req else '❌ Not met, replace unit'} |
    | Rated Voltage met | Rated {pa_Un1:.0f} {'≥' if pa_Un1>=pa_Un_req else '<'} {pa_Un_req:.0f} Vrms | **1 unit sufficient** | {'✅' if pa_Un1>=pa_Un_req else '❌ Not met, replace unit'} |
                    """)

                    if pa_dvdt1 < pa_dvdt_req or pa_Un1 < pa_Un_req:
                        st.error("❌ This capacitor unit does not meet dv/dt or Rated Voltage requirements — **replacing with a higher-spec unit is required**, not more parallels!")
                    else:
                        # (, /)
                        n_rec = n_min
                        # ()
                        if n_rec > 1 and n_rec % 2 != 0:
                            n_rec_even = n_rec + 1
                        else:
                            n_rec_even = n_rec

                        # Parameter
                        Cf_actual = n_rec_even * pa_Cn
                        Ir_actual = n_rec_even * pa_Ir1
                        ESR_actual = pa_ESR1 / n_rec_even
                        margin_Cf = (Cf_actual / pa_Cf_total - 1) * 100
                        margin_Ir = (Ir_actual / pa_Ir_total - 1) * 100

                        if n_rec_even != n_rec:
                            st.info(f"💡 **{n_min} units** is mathematically minimum, engineering recommended → **{n_rec_even} units** (even count for symmetric layout)")
                        else:
                            st.success(f"✅ Recommended parallel count: **{n_rec_even} units**")

                        res_col1, res_col2, res_col3 = st.columns(3)
                        res_col1.metric("Parallel capacitance", f"{Cf_actual:.0f} μF",
                                        delta=f"Margin +{margin_Cf:.0f}%", delta_color="normal")
                        res_col2.metric("Parallel ripple rating", f"{Ir_actual:.0f} Arms",
                                        delta=f"Margin +{margin_Ir:.0f}%", delta_color="normal")
                        res_col3.metric("Parallel ESR", f"{ESR_actual:.2f} mΩ",
                                        delta=f"{pa_ESR1:.1f}mΩ ÷ {n_rec_even} units", delta_color="off")

                        st.markdown(f"""
    **📋 selection(per phase)**

    | | value |
    |-----|------|
    | Parallel Count(per phase) | **{n_rec_even} ** |
    | 3-phase | **{n_rec_even * 3} ** |
    | requirement | Un ≥ {pa_Un_req:.0f} Vrms, Cn = {pa_Cn:.0f} μF, Ir ≥ {pa_Ir1:.0f} A, dv/dt ≥ {pa_dvdt_req:.1f} V/μs |
    | parallel(per phase) | {Cf_actual:.0f} μF(exceedsdesign +{margin_Cf:.0f}%) |
    | parallelrated(per phase) | {Ir_actual:.0f} Arms(exceeds +{margin_Ir:.0f}%) |
    | parallel ESR(per phase) | {ESR_actual:.2f} mΩ |
    | parallelthermal dissipation(single-phase, {pa_Ir_total:.0f}A) | {(pa_Ir_total**2 * ESR_actual * 1e-3):.1f} W |

    > ⚠ **parallel**:
    > 1. parallelcapacitor****, impedanceripple current
    > 2. symmetric,
    > 3. units in parallel ±5% , suggestion
    > 4. power(≥6)suggestion,
                        """)

                        # : Inductive
                        n_range = np.arange(1, min(n_min + 6, 21))
                        margin_Cf_arr = (n_range * pa_Cn / pa_Cf_total - 1) * 100
                        margin_Ir_arr = (n_range * pa_Ir1 / pa_Ir_total - 1) * 100

                        fig_pa = go.Figure()
                        fig_pa.add_trace(go.Bar(x=n_range, y=margin_Cf_arr, name="Capacitance Margin (%)",
                                                marker_color="#4dabf7", opacity=0.75))
                        fig_pa.add_trace(go.Bar(x=n_range, y=margin_Ir_arr, name="Ripple Current Margin (%)",
                                                marker_color="#40c057", opacity=0.75))
                        fig_pa.add_vline(x=n_rec_even, line_dash="dash", line_color="#e03131",
                                         annotation_text=f"Recommended {n_rec_even} units")
                        fig_pa.add_hline(y=0, line_color="#868e96", line_width=1)
                        fig_pa.update_layout(
                            title="Parallel Count vs ParameterMargin",
                            xaxis_title="Parallel Count(per phase)", yaxis_title="Margin(%)",
                            barmode="group", height=280, template="plotly_white",
                            legend=dict(orientation="h", y=1.1)
                        )
                        st.plotly_chart(fig_pa, use_container_width=True)

        # ══════════════════════════════════════════════════════════════════════════
        # TAB 3 — Ripple Current Analysis
        # ══════════════════════════════════════════════════════════════════════════
    # Pre-compute ripple chart for trial mode blurred preview
    _cf_scan_p = np.linspace(10, max(400, _Cf_uF * 2), 300)
    _dI_p = _Vdc / (4 * L1_base * _fsw) if L1_base > 0 else 0.0
    _Ic_f_p, _Ic_sw_p, _Ic_tot_p, _fres_p = [], [], [], []
    for _c in _cf_scan_p * 1e-6:
        _if_c = Ug_phase * omega1 * _c
        _ZL_p = omega_sw * L2_base
        _ZC_p = 1.0 / (omega_sw * _c) if _c > 0 else 1e9
        _d_p = abs(_ZL_p - _ZC_p)
        _sw_p = ((_dI_p / 2) * _ZL_p / _d_p / np.sqrt(2)) if _d_p > 1e-9 else 0.0
        _fr_p = (1/(2*np.pi))*np.sqrt((L1_base+L2_base)/(L1_base*L2_base*_c)) if (_c > 0 and L1_base > 0 and L2_base > 0) else 0.0
        _Ic_f_p.append(_if_c)
        _Ic_sw_p.append(_sw_p)
        _Ic_tot_p.append(np.sqrt(_if_c**2 + _sw_p**2))
        _fres_p.append(_fr_p)
    _Ic_f_p = np.array(_Ic_f_p); _Ic_sw_p = np.array(_Ic_sw_p)
    _Ic_tot_p = np.array(_Ic_tot_p); _fres_p = np.array(_fres_p)
    _fig_rip_p = make_subplots(rows=1, cols=2,
                               subplot_titles=("Ripple Current vs Cf", "Ripple Current vs fres"))
    _fig_rip_p.add_trace(go.Scatter(x=_cf_scan_p, y=_Ic_f_p, name="Fundamental",
                                    line=dict(color="#0d6efd", width=2)), row=1, col=1)
    _fig_rip_p.add_trace(go.Scatter(x=_cf_scan_p, y=_Ic_sw_p, name="Switching Frequency Ripple",
                                    line=dict(color="#fd7e14", width=2)), row=1, col=1)
    _fig_rip_p.add_trace(go.Scatter(x=_cf_scan_p, y=_Ic_tot_p, name="Total (RSS)",
                                    line=dict(color="#6f42c1", width=2.5)), row=1, col=1)
    _fig_rip_p.add_trace(go.Scatter(x=_fres_p, y=_Ic_tot_p, name="Total vs fres",
                                    line=dict(color="#6f42c1", width=2.5), showlegend=False), row=1, col=2)
    _fig_rip_p.update_layout(height=380, template="plotly_white", hovermode=False,
                             legend=dict(yanchor="top", y=0.98, xanchor="left", x=0.01))

    with tab_ripple:
        if is_trial:
            _fig_t = go.Figure(_fig_rip_p)
            _fig_t.update_traces(hoverinfo='none', hovertemplate=None)
            _fig_t.update_yaxes(showticklabels=False)
            st.plotly_chart(_fig_t, use_container_width=True)
            trial_lock("Ripple current analysis and harmonic spectrum require paid access", compact=True)
        else:
            st.markdown("### 〰️ Ripple Current Analysis")
            st.info("""
    **Note**: Ripple current causes capacitor thermal dissipation ($P = I_{ripple}^2 \\times ESR$).
    This tool has two functions:
    ① **Ripple current vs Cf**: Sweep analysis across capacitance, helping verify current at the Design Point and optimize the design;
    ② **Harmonic spectrum**: Decompose ripple current by harmonic order, helping evaluate total RSS ripple current.
            """)


            # ParameterTab2, (), 
            st.markdown("#### Parameter Settings (from Tab 2, adjustable)")
            col_r1, col_r2, col_r3, col_r4 = st.columns(4)
            with col_r1:
                r_L1 = st.number_input("L1 (mH)", value=float(s_L1), step=0.05, format="%.3f", key="s10r_L1")
            with col_r2:
                r_Cf = st.number_input("Cf (μF)", value=float(s_Cf), step=5.0, key="s10r_Cf")
            with col_r3:
                r_L2 = st.number_input("L2 (mH)", value=float(s_L2), step=0.05, format="%.3f", key="s10r_L2")
            with col_r4:
                r_Lg = st.number_input("Lg (mH)", value=float(s_Lg), step=0.01, format="%.3f", key="s10r_Lg",
                help="Grid leakage inductance affects Cf current at switching frequency")

            rL1 = r_L1 * 1e-3
            rL2 = r_L2 * 1e-3
            rLg = r_Lg * 1e-3
            rCf = r_Cf * 1e-6

            # Cf 
            cf_scan_uF = np.linspace(10, max(400, r_Cf * 2), 300)
            cf_scan    = cf_scan_uF * 1e-6
            Ic_fund_s, Ic_sw_s, Ic_tot_s, fres_s = [], [], [], []
            delta_I_base = s_vdc / (ripple_denom * rL1 * s_fsw) if rL1 > 0 else 0.0

            for cf_i in cf_scan:
                ic_f  = (Ug_phase * omega1 * cf_i)
                ZL2_i = omega_sw * (rL2 + rLg)
                ZCf_i = 1.0 / (omega_sw * cf_i) if cf_i > 0 else 1e9
                den_i = abs(ZL2_i - ZCf_i)
                ic_sw = ((delta_I_base / 2) * ZL2_i / den_i / np.sqrt(2)) if den_i > 1e-9 else 0.0
                ic_tot = np.sqrt(ic_f**2 + ic_sw**2)
                fres_i = (1 / (2 * np.pi)) * np.sqrt((rL1 + rL2) / (rL1 * rL2 * cf_i)) if (cf_i > 0 and rL1 > 0 and rL2 > 0) else 0.0
                Ic_fund_s.append(ic_f)
                Ic_sw_s.append(ic_sw)
                Ic_tot_s.append(ic_tot)
                fres_s.append(fres_i)

            Ic_fund_s = np.array(Ic_fund_s)
            Ic_sw_s   = np.array(Ic_sw_s)
            Ic_tot_s  = np.array(Ic_tot_s)
            fres_s    = np.array(fres_s)

            # Design Point
            cur_ic_fund = Ug_phase * omega1 * rCf
            ZL2_cur = omega_sw * (rL2 + rLg)
            ZCf_cur = 1.0 / (omega_sw * rCf) if rCf > 0 else 1e9
            den_cur = abs(ZL2_cur - ZCf_cur)
            cur_ic_sw_pk  = (delta_I_base / 2) * ZL2_cur / den_cur if den_cur > 1e-9 else 0.0
            cur_ic_sw_rms = cur_ic_sw_pk / np.sqrt(2)
            cur_ic_tot = np.sqrt(cur_ic_fund**2 + cur_ic_sw_rms**2)
            cur_fres = (1/(2*np.pi)) * np.sqrt((rL1+rL2)/(rL1*rL2*rCf)) if (rCf > 0 and rL1 > 0 and rL2 > 0) else 0.0

            # 1+2
            st.markdown("#### 📈 Ripple Current vs Capacitance / Resonance Frequency")
            st.caption("""
    **How to read**: Left chart X-axis is Cf, Y-axis is RMS current at each frequency component.
    **Blue (Fundamental)** increases linearly with Cf; **Orange (Switching Frequency Ripple)** decreases as Cf increases (larger Cf → lower impedance → more ripple shunted, but lower $Z_{Cf}$);
    **Purple (Total RSS)** is the combined result — its minimum point corresponds to the "optimal ripple design point".
    Right chart uses Resonance Frequency $f_{res}$ as X-axis, showing how fres affects total ripple — as fres approaches fsw, ripple amplifies sharply and must be avoided.
            """)

            fig_rip = make_subplots(rows=1, cols=2,
                                    subplot_titles=("Ripple Current vs Cf (sweep)", "Ripple Current vs Resonance Frequency"))
            fig_rip.add_trace(go.Scatter(x=cf_scan_uF, y=Ic_fund_s, name="Fundamental current",
                                         line=dict(color="#0d6efd", width=2)), row=1, col=1)
            fig_rip.add_trace(go.Scatter(x=cf_scan_uF, y=Ic_sw_s, name="Switching Frequency current",
                                         line=dict(color="#fd7e14", width=2)), row=1, col=1)
            fig_rip.add_trace(go.Scatter(x=cf_scan_uF, y=Ic_tot_s, name="Total (RSS)",
                                         line=dict(color="#6f42c1", width=2.5)), row=1, col=1)
            fig_rip.add_vline(x=r_Cf, line_dash="dot", line_color="#e03131",
                              annotation_text=f"Current Cf={r_Cf:.0f}μF, Total={cur_ic_tot:.1f}A",
                              row=1, col=1)
            fig_rip.add_trace(go.Scatter(x=fres_s, y=Ic_tot_s, name="Total vs fres",
                                         line=dict(color="#6f42c1", width=2.5), showlegend=False), row=1, col=2)
            if cur_fres > 0:
                fig_rip.add_vline(x=cur_fres, line_dash="dot", line_color="#e03131",
                                  annotation_text=f"Current fres={cur_fres:.0f}Hz", row=1, col=2)
                fig_rip.add_vline(x=s_fsw, line_dash="dash", line_color="#868e96",
                                  annotation_text=f"fsw={s_fsw:.0f}Hz (switching)", row=1, col=2)
            fig_rip.update_xaxes(title_text="Cf (μF)", row=1, col=1)
            fig_rip.update_xaxes(title_text="Resonance Frequency fres (Hz)", row=1, col=2)
            fig_rip.update_yaxes(title_text="current (Arms)", row=1, col=1)
            fig_rip.update_layout(height=380, template="plotly_white",
                                  legend=dict(yanchor="top", y=0.98, xanchor="left", x=0.01))
            st.plotly_chart(fig_rip, use_container_width=True)

            # Result()
            fsw_margin = cur_fres / s_fsw if s_fsw > 0 else 0
            if fsw_margin > 0:
                if fsw_margin < 0.3 or fsw_margin > 3.0:
                    fres_verdict = f"❌ **Resonance frequency {cur_fres:.0f} Hz is too close to switching frequency {s_fsw:.0f} Hz**, filtering performance degraded — suggest adjusting L or Cf!"
                elif 0.3 <= fsw_margin < 0.5:
                    fres_verdict = f"⚠ Resonance frequency {cur_fres:.0f} Hz is {fsw_margin:.0%} of fsw — marginal, recommend checking damping design."
                else:
                    fres_verdict = f"✅ Resonance frequency {cur_fres:.0f} Hz is {fsw_margin:.0%} of fsw — adequate separation, no resonance concern, design is acceptable."
            else:
                fres_verdict = "⚠ L2 parameters not set (L2=0 cannot calculate resonance frequency)."

            st.markdown(f"**📋 Stability Result**: {fres_verdict}")
            st.markdown(f"**📋 Design Point**: Fundamental **{cur_ic_fund:.1f} A**, Switching component **{cur_ic_sw_rms:.1f} Arms**, Total RSS **{cur_ic_tot:.1f} Arms**. "
                        f"{'Fundamental dominates — Cf may be conservative; consider reducing capacitance to lower reactive power. ' if cur_ic_fund > cur_ic_sw_rms else 'High-frequency component dominates — Cf is near the limit; consider increasing capacitance or using a lower-loss unit. '}"
                        f"Select Rated Ripple Current ≥ **{cur_ic_tot * 1.2:.1f} Arms** (1.2× Margin).")

            st.divider()

            #
            st.markdown("#### 🎛️ Harmonic Spectrum of Capacitor Ripple Current")
            st.caption("""
    **How to read**: X-axis is frequency; Y-axis is capacitor RMS current at each frequency component.
    **Blue bar** = fundamental (50/60 Hz); **Red bars** = PWM switching harmonics (Switching Frequency sidebands).
    RSS combines all components. Capacitor thermal dissipation is driven by total RSS — select Rated Ripple Current ≥ RSS.
            """)
            from scipy.special import jv as bessel_j
            r_ma = st.number_input("modulation ma", min_value=0.5, max_value=1.15,
                                   value=0.9, step=0.01, key="s10r_ma",
                                   help="SPWM: 0.85~1.0; SVPWM: up to 1.15 with 3rd harmonic injection")

            harm_f_list, ic_h_list = [s_fg], [Ug_phase * omega1 * rCf]
            for m in range(1, 5):
                for n in range(-3, 4):
                    f_h = m * s_fsw + n * s_fg
                    if f_h < s_fg * 2:
                        continue
                    omega_h = 2 * np.pi * f_h
                    try:
                        J_val    = abs(bessel_j(n, m * np.pi * r_ma / 2.0))
                        i_inv_pk = (4 * s_vdc * J_val) / (m * np.pi**2 * omega_h * rL1) if rL1 > 0 else 0.0
                    except Exception:
                        i_inv_pk = 0.0
                    ZL2_h = omega_h * (rL2 + rLg)
                    ZCf_h = 1.0 / (omega_h * rCf) if rCf > 0 else 1e9
                    den_h = abs(ZL2_h - ZCf_h)
                    ic_h  = (i_inv_pk * ZL2_h / den_h / np.sqrt(2)) if den_h > 1e-9 else 0.0
                    harm_f_list.append(f_h)
                    ic_h_list.append(ic_h)

            harm_f_arr = np.array(harm_f_list)
            ic_h_arr   = np.array(ic_h_list)
            Ic_rss     = np.sqrt(np.sum(ic_h_arr**2))

            colors_bar = ["#0d6efd" if f == s_fg else "#dc3545" for f in harm_f_arr]
            fig_spec = go.Figure()
            fig_spec.add_trace(go.Bar(x=harm_f_arr, y=ic_h_arr, marker_color=colors_bar,
                                      hovertemplate="%{x:.0f}Hz<br>%{y:.3f}Arms<extra></extra>"))
            fig_spec.add_hline(y=Ic_rss, line_dash="dash", line_color="#fd7e14",
                               annotation_text=f"RSS={Ic_rss:.2f} Arms (Select rated ≥ {Ic_rss*1.2:.1f} A)")
            fig_spec.update_layout(title="Filter Capacitor Ripple Current Spectrum (Individual Harmonics)",
                                   xaxis_title="Frequency (Hz)", yaxis_title="Capacitor Current (Arms)",
                                   height=360, template="plotly_white")
            st.plotly_chart(fig_spec, use_container_width=True)

            # Result
            ic_fund_val = ic_h_arr[0]
            ic_hf_arr   = ic_h_arr[1:]
            ic_hf_rss   = np.sqrt(np.sum(ic_hf_arr**2))
            dominant_idx = np.argmax(ic_hf_arr) if len(ic_hf_arr) > 0 else 0
            dominant_f   = harm_f_arr[1:][dominant_idx] if len(ic_hf_arr) > 0 else 0
            dominant_ic  = ic_hf_arr[dominant_idx] if len(ic_hf_arr) > 0 else 0

            col_s1, col_s2, col_s3 = st.columns(3)
            col_s1.metric("Fundamental current", f"{ic_fund_val:.2f} Arms")
            col_s2.metric("High-frequency harmonic RSS", f"{ic_hf_rss:.2f} Arms")
            col_s3.metric("Total (RSS)", f"{Ic_rss:.2f} Arms",
                          delta=f"Selection rated ≥ {Ic_rss*1.2:.1f} A", delta_color="off")

            if dominant_ic > 0:
                hf_ratio = ic_hf_rss / Ic_rss
                st.markdown(f"""
    **📋 Analysis Result**:

    - **Dominant harmonic**: **{dominant_f:.0f} Hz** ({dominant_ic:.2f} Arms), located at {int(round(dominant_f/s_fsw))}×fsw±{int(round(abs(dominant_f - round(dominant_f/s_fsw)*s_fsw)/s_fg))}×fg sideband.
    - **Thermal dissipation source**: High-frequency harmonics account for **{hf_ratio:.0%}** of total — {'(<50%) fundamental reactive current dominates; consider reducing Cf to lower reactive power.' if hf_ratio < 0.5 else '(>50%) switching harmonics dominate capacitor heating; consider increasing Cf or reducing L1 to shift the design point.'}
    - **Selection requirement**: Capacitor Rated Ripple Current ≥ **{Ic_rss*1.2:.1f} Arms** (RSS × 1.2 Margin, IEC 61071 §6.5).

    > 📌 The Cf design result gives the minimum Rated Ripple Current for capacitor selection. This harmonic spectrum is for verification and thermal estimation reference only. If a harmonic component exceeds 30%, set modulation index $m_a$ to the actual operating value for accuracy.
                """)

        # ══════════════════════════════════════════════════════════════════════════
        # TAB 4 — ··(v2, )
        # ══════════════════════════════════════════════════════════════════════════
    with tab_life:
        st.markdown("### 🌡️ Lifetime Prediction & Evaluation")

        st.markdown(r"""
#### Lifetime Model (Arrhenius + Voltage Acceleration)

**Arrhenius + voltage acceleration** (standard engineering model): 

$$L = L_0 \times 2^{\frac{T_{max}-T_h}{10}} \times \left(\frac{U_{rated}}{U_{op}}\right)^n$$

**Formula Notation**:


| Symbol | Meaning | Typical Range |
|------|------|---------|
| $L$ | Predicted lifetime (h) | Calculated result |
| $L_0$ | Rated Lifetime at rated conditions ($T_{max}$, $U_{rated}$) | Datasheet value, typically 50000~200000 h |
| $T_{max}$ | Capacitor maximum rated temperature (°C) | 85°C or 105°C |
| $T_h$ | Hotspot temperature (°C), formula: $T_h = T_{amb} + P \cdot R_{th}$ | Calculated |
| $U_{rated}$ | Capacitor Rated Voltage (Vrms) | Datasheet value |
| $U_{op}$ | Actual operating voltage (Vrms) | System value |
| $n$ | Voltage lifetime exponent | MKP typically 7 (electrolytic 3~5, much lower!) |

- **Temperature rule**: Every +10°C reduces lifetime by half (Arrhenius model)
- **Voltage exponent**: $n = 7$ (MKP standard, per IEC 61071 & TDK general tech info) — electrolytic uses 3~5, MKP 7~9 — do not confuse!

> 📚 TDK EPCOS Film Capacitors General Technical Information Rev.2023, §4; IEC 61071:2017 §6.5
        """)

        st.divider()

        life_mode = st.radio(
            "Evaluation Mode",
            ["Single-point evaluation", "Mission Profile evaluation (Miner's Rule)"],
            horizontal=True, key="s10_lifemode"
        )

        # ──  ───────────────────────────────────────────────────
        with st.expander("📋 Capacitor Parameters (from datasheet)", expanded=True):
            lc1, lc2, lc3 = st.columns(3)
            with lc1:
                lc_L0   = st.number_input("Rated Lifetime L0 (h)",          value=100000, step=10000, key="s10_L0")
                lc_Tmax = st.number_input("Rated max temperature Tmax (°C)", value=85.0, step=5.0, key="s10_Tmax")
            with lc2:
                lc_ESR  = st.number_input("ESR @ fsw (mΩ)",           value=3.0,    step=0.5,   key="s10_ESR")
                lc_Rth = st.number_input("Thermal resistance Rth (°C/W)", value=2.0, step=0.1, key="s10_Rth")
            with lc3:
                lc_Ir   = st.number_input("Rated Ripple Current Ir_rated (A)", value=50.0,   step=1.0,   key="s10_Ir")
                lc_Un   = st.number_input("Rated Voltage Un (Vrms)",        value=float(int(np.ceil(s_Uc*1.5/50)*50)),
                                           step=50.0, key="s10_Un")
        n_voltage = 7.0

        def calc_life(Tamb, Irms, Uc_op):
            """Calculate lifetime with temperature and voltage acceleration factors"""
            P_rip  = (Irms ** 2) * (lc_ESR * 1e-3)
            P_fund = s_Uc**2 * s_omega1 * sCf * 0.001
            Th     = Tamb + (P_rip + P_fund) * lc_Rth + 3.0
            L_T    = lc_L0 * (2 ** ((lc_Tmax - Th) / 10.0))
            V_rat  = Uc_op / lc_Un if lc_Un > 0 else 1.0
            L_V    = (1.0 / V_rat) ** n_voltage if V_rat > 0 else 1.0
            L_tot  = L_T * L_V
            return Th, L_tot, P_rip

        if is_trial:
            trial_lock("Capacitor lifetime prediction results require paid access")
        elif life_mode == "Single-point evaluation":
            col_lc1, col_lc2 = st.columns(2)
            with col_lc1:
                lc_Tamb = st.number_input("Ambient temperature Tamb (°C)", value=40.0, step=1.0, key="s10_Tamb")
                lc_Irms = st.number_input("Actual ripple current (Arms)", value=float(Ic_total), step=1.0, key="s10_Iact")
                lc_Ucop = st.number_input("Actual operating voltage (Vrms)", value=float(s_Uc), step=10.0, key="s10_Ucop")

            with col_lc2:
                Th, L_tot, P_rip = calc_life(lc_Tamb, lc_Irms, lc_Ucop)
                L_years = L_tot / 8760
                ripple_ratio = lc_Irms / lc_Ir if lc_Ir > 0 else 0.0
                st.metric("Ripple thermal dissipation", f"{P_rip:.2f} W")
                st.metric("Hotspot Temp Th", f"{Th:.1f} °C",
                          delta="✅ Safe" if Th < lc_Tmax else f"❌ Exceeds Tmax={lc_Tmax}°C",
                          delta_color="normal" if Th < lc_Tmax else "inverse")
                st.metric("Ripple current utilization", f"{ripple_ratio*100:.1f}%",
                          delta="✅ Within rated" if ripple_ratio <= 1.0 else "❌ Exceeds rated",
                          delta_color="normal" if ripple_ratio <= 1.0 else "inverse")
                st.metric("Predicted lifetime", f"{L_tot/1e4:.1f} ×10⁴ h ({L_years:.1f} yr)",
                          delta="✅ >20 yr" if L_years >= 20 else "⚠ <20 yr, optimize design",
                          delta_color="normal" if L_years >= 20 else "inverse")

        else:
            # ── [Add2] Miner  ──────────────────────────────────
            st.markdown("#### 🔄 Mission Profile Input (Miner's linear damage rule)")
            st.markdown("""
**Miner's Rule** (IEC 60300-3-1:2003):

$$\\frac{1}{L_{total}} = \\sum_i \\frac{p_i}{L_i}$$

Where $p_i$ is the time fraction of operating condition $i$, and $L_i$ is the lifetime under that condition.
            """)
            st.caption("All condition time fractions must sum to 100%.")

            mission_data = []
            col_m = st.columns(3)
            labels = ["Full Load", "Half Load", "Standby"]
            default_pct   = [20.0, 50.0, 30.0]
            default_tamb  = [50.0, 40.0, 30.0]
            default_irms  = [float(Ic_total), float(Ic_total*0.55), float(Ic_total*0.1)]
            default_uc    = [float(s_Uc), float(s_Uc), float(s_Uc*0.9)]

            for i, (col, lbl) in enumerate(zip(col_m, labels)):
                with col:
                    st.markdown(f"**{lbl}**")
                    pct = st.number_input(f"Time fraction (%)", value=default_pct[i], step=1.0, key=f"s10m_pct{i}")
                    tamb = st.number_input(f"Ambient temp (°C)", value=default_tamb[i], step=1.0, key=f"s10m_T{i}")
                    irms = st.number_input(f"Current (A)", value=default_irms[i], step=1.0, key=f"s10m_I{i}")
                    ucop = st.number_input(f"Voltage (V)", value=default_uc[i], step=5.0, key=f"s10m_U{i}")
                    mission_data.append((pct/100.0, tamb, irms, ucop, lbl))

            total_pct = sum(d[0] for d in mission_data)
            if abs(total_pct - 1.0) > 0.01:
                st.error(f"⚠ Total time fractions sum to {total_pct*100:.1f}%, must equal 100% — please adjust!")
            else:
                miner_sum = 0.0
                rows = []
                for pct, tamb, irms, ucop, lbl in mission_data:
                    Th_i, L_i, P_i = calc_life(tamb, irms, ucop)
                    miner_sum += pct / (L_i / 8760)
                    rows.append({
                        "Condition": lbl,
                        "Time Share": f"{pct*100:.0f}%",
                        "Ambient Temp.": f"{tamb:.0f}°C",
                        "Hotspot Temp.": f"{Th_i:.1f}°C",
                        "Ripple Loss (W)": f"{P_i:.2f}W",
                        "Lifetime (yr)": f"{L_i/8760:.1f}",
                        "Damage pi/Li": f"{pct/(L_i/8760):.4f}/yr"
                    })

                L_miner = 1.0 / miner_sum if miner_sum > 0 else float("inf")

                st.dataframe(pd.DataFrame(rows), width="stretch", hide_index=True)

                col_mn1, col_mn2 = st.columns(2)
                col_mn1.metric("Lifetime (Miner's Rule)",
                               f"{L_miner:.1f} yr",
                               delta="✅ >20 yr, design acceptable" if L_miner >= 20 else "⚠ <20 yr, optimize design",
                               delta_color="normal" if L_miner >= 20 else "inverse")
                col_mn2.metric("vs Single-point estimate",
                               f"{calc_life(mission_data[0][1], mission_data[0][2], mission_data[0][3])[1]/8760:.1f} yr",
                               delta=f"Miner result is {L_miner/max(calc_life(mission_data[0][1], mission_data[0][2], mission_data[0][3])[1]/8760,0.1):.1f}× single-point estimate",
                               delta_color="normal")

                if L_miner < 20:
                    st.warning("The selected capacitor's actual lifetime is below 20 years under this mission profile. Recommend increasing rated voltage or improving thermal path.")
                else:
                    st.success("✅ Capacitor selection passes lifetime verification — predicted lifetime exceeds 20 years.")

        # Lifetime vs Hotspot Temperature chart (single-point mode only, hidden in trial)
        if life_mode == "Single-point evaluation" and not is_trial:
            T_rng = np.linspace(20, lc_Tmax + 5, 200)
            V_fac = (lc_Un / lc_Ucop) ** n_voltage if lc_Ucop > 0 else 1.0
            L_rng = lc_L0 * (2 ** ((lc_Tmax - T_rng) / 10.0)) * V_fac / 8760
            Th_cur = Th
            L_cur  = L_years

            fig_lf = go.Figure()
            fig_lf.add_trace(go.Scatter(x=T_rng, y=L_rng, name="Predicted lifetime (with voltage derating)",
                                         line=dict(color="#0d6efd", width=2.5),
                                         fill="tozeroy", fillcolor="rgba(13,110,253,0.07)"))
            fig_lf.add_vline(x=Th_cur, line_dash="dash", line_color="#dc3545",
                              annotation_text=f"Current Th={Th_cur:.1f}°C")
            fig_lf.add_hline(y=20, line_dash="dot", line_color="#198754",
            annotation_text="20 yr target")
            fig_lf.add_trace(go.Scatter(x=[Th_cur], y=[L_cur], mode="markers",
                                         marker=dict(size=14, color="#dc3545", symbol="star"),
                                         name=f"Current design ({L_cur:.1f} yr)"))
            fig_lf.update_layout(title="Lifetime vs Hotspot Temperature (n=7, MKP)",
                                  xaxis_title="Hotspot Temp Th (°C)", yaxis_title="Predicted lifetime (yr)",
                                  height=360, template="plotly_white")
            st.plotly_chart(fig_lf, use_container_width=True)

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 5 — AI & (v2, )
    # ══════════════════════════════════════════════════════════════════════════
    with tab_upload:

        # ── [Add3]  ─────────────────────────────────────────────
        st.markdown("### 🏭 Selection Recommendation")
        st.caption("Recommended capacitor series matching current system calculations (reference data — always verify with manufacturer datasheet)")

        cap_db = pd.DataFrame([
 {"Manufacturer": "EPCOS/TDK", "Series": "B32373", "Un(Vrms)": 305, "Cfrange(μF)": "5~120", "Ir(A)": 30, "dvdt(V/μs)": 40, "ESR(mΩ)": 5, "Notes": "Standard size, cost-effective"},
 {"Manufacturer": "EPCOS/TDK", "Series": "B32776", "Un(Vrms)": 450, "Cfrange(μF)": "10~250", "Ir(A)": 55, "dvdt(V/μs)": 50, "ESR(mΩ)": 4, "Notes": "High ripple current, DC-LINK"},
 {"Manufacturer": "EPCOS/TDK", "Series": "B32778", "Un(Vrms)": 700, "Cfrange(μF)": "10~500", "Ir(A)": 80, "dvdt(V/μs)": 50, "ESR(mΩ)": 3, "Notes": "Large capacity range"},
 {"Manufacturer": "WIMA", "Series": "MKP10", "Un(Vrms)": 400, "Cfrange(μF)": "0.1~100", "Ir(A)": 40, "dvdt(V/μs)": 100,"ESR(mΩ)": 3, "Notes": "High dv/dt, compact size"},
 {"Manufacturer": "WIMA", "Series": "FKP1", "Un(Vrms)": 630, "Cfrange(μF)": "0.1~47", "Ir(A)": 35, "dvdt(V/μs)": 200,"ESR(mΩ)": 2, "Notes": "Very high dv/dt, ultra-low loss"},
 {"Manufacturer": "Vishay", "Series": "MKP1848","Un(Vrms)": 480, "Cfrange(μF)": "1~200", "Ir(A)": 60, "dvdt(V/μs)": 50, "ESR(mΩ)": 4, "Notes": "UL/IEC certified, mainstream inverter"},
 {"Manufacturer": "Vishay", "Series": "MKP338x","Un(Vrms)": 630, "Cfrange(μF)": "1~150", "Ir(A)": 50, "dvdt(V/μs)": 50, "ESR(mΩ)": 3, "Notes": "High voltage, wide range"},
 {"Manufacturer": "Domestic CN", "Series": "CL21/CBB61","Un(Vrms)": 450,"Cfrange(μF)":"1~150", "Ir(A)": 35, "dvdt(V/μs)": 25, "ESR(mΩ)": 6, "Notes": "Cost-effective, lower dv/dt"},
 {"Manufacturer": "Domestic CN", "Series": "CBB65", "Un(Vrms)": 700, "Cfrange(μF)": "5~300", "Ir(A)": 50, "dvdt(V/μs)": 30, "ESR(mΩ)": 5, "Notes": "Medium/high voltage filter"},
        ])

        req_Un = int(np.ceil(s_Uc * 1.5 / 50) * 50)
        req_Ir = Ic_total * 1.2
        req_dvdt = dvdt_design

        matched = cap_db[
            (cap_db["Un(Vrms)"] >= req_Un) &
            (cap_db["Ir(A)"] >= req_Ir) &
            (cap_db["dvdt(V/μs)"] >= req_dvdt)
        ].copy()

        st.markdown(f"**Current Design Requirements**: Rated Voltage ≥ {req_Un}V | ripple current ≥ {req_Ir:.1f}A | dv/dt ≥ {req_dvdt:.1f}V/μs")

        if len(matched) > 0:
            st.success(f"✅ Matched {len(matched)} recommended part(s):")
            def highlight_match(row):
                return ["background-color:#d4edda"] * len(row)
            st.dataframe(matched.style.apply(highlight_match, axis=1),
                          width="stretch", hide_index=True)
        else:
            st.warning("⚠ No capacitor series matched all constraints — please adjust the parameter requirements.")
            st.dataframe(cap_db, width="stretch", hide_index=True)

        st.caption("Note: Values above are typical series parameters. Always verify against the latest manufacturer datasheet.")

        st.divider()

        # ── Manual Datasheet Entry & Verification ─────────────────────────────────────────
        st.markdown("### 📋 Manual Datasheet Entry & Selection Verification")
        st.info("""
Enter your target capacitor's parameters below; the tool will automatically compare them against the system design requirements and show pass/fail results. For each parameter, refer to the Datasheet Keyword table below.
        """)

        # ()
        with st.expander("📖 Datasheet Keyword Reference (click to expand)", expanded=False):
            st.markdown("""
| Parameter | Datasheet Keyword | Typical Location | Notes |
|------|---------------|---------|------|
| Rated Voltage | Rated voltage / Rated AC voltage / Un / VAC | First row of spec table | Note: distinguish between AC and DC ratings |
| Nominal Capacitance | Capacitance / Nominal capacitance / Cn | Spec table on first page | Unit: μF, note tolerance class |
| Rated Ripple Current | Rated ripple current / Max ripple current / Ir | Derating curve / thermal characteristics table | Unit: Arms, note temperature and frequency conditions |
| dv/dt Rating | Max dv/dt / Max du/dt / Surge dv/dt | Pulse characteristics table | Unit: V/μs — **some types do not specify this → N/A** |
| ESR | ESR / Equivalent series resistance / tan δ | Electrical characteristics table | If tan δ is given, convert: $ESR = \\tan\\delta/(2\\pi f C)$, f typically 1kHz |
| Maximum Operating Temperature | Max operating temperature / Tmax / Upper category temperature | Operating conditions table | Unit: °C |
| Rated Lifetime | Expected lifetime / Rated lifetime / L0 | Reliability / lifetime section | Unit: h, note the corresponding temperature condition |
            """)

        st.markdown("#### Capacitor Parameter Entry")
        st.caption("Parameters left at 0 will be skipped during verification — fill in as many as possible for a complete check")

        # ──  ───────────────────────────────────────────────────────
        mi_col1, mi_col2 = st.columns(2, gap="large")
        with mi_col1:
            st.markdown("**① Electrical Rated Parameters**")
            v_Un = st.number_input(
                "Rated Voltage Un(Vrms AC)",
                min_value=0.0, value=0.0,
                step=50.0, key="s10m_Un",
                help="Datasheet keyword: Rated voltage / VAC / Un"
            )
            v_Cf = st.number_input(
                "Nominal Capacitance Cn(μF)",
                min_value=0.0, value=0.0,
                step=1.0, key="s10m_Cn",
                help="Datasheet keyword: Capacitance / Cn"
            )
            v_Ir = st.number_input(
                "Rated Ripple Current Ir(Arms)",
                min_value=0.0, value=0.0,
                step=1.0, key="s10m_Ir",
                help="Datasheet keyword: Ripple current / Ir (note temperature and frequency conditions)"
            )
            v_dvdt = st.number_input(
                "dv/dt Rating (V/μs) — leave 0 if N/A",
                min_value=0.0, value=0.0,
                step=5.0, key="s10m_dvdt",
                help="Datasheet keyword: Max dv/dt / Max du/dt (some types do not specify this)"
            )

        with mi_col2:
            st.markdown("**② Thermal & Lifetime Parameters**")
            v_ESR = st.number_input(
                "ESR @ High Frequency (mΩ)",
                min_value=0.0, value=0.0,
                step=0.5, key="s10m_ESR",
                help="Datasheet keyword: ESR / tan δ (if tan δ is given, see conversion table above)"
            )
            v_Tmax = st.number_input(
                "Maximum Operating Temperature Tmax(°C)",
                min_value=0.0, value=0.0,
                step=5.0, key="s10m_Tmax",
                help="Datasheet keyword: Max operating temperature / Upper category temperature"
            )
            v_L0 = st.number_input(
                "Rated Lifetime L0(h)",
                min_value=0.0, value=0.0,
                step=10000.0, key="s10m_L0",
                help="Datasheet keyword: Expected lifetime / Rated lifetime / L0"
            )
            st.markdown("")
            st.markdown("**③ Parallel Configuration (Optional)**")
            v_npar = st.number_input(
                "Parallel Count(per phase)",
                min_value=1, value=1, step=1, key="s10m_npar",
                help="If satisfied only by parallel units, enter Parallel Count here; tool auto-calculates combined parameters"
            )

        # ──  ─────────────────────────────────────────────────────
        st.divider()
        if is_trial:
            st.button("🔍 Verify Selection", key="s10_verify_trial", type="primary", disabled=True)
            trial_lock("Unlock with paid access: one-click verification (Rated Voltage / Ripple Current / dv/dt / Lifetime auto-compare)", compact=True)
        elif st.button("🔍 Verify Selection", key="s10_verify", type="primary"):

            # Parameter
            eff_Cn   = v_Cf   * v_npar
            eff_Ir   = v_Ir   * v_npar
            eff_ESR  = v_ESR  / v_npar if v_npar > 0 else v_ESR

            st.markdown("#### 📊 Verification Results")

            if v_npar > 1:
                st.info(f"Calculated with **{v_npar} units in parallel** — combined parameters: capacitance {eff_Cn:.1f} μF, ripple current {eff_Ir:.1f} Arms, ESR {eff_ESR:.2f} mΩ")

            # : (, Datasheet Value, Design Requirement, , , )
            checks = [
                ("Rated Voltage",       v_Un,    req_Un,       "Vrms",  "≥",
                 "Rated Voltage must be at least 1.5× the actual operating voltage (67% derating), otherwise voltage overstress occurs"),
                ("Nominal Capacitance",       eff_Cn,  s_Cf,         "μF",    "≈±5%",
                 "Capacitance should match design value (±5%), as it affects LCL resonance frequency and filter performance"),
                ("Rated Ripple Current",   eff_Ir,  req_Ir,       "Arms",  "≥",
                 "Rated Ripple Current must exceed the actual ripple current (with 1.2× margin), otherwise thermal runaway"),
                ("dv/dt Rating",   v_dvdt,  dvdt_design,  "V/μs",  "≥",
                 "dv/dt rating must exceed the actual dv/dt from PWM switching, otherwise the metallized end contacts may delaminate."),
                ("ESR",            eff_ESR, 10.0,         "mΩ",    "≤",
                 "Lower ESR means less heat dissipation. MKP film capacitors are typically 1–5 mΩ; values above 10 mΩ require thermal evaluation."),
                ("Maximum Operating Temperature",   v_Tmax,  85.0,         "°C",    "≥",
                 "Rated temperature must exceed the maximum ambient installation temperature with adequate margin."),
                ("Rated Lifetime L0",   v_L0,    100000.0,     "h",     "≥",
                 "Rated Lifetime ≥ 100000h (11.4 yr); with voltage derating ×4.5, this yields >20 yr system lifetime"),
            ]

            items = []
            skip_count = 0
            for name, val, req, unit, op, reason in checks:
                if val <= 0:
                    skip_count += 1
                    continue

                if op == "≥":
                    passed = val >= req
                    margin = (val / req - 1) * 100 if req > 0 else 0
                    req_str = f"≥ {req:.1f} {unit}"
                    val_str = f"{val:.1f} {unit}"
                elif op == "≤":
                    passed = val <= req
                    margin = (req / val - 1) * 100 if val > 0 else 0
                    req_str = f"≤ {req:.1f} {unit}"
                    val_str = f"{val:.1f} {unit}"
                else:  # ≈±5%
                    passed = abs(val - req) / req <= 0.05 if req > 0 else True
                    margin = (val / req - 1) * 100 if req > 0 else 0
                    req_str = f"{req:.1f} ±5% {unit}"
                    val_str = f"{val:.1f} {unit}"

                items.append({
                    "Check Item": name,
                    "Design Requirement":  req_str,
                    "Datasheet Value":  val_str,
                    "Margin":      f"{margin:+.1f}%",
 "Result": "✅ Pass" if passed else "❌ Fail",
                    "Failure Reason": "" if passed else reason,
                })

            if not items:
                st.warning("No parameters entered for verification — please fill in at least one field and click Verify.")
            else:
                df_v = pd.DataFrame(items)

                def color_v(row):
                    color = "#d4edda" if "✅" in str(row["Result"]) else "#f8d7da"
                    return [f"background-color:{color}"] * len(row)

                # ()
                st.dataframe(
                    df_v[["Check Item","Design Requirement","Datasheet Value","Margin","Result"]]
                    .style.apply(color_v, axis=1),
                    width="stretch", hide_index=True
                )

                n_pass = sum(1 for it in items if "✅" in it["Result"])
                n_fail = len(items) - n_pass

                # Result
                if n_fail == 0:
                    st.success(f"✅ All {len(items)} checks passed! This capacitor meets all design requirements."
                               + (f" ({skip_count} items were skipped — not filled in)" if skip_count > 0 else ""))
                else:
                    st.error(f"❌ {n_fail} item(s) failed design requirements — please review and adjust capacitor or design parameters.")
                    st.markdown("**Failed Check Details**: ")
                    for it in items:
                        if "❌" in it["Result"]:
                            st.markdown(f"- **{it['Check Item']}**: {it['Failure Reason']}")

                if skip_count > 0:
                    st.caption(f"Note: {skip_count} items were left at 0 and skipped. It is recommended to complete all fields and re-verify.")

                # : ESR Thermal Dissipation Estimate
                if eff_ESR > 0 and req_Ir > 0:
                    P_loss = (req_Ir ** 2) * (eff_ESR * 1e-3)
                    st.markdown(f"""
> 📌 **Thermal Dissipation Estimate**(per phase): $P = I_{{ripple}}^2 \\times ESR = {req_Ir:.1f}^2 \\times {eff_ESR:.2f}\\text{{mΩ}} = {P_loss:.2f}\\text{{ W}}$
> — If thermal dissipation is high, consider increasing parallel count to lower equivalent ESR.
                    """)

